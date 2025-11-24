
from fpdf import FPDF
from docx import Document
from docx.shared import Pt

# Content for the document
content = {
    "title": "Etapa 1. Proceso productivo de una organización",
    "header": [
        "Nombre del Alumno: Adán Muñoz Pablo",
        "Materia: Administración de Operaciones",
        "Actividad: Etapa 1 - Proyecto"
    ],
    "sections": [
        {
            "title": "a) Información general de la organización",
            "body": """Nombre de la empresa: Muebles Rústicos "El Roble"

Breve historia:
La empresa fue fundada hace 8 años por el Sr. Juan Pérez, un carpintero de oficio que comenzó fabricando mesas de centro en el garaje de su casa. Debido a la calidad de su trabajo, la demanda creció rápidamente, permitiéndole alquilar una pequeña nave industrial. Actualmente, la empresa cuenta con 15 empleados operativos. Sin embargo, el crecimiento ha sido desordenado; la maquinaria se ha ido acomodando donde había espacio disponible sin un layout estratégico, y no existe una planificación formal de la producción, lo que ocasiona retrasos frecuentes en las entregas y acumulación de materiales.

Actividades que realiza:
La empresa se dedica a la fabricación y venta de muebles de madera de pino para el hogar, especializándose en comedores, sillas, alacenas y bases de cama con acabados rústicos y entintados."""
        },
        {
            "title": "b) Determinación del enfoque del análisis",
            "body": """Debido a que "Muebles Rústicos El Roble" fabrica una gran variedad de productos (mesas, sillas, camas, armarios), el análisis de este proyecto se centrará únicamente en un proceso específico para asegurar la profundidad requerida.

Proceso seleccionado: Fabricación de la "Silla de Comedor Modelo Clásico".

Justificación:
Este es el producto con mayor volumen de venta (se venden en juegos de 4 o 6 unidades) y es donde se presentan los mayores problemas de acumulación de material en proceso (WIP) y tiempos de espera."""
        },
        {
            "title": "c) Diagramas de Flujo de Procesos",
            "body": """Diagrama 1: Operación General de la Empresa

1. Recepción de Pedidos: El cliente solicita muebles (taller o tienda).
2. Compra de Materia Prima: Se adquiere madera, barnices y herrajes.
3. Almacenamiento: Se guarda la madera en el patio y bodegas.
4. Programación: El jefe de taller decide qué fabricar según urgencia.
5. Corte y Habilitado: La madera se corta a medidas generales.
6. Maquinado y Ensamble: Se da forma y se arman los muebles.
7. Acabado: Lijado fino, entintado y barnizado.
8. Empaquetado: Se protege el mueble para transporte.
9. Entrega: Distribución al cliente final.

Diagrama 2: Proceso Específico (Silla Modelo Clásico)

1. Recepción de madera: Llegada de tablones de pino al área de corte.
2. Trazo y Corte: Corte de patas, asiento y respaldo con plantilla.
3. Cepillado: Calibración del grosor de las piezas.
4. Lijado de piezas: Eliminación de astillas en componentes sueltos.
5. Ensamble: Unión de piezas con pegamento y pijas.
6. Resanado: Ocultamiento de orificios de tornillos con pasta.
7. Lijado final: Suavizado de la silla armada.
8. Barnizado y Secado: (CUELLO DE BOTELLA - ÁREA CRÍTICA) Aplicación de tinte y laca; tiempo prolongado de secado.
9. Inspección: Revisión visual de defectos.
10. Almacenaje: Traslado a bodega de producto terminado."""
        },
        {
            "title": "d) Desarrollo y Análisis (Enfoque Justo a Tiempo)",
            "body": """Analizando la operación de "Muebles Rústicos El Roble" bajo la filosofía Justo a Tiempo (JIT), se detecta que la empresa opera bajo un sistema tradicional "Push" (empujar). Se producen grandes lotes de sillas sin tener pedidos confirmados inmediatos, simplemente para "mantener ocupados" a los carpinteros y aprovechar la madera cortada. Esto genera un inventario excesivo de producto en proceso (WIP). Es común observar montañas de sillas armadas ("en blanco") esperando días en los pasillos antes de poder entrar al área de barnizado, lo que obstruye el paso y daña el material por golpes accidentales.

El cuello de botella principal se encuentra claramente en el área de Barnizado y Secado (marcado en el diagrama de flujo). Mientras que el proceso de corte y ensamble es relativamente rápido (aprox. 20 minutos de trabajo hombre por silla), el proceso de barnizado y secado toma hasta 4 horas debido a la falta de espacio, la dependencia del clima para el secado y la falta de equipos de extracción de aire. Esto crea un inventario de "desperdicio" (espera) acumulado antes de esta estación. Adicionalmente, el inventario de materia prima es alto (se compra madera para 3 meses) por miedo a la escasez y variaciones de precio, lo cual ocupa espacio vital y capital de trabajo, contradiciendo el principio JIT de reducir inventarios a lo mínimo necesario."""
        },
        {
            "title": "e) Propuestas de mejora",
            "body": """Tras realizar una lluvia de ideas con el equipo de trabajo para solucionar los cuellos de botella y problemas detectados, seleccionamos las siguientes 3 mejores opciones:

1. Implementación de un sistema Kanban (Tarjeta visual):
Establecer un límite de inventario máximo ("Supermercado") entre el área de Ensamble y el área de Barnizado. El área de corte y ensamble solo podrá trabajar si hay una tarjeta Kanban disponible o un espacio vacío en la zona designada. Esto detendrá la sobreproducción y evitará que se acumulen sillas en los pasillos cuando el área de barnizado esté saturada, alineando la producción a la capacidad real del cuello de botella.

2. Redistribución de la planta en Célula de Manufactura (Layout en "U"):
Actualmente, el recorrido del material es lineal y extenso. Se propone reorganizar las mesas de trabajo de la línea de sillas en forma de "U". Esto permitirá que los operarios de corte, lijado y ensamble estén físicamente más cerca, facilitando la comunicación inmediata sobre defectos de calidad, reduciendo los tiempos de traslado de material (desperdicio de transporte) y permitiendo que los operarios puedan ayudarse entre sí (polivalencia) si una estación se retrasa.

3. Optimización técnica del Cuello de Botella (Secado):
Para atacar la restricción principal en el Barnizado, se propone cambiar el insumo actual por lacas de poliuretano de secado rápido y altos sólidos, además de instalar un estante vertical de secado con ventilación forzada. Esto reducirá drásticamente el tiempo de ciclo de esta etapa, permitiendo un flujo más continuo y reduciendo el tiempo total de entrega (Lead Time)."""
        }
    ]
}

# Generate DOCX
doc = Document()
doc.add_heading(content["title"], 0)

for line in content["header"]:
    p = doc.add_paragraph(line)
    p.style = 'Normal'

doc.add_page_break()

for section in content["sections"]:
    doc.add_heading(section["title"], level=1)
    doc.add_paragraph(section["body"])

doc.save("A1_ADZ.docx")

# Generate PDF
class PDF(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 12)
        self.cell(0, 10, 'Etapa 1. Proceso productivo', 0, 1, 'C')
        self.ln(5)

    def chapter_title(self, title):
        self.set_font('Arial', 'B', 12)
        self.cell(0, 10, title, 0, 1, 'L')
        self.ln(4)

    def chapter_body(self, body):
        self.set_font('Arial', '', 11)
        self.multi_cell(0, 6, body)
        self.ln()

pdf = PDF()
pdf.add_page()

# Add header info
pdf.set_font('Arial', '', 11)
for line in content["header"]:
    pdf.cell(0, 6, line, 0, 1)
pdf.ln(10)

# Add sections
for section in content["sections"]:
    pdf.chapter_title(section["title"])
    # Replace unicode characters not supported by default font
    safe_body = section["body"].encode('latin-1', 'replace').decode('latin-1')
    pdf.chapter_body(safe_body)

pdf.output("A1_ADZ.pdf")


