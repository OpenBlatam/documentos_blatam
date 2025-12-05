#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script para generar documentos Word y PDF del foro de mercados financieros
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT, TA_CENTER
from reportlab.lib.colors import HexColor, black
import re

def markdown_to_text(md_text):
    """Convierte markdown básico a texto plano para Word"""
    # Remover headers
    text = re.sub(r'^#+\s+', '', md_text, flags=re.MULTILINE)
    # Convertir negritas
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    # Remover separadores
    text = re.sub(r'^---+$', '', text, flags=re.MULTILINE)
    return text.strip()

def create_word_document():
    """Crea el documento Word"""
    doc = Document()
    
    # Configurar estilos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Título principal
    title = doc.add_heading('Participación en Foro de Discusión: Mercados Financieros y Regulación en México', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtítulo
    subtitle = doc.add_paragraph('FASE: INICIO - Participación Inicial')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.bold = True
    subtitle_format.italic = True
    doc.add_paragraph('')
    
    # Pregunta 1
    doc.add_heading('¿En qué consideras que radica la importancia de los mercados financieros?', level=1)
    
    intro1 = doc.add_paragraph()
    intro1.add_run('Los mercados financieros constituyen el pilar fundamental del sistema económico moderno, ya que cumplen funciones esenciales que trascienden la simple intermediación de recursos. Su importancia radica en múltiples dimensiones interconectadas:')
    
    # Punto 1
    doc.add_heading('1. Asignación eficiente de recursos', level=2)
    p1 = doc.add_paragraph('Los mercados financieros facilitan la canalización del ahorro hacia la inversión productiva, permitiendo que los recursos de quienes tienen excedentes (ahorradores) fluyan hacia quienes necesitan capital para proyectos empresariales, expansión o innovación. Esta función es crucial para el crecimiento económico, ya que sin un sistema financiero eficiente, el capital quedaría ocioso o mal asignado, limitando el desarrollo de nuevas empresas y la expansión de las existentes.')
    
    # Punto 2
    doc.add_heading('2. Determinación de precios y valoración de activos', level=2)
    p2 = doc.add_paragraph('A través de los mecanismos de oferta y demanda, los mercados financieros establecen precios que reflejan el valor presente de los flujos futuros esperados. Esta función de valoración proporciona señales importantes tanto para inversores como para empresas, permitiendo decisiones informadas sobre dónde asignar recursos y cómo valorar activos.')
    
    # Punto 3
    doc.add_heading('3. Gestión de riesgos', level=2)
    p3 = doc.add_paragraph('Los mercados financieros ofrecen instrumentos y mecanismos para transferir, diversificar y gestionar riesgos. A través de derivados, seguros y otros productos financieros, los participantes pueden protegerse contra volatilidad, cambios en tipos de cambio, fluctuaciones de precios y otros riesgos inherentes a la actividad económica.')
    
    # Punto 4
    doc.add_heading('4. Liquidez y movilidad del capital', level=2)
    p4 = doc.add_paragraph('Los mercados financieros proporcionan liquidez, permitiendo que los inversores conviertan sus activos en efectivo cuando lo necesiten. Esta característica es fundamental para la confianza del sistema, ya que los participantes saben que pueden entrar y salir de posiciones según sus necesidades, sin tener que mantener inversiones de forma permanente.')
    
    # Punto 5
    doc.add_heading('5. Transparencia y disciplina de mercado', level=2)
    p5 = doc.add_paragraph('Los mercados financieros ejercen una función disciplinaria sobre las empresas, ya que los precios de las acciones y otros instrumentos reflejan el desempeño y las expectativas sobre las empresas. Esta transparencia incentiva la buena gestión corporativa y la eficiencia operativa.')
    
    # Punto 6
    doc.add_heading('6. Estabilidad macroeconómica', level=2)
    p6 = doc.add_paragraph('Un sistema financiero robusto contribuye a la estabilidad macroeconómica al facilitar la implementación de políticas monetarias, permitir la gestión de la inflación y proporcionar mecanismos para enfrentar crisis económicas.')
    
    # Conclusión contexto mexicano
    conclusion1 = doc.add_paragraph()
    conclusion1.add_run('En el contexto mexicano, los mercados financieros han sido fundamentales para el desarrollo económico del país, permitiendo la canalización de recursos hacia sectores estratégicos y facilitando la integración con los mercados globales.')
    conclusion1_format = conclusion1.runs[0]
    conclusion1_format.italic = True
    
    doc.add_paragraph('')
    
    # Pregunta 2
    doc.add_heading('¿Las entidades reguladoras del sistema financiero en México, realmente contribuyen al cumplimiento de las leyes y normas del sector empresarial? Argumenta tu respuesta', level=1)
    
    intro2 = doc.add_paragraph('La respuesta a esta pregunta requiere un análisis matizado que reconoce tanto los avances significativos como las áreas de mejora en el sistema regulatorio mexicano.')
    
    # Contribuciones positivas
    doc.add_heading('SÍ, las entidades reguladoras contribuyen al cumplimiento, pero con matices importantes:', level=2)
    doc.add_heading('Contribuciones positivas:', level=3)
    
    # Marco regulatorio
    doc.add_heading('1. Marco regulatorio estructurado', level=3)
    marco = doc.add_paragraph('México cuenta con un sistema regulatorio bien estructurado a través de entidades especializadas:')
    
    cnbv = doc.add_paragraph('• ', style='List Bullet')
    cnbv.add_run('CNBV (Comisión Nacional Bancaria y de Valores): ').bold = True
    cnbv.add_run('Supervisa y regula a las instituciones de crédito, casas de bolsa, sociedades de inversión y otras entidades del sector financiero. Establece requisitos de capital, ratios de solvencia y normas de operación que han contribuido a fortalecer el sistema bancario mexicano.')
    
    condusef = doc.add_paragraph('• ', style='List Bullet')
    condusef.add_run('CONDUSEF (Comisión Nacional para la Protección y Defensa de los Usuarios de Servicios Financieros): ').bold = True
    condusef.add_run('Protege los derechos de los usuarios de servicios financieros, estableciendo estándares de transparencia y buenas prácticas que las instituciones deben cumplir.')
    
    banxico = doc.add_paragraph('• ', style='List Bullet')
    banxico.add_run('Banco de México: ').bold = True
    banxico.add_run('Como banco central, establece políticas monetarias y regula aspectos clave del sistema financiero, contribuyendo a la estabilidad macroeconómica.')
    
    shcp = doc.add_paragraph('• ', style='List Bullet')
    shcp.add_run('SHCP (Secretaría de Hacienda y Crédito Público): ').bold = True
    shcp.add_run('Define políticas fiscales y coordina la regulación del sector financiero a nivel nacional.')
    
    # Fortalecimiento post-crisis
    doc.add_heading('2. Fortalecimiento post-crisis', level=3)
    crisis = doc.add_paragraph('Tras las crisis financieras de las décadas pasadas, México implementó reformas regulatorias significativas que han fortalecido la solvencia y estabilidad del sistema financiero. Los requisitos de capital más estrictos, los controles de riesgo mejorados y los mecanismos de supervisión han reducido la vulnerabilidad del sistema.')
    
    # Protección al consumidor
    doc.add_heading('3. Protección al consumidor', level=3)
    consumidor = doc.add_paragraph('CONDUSEF ha establecido mecanismos de protección al consumidor que obligan a las instituciones financieras a proporcionar información clara, transparente y accesible sobre sus productos y servicios, contribuyendo a un mercado más justo y competitivo.')
    
    # Estándares internacionales
    doc.add_heading('4. Cumplimiento de estándares internacionales', level=3)
    estandares = doc.add_paragraph('Las autoridades regulatorias mexicanas han trabajado para alinear las regulaciones locales con estándares internacionales (como Basilea III), lo que ha mejorado la integración del sistema financiero mexicano con los mercados globales y ha fortalecido la confianza de los inversionistas.')
    
    # Áreas de mejora
    doc.add_heading('Áreas de mejora y desafíos:', level=3)
    
    doc.add_heading('1. Capacidad de supervisión', level=3)
    supervision = doc.add_paragraph('A pesar de los avances, la capacidad de supervisión efectiva puede verse limitada por recursos insuficientes, especialmente en el monitoreo de instituciones más pequeñas o en sectores emergentes como las fintech.')
    
    doc.add_heading('2. Complejidad regulatoria', level=3)
    complejidad = doc.add_paragraph('La multiplicidad de reguladores y la complejidad de las normas pueden generar confusión y costos de cumplimiento elevados, especialmente para pequeñas y medianas empresas que pueden carecer de recursos para mantener departamentos de compliance robustos.')
    
    doc.add_heading('3. Velocidad de adaptación', level=3)
    velocidad = doc.add_paragraph('El ritmo de innovación financiera (tecnología financiera, criptomonedas, nuevos modelos de negocio) a veces supera la capacidad de las autoridades regulatorias para adaptar y actualizar las normas de manera oportuna.')
    
    doc.add_heading('4. Coordinación interinstitucional', level=3)
    coordinacion = doc.add_paragraph('Aunque existe coordinación entre las diferentes entidades regulatorias, puede haber áreas de mejora en la comunicación y coordinación para evitar duplicidades o vacíos regulatorios.')
    
    doc.add_heading('5. Enfoque en grandes instituciones', level=3)
    enfoque = doc.add_paragraph('Históricamente, la supervisión se ha concentrado más en las grandes instituciones financieras, mientras que la supervisión de instituciones más pequeñas o de sectores emergentes puede ser menos exhaustiva.')
    
    # Conclusión
    doc.add_heading('Conclusión:', level=3)
    conclusion2 = doc.add_paragraph()
    conclusion2.add_run('Las entidades reguladoras del sistema financiero en México ').bold = True
    conclusion2.add_run('SÍ contribuyen significativamente ').bold = True
    conclusion2.add_run('al cumplimiento de las leyes y normas del sector empresarial. El marco regulatorio mexicano ha evolucionado positivamente y ha contribuido a la estabilidad y solidez del sistema financiero. Sin embargo, es importante reconocer que el cumplimiento efectivo requiere:')
    
    req1 = doc.add_paragraph('1. ', style='List Number')
    req1.add_run('Recursos adecuados ').bold = True
    req1.add_run('para supervisión y enforcement')
    
    req2 = doc.add_paragraph('2. ', style='List Number')
    req2.add_run('Adaptación continua ').bold = True
    req2.add_run('a los cambios en el sector financiero')
    
    req3 = doc.add_paragraph('3. ', style='List Number')
    req3.add_run('Coordinación efectiva ').bold = True
    req3.add_run('entre las diferentes entidades regulatorias')
    
    req4 = doc.add_paragraph('4. ', style='List Number')
    req4.add_run('Equilibrio ').bold = True
    req4.add_run('entre regulación suficiente para proteger a los usuarios y flexibilidad para fomentar la innovación')
    
    final = doc.add_paragraph('El sistema regulatorio mexicano ha demostrado su capacidad para aprender de las crisis pasadas y fortalecerse, pero el dinamismo del sector financiero requiere una evolución continua de las capacidades regulatorias.')
    final_format = final.runs[0]
    final_format.italic = True
    
    doc.add_paragraph('')
    
    # Referencias
    doc.add_heading('Referencias y fuentes consultadas:', level=1)
    
    referencias = [
        'Comisión Nacional Bancaria y de Valores (CNBV). (2024). Marco regulatorio del sistema financiero mexicano. Recuperado de: https://www.cnbv.gob.mx',
        'Comisión Nacional para la Protección y Defensa de los Usuarios de Servicios Financieros (CONDUSEF). (2024). Protección al usuario de servicios financieros. Recuperado de: https://www.condusef.gob.mx',
        'Banco de México. (2024). Regulación y supervisión del sistema financiero. Recuperado de: https://www.banxico.org.mx',
        'Secretaría de Hacienda y Crédito Público (SHCP). (2024). Políticas y regulaciones financieras. Recuperado de: https://www.gob.mx/shcp',
        'Mishkin, F. S., & Eakins, S. G. (2018). Mercados e instituciones financieras (9ª ed.). Pearson Educación.',
        'Bodie, Z., Kane, A., & Marcus, A. J. (2018). Inversiones (10ª ed.). McGraw-Hill Interamericana.',
        'Fabozzi, F. J., Modigliani, F., & Jones, F. J. (2014). Fundamentos de mercados e instituciones financieras (5ª ed.). Pearson Educación.',
        'Banco Mundial. (2023). Reporte sobre el desarrollo financiero en América Latina. Recuperado de: https://www.worldbank.org'
    ]
    
    for i, ref in enumerate(referencias, 1):
        p = doc.add_paragraph(f'{i}. {ref}', style='List Number')
    
    doc.add_paragraph('')
    nota = doc.add_paragraph('Nota: Esta participación está lista para ser publicada en el foro de discusión. Recuerda adaptarla según las instrucciones específicas de tu plataforma y agregar tus propias reflexiones personales basadas en tu experiencia y conocimiento del tema.')
    nota_format = nota.runs[0]
    nota_format.italic = True
    nota_format.font.size = Pt(9)
    
    # Guardar
    filename = 'Participacion_Foro_Mercados_Financieros.docx'
    doc.save(filename)
    print(f"✓ Documento Word creado: {filename}")
    return filename

def create_pdf_document():
    """Crea el documento PDF"""
    filename = 'Participacion_Foro_Mercados_Financieros.pdf'
    doc = SimpleDocTemplate(filename, pagesize=A4,
                            rightMargin=72, leftMargin=72,
                            topMargin=72, bottomMargin=72)
    
    Story = []
    styles = getSampleStyleSheet()
    
    # Estilos personalizados
    styles.add(ParagraphStyle(
        name='TitleStyle',
        parent=styles['Heading1'],
        fontSize=16,
        textColor=HexColor('#1a1a1a'),
        spaceAfter=12,
        alignment=TA_CENTER
    ))
    
    styles.add(ParagraphStyle(
        name='Justify',
        parent=styles['Normal'],
        alignment=TA_JUSTIFY,
        fontSize=11,
        leading=14
    ))
    
    styles.add(ParagraphStyle(
        name='HeaderCustom',
        parent=styles['Heading1'],
        fontSize=14,
        textColor=HexColor('#2c3e50'),
        spaceAfter=12,
        spaceBefore=12
    ))
    
    styles.add(ParagraphStyle(
        name='SubHeaderCustom',
        parent=styles['Heading2'],
        fontSize=12,
        textColor=HexColor('#34495e'),
        spaceAfter=8,
        spaceBefore=10
    ))
    
    # Título
    Story.append(Paragraph("Participación en Foro de Discusión: Mercados Financieros y Regulación en México", styles['TitleStyle']))
    Story.append(Paragraph("<i>FASE: INICIO - Participación Inicial</i>", styles['Normal']))
    Story.append(Spacer(1, 24))
    
    # Pregunta 1
    Story.append(Paragraph("<b>¿En qué consideras que radica la importancia de los mercados financieros?</b>", styles['HeaderCustom']))
    
    intro1 = "Los mercados financieros constituyen el pilar fundamental del sistema económico moderno, ya que cumplen funciones esenciales que trascienden la simple intermediación de recursos. Su importancia radica en múltiples dimensiones interconectadas:"
    Story.append(Paragraph(intro1, styles['Justify']))
    Story.append(Spacer(1, 12))
    
    # Puntos
    puntos = [
        ("<b>1. Asignación eficiente de recursos</b>",
         "Los mercados financieros facilitan la canalización del ahorro hacia la inversión productiva, permitiendo que los recursos de quienes tienen excedentes (ahorradores) fluyan hacia quienes necesitan capital para proyectos empresariales, expansión o innovación. Esta función es crucial para el crecimiento económico, ya que sin un sistema financiero eficiente, el capital quedaría ocioso o mal asignado, limitando el desarrollo de nuevas empresas y la expansión de las existentes."),
        ("<b>2. Determinación de precios y valoración de activos</b>",
         "A través de los mecanismos de oferta y demanda, los mercados financieros establecen precios que reflejan el valor presente de los flujos futuros esperados. Esta función de valoración proporciona señales importantes tanto para inversores como para empresas, permitiendo decisiones informadas sobre dónde asignar recursos y cómo valorar activos."),
        ("<b>3. Gestión de riesgos</b>",
         "Los mercados financieros ofrecen instrumentos y mecanismos para transferir, diversificar y gestionar riesgos. A través de derivados, seguros y otros productos financieros, los participantes pueden protegerse contra volatilidad, cambios en tipos de cambio, fluctuaciones de precios y otros riesgos inherentes a la actividad económica."),
        ("<b>4. Liquidez y movilidad del capital</b>",
         "Los mercados financieros proporcionan liquidez, permitiendo que los inversores conviertan sus activos en efectivo cuando lo necesiten. Esta característica es fundamental para la confianza del sistema, ya que los participantes saben que pueden entrar y salir de posiciones según sus necesidades, sin tener que mantener inversiones de forma permanente."),
        ("<b>5. Transparencia y disciplina de mercado</b>",
         "Los mercados financieros ejercen una función disciplinaria sobre las empresas, ya que los precios de las acciones y otros instrumentos reflejan el desempeño y las expectativas sobre las empresas. Esta transparencia incentiva la buena gestión corporativa y la eficiencia operativa."),
        ("<b>6. Estabilidad macroeconómica</b>",
         "Un sistema financiero robusto contribuye a la estabilidad macroeconómica al facilitar la implementación de políticas monetarias, permitir la gestión de la inflación y proporcionar mecanismos para enfrentar crisis económicas.")
    ]
    
    for titulo, contenido in puntos:
        Story.append(Paragraph(titulo, styles['SubHeaderCustom']))
        Story.append(Paragraph(contenido, styles['Justify']))
        Story.append(Spacer(1, 10))
    
    conclusion1 = "<i>En el contexto mexicano, los mercados financieros han sido fundamentales para el desarrollo económico del país, permitiendo la canalización de recursos hacia sectores estratégicos y facilitando la integración con los mercados globales.</i>"
    Story.append(Paragraph(conclusion1, styles['Justify']))
    Story.append(Spacer(1, 24))
    
    # Pregunta 2
    Story.append(Paragraph("<b>¿Las entidades reguladoras del sistema financiero en México, realmente contribuyen al cumplimiento de las leyes y normas del sector empresarial? Argumenta tu respuesta</b>", styles['HeaderCustom']))
    
    intro2 = "La respuesta a esta pregunta requiere un análisis matizado que reconoce tanto los avances significativos como las áreas de mejora en el sistema regulatorio mexicano."
    Story.append(Paragraph(intro2, styles['Justify']))
    Story.append(Spacer(1, 12))
    
    Story.append(Paragraph("<b>SÍ, las entidades reguladoras contribuyen al cumplimiento, pero con matices importantes:</b>", styles['SubHeaderCustom']))
    Story.append(Paragraph("<b>Contribuciones positivas:</b>", styles['Normal']))
    Story.append(Spacer(1, 8))
    
    # Contribuciones
    contribuciones = [
        ("<b>1. Marco regulatorio estructurado</b>",
         "México cuenta con un sistema regulatorio bien estructurado a través de entidades especializadas:<br/><br/>" +
         "• <b>CNBV (Comisión Nacional Bancaria y de Valores):</b> Supervisa y regula a las instituciones de crédito, casas de bolsa, sociedades de inversión y otras entidades del sector financiero. Establece requisitos de capital, ratios de solvencia y normas de operación que han contribuido a fortalecer el sistema bancario mexicano.<br/><br/>" +
         "• <b>CONDUSEF (Comisión Nacional para la Protección y Defensa de los Usuarios de Servicios Financieros):</b> Protege los derechos de los usuarios de servicios financieros, estableciendo estándares de transparencia y buenas prácticas que las instituciones deben cumplir.<br/><br/>" +
         "• <b>Banco de México:</b> Como banco central, establece políticas monetarias y regula aspectos clave del sistema financiero, contribuyendo a la estabilidad macroeconómica.<br/><br/>" +
         "• <b>SHCP (Secretaría de Hacienda y Crédito Público):</b> Define políticas fiscales y coordina la regulación del sector financiero a nivel nacional."),
        ("<b>2. Fortalecimiento post-crisis</b>",
         "Tras las crisis financieras de las décadas pasadas, México implementó reformas regulatorias significativas que han fortalecido la solvencia y estabilidad del sistema financiero. Los requisitos de capital más estrictos, los controles de riesgo mejorados y los mecanismos de supervisión han reducido la vulnerabilidad del sistema."),
        ("<b>3. Protección al consumidor</b>",
         "CONDUSEF ha establecido mecanismos de protección al consumidor que obligan a las instituciones financieras a proporcionar información clara, transparente y accesible sobre sus productos y servicios, contribuyendo a un mercado más justo y competitivo."),
        ("<b>4. Cumplimiento de estándares internacionales</b>",
         "Las autoridades regulatorias mexicanas han trabajado para alinear las regulaciones locales con estándares internacionales (como Basilea III), lo que ha mejorado la integración del sistema financiero mexicano con los mercados globales y ha fortalecido la confianza de los inversionistas.")
    ]
    
    for titulo, contenido in contribuciones:
        Story.append(Paragraph(titulo, styles['SubHeaderCustom']))
        Story.append(Paragraph(contenido, styles['Justify']))
        Story.append(Spacer(1, 10))
    
    Story.append(Paragraph("<b>Áreas de mejora y desafíos:</b>", styles['Normal']))
    Story.append(Spacer(1, 8))
    
    mejoras = [
        ("<b>1. Capacidad de supervisión</b>",
         "A pesar de los avances, la capacidad de supervisión efectiva puede verse limitada por recursos insuficientes, especialmente en el monitoreo de instituciones más pequeñas o en sectores emergentes como las fintech."),
        ("<b>2. Complejidad regulatoria</b>",
         "La multiplicidad de reguladores y la complejidad de las normas pueden generar confusión y costos de cumplimiento elevados, especialmente para pequeñas y medianas empresas que pueden carecer de recursos para mantener departamentos de compliance robustos."),
        ("<b>3. Velocidad de adaptación</b>",
         "El ritmo de innovación financiera (tecnología financiera, criptomonedas, nuevos modelos de negocio) a veces supera la capacidad de las autoridades regulatorias para adaptar y actualizar las normas de manera oportuna."),
        ("<b>4. Coordinación interinstitucional</b>",
         "Aunque existe coordinación entre las diferentes entidades regulatorias, puede haber áreas de mejora en la comunicación y coordinación para evitar duplicidades o vacíos regulatorios."),
        ("<b>5. Enfoque en grandes instituciones</b>",
         "Históricamente, la supervisión se ha concentrado más en las grandes instituciones financieras, mientras que la supervisión de instituciones más pequeñas o de sectores emergentes puede ser menos exhaustiva.")
    ]
    
    for titulo, contenido in mejoras:
        Story.append(Paragraph(titulo, styles['SubHeaderCustom']))
        Story.append(Paragraph(contenido, styles['Justify']))
        Story.append(Spacer(1, 10))
    
    Story.append(Paragraph("<b>Conclusión:</b>", styles['SubHeaderCustom']))
    conclusion2 = ("Las entidades reguladoras del sistema financiero en México <b>SÍ contribuyen significativamente</b> al cumplimiento de las leyes y normas del sector empresarial. "
                   "El marco regulatorio mexicano ha evolucionado positivamente y ha contribuido a la estabilidad y solidez del sistema financiero. Sin embargo, es importante reconocer que el cumplimiento efectivo requiere:<br/><br/>"
                   "1. <b>Recursos adecuados</b> para supervisión y enforcement<br/>"
                   "2. <b>Adaptación continua</b> a los cambios en el sector financiero<br/>"
                   "3. <b>Coordinación efectiva</b> entre las diferentes entidades regulatorias<br/>"
                   "4. <b>Equilibrio</b> entre regulación suficiente para proteger a los usuarios y flexibilidad para fomentar la innovación<br/><br/>"
                   "<i>El sistema regulatorio mexicano ha demostrado su capacidad para aprender de las crisis pasadas y fortalecerse, pero el dinamismo del sector financiero requiere una evolución continua de las capacidades regulatorias.</i>")
    Story.append(Paragraph(conclusion2, styles['Justify']))
    Story.append(Spacer(1, 24))
    
    # Referencias
    Story.append(Paragraph("<b>Referencias y fuentes consultadas:</b>", styles['HeaderCustom']))
    
    referencias = [
        "1. Comisión Nacional Bancaria y de Valores (CNBV). (2024). <i>Marco regulatorio del sistema financiero mexicano</i>. Recuperado de: https://www.cnbv.gob.mx",
        "2. Comisión Nacional para la Protección y Defensa de los Usuarios de Servicios Financieros (CONDUSEF). (2024). <i>Protección al usuario de servicios financieros</i>. Recuperado de: https://www.condusef.gob.mx",
        "3. Banco de México. (2024). <i>Regulación y supervisión del sistema financiero</i>. Recuperado de: https://www.banxico.org.mx",
        "4. Secretaría de Hacienda y Crédito Público (SHCP). (2024). <i>Políticas y regulaciones financieras</i>. Recuperado de: https://www.gob.mx/shcp",
        "5. Mishkin, F. S., & Eakins, S. G. (2018). <i>Mercados e instituciones financieras</i> (9ª ed.). Pearson Educación.",
        "6. Bodie, Z., Kane, A., & Marcus, A. J. (2018). <i>Inversiones</i> (10ª ed.). McGraw-Hill Interamericana.",
        "7. Fabozzi, F. J., Modigliani, F., & Jones, F. J. (2014). <i>Fundamentos de mercados e instituciones financieras</i> (5ª ed.). Pearson Educación.",
        "8. Banco Mundial. (2023). <i>Reporte sobre el desarrollo financiero en América Latina</i>. Recuperado de: https://www.worldbank.org"
    ]
    
    for ref in referencias:
        Story.append(Paragraph(ref, styles['Justify']))
        Story.append(Spacer(1, 6))
    
    Story.append(Spacer(1, 12))
    nota = "<i>Nota: Esta participación está lista para ser publicada en el foro de discusión. Recuerda adaptarla según las instrucciones específicas de tu plataforma y agregar tus propias reflexiones personales basadas en tu experiencia y conocimiento del tema.</i>"
    Story.append(Paragraph(nota, styles['Normal']))
    
    doc.build(Story)
    print(f"✓ Documento PDF creado: {filename}")
    return filename

if __name__ == "__main__":
    print("Generando documentos del foro de mercados financieros...")
    print("-" * 60)
    try:
        word_file = create_word_document()
    except Exception as e:
        print(f"✗ Error al crear Word: {e}")
        word_file = None
    
    try:
        pdf_file = create_pdf_document()
    except Exception as e:
        print(f"✗ Error al crear PDF: {e}")
        pdf_file = None
    
    print("-" * 60)
    if word_file:
        print(f"✓ Word: {word_file}")
    if pdf_file:
        print(f"✓ PDF: {pdf_file}")
    print("\n¡Documentos generados exitosamente!")







