#!/usr/bin/env python3
"""
Script completo para convertir los documentos más importantes a PDF, Word y Excel
con gráficas de alta calidad y formato profesional.

Requisitos:
    pip install python-docx openpyxl matplotlib seaborn markdown reportlab pandas pillow
"""

import os
import re
import json
import statistics
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Tuple, Any, Optional
import markdown
from markdown.extensions import codehilite, tables, fenced_code

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import openpyxl
    from openpyxl import Workbook
    from openpyxl.chart import BarChart, LineChart, PieChart, Reference, ScatterChart
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    from openpyxl.drawing.image import Image as XLImage
    import matplotlib.pyplot as plt
    import matplotlib
    matplotlib.use('Agg')
    import numpy as np
    import seaborn as sns
    sns.set_style("whitegrid")
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image as RLImage
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    from PIL import Image as PILImage
    import pandas as pd
except ImportError as e:
    print(f"❌ Error: Faltan dependencias. Instala con:")
    print(f"pip install python-docx openpyxl matplotlib seaborn markdown reportlab pandas pillow")
    print(f"Error específico: {e}")
    exit(1)


class DocumentConverter:
    """Convierte documentos Markdown a PDF, Word y Excel con gráficas profesionales."""
    
    def __init__(self, output_dir: str = "exports_docs_importantes"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        self.stats = {}
        self.charts_dir = self.output_dir / "charts"
        self.charts_dir.mkdir(exist_ok=True)
        
    def parse_markdown(self, file_path: str) -> Dict[str, Any]:
        """Parsea un archivo Markdown y extrae información estructurada."""
        print(f"  📖 Parseando: {Path(file_path).name}")
        
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Parsear con markdown
        md = markdown.Markdown(extensions=['codehilite', 'tables', 'fenced_code', 'toc'])
        html = md.convert(content)
        
        # Extraer estadísticas detalladas
        lines = content.split('\n')
        stats = {
            'total_lines': len(lines),
            'total_words': len(content.split()),
            'total_chars': len(content),
            'headers': [],
            'code_blocks': [],
            'tables': [],
            'links': [],
            'images': [],
            'sections': [],
            'lists': [],
            'emphasis': []
        }
        
        # Extraer headers con contexto
        for i, line in enumerate(lines):
            if line.strip().startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                text = line.lstrip('#').strip()
                stats['headers'].append({
                    'level': level,
                    'text': text,
                    'line': i + 1
                })
        
        # Extraer código
        code_pattern = r'```[\s\S]*?```'
        code_blocks = re.findall(code_pattern, content)
        stats['code_blocks'] = [cb[:100] for cb in code_blocks]  # Primeros 100 chars
        
        # Extraer tablas
        table_pattern = r'\|.*\|'
        tables = re.findall(table_pattern, content, re.MULTILINE)
        stats['tables'] = tables[:50]  # Limitar
        
        # Extraer links
        link_pattern = r'\[([^\]]+)\]\(([^\)]+)\)'
        links = re.findall(link_pattern, content)
        stats['links'] = links[:100]  # Limitar
        
        # Extraer imágenes
        img_pattern = r'!\[([^\]]*)\]\(([^\)]+)\)'
        images = re.findall(img_pattern, content)
        stats['images'] = images
        
        # Identificar secciones principales
        current_section = None
        for i, line in enumerate(lines):
            if line.strip().startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                if level <= 2:  # Solo secciones principales
                    current_section = line.lstrip('#').strip()
                    stats['sections'].append({
                        'title': current_section,
                        'level': level,
                        'line': i + 1,
                        'word_count': 0  # Se calculará después
                    })
        
        # Calcular palabras por sección
        for i, section in enumerate(stats['sections']):
            start_line = section['line'] - 1
            end_line = stats['sections'][i + 1]['line'] - 1 if i + 1 < len(stats['sections']) else len(lines)
            section_text = ' '.join(lines[start_line:end_line])
            section['word_count'] = len(section_text.split())
        
        return {
            'content': content,
            'html': html,
            'stats': stats,
            'filename': Path(file_path).stem,
            'filepath': file_path
        }
    
    def create_visualizations(self, parsed_data: Dict[str, Any]) -> List[str]:
        """Crea visualizaciones profesionales y las guarda como imágenes."""
        stats = parsed_data['stats']
        filename = parsed_data['filename']
        chart_paths = []
        
        # Configurar estilo mejorado
        plt.style.use('seaborn-v0_8-darkgrid')
        sns.set_palette("husl")
        
        # Gráfica 1: Dashboard completo (2x2) - MEJORADO
        fig, axes = plt.subplots(2, 2, figsize=(16, 14))
        fig.patch.set_facecolor('white')
        fig.suptitle(f'📊 Análisis Completo: {filename}', fontsize=20, fontweight='bold', y=0.995, color='#003366')
        
        # Subplot 1: Distribución de niveles de encabezados - MEJORADO
        if stats['headers']:
            level_counts = {}
            for h in stats['headers']:
                level_counts[h['level']] = level_counts.get(h['level'], 0) + 1
            
            if level_counts:
                levels = sorted(level_counts.keys())
                counts = [level_counts[l] for l in levels]
                max_level = max(levels) if levels else 1
                colors_viridis = plt.cm.viridis([l/max_level if max_level > 0 else 0 for l in levels])
                bars = axes[0, 0].bar([f'Nivel {l}' for l in levels], counts, 
                                     color=colors_viridis, edgecolor='black', linewidth=1.2, alpha=0.85)
                axes[0, 0].set_title('📊 Distribución de Niveles de Encabezados', 
                                    fontsize=13, fontweight='bold', pad=12)
                axes[0, 0].set_xlabel('Nivel', fontsize=11, fontweight='bold')
                axes[0, 0].set_ylabel('Cantidad', fontsize=11, fontweight='bold')
                axes[0, 0].grid(True, alpha=0.3, axis='y', linestyle='--', linewidth=0.8)
                axes[0, 0].set_axisbelow(True)
                # Agregar valores en las barras con mejor formato
                for bar in bars:
                    height = bar.get_height()
                    axes[0, 0].text(bar.get_x() + bar.get_width()/2., height,
                                   f'{int(height)}', ha='center', va='bottom', 
                                   fontsize=10, fontweight='bold',
                                   bbox=dict(boxstyle='round,pad=0.3', facecolor='white', alpha=0.8))
        
        # Subplot 2: Métricas principales (horizontal bar) - MEJORADO
        metrics = ['Líneas', 'Palabras', 'Encabezados', 'Enlaces', 'Código']
        values = [
            stats['total_lines'],
            stats['total_words'],
            len(stats['headers']),
            len(stats['links']),
            len(stats['code_blocks'])
        ]
        
        # Usar escala logarítmica si hay mucha diferencia
        max_val = max(values) if values else 1
        if max_val > 1000:
            # Normalizar pero mostrar valores reales
            normalized = [v/max_val * 100 for v in values]
            bars = axes[0, 1].barh(metrics, normalized, color=plt.cm.viridis([i/len(metrics) for i in range(len(metrics))]))
            axes[0, 1].set_xlabel('Valor Normalizado (%)', fontsize=10)
        else:
            bars = axes[0, 1].barh(metrics, values, color=plt.cm.viridis([i/len(metrics) for i in range(len(metrics))]))
            axes[0, 1].set_xlabel('Cantidad', fontsize=10)
        
        axes[0, 1].set_title('📈 Métricas Principales', fontsize=13, fontweight='bold', pad=15)
        axes[0, 1].grid(True, alpha=0.3, axis='x', linestyle='--')
        # Agregar valores con mejor formato
        for i, (bar, val) in enumerate(zip(bars, values)):
            x_pos = bar.get_width() if max_val <= 1000 else bar.get_width()
            axes[0, 1].text(x_pos, bar.get_y() + bar.get_height()/2,
                           f'{val:,}', ha='left', va='center', fontsize=10, fontweight='bold',
                           bbox=dict(boxstyle='round,pad=0.3', facecolor='white', alpha=0.7))
        
        # Subplot 3: Tipos de contenido (pie chart mejorado) - MEJORADO
        content_types = ['Código', 'Tablas', 'Enlaces', 'Imágenes']
        content_counts = [
            len(stats['code_blocks']),
            len(stats['tables']),
            len(stats['links']),
            len(stats['images'])
        ]
        
        # Filtrar valores cero
        filtered_types = []
        filtered_counts = []
        for ct, cc in zip(content_types, content_counts):
            if cc > 0:
                filtered_types.append(ct)
                filtered_counts.append(cc)
        
        if filtered_counts:
            colors_pie = plt.cm.Set3(range(len(filtered_types)))
            wedges, texts, autotexts = axes[1, 0].pie(filtered_counts, labels=filtered_types, 
                                                     autopct='%1.1f%%', startangle=90,
                                                     colors=colors_pie, textprops={'fontsize': 11, 'fontweight': 'bold'},
                                                     explode=[0.05] * len(filtered_types),
                                                     shadow=True, wedgeprops={'edgecolor': 'black', 'linewidth': 1.5})
            axes[1, 0].set_title('📦 Distribución de Tipos de Contenido', 
                               fontsize=13, fontweight='bold', pad=15)
            for autotext in autotexts:
                autotext.set_color('black')
                autotext.set_fontweight('bold')
                autotext.set_fontsize(11)
        else:
            axes[1, 0].text(0.5, 0.5, 'Sin contenido adicional', ha='center', va='center',
                           transform=axes[1, 0].transAxes, fontsize=12, style='italic', color='gray')
            axes[1, 0].set_title('📦 Distribución de Tipos de Contenido', 
                               fontsize=13, fontweight='bold')
        
        # Subplot 4: Palabras por sección (top 10) - MEJORADO
        if stats['sections']:
            sections_sorted = sorted(stats['sections'], key=lambda x: x['word_count'], reverse=True)[:10]
            section_names = [s['title'][:30] + '...' if len(s['title']) > 30 else s['title'] 
                           for s in sections_sorted]
            section_sizes = [s['word_count'] for s in sections_sorted]
            
            if section_sizes:
                colors_coolwarm = plt.cm.coolwarm([i/len(section_names) for i in range(len(section_names))])
                bars = axes[1, 1].barh(range(len(section_names)), section_sizes, 
                                      color=colors_coolwarm, edgecolor='black', linewidth=0.8, alpha=0.85)
                axes[1, 1].set_yticks(range(len(section_names)))
                axes[1, 1].set_yticklabels(section_names, fontsize=9)
                axes[1, 1].set_title('📝 Top 10 Secciones por Palabras', 
                                   fontsize=13, fontweight='bold', pad=12)
                axes[1, 1].set_xlabel('Palabras', fontsize=11, fontweight='bold')
                axes[1, 1].grid(True, alpha=0.3, axis='x', linestyle='--', linewidth=0.8)
                axes[1, 1].set_axisbelow(True)
                # Agregar valores con mejor formato
                for bar, val in zip(bars, section_sizes):
                    axes[1, 1].text(bar.get_width(), bar.get_y() + bar.get_height()/2,
                                   f'{val:,}', ha='left', va='center', fontsize=9, fontweight='bold',
                                   bbox=dict(boxstyle='round,pad=0.3', facecolor='white', alpha=0.7))
        else:
            axes[1, 1].text(0.5, 0.5, 'Sin secciones disponibles', ha='center', va='center',
                           transform=axes[1, 1].transAxes, fontsize=12, style='italic', color='gray')
            axes[1, 1].set_title('📝 Top 10 Secciones por Palabras', fontsize=13, fontweight='bold')
        
        plt.tight_layout(rect=[0, 0, 1, 0.98])
        chart_path1 = self.charts_dir / f"{filename}_dashboard.png"
        plt.savefig(chart_path1, dpi=300, bbox_inches='tight', facecolor='white', 
                   edgecolor='none', pad_inches=0.2)
        plt.close()
        chart_paths.append(str(chart_path1))
        print(f"    ✅ Dashboard creado: {chart_path1.name}")
        
        # Gráfica 2: Análisis de estructura (si hay suficientes headers) - MEJORADO
        if len(stats['headers']) > 5:
            fig, ax = plt.subplots(figsize=(14, 9))
            fig.patch.set_facecolor('white')
            
            # Crear árbol de estructura
            levels = {}
            for h in stats['headers']:
                if h['level'] not in levels:
                    levels[h['level']] = []
                levels[h['level']].append(h)
            
            # Gráfico de barras mejorado por nivel
            level_data = {}
            for level in sorted(levels.keys()):
                level_data[f'Nivel {level}'] = len(levels[level])
            
            if level_data:
                colors_map = plt.cm.plasma([i/len(level_data) for i in range(len(level_data))])
                bars = ax.bar(level_data.keys(), level_data.values(), 
                             color=colors_map, edgecolor='black', linewidth=1.5, alpha=0.8)
                ax.set_title(f'📑 Estructura de Encabezados: {filename}', fontsize=16, fontweight='bold', 
                           color='#003366', pad=20)
                ax.set_xlabel('Nivel de Encabezado', fontsize=12, fontweight='bold')
                ax.set_ylabel('Cantidad', fontsize=12, fontweight='bold')
                ax.grid(True, alpha=0.3, axis='y', linestyle='--', linewidth=0.8)
                ax.set_axisbelow(True)
                
                # Agregar valores con mejor formato
                for bar in bars:
                    height = bar.get_height()
                    ax.text(bar.get_x() + bar.get_width()/2., height,
                           f'{int(height)}', ha='center', va='bottom', fontsize=11, fontweight='bold',
                           bbox=dict(boxstyle='round,pad=0.4', facecolor='white', alpha=0.8, edgecolor='gray'))
            
            plt.tight_layout()
            chart_path2 = self.charts_dir / f"{filename}_estructura.png"
            plt.savefig(chart_path2, dpi=300, bbox_inches='tight', facecolor='white', 
                       edgecolor='none', pad_inches=0.2)
            plt.close()
            chart_paths.append(str(chart_path2))
            print(f"    ✅ Gráfica de estructura creada: {chart_path2.name}")
        
        # Gráfica 3: Análisis de complejidad (NUEVA)
        if stats['sections'] and len(stats['sections']) > 3:
            fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(16, 7))
            fig.patch.set_facecolor('white')
            fig.suptitle(f'📊 Análisis de Complejidad: {filename}', fontsize=16, fontweight='bold', 
                        color='#003366', y=1.02)
            
            # Subplot 1: Distribución de palabras por sección
            sections_sorted = sorted(stats['sections'], key=lambda x: x['word_count'], reverse=True)[:15]
            section_names = [s['title'][:25] + '...' if len(s['title']) > 25 else s['title'] 
                           for s in sections_sorted]
            section_sizes = [s['word_count'] for s in sections_sorted]
            
            bars = ax1.barh(range(len(section_names)), section_sizes, 
                           color=plt.cm.coolwarm([i/len(section_names) for i in range(len(section_names))]),
                           edgecolor='black', linewidth=0.5)
            ax1.set_yticks(range(len(section_names)))
            ax1.set_yticklabels(section_names, fontsize=9)
            ax1.set_title('Top 15 Secciones por Palabras', fontsize=12, fontweight='bold')
            ax1.set_xlabel('Palabras', fontsize=11, fontweight='bold')
            ax1.grid(True, alpha=0.3, axis='x', linestyle='--')
            ax1.set_axisbelow(True)
            
            for bar, val in zip(bars, section_sizes):
                ax1.text(bar.get_width(), bar.get_y() + bar.get_height()/2,
                       f'{val:,}', ha='left', va='center', fontsize=9, fontweight='bold')
            
            # Subplot 2: Relación nivel vs palabras
            level_word_data = {}
            for section in stats['sections']:
                level = section['level']
                if level not in level_word_data:
                    level_word_data[level] = []
                level_word_data[level].append(section['word_count'])
            
            levels_avg = {level: sum(words)/len(words) for level, words in level_word_data.items()}
            levels_sorted = sorted(levels_avg.keys())
            
            bars2 = ax2.bar([f'Nivel {l}' for l in levels_sorted], 
                           [levels_avg[l] for l in levels_sorted],
                           color=plt.cm.viridis([l/max(levels_sorted) if max(levels_sorted) > 0 else 0 
                                               for l in levels_sorted]),
                           edgecolor='black', linewidth=1)
            ax2.set_title('Promedio de Palabras por Nivel', fontsize=12, fontweight='bold')
            ax2.set_xlabel('Nivel', fontsize=11, fontweight='bold')
            ax2.set_ylabel('Palabras Promedio', fontsize=11, fontweight='bold')
            ax2.grid(True, alpha=0.3, axis='y', linestyle='--')
            ax2.set_axisbelow(True)
            
            for bar in bars2:
                height = bar.get_height()
                ax2.text(bar.get_x() + bar.get_width()/2., height,
                       f'{int(height):,}', ha='center', va='bottom', fontsize=10, fontweight='bold')
            
            plt.tight_layout()
            chart_path3 = self.charts_dir / f"{filename}_complejidad.png"
            plt.savefig(chart_path3, dpi=300, bbox_inches='tight', facecolor='white', 
                       edgecolor='none', pad_inches=0.2)
            plt.close()
            chart_paths.append(str(chart_path3))
            print(f"    ✅ Gráfica de complejidad creada: {chart_path3.name}")
        
        # Gráfica 4: Análisis temporal y distribución (NUEVA)
        if stats['headers'] and len(stats['headers']) > 10:
            fig, axes = plt.subplots(1, 2, figsize=(16, 7))
            fig.patch.set_facecolor('white')
            fig.suptitle(f'📈 Análisis Temporal y Distribución: {filename}', fontsize=16, 
                        fontweight='bold', color='#003366', y=1.02)
            
            # Subplot 1: Distribución de headers por posición en documento
            header_positions = [h['line'] for h in stats['headers']]
            if header_positions:
                axes[0].hist(header_positions, bins=min(20, len(header_positions)//2), 
                            color=plt.cm.viridis(0.5), edgecolor='black', alpha=0.7)
                axes[0].set_title('Distribución de Encabezados en el Documento', 
                                 fontsize=12, fontweight='bold')
                axes[0].set_xlabel('Línea en el Documento', fontsize=11, fontweight='bold')
                axes[0].set_ylabel('Frecuencia', fontsize=11, fontweight='bold')
                axes[0].grid(True, alpha=0.3, axis='y', linestyle='--')
                axes[0].set_axisbelow(True)
            
            # Subplot 2: Relación nivel vs posición (scatter)
            if len(stats['headers']) > 5:
                levels = [h['level'] for h in stats['headers']]
                positions = [h['line'] for h in stats['headers']]
                scatter = axes[1].scatter(positions, levels, c=levels, cmap='plasma', 
                                         s=100, alpha=0.6, edgecolors='black', linewidth=0.5)
                axes[1].set_title('Relación Nivel vs Posición', fontsize=12, fontweight='bold')
                axes[1].set_xlabel('Posición en Documento (línea)', fontsize=11, fontweight='bold')
                axes[1].set_ylabel('Nivel de Encabezado', fontsize=11, fontweight='bold')
                axes[1].grid(True, alpha=0.3, linestyle='--')
                axes[1].set_axisbelow(True)
                plt.colorbar(scatter, ax=axes[1], label='Nivel')
            
            plt.tight_layout()
            chart_path4 = self.charts_dir / f"{filename}_temporal.png"
            plt.savefig(chart_path4, dpi=300, bbox_inches='tight', facecolor='white', 
                       edgecolor='none', pad_inches=0.2)
            plt.close()
            chart_paths.append(str(chart_path4))
            print(f"    ✅ Gráfica temporal creada: {chart_path4.name}")
        
        # Gráfica 5: Análisis de contenido (NUEVA)
        if stats['code_blocks'] or stats['tables'] or stats['links']:
            fig, axes = plt.subplots(1, 3, figsize=(18, 6))
            fig.patch.set_facecolor('white')
            fig.suptitle(f'📊 Análisis Detallado de Contenido: {filename}', fontsize=16, 
                        fontweight='bold', color='#003366', y=1.02)
            
            # Subplot 1: Longitud de bloques de código
            if stats['code_blocks']:
                code_lengths = [len(cb) for cb in stats['code_blocks'][:50]]
                axes[0].hist(code_lengths, bins=15, color=plt.cm.Set2(0), edgecolor='black', alpha=0.7)
                axes[0].set_title('Distribución de Longitud de Bloques de Código', 
                                fontsize=11, fontweight='bold')
                axes[0].set_xlabel('Longitud (caracteres)', fontsize=10)
                axes[0].set_ylabel('Frecuencia', fontsize=10)
                axes[0].grid(True, alpha=0.3, axis='y', linestyle='--')
            else:
                axes[0].text(0.5, 0.5, 'Sin bloques de código', ha='center', va='center',
                           transform=axes[0].transAxes, fontsize=12, style='italic')
                axes[0].set_title('Bloques de Código', fontsize=11, fontweight='bold')
            
            # Subplot 2: Distribución de enlaces
            if stats['links']:
                link_text_lengths = [len(link[0]) for link in stats['links'][:100]]
                axes[1].hist(link_text_lengths, bins=15, color=plt.cm.Set2(1), edgecolor='black', alpha=0.7)
                axes[1].set_title('Distribución de Longitud de Texto de Enlaces', 
                                fontsize=11, fontweight='bold')
                axes[1].set_xlabel('Longitud del texto (caracteres)', fontsize=10)
                axes[1].set_ylabel('Frecuencia', fontsize=10)
                axes[1].grid(True, alpha=0.3, axis='y', linestyle='--')
            else:
                axes[1].text(0.5, 0.5, 'Sin enlaces', ha='center', va='center',
                           transform=axes[1].transAxes, fontsize=12, style='italic')
                axes[1].set_title('Enlaces', fontsize=11, fontweight='bold')
            
            # Subplot 3: Resumen de tipos de contenido
            content_summary = {
                'Código': len(stats['code_blocks']),
                'Tablas': len(stats['tables']),
                'Enlaces': len(stats['links']),
                'Imágenes': len(stats['images']),
                'Secciones': len(stats['sections'])
            }
            filtered_summary = {k: v for k, v in content_summary.items() if v > 0}
            if filtered_summary:
                bars = axes[2].bar(filtered_summary.keys(), filtered_summary.values(),
                                  color=plt.cm.Pastel1(range(len(filtered_summary))),
                                  edgecolor='black', linewidth=1)
                axes[2].set_title('Resumen de Tipos de Contenido', fontsize=11, fontweight='bold')
                axes[2].set_ylabel('Cantidad', fontsize=10)
                axes[2].tick_params(axis='x', rotation=45)
                axes[2].grid(True, alpha=0.3, axis='y', linestyle='--')
                for bar in bars:
                    height = bar.get_height()
                    axes[2].text(bar.get_x() + bar.get_width()/2., height,
                               f'{int(height)}', ha='center', va='bottom', fontsize=9, fontweight='bold')
            else:
                axes[2].text(0.5, 0.5, 'Sin contenido adicional', ha='center', va='center',
                           transform=axes[2].transAxes, fontsize=12, style='italic')
                axes[2].set_title('Resumen', fontsize=11, fontweight='bold')
            
            plt.tight_layout()
            chart_path5 = self.charts_dir / f"{filename}_contenido.png"
            plt.savefig(chart_path5, dpi=300, bbox_inches='tight', facecolor='white', 
                       edgecolor='none', pad_inches=0.2)
            plt.close()
            chart_paths.append(str(chart_path5))
            print(f"    ✅ Gráfica de contenido creada: {chart_path5.name}")
        
        # Gráfica 6: Análisis de densidad y frecuencia (NUEVA)
        if stats['headers'] and len(stats['headers']) > 15:
            fig, axes = plt.subplots(2, 1, figsize=(14, 10))
            fig.patch.set_facecolor('white')
            fig.suptitle(f'📊 Análisis de Densidad y Frecuencia: {filename}', fontsize=16, 
                        fontweight='bold', color='#003366', y=0.995)
            
            # Subplot 1: Densidad de encabezados por sección del documento
            header_positions = [h['line'] for h in stats['headers']]
            total_lines = stats['total_lines']
            
            if header_positions and total_lines > 0:
                # Dividir en 10 secciones
                bins = 10
                hist, bin_edges = np.histogram(header_positions, bins=bins)
                bin_centers = (bin_edges[:-1] + bin_edges[1:]) / 2
                bin_labels = [f'{int(bin_edges[i])}-{int(bin_edges[i+1])}' for i in range(len(bin_edges)-1)]
                
                bars = axes[0].bar(range(len(bin_labels)), hist, 
                                  color=plt.cm.viridis([i/len(bin_labels) for i in range(len(bin_labels))]),
                                  edgecolor='black', linewidth=1, alpha=0.8)
                axes[0].set_xticks(range(len(bin_labels)))
                axes[0].set_xticklabels(bin_labels, rotation=45, ha='right', fontsize=9)
                axes[0].set_title('Densidad de Encabezados por Sección del Documento', 
                                fontsize=12, fontweight='bold', pad=15)
                axes[0].set_xlabel('Rango de Líneas', fontsize=11, fontweight='bold')
                axes[0].set_ylabel('Cantidad de Encabezados', fontsize=11, fontweight='bold')
                axes[0].grid(True, alpha=0.3, axis='y', linestyle='--')
                axes[0].set_axisbelow(True)
                
                for bar in bars:
                    height = bar.get_height()
                    if height > 0:
                        axes[0].text(bar.get_x() + bar.get_width()/2., height,
                                   f'{int(height)}', ha='center', va='bottom', 
                                   fontsize=9, fontweight='bold')
            
            # Subplot 2: Frecuencia de palabras clave (top palabras más comunes)
            if stats['sections']:
                # Extraer palabras más comunes de los títulos
                from collections import Counter
                import re
                all_words = []
                for section in stats['sections']:
                    words = re.findall(r'\b\w+\b', section['title'].lower())
                    all_words.extend(words)
                
                # Filtrar palabras comunes
                stop_words = {'de', 'la', 'el', 'en', 'y', 'a', 'los', 'las', 'un', 'una', 
                            'del', 'que', 'por', 'con', 'para', 'se', 'es', 'al', 'lo'}
                filtered_words = [w for w in all_words if w not in stop_words and len(w) > 3]
                
                if filtered_words:
                    word_counts = Counter(filtered_words)
                    top_words = word_counts.most_common(15)
                    
                    words, counts = zip(*top_words) if top_words else ([], [])
                    if words:
                        bars = axes[1].barh(range(len(words)), counts,
                                          color=plt.cm.plasma([i/len(words) for i in range(len(words))]),
                                          edgecolor='black', linewidth=0.8, alpha=0.85)
                        axes[1].set_yticks(range(len(words)))
                        axes[1].set_yticklabels(words, fontsize=10)
                        axes[1].set_title('Top 15 Palabras Clave en Títulos de Secciones', 
                                        fontsize=12, fontweight='bold', pad=15)
                        axes[1].set_xlabel('Frecuencia', fontsize=11, fontweight='bold')
                        axes[1].grid(True, alpha=0.3, axis='x', linestyle='--')
                        axes[1].set_axisbelow(True)
                        
                        for bar, count in zip(bars, counts):
                            axes[1].text(bar.get_width(), bar.get_y() + bar.get_height()/2,
                                       f'{count}', ha='left', va='center', 
                                       fontsize=9, fontweight='bold',
                                       bbox=dict(boxstyle='round,pad=0.3', facecolor='white', alpha=0.7))
            
            plt.tight_layout(rect=[0, 0, 1, 0.98])
            chart_path6 = self.charts_dir / f"{filename}_densidad.png"
            plt.savefig(chart_path6, dpi=300, bbox_inches='tight', facecolor='white', 
                       edgecolor='none', pad_inches=0.2)
            plt.close()
            chart_paths.append(str(chart_path6))
            print(f"    ✅ Gráfica de densidad creada: {chart_path6.name}")
        
        return chart_paths

    def create_master_dashboard(self, conversion_results: List[Dict[str, Any]]):
        """Crea un dashboard maestro en Excel y un resumen JSON global."""
        if not self.stats:
            print("⚠️  No hay estadísticas para el dashboard maestro.")
            return

        summary_data = []
        result_lookup = {}
        for result in conversion_results:
            doc_stem = Path(result['word']).stem if 'word' in result else None
            if doc_stem:
                result_lookup[doc_stem] = result

        for doc_name, stats in self.stats.items():
            doc_stats = {
                'documento': doc_name,
                'lineas': stats.get('total_lines', 0),
                'palabras': stats.get('total_words', 0),
                'caracteres': stats.get('total_chars', 0),
                'encabezados': len(stats.get('headers', [])),
                'secciones': len(stats.get('sections', [])),
                'codigo': len(stats.get('code_blocks', [])),
                'tablas': len(stats.get('tables', [])),
                'enlaces': len(stats.get('links', [])),
                'imagenes': len(stats.get('images', []))
            }
            result_info = result_lookup.get(doc_name)
            doc_stats['word_path'] = result_info['word'] if result_info and 'word' in result_info else ''
            doc_stats['excel_path'] = result_info['excel'] if result_info and 'excel' in result_info else ''
            doc_stats['pdf_path'] = result_info.get('pdf', '') if result_info else ''
            summary_data.append(doc_stats)

        if not summary_data:
            print("⚠️  No se encontraron datos para el dashboard maestro.")
            return

        summary_data = sorted(summary_data, key=lambda x: x['palabras'], reverse=True)

        master_path = self.output_dir / "master_dashboard.xlsx"
        wb = Workbook()

        # Hoja 1: Resumen Global
        ws = wb.active
        ws.title = "Resumen Global"
        headers = [
            "Documento", "Líneas", "Palabras", "Caracteres", "Encabezados",
            "Secciones", "Bloques de Código", "Tablas", "Enlaces", "Imágenes",
            "Archivo Word", "Archivo Excel", "Archivo PDF"
        ]
        ws.append(headers)
        header_fill = PatternFill(start_color="003366", end_color="003366", fill_type="solid")
        for col_idx, header in enumerate(headers, start=1):
            cell = ws.cell(row=1, column=col_idx)
            cell.font = Font(color="FFFFFF", bold=True)
            cell.fill = header_fill
            ws.column_dimensions[get_column_letter(col_idx)].width = 18

        for entry in summary_data:
            ws.append([
                entry['documento'],
                entry['lineas'],
                entry['palabras'],
                entry['caracteres'],
                entry['encabezados'],
                entry['secciones'],
                entry['codigo'],
                entry['tablas'],
                entry['enlaces'],
                entry['imagenes'],
                Path(entry['word_path']).name if entry['word_path'] else "",
                Path(entry['excel_path']).name if entry['excel_path'] else "",
                Path(entry['pdf_path']).name if entry['pdf_path'] else ""
            ])

        # Formato alternado
        for row in range(2, ws.max_row + 1):
            fill_color = "F7F9FC" if row % 2 == 0 else "FFFFFF"
            for col in range(1, ws.max_column + 1):
                ws.cell(row=row, column=col).fill = PatternFill(start_color=fill_color, end_color=fill_color, fill_type="solid")

        # Gráfica Top palabras
        top_count = min(10, len(summary_data))
        if top_count > 0:
            chart = BarChart()
            chart.title = "Top Documentos por Palabras"
            chart.y_axis.title = "Palabras"
            chart.x_axis.title = "Documento"
            data = Reference(ws, min_col=3, min_row=1, max_row=top_count + 1)
            cats = Reference(ws, min_col=1, min_row=2, max_row=top_count + 1)
            chart.add_data(data, titles_from_data=True)
            chart.set_categories(cats)
            chart.height = 10
            chart.width = 20
            ws.add_chart(chart, "N2")

        # Hoja 2: KPIs
        ws_kpi = wb.create_sheet("KPIs")
        kpi_data = [
            ("Documentos procesados", len(summary_data)),
            ("Palabras totales", sum(entry['palabras'] for entry in summary_data)),
            ("Promedio de palabras por documento", round(statistics.mean(entry['palabras'] for entry in summary_data), 2)),
            ("Encabezados totales", sum(entry['encabezados'] for entry in summary_data)),
            ("Secciones totales", sum(entry['secciones'] for entry in summary_data)),
            ("Enlaces totales", sum(entry['enlaces'] for entry in summary_data)),
            ("Imágenes totales", sum(entry['imagenes'] for entry in summary_data)),
        ]
        ws_kpi['A1'] = "Indicador"
        ws_kpi['B1'] = "Valor"
        ws_kpi['A1'].font = Font(bold=True)
        ws_kpi['B1'].font = Font(bold=True)
        for idx, (metric, value) in enumerate(kpi_data, start=2):
            ws_kpi[f"A{idx}"] = metric
            ws_kpi[f"B{idx}"] = value
            ws_kpi[f"A{idx}"].font = Font(bold=True)
        ws_kpi.column_dimensions['A'].width = 45
        ws_kpi.column_dimensions['B'].width = 20

        # Hoja 3: Distribución de contenido
        ws_dist = wb.create_sheet("Distribución")
        dist_headers = ["Tipo", "Cantidad"]
        ws_dist.append(dist_headers)
        total_content = {
            "Bloques de código": sum(entry['codigo'] for entry in summary_data),
            "Tablas": sum(entry['tablas'] for entry in summary_data),
            "Enlaces": sum(entry['enlaces'] for entry in summary_data),
            "Imágenes": sum(entry['imagenes'] for entry in summary_data)
        }
        for col_idx, header in enumerate(dist_headers, start=1):
            cell = ws_dist.cell(row=1, column=col_idx)
            cell.font = Font(color="FFFFFF", bold=True)
            cell.fill = header_fill

        for content_type, count in total_content.items():
            ws_dist.append([content_type, count])

        pie_chart = PieChart()
        pie_chart.title = "Distribución de Contenido"
        pie_data = Reference(ws_dist, min_col=2, min_row=1, max_row=len(total_content) + 1)
        pie_labels = Reference(ws_dist, min_col=1, min_row=2, max_row=len(total_content) + 1)
        pie_chart.add_data(pie_data, titles_from_data=True)
        pie_chart.set_categories(pie_labels)
        pie_chart.height = 12
        pie_chart.width = 12
        ws_dist.add_chart(pie_chart, "D2")

        # Guardar workbook y JSON
        wb.save(master_path)
        summary_json_path = self.output_dir / "conversion_summary.json"
        with open(summary_json_path, 'w', encoding='utf-8') as json_file:
            json.dump(summary_data, json_file, ensure_ascii=False, indent=2)

        print(f"✅ Dashboard maestro creado: {master_path}")
        print(f"✅ Resumen JSON global: {summary_json_path}")
    
    def create_word_document(self, parsed_data: Dict[str, Any], chart_paths: List[str], output_path: str):
        """Crea un documento Word profesional desde datos parseados."""
        print(f"  📝 Creando documento Word...")
        
        doc = Document()
        
        # Configurar márgenes
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3)
            section.right_margin = Cm(3)
        
        # Estilo base
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.space_after = Pt(6)
        
        # Título principal
        title = doc.add_heading(parsed_data['filename'].replace('_', ' ').title(), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Subtítulo
        subtitle = doc.add_paragraph(f"Análisis y Documentación Completa")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.size = Pt(12)
        subtitle_run.font.italic = True
        subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
        
        # Fecha
        date_para = doc.add_paragraph(f"Generado el: {datetime.now().strftime('%d de %B de %Y a las %H:%M')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.runs[0].font.size = Pt(10)
        date_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_paragraph()  # Espacio
        
        # Agregar gráficas con mejor organización
        if chart_paths:
            charts_heading = doc.add_heading('📊 Análisis Visual Completo', 1)
            charts_heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
            
            # Agrupar gráficas por tipo
            dashboard_charts = [p for p in chart_paths if 'dashboard' in p]
            other_charts = [p for p in chart_paths if 'dashboard' not in p]
            
            if dashboard_charts:
                doc.add_heading('Dashboard Principal', 2)
                for chart_path in dashboard_charts:
                    if os.path.exists(chart_path):
                        try:
                            doc.add_picture(chart_path, width=Inches(6.5))
                            last_paragraph = doc.paragraphs[-1]
                            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            doc.add_paragraph()  # Espacio
                        except Exception as e:
                            print(f"    ⚠️  No se pudo agregar imagen {chart_path}: {e}")
            
            if other_charts:
                doc.add_heading('Análisis Adicionales', 2)
                for chart_path in other_charts:
                    if os.path.exists(chart_path):
                        try:
                            doc.add_picture(chart_path, width=Inches(6))
                            last_paragraph = doc.paragraphs[-1]
                            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            doc.add_paragraph()  # Espacio
                        except Exception as e:
                            print(f"    ⚠️  No se pudo agregar imagen {chart_path}: {e}")
        
        doc.add_page_break()
        
        # Estadísticas mejoradas
        stats = parsed_data['stats']
        stats_heading = doc.add_heading('📊 Estadísticas del Documento', 1)
        stats_heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        
        # Agregar resumen ejecutivo de estadísticas
        summary_para = doc.add_paragraph()
        summary_para.add_run('Resumen Ejecutivo: ').bold = True
        summary_para.add_run(
            f"Este documento contiene {stats['total_words']:,} palabras distribuidas en "
            f"{stats['total_lines']:,} líneas, con {len(stats['headers']):,} encabezados organizados en "
            f"{len(stats['sections']):,} secciones principales. "
        )
        summary_para.runs[0].font.size = Pt(11)
        summary_para.runs[1].font.size = Pt(11)
        doc.add_paragraph()  # Espacio
        
        stats_table = doc.add_table(rows=1, cols=2)
        stats_table.style = 'Light Grid Accent 1'
        hdr_cells = stats_table.rows[0].cells
        hdr_cells[0].text = 'Métrica'
        hdr_cells[1].text = 'Valor'
        hdr_cells[0].paragraphs[0].runs[0].bold = True
        hdr_cells[1].paragraphs[0].runs[0].bold = True
        
        metrics = [
            ('Total de líneas', f"{stats['total_lines']:,}"),
            ('Total de palabras', f"{stats['total_words']:,}"),
            ('Total de caracteres', f"{stats['total_chars']:,}"),
            ('Encabezados', f"{len(stats['headers']):,}"),
            ('Bloques de código', f"{len(stats['code_blocks']):,}"),
            ('Tablas', f"{len(stats['tables']):,}"),
            ('Enlaces', f"{len(stats['links']):,}"),
            ('Imágenes', f"{len(stats['images']):,}"),
            ('Secciones principales', f"{len(stats['sections']):,}")
        ]
        
        for metric, value in metrics:
            row_cells = stats_table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = str(value)
        
        doc.add_paragraph()  # Espacio
        
        # Índice de secciones
        if stats['sections']:
            doc.add_page_break()
            sections_heading = doc.add_heading('📑 Índice de Secciones', 1)
            sections_heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
            
            sections_table = doc.add_table(rows=1, cols=3)
            sections_table.style = 'Light List Accent 1'
            hdr_cells = sections_table.rows[0].cells
            hdr_cells[0].text = 'Nivel'
            hdr_cells[1].text = 'Título'
            hdr_cells[2].text = 'Línea'
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].bold = True
            
            for section in stats['sections'][:50]:  # Limitar a 50
                row_cells = sections_table.add_row().cells
                row_cells[0].text = str(section['level'])
                row_cells[1].text = section['title'][:80]  # Limitar longitud
                row_cells[2].text = str(section['line'])
        
        doc.add_page_break()
        
        # Contenido principal (primeras 2000 líneas para no hacer el doc muy grande)
        content_heading = doc.add_heading('📄 Contenido', 1)
        content_heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        
        lines = parsed_data['content'].split('\n')
        processed_lines = 0
        max_lines = 2000
        
        for i, line in enumerate(lines[:max_lines]):
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
            
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                text = line.lstrip('#').strip()
                if level <= 6:
                    heading = doc.add_heading(text, min(level, 9))
                    if level <= 2:
                        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
            elif line.startswith('-') or line.startswith('*'):
                p = doc.add_paragraph(line, style='List Bullet')
            elif line.startswith('```'):
                # Código - omitir delimitadores
                continue
            elif '|' in line and line.count('|') >= 2:
                # Tabla - procesar básicamente (simplificado)
                continue
            else:
                doc.add_paragraph(line)
            
            processed_lines += 1
        
        if len(lines) > max_lines:
            doc.add_paragraph()
            note = doc.add_paragraph(f"[Nota: El documento original tiene {len(lines):,} líneas. "
                                    f"Se muestran las primeras {max_lines:,} líneas en este resumen.]")
            note.runs[0].font.italic = True
            note.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        
        # Guardar
        doc.save(output_path)
        print(f"    ✅ Documento Word creado: {Path(output_path).name}")
    
    def create_excel_workbook(self, parsed_data: Dict[str, Any], chart_paths: List[str], output_path: str):
        """Crea un libro Excel con datos estructurados y gráficas."""
        print(f"  📊 Creando libro Excel...")
        
        # Obtener estadísticas
        stats = parsed_data['stats']
        
        wb = Workbook()
        
        # Hoja 1: Resumen y Estadísticas
        ws_stats = wb.active
        ws_stats.title = "Resumen"
        
        # Título
        ws_stats['A1'] = parsed_data['filename'].replace('_', ' ').title()
        ws_stats['A1'].font = Font(size=18, bold=True, color="003366")
        ws_stats.merge_cells('A1:C1')
        ws_stats['A1'].alignment = Alignment(horizontal='center', vertical='center')
        
        ws_stats['A2'] = f"Generado el: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
        ws_stats['A2'].font = Font(size=10, italic=True)
        ws_stats.merge_cells('A2:C2')
        
        # Resumen ejecutivo en Excel
        ws_stats['A3'] = 'Resumen Ejecutivo'
        ws_stats['A3'].font = Font(bold=True, size=12, color="003366")
        ws_stats.merge_cells('A3:C3')
        summary_text = (f"Documento con {stats['total_words']:,} palabras, {stats['total_lines']:,} líneas, "
                       f"{len(stats['headers']):,} encabezados y {len(stats['sections']):,} secciones principales.")
        ws_stats['A4'] = summary_text
        ws_stats['A4'].font = Font(size=10, italic=True)
        ws_stats.merge_cells('A4:C4')
        ws_stats.row_dimensions[4].height = 30
        
        # Datos estadísticos
        row = 6
        ws_stats[f'A{row}'] = 'Métrica'
        ws_stats[f'B{row}'] = 'Valor'
        ws_stats[f'A{row}'].font = Font(bold=True, size=12)
        ws_stats[f'B{row}'].font = Font(bold=True, size=12)
        
        # Color de fondo para encabezados
        fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        ws_stats[f'A{row}'].fill = fill
        ws_stats[f'B{row}'].fill = fill
        ws_stats[f'A{row}'].font = Font(bold=True, color="FFFFFF", size=12)
        ws_stats[f'B{row}'].font = Font(bold=True, color="FFFFFF", size=12)
        
        metrics_data = [
            ('Total de líneas', stats['total_lines']),
            ('Total de palabras', stats['total_words']),
            ('Total de caracteres', stats['total_chars']),
            ('Encabezados', len(stats['headers'])),
            ('Bloques de código', len(stats['code_blocks'])),
            ('Tablas', len(stats['tables'])),
            ('Enlaces', len(stats['links'])),
            ('Imágenes', len(stats['images'])),
            ('Secciones principales', len(stats['sections']))
        ]
        
        for metric, value in metrics_data:
            row += 1
            ws_stats[f'A{row}'] = metric
            ws_stats[f'B{row}'] = value
            ws_stats[f'A{row}'].font = Font(size=11)
            ws_stats[f'B{row}'].font = Font(size=11)
            # Filas alternadas
            if row % 2 == 0:
                fill_alt = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
                ws_stats[f'A{row}'].fill = fill_alt
                ws_stats[f'B{row}'].fill = fill_alt
        
        # Ajustar columnas
        ws_stats.column_dimensions['A'].width = 25
        ws_stats.column_dimensions['B'].width = 20
        
        # Gráfica de barras
        chart = BarChart()
        chart.type = "col"
        chart.style = 10
        chart.title = "Estadísticas del Documento"
        chart.y_axis.title = 'Cantidad'
        chart.x_axis.title = 'Métricas'
        chart.height = 10
        chart.width = 15
        
        data = Reference(ws_stats, min_col=2, min_row=6, max_row=row)
        cats = Reference(ws_stats, min_col=1, min_row=7, max_row=row)
        chart.add_data(data, titles_from_data=False)
        chart.set_categories(cats)
        ws_stats.add_chart(chart, "D6")
        
        # Hoja 2: Encabezados
        ws_headers = wb.create_sheet("Encabezados")
        ws_headers['A1'] = 'Nivel'
        ws_headers['B1'] = 'Texto'
        ws_headers['C1'] = 'Línea'
        for cell in ['A1', 'B1', 'C1']:
            ws_headers[cell].font = Font(bold=True, size=12, color="FFFFFF")
            ws_headers[cell].fill = fill
            ws_headers[cell].alignment = Alignment(horizontal='center')
        
        for i, header in enumerate(stats['headers'][:500], start=2):  # Limitar a 500
            ws_headers[f'A{i}'] = header['level']
            ws_headers[f'B{i}'] = header['text'][:100]  # Limitar longitud
            ws_headers[f'C{i}'] = header['line']
            if i % 2 == 0:
                fill_alt = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
                ws_headers[f'A{i}'].fill = fill_alt
                ws_headers[f'B{i}'].fill = fill_alt
                ws_headers[f'C{i}'].fill = fill_alt
        
        ws_headers.column_dimensions['A'].width = 10
        ws_headers.column_dimensions['B'].width = 60
        ws_headers.column_dimensions['C'].width = 10
        
        # Gráfica de distribución de niveles
        if stats['headers']:
            level_counts = {}
            for h in stats['headers']:
                level_counts[h['level']] = level_counts.get(h['level'], 0) + 1
            
            ws_levels = wb.create_sheet("Distribución Niveles")
            ws_levels['A1'] = 'Nivel'
            ws_levels['B1'] = 'Cantidad'
            ws_levels['A1'].font = Font(bold=True, size=12, color="FFFFFF")
            ws_levels['B1'].font = Font(bold=True, size=12, color="FFFFFF")
            ws_levels['A1'].fill = fill
            ws_levels['B1'].fill = fill
            
            row = 2
            for level, count in sorted(level_counts.items()):
                ws_levels[f'A{row}'] = f"Nivel {level}"
                ws_levels[f'B{row}'] = count
                row += 1
            
            # Gráfica de pastel
            pie_chart = PieChart()
            pie_chart.title = "Distribución de Niveles de Encabezados"
            pie_chart.height = 10
            pie_chart.width = 10
            data = Reference(ws_levels, min_col=2, min_row=1, max_row=row-1)
            labels = Reference(ws_levels, min_col=1, min_row=2, max_row=row-1)
            pie_chart.add_data(data, titles_from_data=False)
            pie_chart.set_categories(labels)
            ws_levels.add_chart(pie_chart, "D2")
        
        # Hoja 3: Enlaces
        if stats['links']:
            ws_links = wb.create_sheet("Enlaces")
            ws_links['A1'] = 'Texto'
            ws_links['B1'] = 'URL'
            for cell in ['A1', 'B1']:
                ws_links[cell].font = Font(bold=True, size=12, color="FFFFFF")
                ws_links[cell].fill = fill
            
            for i, (text, url) in enumerate(stats['links'][:300], start=2):  # Limitar a 300
                ws_links[f'A{i}'] = text[:50]
                ws_links[f'B{i}'] = url[:100]  # Limitar URL
                if i % 2 == 0:
                    fill_alt = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
                    ws_links[f'A{i}'].fill = fill_alt
                    ws_links[f'B{i}'].fill = fill_alt
            
            ws_links.column_dimensions['A'].width = 30
            ws_links.column_dimensions['B'].width = 80
        
        # Hoja 4: Secciones
        if stats['sections']:
            ws_sections = wb.create_sheet("Secciones")
            ws_sections['A1'] = 'Nivel'
            ws_sections['B1'] = 'Título'
            ws_sections['C1'] = 'Línea'
            ws_sections['D1'] = 'Palabras'
            for cell in ['A1', 'B1', 'C1', 'D1']:
                ws_sections[cell].font = Font(bold=True, size=12, color="FFFFFF")
                ws_sections[cell].fill = fill
            
            for i, section in enumerate(stats['sections'], start=2):
                ws_sections[f'A{i}'] = section['level']
                ws_sections[f'B{i}'] = section['title'][:80]
                ws_sections[f'C{i}'] = section['line']
                ws_sections[f'D{i}'] = section['word_count']
                if i % 2 == 0:
                    fill_alt = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
                    for col in ['A', 'B', 'C', 'D']:
                        ws_sections[f'{col}{i}'].fill = fill_alt
            
            ws_sections.column_dimensions['A'].width = 10
            ws_sections.column_dimensions['B'].width = 50
            ws_sections.column_dimensions['C'].width = 10
            ws_sections.column_dimensions['D'].width = 12
            
            # Gráfica de palabras por sección
            if len(stats['sections']) > 1:
                chart_words = BarChart()
                chart_words.type = "col"
                chart_words.style = 10
                chart_words.title = "Palabras por Sección (Top 15)"
                chart_words.y_axis.title = 'Palabras'
                chart_words.height = 10
                chart_words.width = 15
                
                # Top 15 secciones
                top_sections = sorted(stats['sections'], key=lambda x: x['word_count'], reverse=True)[:15]
                start_row = 2
                end_row = start_row + len(top_sections) - 1
                
                data = Reference(ws_sections, min_col=4, min_row=start_row, max_row=end_row)
                cats = Reference(ws_sections, min_col=2, min_row=start_row, max_row=end_row)
                chart_words.add_data(data, titles_from_data=False)
                chart_words.set_categories(cats)
                ws_sections.add_chart(chart_words, "F2")
        
        # Guardar
        wb.save(output_path)
        print(f"    ✅ Libro Excel creado: {Path(output_path).name}")
    
    def create_pdf_document(self, parsed_data: Dict[str, Any], chart_paths: List[str], output_path: str):
        """Crea un documento PDF profesional."""
        print(f"  📄 Creando documento PDF...")
        
        doc = SimpleDocTemplate(output_path, pagesize=A4,
                               rightMargin=2*inch, leftMargin=2*inch,
                               topMargin=1.5*inch, bottomMargin=1.5*inch)
        
        # Estilos mejorados
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=26,
            textColor=colors.HexColor('#003366'),
            spaceAfter=30,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold',
            leading=32
        )
        
        # Agregar estilo para resumen ejecutivo
        summary_style = ParagraphStyle(
            'Summary',
            parent=styles['Normal'],
            fontSize=11,
            textColor=colors.HexColor('#333333'),
            spaceAfter=12,
            alignment=TA_LEFT,
            fontName='Helvetica',
            leading=14
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading1'],
            fontSize=16,
            textColor=colors.HexColor('#003366'),
            spaceAfter=12,
            spaceBefore=12
        )
        
        # Contenido
        story = []
        
        # Título
        title = Paragraph(parsed_data['filename'].replace('_', ' ').title(), title_style)
        story.append(title)
        story.append(Spacer(1, 0.2*inch))
        
        # Subtítulo
        subtitle = Paragraph(f"Análisis y Documentación Completa", 
                           ParagraphStyle('Subtitle', parent=styles['Normal'], 
                                         fontSize=12, textColor=colors.grey,
                                         alignment=TA_CENTER, fontStyle='Italic'))
        story.append(subtitle)
        story.append(Spacer(1, 0.1*inch))
        
        # Fecha
        date_text = f"Generado el: {datetime.now().strftime('%d de %B de %Y a las %H:%M')}"
        date_para = Paragraph(date_text, 
                            ParagraphStyle('Date', parent=styles['Normal'],
                                         fontSize=10, textColor=colors.grey,
                                         alignment=TA_CENTER))
        story.append(date_para)
        story.append(Spacer(1, 0.3*inch))
        
        # Agregar gráficas con mejor organización
        if chart_paths:
            charts_heading = Paragraph('Análisis Visual Completo', heading_style)
            story.append(charts_heading)
            story.append(Spacer(1, 0.2*inch))
            
            # Agrupar gráficas
            dashboard_charts = [p for p in chart_paths if 'dashboard' in p]
            other_charts = [p for p in chart_paths if 'dashboard' not in p]
            
            if dashboard_charts:
                sub_heading = Paragraph('Dashboard Principal', 
                                     ParagraphStyle('SubHeading', parent=styles['Heading2'],
                                                   fontSize=14, textColor=colors.HexColor('#003366')))
                story.append(sub_heading)
                story.append(Spacer(1, 0.15*inch))
                
                for chart_path in dashboard_charts:
                    if os.path.exists(chart_path):
                        try:
                            img = RLImage(chart_path, width=6*inch, height=5*inch)
                            story.append(img)
                            story.append(Spacer(1, 0.3*inch))
                        except Exception as e:
                            print(f"    ⚠️  No se pudo agregar imagen {chart_path}: {e}")
            
            if other_charts:
                sub_heading2 = Paragraph('Análisis Adicionales', 
                                        ParagraphStyle('SubHeading2', parent=styles['Heading2'],
                                                      fontSize=14, textColor=colors.HexColor('#003366')))
                story.append(sub_heading2)
                story.append(Spacer(1, 0.15*inch))
                
                for chart_path in other_charts:
                    if os.path.exists(chart_path):
                        try:
                            img = RLImage(chart_path, width=5.5*inch, height=4.5*inch)
                            story.append(img)
                            story.append(Spacer(1, 0.25*inch))
                        except Exception as e:
                            print(f"    ⚠️  No se pudo agregar imagen {chart_path}: {e}")
        
        story.append(PageBreak())
        
        # Resumen ejecutivo
        stats = parsed_data['stats']
        summary_text = (
            f"<b>Resumen Ejecutivo:</b> Este documento contiene {stats['total_words']:,} palabras "
            f"distribuidas en {stats['total_lines']:,} líneas, con {len(stats['headers']):,} encabezados "
            f"organizados en {len(stats['sections']):,} secciones principales."
        )
        summary_para = Paragraph(summary_text, summary_style)
        story.append(summary_para)
        story.append(Spacer(1, 0.2*inch))
        
        # Estadísticas
        stats_heading = Paragraph('Estadísticas del Documento', heading_style)
        story.append(stats_heading)
        story.append(Spacer(1, 0.2*inch))
        
        stats = parsed_data['stats']
        stats_data = [
            ['Métrica', 'Valor'],
            ['Total de líneas', f"{stats['total_lines']:,}"],
            ['Total de palabras', f"{stats['total_words']:,}"],
            ['Total de caracteres', f"{stats['total_chars']:,}"],
            ['Encabezados', f"{len(stats['headers']):,}"],
            ['Bloques de código', f"{len(stats['code_blocks']):,}"],
            ['Tablas', f"{len(stats['tables']):,}"],
            ['Enlaces', f"{len(stats['links']):,}"],
            ['Imágenes', f"{len(stats['images']):,}"],
            ['Secciones principales', f"{len(stats['sections']):,}"]
        ]
        
        stats_table = Table(stats_data, colWidths=[3*inch, 2*inch])
        stats_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4472C4')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F2F2F2')])
        ]))
        story.append(stats_table)
        story.append(Spacer(1, 0.3*inch))
        
        # Índice de secciones
        if stats['sections']:
            story.append(PageBreak())
            sections_heading = Paragraph('Índice de Secciones', heading_style)
            story.append(sections_heading)
            story.append(Spacer(1, 0.2*inch))
            
            sections_data = [['Nivel', 'Título', 'Línea']]
            for section in stats['sections'][:50]:
                sections_data.append([
                    str(section['level']),
                    section['title'][:60],
                    str(section['line'])
                ])
            
            sections_table = Table(sections_data, colWidths=[0.5*inch, 4*inch, 0.5*inch])
            sections_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4472C4')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F2F2F2')])
            ]))
            story.append(sections_table)
        
        # Contenido (resumen)
        story.append(PageBreak())
        content_heading = Paragraph('Resumen del Contenido', heading_style)
        story.append(content_heading)
        story.append(Spacer(1, 0.2*inch))
        
        # Primeras líneas del contenido
        lines = parsed_data['content'].split('\n')
        preview_lines = lines[:100]  # Primeras 100 líneas
        
        for line in preview_lines:
            line = line.strip()
            if line and not line.startswith('```'):
                if line.startswith('#'):
                    level = len(line) - len(line.lstrip('#'))
                    text = line.lstrip('#').strip()
                    if level <= 3:
                        para = Paragraph(text, styles[f'Heading{min(level, 3)}'])
                        story.append(para)
                else:
                    para = Paragraph(line, styles['Normal'])
                    story.append(para)
                    story.append(Spacer(1, 0.1*inch))
        
        if len(lines) > 100:
            note = Paragraph(
                f"[Nota: El documento original tiene {len(lines):,} líneas. "
                f"Se muestra un resumen de las primeras 100 líneas.]",
                ParagraphStyle('Note', parent=styles['Normal'],
                              fontSize=9, textColor=colors.grey,
                              fontStyle='Italic'))
            story.append(Spacer(1, 0.2*inch))
            story.append(note)
        
        # Construir PDF
        doc.build(story)
        print(f"    ✅ Documento PDF creado: {Path(output_path).name}")
    
    def convert(self, markdown_file: str):
        """Convierte un archivo Markdown a PDF, Word y Excel."""
        print(f"\n🔄 Procesando: {Path(markdown_file).name}")
        
        try:
            # Parsear
            parsed_data = self.parse_markdown(markdown_file)
            filename = parsed_data['filename']
            
            # Crear visualizaciones
            chart_paths = self.create_visualizations(parsed_data)
            
            # Crear Word
            word_path = self.output_dir / f"{filename}.docx"
            self.create_word_document(parsed_data, chart_paths, str(word_path))
            
            # Crear Excel
            excel_path = self.output_dir / f"{filename}.xlsx"
            self.create_excel_workbook(parsed_data, chart_paths, str(excel_path))
            
            # Crear PDF
            pdf_path = self.output_dir / f"{filename}.pdf"
            try:
                self.create_pdf_document(parsed_data, chart_paths, str(pdf_path))
            except Exception as e:
                print(f"    ⚠️  Error creando PDF (continuando...): {e}")
                pdf_path = None
            
            # Guardar estadísticas
            self.stats[filename] = parsed_data['stats']
            
            print(f"✅ Conversión completada para: {filename}")
            result = {
                'word': str(word_path),
                'excel': str(excel_path),
                'charts': chart_paths
            }
            if pdf_path and pdf_path.exists():
                result['pdf'] = str(pdf_path)
            return result
        except Exception as e:
            print(f"❌ Error procesando {markdown_file}: {e}")
            import traceback
            traceback.print_exc()
            return None


def main():
    """Función principal."""
    print("=" * 70)
    print("📚 GENERADOR DE DOCUMENTOS IMPORTANTES - PDF, WORD Y EXCEL")
    print("=" * 70)
    print()
    
    # Directorio base
    base_dir = Path("/Users/adan/Documents/documentos_blatam")
    
        # Documentos más importantes a convertir - ULTRA EXPANDIDO
    important_files = [
        # Documentos principales
        base_dir / "airflow_automation_prompt.md",
        base_dir / "README.md",
        base_dir / "ARCHITECTURE.md",
        base_dir / "BEST_PRACTICES.md",
        
        # Documentos de producción
        base_dir / "truthgpt_collected/integration_code/production_code/ARCHITECTURE_IMPROVEMENTS.md",
        base_dir / "truthgpt_collected/integration_code/production_code/REFACTORING_PLAN.md",
        base_dir / "truthgpt_collected/integration_code/production_code/ARCHITECTURE.md",
        base_dir / "truthgpt_collected/integration_code/production_code/RESUMEN_FINAL_MEJORAS.md",
        base_dir / "truthgpt_collected/integration_code/production_code/MEJORAS_ARQUITECTURA_COMPLETAS.md",
        base_dir / "truthgpt_collected/integration_code/production_code/INDICE_DOCUMENTACION.md",
        base_dir / "truthgpt_collected/integration_code/production_code/QUICK_START.md",
        base_dir / "truthgpt_collected/integration_code/production_code/IMPORT_STANDARDS.md",
        base_dir / "truthgpt_collected/integration_code/production_code/ARCHITECTURE_IMPROVEMENTS_SUMMARY.md",
        
        # Documentación técnica importante
        base_dir / "truthgpt_collected/integration_code/production_code/docs/README.md",
        base_dir / "Docs/README.md",
        base_dir / "Docs/BEST_PRACTICES.md",
        base_dir / "06_documentation/Technical_docs/documentacion_completa.md",
        
        # Guías importantes
        base_dir / "truthgpt_collected/integration_code/production_code/REFACTORING_QUICK_REFERENCE.md",
        
        # Resúmenes importantes
        base_dir / "truthgpt_collected/integration_code/production_code/CODE_CLEANUP_SUMMARY.md",
        base_dir / "truthgpt_collected/integration_code/production_code/CLEANUP_SUMMARY.md",
    ]
    
    # Crear conversor
    converter = DocumentConverter(output_dir=str(base_dir / "exports_docs_importantes"))
    
    # Convertir documentos
    results = []
    for file_path in important_files:
        if file_path.exists():
            result = converter.convert(str(file_path))
            if result:
                results.append(result)
        else:
            print(f"⚠️  Archivo no encontrado: {file_path}")
    
    # Resumen final
    print("\n" + "=" * 70)
    print("📊 RESUMEN DE CONVERSIÓN")
    print("=" * 70)
    
    for result in results:
        filename = Path(result['word']).stem
        print(f"\n✅ {filename}:")
        print(f"   📝 Word:  {result['word']}")
        print(f"   📊 Excel: {result['excel']}")
        if 'pdf' in result:
            print(f"   📄 PDF:   {result['pdf']}")
        else:
            print(f"   📄 PDF:   ⚠️  No generado")
        print(f"   📈 Gráficas: {len(result['charts'])} creadas")
    
    # Dashboard maestro y resumen global
    if results:
        converter.create_master_dashboard(results)

    print("\n" + "=" * 70)
    print("✨ ¡Conversión completada exitosamente!")
    print("=" * 70)
    print(f"\n📁 Todos los archivos se guardaron en: {converter.output_dir}")
    print(f"📈 Gráficas guardadas en: {converter.charts_dir}")


if __name__ == "__main__":
    main()

