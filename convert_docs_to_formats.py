#!/usr/bin/env python3
"""
Script para convertir documentos importantes a PDF, Word y Excel con gráficas
"""

import os
import sys
from pathlib import Path
import markdown
from datetime import datetime
import json

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("Instalando python-docx...")
    os.system("pip install python-docx")
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.chart import BarChart, PieChart, LineChart, Reference
except ImportError:
    print("Instalando openpyxl...")
    os.system("pip install openpyxl")
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.chart import BarChart, PieChart, LineChart, Reference

try:
    import matplotlib.pyplot as plt
    import matplotlib
    matplotlib.use('Agg')  # Para no requerir display
except ImportError:
    print("Instalando matplotlib...")
    os.system("pip install matplotlib")
    import matplotlib.pyplot as plt
    import matplotlib
    matplotlib.use('Agg')

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
except ImportError:
    print("Instalando reportlab...")
    os.system("pip install reportlab")
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

class DocumentConverter:
    def __init__(self, output_dir="converted_docs"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        self.temp_dir = self.output_dir / "temp"
        self.temp_dir.mkdir(exist_ok=True)
        
    def read_markdown(self, file_path):
        """Lee un archivo markdown"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def parse_markdown_sections(self, content):
        """Parsea el markdown en secciones"""
        sections = []
        current_section = {"title": "", "content": "", "level": 0}
        
        lines = content.split('\n')
        for line in lines:
            if line.startswith('#'):
                if current_section["content"]:
                    sections.append(current_section)
                level = len(line) - len(line.lstrip('#'))
                title = line.lstrip('#').strip()
                current_section = {"title": title, "content": "", "level": level}
            else:
                current_section["content"] += line + "\n"
        
        if current_section["content"]:
            sections.append(current_section)
        
        return sections
    
    def create_statistics(self, content):
        """Crea estadísticas del documento"""
        stats = {
            "total_lines": len(content.split('\n')),
            "total_words": len(content.split()),
            "total_chars": len(content),
            "sections": len([l for l in content.split('\n') if l.startswith('#')]),
            "code_blocks": content.count('```'),
            "links": content.count('[') - content.count(']'),
            "images": content.count('!['),
        }
        return stats
    
    def create_charts(self, stats, doc_name):
        """Crea gráficas basadas en estadísticas"""
        charts = []
        
        # Gráfica de barras - Estadísticas generales
        fig, ax = plt.subplots(figsize=(10, 6))
        categories = ['Líneas', 'Palabras', 'Secciones', 'Bloques Código']
        values = [
            stats['total_lines'] / 100,  # Normalizado
            stats['total_words'] / 100,
            stats['sections'] * 10,
            stats['code_blocks'] * 5
        ]
        ax.bar(categories, values, color=['#3498db', '#2ecc71', '#e74c3c', '#f39c12'])
        ax.set_title(f'Estadísticas del Documento: {doc_name}', fontsize=14, fontweight='bold')
        ax.set_ylabel('Valor Normalizado')
        plt.xticks(rotation=45)
        plt.tight_layout()
        chart_path = self.temp_dir / f"{doc_name}_stats_bar.png"
        plt.savefig(chart_path, dpi=300, bbox_inches='tight')
        plt.close()
        charts.append(chart_path)
        
        # Gráfica de pie - Distribución de contenido
        fig, ax = plt.subplots(figsize=(8, 8))
        labels = ['Texto', 'Código', 'Enlaces', 'Imágenes']
        sizes = [
            stats['total_words'],
            stats['code_blocks'] * 50,
            stats['links'] * 10,
            stats['images'] * 20
        ]
        colors_pie = ['#3498db', '#2ecc71', '#e74c3c', '#f39c12']
        ax.pie(sizes, labels=labels, autopct='%1.1f%%', colors=colors_pie, startangle=90)
        ax.set_title('Distribución de Contenido', fontsize=14, fontweight='bold')
        chart_path = self.temp_dir / f"{doc_name}_content_pie.png"
        plt.savefig(chart_path, dpi=300, bbox_inches='tight')
        plt.close()
        charts.append(chart_path)
        
        return charts
    
    def convert_to_word(self, file_path, output_name):
        """Convierte markdown a Word con formato profesional"""
        print(f"Convirtiendo {file_path} a Word...")
        
        content = self.read_markdown(file_path)
        sections = self.parse_markdown_sections(content)
        stats = self.create_statistics(content)
        charts = self.create_charts(stats, output_name)
        
        doc = Document()
        
        # Configurar estilos
        styles = doc.styles
        title_style = styles['Title']
        title_style.font.name = 'Calibri'
        title_style.font.size = Pt(24)
        title_style.font.bold = True
        
        # Título del documento
        title = doc.add_heading(output_name.replace('_', ' ').title(), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Información del documento
        info_para = doc.add_paragraph()
        info_para.add_run(f"Generado el: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        info_para.add_run(f"Archivo original: {Path(file_path).name}\n")
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Espacio
        
        # Agregar gráficas
        for chart_path in charts:
            doc.add_picture(str(chart_path), width=Inches(6))
            doc.add_paragraph()  # Espacio después de gráfica
        
        # Agregar estadísticas en tabla
        doc.add_heading('Estadísticas del Documento', level=1)
        stats_table = doc.add_table(rows=1, cols=2)
        stats_table.style = 'Light Grid Accent 1'
        hdr_cells = stats_table.rows[0].cells
        hdr_cells[0].text = 'Métrica'
        hdr_cells[1].text = 'Valor'
        
        for key, value in stats.items():
            row_cells = stats_table.add_row().cells
            row_cells[0].text = key.replace('_', ' ').title()
            row_cells[1].text = str(value)
        
        doc.add_page_break()
        
        # Agregar contenido
        doc.add_heading('Contenido del Documento', level=1)
        
        for section in sections:
            if section['title']:
                level = min(section['level'], 9)
                doc.add_heading(section['title'], level=level)
            
            # Procesar contenido
            content_lines = section['content'].split('\n')
            for line in content_lines:
                if line.strip():
                    if line.startswith('- ') or line.startswith('* '):
                        # Lista
                        doc.add_paragraph(line[2:], style='List Bullet')
                    elif line.startswith('```'):
                        # Bloque de código
                        continue
                    elif '`' in line:
                        # Código inline
                        para = doc.add_paragraph()
                        parts = line.split('`')
                        for i, part in enumerate(parts):
                            if i % 2 == 0:
                                para.add_run(part)
                            else:
                                run = para.add_run(part)
                                run.font.name = 'Courier New'
                                run.font.size = Pt(10)
                    else:
                        doc.add_paragraph(line)
            
            doc.add_paragraph()  # Espacio entre secciones
        
        # Guardar
        output_path = self.output_dir / f"{output_name}.docx"
        doc.save(str(output_path))
        print(f"✓ Word guardado: {output_path}")
        return output_path
    
    def convert_to_excel(self, file_path, output_name):
        """Convierte markdown a Excel con gráficas"""
        print(f"Convirtiendo {file_path} a Excel...")
        
        content = self.read_markdown(file_path)
        sections = self.parse_markdown_sections(content)
        stats = self.create_statistics(content)
        
        wb = Workbook()
        ws = wb.active
        ws.title = "Resumen"
        
        # Estilos
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF", size=14)
        title_font = Font(bold=True, size=16)
        
        # Título
        ws['A1'] = output_name.replace('_', ' ').title()
        ws['A1'].font = title_font
        ws.merge_cells('A1:D1')
        ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
        
        # Información
        ws['A3'] = f"Generado el: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        ws['A4'] = f"Archivo original: {Path(file_path).name}"
        
        # Estadísticas
        row = 6
        ws[f'A{row}'] = "Estadísticas del Documento"
        ws[f'A{row}'].font = header_font
        ws[f'A{row}'].fill = header_fill
        ws.merge_cells(f'A{row}:B{row}')
        
        row += 1
        ws[f'A{row}'] = "Métrica"
        ws[f'B{row}'] = "Valor"
        ws[f'A{row}'].font = Font(bold=True)
        ws[f'B{row}'].font = Font(bold=True)
        
        for key, value in stats.items():
            row += 1
            ws[f'A{row}'] = key.replace('_', ' ').title()
            ws[f'B{row}'] = value
        
        # Gráfica de barras
        chart = BarChart()
        chart.type = "col"
        chart.style = 10
        chart.title = "Estadísticas del Documento"
        chart.y_axis.title = "Valor"
        chart.x_axis.title = "Métrica"
        
        data = Reference(ws, min_col=2, min_row=7, max_row=row)
        cats = Reference(ws, min_col=1, min_row=8, max_row=row)
        chart.add_data(data, titles_from_data=False)
        chart.set_categories(cats)
        chart.height = 10
        chart.width = 15
        ws.add_chart(chart, "D6")
        
        # Hoja de contenido
        ws2 = wb.create_sheet("Contenido")
        row = 1
        
        for section in sections:
            if section['title']:
                ws2[f'A{row}'] = section['title']
                ws2[f'A{row}'].font = Font(bold=True, size=12 + (3 - section['level']) * 2)
                row += 1
            
            # Agregar contenido
            content_lines = section['content'].split('\n')
            for line in content_lines:
                if line.strip() and not line.startswith('```'):
                    ws2[f'B{row}'] = line[:200]  # Limitar longitud
                    row += 1
                    if row > 1000:  # Limitar filas
                        break
            
            row += 1
            if row > 1000:
                break
        
        # Guardar
        output_path = self.output_dir / f"{output_name}.xlsx"
        wb.save(str(output_path))
        print(f"✓ Excel guardado: {output_path}")
        return output_path
    
    def convert_to_pdf(self, file_path, output_name):
        """Convierte markdown a PDF con formato profesional"""
        print(f"Convirtiendo {file_path} a PDF...")
        
        content = self.read_markdown(file_path)
        sections = self.parse_markdown_sections(content)
        stats = self.create_statistics(content)
        charts = self.create_charts(stats, output_name)
        
        # Crear PDF
        output_path = self.output_dir / f"{output_name}.pdf"
        doc = SimpleDocTemplate(str(output_path), pagesize=A4,
                              rightMargin=72, leftMargin=72,
                              topMargin=72, bottomMargin=18)
        
        # Estilos
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1a1a1a'),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        story = []
        
        # Título
        title = Paragraph(output_name.replace('_', ' ').title(), title_style)
        story.append(title)
        story.append(Spacer(1, 12))
        
        # Información
        info_style = ParagraphStyle(
            'Info',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.HexColor('#666666'),
            alignment=TA_CENTER
        )
        story.append(Paragraph(f"Generado el: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", info_style))
        story.append(Paragraph(f"Archivo original: {Path(file_path).name}", info_style))
        story.append(Spacer(1, 20))
        
        # Agregar gráficas
        for chart_path in charts:
            img = Image(str(chart_path), width=5*inch, height=3*inch)
            story.append(img)
            story.append(Spacer(1, 12))
        
        # Tabla de estadísticas
        story.append(Paragraph("Estadísticas del Documento", styles['Heading2']))
        story.append(Spacer(1, 12))
        
        stats_data = [['Métrica', 'Valor']]
        for key, value in stats.items():
            stats_data.append([key.replace('_', ' ').title(), str(value)])
        
        stats_table = Table(stats_data, colWidths=[3*inch, 2*inch])
        stats_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#366092')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        story.append(stats_table)
        story.append(PageBreak())
        
        # Contenido
        story.append(Paragraph("Contenido del Documento", styles['Heading1']))
        story.append(Spacer(1, 12))
        
        for section in sections:
            if section['title']:
                level = min(section['level'], 6)
                heading_style = styles[f'Heading{level}']
                story.append(Paragraph(section['title'], heading_style))
                story.append(Spacer(1, 6))
            
            # Procesar contenido
            content_lines = section['content'].split('\n')
            for line in content_lines:
                if line.strip() and not line.startswith('```'):
                    # Limitar longitud de líneas para PDF
                    if len(line) > 100:
                        line = line[:97] + "..."
                    # Escapar caracteres especiales para ReportLab
                    line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    # Remover caracteres problemáticos
                    line = ''.join(char for char in line if ord(char) < 128 or char.isprintable())
                    try:
                        story.append(Paragraph(line, styles['Normal']))
                        story.append(Spacer(1, 3))
                    except:
                        # Si falla, agregar como texto simple
                        story.append(Paragraph(line[:200], styles['Normal']))
                        story.append(Spacer(1, 3))
            
            story.append(Spacer(1, 12))
        
        # Construir PDF
        doc.build(story)
        print(f"✓ PDF guardado: {output_path}")
        return output_path

def main():
    """Función principal"""
    # Documentos importantes a convertir
    base_dir = Path("/Users/adan/Documents/documentos_blatam")
    production_dir = base_dir / "truthgpt_collected/integration_code/production_code"
    
    documents = [
        {
            "path": base_dir / "airflow_automation_prompt.md",
            "name": "Automation_Expert_Prompt"
        },
        {
            "path": production_dir / "ARCHITECTURE_IMPROVEMENTS.md",
            "name": "Architecture_Improvements"
        },
        {
            "path": production_dir / "REFACTORING_PLAN.md",
            "name": "Refactoring_Plan"
        }
    ]
    
    converter = DocumentConverter()
    
    print("=" * 60)
    print("CONVERSIÓN DE DOCUMENTOS A PDF, WORD Y EXCEL")
    print("=" * 60)
    print()
    
    results = []
    
    for doc in documents:
        if not doc["path"].exists():
            print(f"⚠ Archivo no encontrado: {doc['path']}")
            continue
        
        print(f"\n📄 Procesando: {doc['name']}")
        print("-" * 60)
        
        try:
            # Convertir a Word
            word_path = converter.convert_to_word(doc["path"], doc["name"])
            
            # Convertir a Excel
            excel_path = converter.convert_to_excel(doc["path"], doc["name"])
            
            # Convertir a PDF
            pdf_path = converter.convert_to_pdf(doc["path"], doc["name"])
            
            results.append({
                "document": doc["name"],
                "word": str(word_path),
                "excel": str(excel_path),
                "pdf": str(pdf_path)
            })
            
        except Exception as e:
            print(f"❌ Error procesando {doc['name']}: {str(e)}")
            import traceback
            traceback.print_exc()
    
    # Resumen
    print("\n" + "=" * 60)
    print("RESUMEN DE CONVERSIÓN")
    print("=" * 60)
    print(f"\nDocumentos procesados: {len(results)}")
    print(f"Directorio de salida: {converter.output_dir}")
    print("\nArchivos generados:")
    for result in results:
        print(f"\n  📄 {result['document']}:")
        print(f"     - Word: {result['word']}")
        print(f"     - Excel: {result['excel']}")
        print(f"     - PDF: {result['pdf']}")
    
    print("\n✅ Conversión completada!")

if __name__ == "__main__":
    main()

