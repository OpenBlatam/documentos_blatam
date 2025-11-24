import docx
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_investor_faq(filename):
    doc = docx.Document()
    
    # Styles
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(16)
    h1.font.color.rgb = RGBColor(32, 55, 100)
    h1.paragraph_format.space_before = Pt(18)

    # Title
    title = doc.add_paragraph("VENTURA CAPITAL - INVESTOR FAQ & DEFENSE PLAYBOOK")
    title.style = 'Title'
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Confidential | Strategic Responses to Hard Questions").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("_" * 80)

    # Q&A Section
    
    def add_qa(question, answer, bullet_points=None):
        doc.add_heading(question, level=1)
        p = doc.add_paragraph()
        run = p.add_run("Strategic Response:")
        run.bold = True
        run.font.color.rgb = RGBColor(47, 117, 181)
        
        doc.add_paragraph(answer)
        
        if bullet_points:
            for point in bullet_points:
                doc.add_paragraph(point, style='List Bullet')

    # Q1: Competition
    add_qa(
        "1. Why won't GPT-5 or Google Gemini kill your business?",
        "Los modelos fundacionales (Foundation Models) son generalistas. Nosotros somos verticalistas. Nuestra ventaja no es el modelo base, sino la capa de aplicación y los datos propietarios.",
        [
            "Data Moat: Tenemos 50M+ de ejemplos de copy calibrados culturalmente para LATAM.",
            "Workflow Integration: Estamos integrados en el flujo de trabajo de las agencias (Slack, HubSpot, Meta Ads), no somos solo un chat.",
            "Brand Voice: Nuestro fine-tuning permite mantener la voz de marca consistente, algo que ChatGPT 'out of the box' no logra bien."
        ]
    )

    # Q2: CAC Saturation
    add_qa(
        "2. Your CAC is low ($150). Won't it skyrocket as you scale?",
        "Esperamos un aumento natural del CAC en canales pagados, pero nuestra estrategia de blended CAC compensa esto.",
        [
            "Viral Loops: El 30% de nuestros usuarios nuevos vienen por referidos (Product-Led Growth).",
            "Content SEO: Dominamos keywords de 'Marketing IA' en español.",
            "Partnerships: El costo de adquisición a través de partners (e.g., TiendaNube) es cercano a cero."
        ]
    )

    # Q3: Market Size
    add_qa(
        "3. Is the LATAM market deep enough for a Venture outcome?",
        "Absolutamente. El mercado digital de LATAM es el de mayor crecimiento mundial (CAGR 30%).",
        [
            "TAM Expansion: No solo vendemos a agencias. Nuestro roadmap incluye Enterprise (Bancos, Retail) y SMBs.",
            "Cross-border: Nuestra tecnología de dialectos es exportable a mercado US Hispanic ($2T GDP) y España.",
            "Exit Path: Adquirentes estratégicos como HubSpot, Salesforce o Globant buscan activamente tecnología nativa de la región."
        ]
    )

    # Q4: Team
    add_qa(
        "4. Why is this the right team to execute?",
        "Combinamos experiencia técnica profunda con ejecución comercial probada.",
        [
            "Founder/CEO: Ex-VP de Ventas en Unicornio SaaS. Sabe escalar equipos comerciales.",
            "CTO: PhD en NLP. Ha publicado papers sobre 'Low-resource language fine-tuning'.",
            "Advisors: VP de Marketing de MercadoLibre y Ex-Partner de Kaszek."
        ]
    )

    doc.save(filename)
    print(f"Investor FAQ created: {filename}")

if __name__ == "__main__":
    create_investor_faq("Ventura_Capital_Investor_FAQ.docx")


