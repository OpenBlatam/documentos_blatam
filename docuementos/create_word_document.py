#!/usr/bin/env python3
"""
Script para crear documento Word del Curso Premium de IA en Marketing
Usa python-docx para generar un documento Word profesional
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def read_md_file(file_path):
    """Lee un archivo Markdown y retorna su contenido"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except FileNotFoundError:
        print(f"Error: No se encontró el archivo {file_path}")
        return ""
    except Exception as e:
        print(f"Error leyendo {file_path}: {e}")
        return ""

def add_heading_with_style(doc, text, level):
    """Añade un encabezado con estilo personalizado"""
    heading = doc.add_heading(text, level)
    
    # Personalizar estilo según el nivel
    if level == 1:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.runs[0].font.size = Pt(24)
        heading.runs[0].font.color.rgb = None  # Color por defecto
    elif level == 2:
        heading.runs[0].font.size = Pt(18)
    elif level == 3:
        heading.runs[0].font.size = Pt(14)
    elif level == 4:
        heading.runs[0].font.size = Pt(12)
    
    return heading

def add_formatted_paragraph(doc, text):
    """Añade un párrafo con formato básico"""
    if not text.strip():
        return doc.add_paragraph()
    
    # Convertir markdown básico a formato Word
    text = text.replace('**', '')  # Remover negritas de markdown
    text = text.replace('__', '')  # Remover negritas de markdown
    text = text.replace('*', '')   # Remover cursivas de markdown
    text = text.replace('_', '')   # Remover cursivas de markdown
    text = text.replace('`', '')   # Remover código de markdown
    
    paragraph = doc.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return paragraph

def add_list_item(doc, text, is_numbered=False):
    """Añade un elemento de lista"""
    if not text.strip():
        return
    
    # Limpiar texto
    text = text.strip()
    if text.startswith('- ') or text.startswith('* '):
        text = text[2:]
    elif text.startswith(('1. ', '2. ', '3. ', '4. ', '5. ', '6. ', '7. ', '8. ', '9. ')):
        text = text[3:]
    
    # Remover formato markdown
    text = text.replace('**', '')
    text = text.replace('__', '')
    text = text.replace('*', '')
    text = text.replace('_', '')
    text = text.replace('`', '')
    
    if is_numbered:
        doc.add_paragraph(text, style='List Number')
    else:
        doc.add_paragraph(text, style='List Bullet')

def process_markdown_content(doc, content):
    """Procesa el contenido Markdown y lo convierte a Word"""
    lines = content.split('\n')
    in_list = False
    in_numbered_list = False
    
    for line in lines:
        line = line.strip()
        
        if not line:
            if in_list:
                in_list = False
            if in_numbered_list:
                in_numbered_list = False
            continue
        
        # Títulos
        if line.startswith('# '):
            if in_list or in_numbered_list:
                in_list = False
                in_numbered_list = False
            add_heading_with_style(doc, line[2:], 1)
            
        elif line.startswith('## '):
            if in_list or in_numbered_list:
                in_list = False
                in_numbered_list = False
            add_heading_with_style(doc, line[3:], 2)
            
        elif line.startswith('### '):
            if in_list or in_numbered_list:
                in_list = False
                in_numbered_list = False
            add_heading_with_style(doc, line[4:], 3)
            
        elif line.startswith('#### '):
            if in_list or in_numbered_list:
                in_list = False
                in_numbered_list = False
            add_heading_with_style(doc, line[5:], 4)
            
        elif line.startswith('##### '):
            if in_list or in_numbered_list:
                in_list = False
                in_numbered_list = False
            add_heading_with_style(doc, line[6:], 5)
            
        elif line.startswith('###### '):
            if in_list or in_numbered_list:
                in_list = False
                in_numbered_list = False
            add_heading_with_style(doc, line[7:], 6)
            
        # Listas
        elif line.startswith('- ') or line.startswith('* '):
            if not in_list:
                in_list = True
                in_numbered_list = False
            add_list_item(doc, line, False)
            
        elif line.startswith(('1. ', '2. ', '3. ', '4. ', '5. ', '6. ', '7. ', '8. ', '9. ')):
            if not in_numbered_list:
                in_numbered_list = True
                in_list = False
            add_list_item(doc, line, True)
            
        # Separadores
        elif line.startswith('---'):
            if in_list or in_numbered_list:
                in_list = False
                in_numbered_list = False
            doc.add_paragraph('─' * 50)
            
        # Párrafos normales
        else:
            if in_list or in_numbered_list:
                in_list = False
                in_numbered_list = False
            add_formatted_paragraph(doc, line)

def create_word_document(content1, content2, output_path):
    """Crea el documento Word completo"""
    
    # Crear nuevo documento
    doc = Document()
    
    # Configurar estilos
    styles = doc.styles
    
    # Estilo para títulos principales
    title_style = styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Calibri'
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Estilo para subtítulos
    subtitle_style = styles.add_style('Custom Subtitle', WD_STYLE_TYPE.PARAGRAPH)
    subtitle_style.font.name = 'Calibri'
    subtitle_style.font.size = Pt(16)
    subtitle_style.font.bold = True
    subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Añadir portada
    title = doc.add_paragraph('🚀 CURSO PREMIUM: INTELIGENCIA ARTIFICIAL APLICADA AL MARKETING', style='Custom Title')
    subtitle = doc.add_paragraph('Programa de Certificación Profesional Avanzado', style='Custom Subtitle')
    subtitle2 = doc.add_paragraph('Transforma tu Carrera con IA: De Principiante a Experto en 8 Semanas', style='Custom Subtitle')
    
    # Añadir información del curso
    doc.add_paragraph()
    info_para = doc.add_paragraph('📋 INFORMACIÓN DEL CURSO')
    info_para.runs[0].font.bold = True
    info_para.runs[0].font.size = Pt(14)
    
    doc.add_paragraph('• 8 Módulos Principales + 12 Módulos Especializados Opcionales')
    doc.add_paragraph('• 100+ Herramientas de IA premium incluidas')
    doc.add_paragraph('• 50+ Casos de Estudio reales por industria')
    doc.add_paragraph('• 5 Especializaciones por industria')
    doc.add_paragraph('• Certificaciones verificables en blockchain')
    doc.add_paragraph('• Valor total: $15,000+ en herramientas y recursos')
    doc.add_paragraph('• Precio del curso: $497 (97% de descuento)')
    
    doc.add_paragraph()
    doc.add_paragraph('─' * 50)
    doc.add_paragraph()
    
    # Procesar contenido del primer archivo
    print("📝 Procesando estructura del curso...")
    process_markdown_content(doc, content1)
    
    # Añadir separador entre secciones
    doc.add_paragraph()
    doc.add_paragraph('─' * 50)
    doc.add_paragraph()
    
    # Procesar contenido del segundo archivo
    print("📝 Procesando recursos y marketing...")
    process_markdown_content(doc, content2)
    
    # Añadir pie de página
    doc.add_paragraph()
    doc.add_paragraph('─' * 50)
    doc.add_paragraph()
    
    footer_para = doc.add_paragraph('© 2024 - Blatam AI Marketing. Todos los derechos reservados.')
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.runs[0].font.italic = True
    
    # Guardar documento
    try:
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Error guardando documento Word: {e}")
        return False

def main():
    """Función principal"""
    print("🚀 Generando documento Word del Curso Premium de IA en Marketing...")
    
    # Rutas de archivos
    base_dir = Path(".")
    file1 = base_dir / "AI_Marketing_Course_Complete_Structure.md"
    file2 = base_dir / "AI_Marketing_Course_Resources_Delivery_Marketing_Feedback.md"
    output_file = base_dir / "Curso_Premium_IA_Marketing_Completo.docx"
    
    # Verificar que los archivos existen
    if not file1.exists():
        print(f"❌ Error: No se encontró {file1}")
        return False
    
    if not file2.exists():
        print(f"❌ Error: No se encontró {file2}")
        return False
    
    print("📖 Leyendo archivos Markdown...")
    
    # Leer contenido de ambos archivos
    content1 = read_md_file(file1)
    content2 = read_md_file(file2)
    
    if not content1 or not content2:
        print("❌ Error: No se pudo leer el contenido de los archivos")
        return False
    
    # Crear documento Word
    print("📄 Creando documento Word...")
    if create_word_document(content1, content2, output_file):
        print(f"✅ Documento Word generado exitosamente: {output_file}")
        if output_file.exists():
            file_size = output_file.stat().st_size / 1024 / 1024
            print(f"📊 Tamaño del archivo: {file_size:.2f} MB")
        return True
    else:
        print("❌ Error generando documento Word")
        return False

if __name__ == "__main__":
    try:
        success = main()
        if success:
            print("\n🎉 ¡Documento Word generado exitosamente!")
            print("📧 El archivo está listo para compartir y distribuir.")
            print("💡 Puedes abrirlo con Microsoft Word, Google Docs, o LibreOffice.")
        else:
            print("\n💥 Error en la generación del documento Word")
            sys.exit(1)
    except ImportError as e:
        print(f"❌ Error: Falta instalar una dependencia: {e}")
        print("💡 Instala las dependencias con: pip install python-docx")
        sys.exit(1)
    except Exception as e:
        print(f"❌ Error inesperado: {e}")
        sys.exit(1)
