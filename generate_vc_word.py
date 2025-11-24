import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_word_document(filename):
    doc = docx.Document()

    # --- Styles ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    heading_style = doc.styles['Heading 1']
    heading_font = heading_style.font
    heading_font.name = 'Calibri'
    heading_font.size = Pt(24)
    heading_font.color.rgb = RGBColor(31, 78, 120) # Dark Blue

    heading2_style = doc.styles['Heading 2']
    heading2_font = heading2_style.font
    heading2_font.name = 'Calibri'
    heading2_font.size = Pt(16)
    heading2_font.color.rgb = RGBColor(47, 117, 181) # Lighter Blue

    # --- Title Page ---
    doc.add_paragraph('\n\n\n\n')
    title = doc.add_paragraph('Ventura Capital\nDocumentación Legal y Estratégica')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style = 'Title'
    
    subtitle = doc.add_paragraph('\nSaaS IA Marketing/Copywriting\nLATAM Market Entry')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.style = 'Subtitle'
    
    doc.add_page_break()

    # --- Section 1: Due Diligence & Risk Analysis ---
    doc.add_heading('1. Due Diligence & Risk Analysis', 0)
    doc.add_paragraph('Este análisis cubre los riesgos financieros, técnicos y legales identificados durante el proceso de Due Diligence.')

    risks = [
        ("Riesgos Financieros", [
            "Liquidez: Falta de flujo de caja positivo. Mitigación: Línea de crédito $500K, reducción de costos.",
            "Crecimiento: Crecimiento lento. Mitigación: Diversificación de canales.",
            "Competencia: Entrada de competidores fuertes. Mitigación: Diferenciación de producto."
        ]),
        ("Riesgos Técnicos", [
            "Tecnología: Obsolescencia. Mitigación: Inversión continua en R&D.",
            "Seguridad: Brechas de seguridad. Mitigación: Auditorías regulares, Seguros.",
            "Escalabilidad: Limitaciones. Mitigación: Arquitectura escalable."
        ]),
        ("Riesgos Legales", [
            "Regulatorios: Cambios en regulaciones. Mitigación: Compliance proactivo.",
            "Propiedad Intelectual: Violación de IP. Mitigación: Due diligence, Patentes."
        ])
    ]

    for category, items in risks:
        doc.add_heading(category, level=2)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # --- Section 2: Legal Templates ---
    doc.add_heading('2. Templates Legales VC', 0)
    
    # Term Sheet
    doc.add_heading('2.1 Investment Term Sheet', level=2)
    doc.add_paragraph('Investment Amount: $2,000,000 | Pre-money Valuation: $8,000,000')
    
    doc.add_heading('Investment Terms', level=3)
    terms = [
        "Security Type: Series A Preferred Stock",
        "Liquidation Preference: 1x participating",
        "Anti-dilution: Weighted average broad-based",
        "Board of Directors: 5 members (2 Founders, 2 Investors, 1 Independent)",
        "Vesting: 4 years, 1 year cliff"
    ]
    for term in terms:
        doc.add_paragraph(term, style='List Bullet')

    # Board Resolution
    doc.add_heading('2.2 Board Resolution Template', level=2)
    p = doc.add_paragraph()
    p.add_run("RESOLVED, that:").bold = True
    doc.add_paragraph("1. Investment Approval: The Company hereby approves the Series A investment of $2,000,000.", style='List Number')
    doc.add_paragraph("2. Board Composition: The Board shall consist of 5 members.", style='List Number')
    doc.add_paragraph("3. Use of Proceeds: Product Development (40%), Marketing (35%), Team (20%).", style='List Number')

    doc.add_page_break()

    # --- Section 3: Negotiation Scripts ---
    doc.add_heading('3. Scripts de Negociación', 0)
    
    doc.add_heading('3.1 Apertura de Reunión', level=2)
    p = doc.add_paragraph("Enfoque en Oportunidad:")
    p.style = 'Heading 3'
    script_box = doc.add_paragraph()
    run = script_box.add_run('"Gracias por su interés. Estamos construyendo la plataforma de IA copywriting líder en LATAM, un mercado de $2.8B con solo 2 competidores principales. Nuestro MRR ha crecido 20% mensual. Estamos buscando un partner estratégico."')
    run.italic = True
    
    doc.add_heading('3.2 Negociación de Valuación', level=2)
    p = doc.add_paragraph("Justificación con Comparables:")
    p.style = 'Heading 3'
    script_box = doc.add_paragraph()
    run = script_box.add_run('"Mirando los comparables del mercado (Copy.ai @ 24x, Jasper @ 21x), y ajustando por mercado LATAM y etapa temprana, creemos que 12x nuestro ARR proyectado es razonable."')
    run.italic = True

    doc.add_heading('3.3 Manejo de Objeciones', level=2)
    doc.add_paragraph("Objeción: 'La valuación es muy alta'", style='Heading 3')
    doc.add_paragraph('"Entiendo su preocupación. Permíteme explicar por qué creemos que es justa: 1) Crecimiento mensual 20%, 2) LTV/CAC 9:1, 3) Churn 5%. ¿Qué múltiplo considera apropiado con estas métricas?"')

    doc.add_heading('3.4 Cierre', level=2)
    doc.add_paragraph('"Perfecto, hemos llegado a un acuerdo que funciona para ambas partes. ¿Están listos para proceder con la documentación legal? Podemos tener el term sheet listo para mañana."')

    doc.save(filename)
    print(f"Word document created: {filename}")

if __name__ == "__main__":
    create_word_document("Ventura_Capital_Documentacion.docx")


