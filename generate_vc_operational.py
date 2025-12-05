import docx
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_case_studies(filename):
    doc = docx.Document()
    
    # Styles
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(16)
    h1.font.color.rgb = RGBColor(32, 55, 100)
    h1.font.bold = True
    
    # Header
    doc.add_paragraph("VENTURA CAPITAL", style='Title').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("CUSTOMER SUCCESS STORIES & ROI REPORT", style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("_" * 80).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Case Study 1: Marketing Agency
    doc.add_heading("CASE STUDY 1: Agencia Digital 'Creativa MX'", 1)
    
    doc.add_paragraph("Client Profile:", style='Heading 3')
    doc.add_paragraph("Leading digital agency in Mexico City with 50+ employees managing 20 enterprise accounts.")
    
    doc.add_paragraph("The Challenge:", style='Heading 3')
    doc.add_paragraph("High churn due to slow content turnaround (5 days avg) and inconsistent brand voice across junior copywriters.")
    
    doc.add_paragraph("The Solution:", style='Heading 3')
    doc.add_paragraph("Implemented Ventura AI 'Enterprise Workflows' to automate first-draft generation for social media calendars.")
    
    doc.add_paragraph("Key Results (validated 6 months):", style='Heading 3')
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run("40% Reduction").bold = True
    p.add_run(" in production time (5 days -> 3 days).")
    
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run("2x Volume").bold = True
    p.add_run(" of assets produced without adding headcount.")
    
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run("$120k Annual Savings").bold = True
    p.add_run(" in freelance costs.")

    doc.add_paragraph("_" * 40).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Case Study 2: E-commerce Brand
    doc.add_heading("CASE STUDY 2: 'ModaLatina' (E-commerce)", 1)
    
    doc.add_paragraph("Client Profile:", style='Heading 3')
    doc.add_paragraph("Regional fashion retailer scaling to Brazil and Colombia.")
    
    doc.add_paragraph("The Challenge:", style='Heading 3')
    doc.add_paragraph("Need for thousands of unique product descriptions localized for distinct Spanish dialects and Portuguese.")
    
    doc.add_paragraph("The Solution:", style='Heading 3')
    doc.add_paragraph("Used Ventura AI's 'Bulk Product Generator' with dialect-specific fine-tuning.")
    
    doc.add_paragraph("Key Results:", style='Heading 3')
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run("30% Uplift in SEO Traffic").bold = True
    p.add_run(" due to unique, non-duplicate content.")
    
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run("95% Faster Time-to-Market").bold = True
    p.add_run(" for new seasonal collections.")

    doc.save(filename)
    print(f"Case Studies created: {filename}")

def create_esg_policy(filename):
    doc = docx.Document()
    
    doc.add_paragraph("VENTURA CAPITAL", style='Title').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("AI ETHICS, GOVERNANCE & ESG POLICY", style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Effective Date: December 2024").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. Responsible AI
    doc.add_heading("1. RESPONSIBLE AI PRINCIPLES", 1)
    doc.add_paragraph("Ventura Capital is committed to developing AI systems that are safe, fair, and transparent.")
    
    doc.add_paragraph("Bias Mitigation:", style='Heading 3')
    doc.add_paragraph("We actively monitor our training datasets to ensure fair representation of Latin American diversity (indigenous languages, gender neutrality where appropriate).")
    
    doc.add_paragraph("Content Safety:", style='Heading 3')
    doc.add_paragraph("Implementation of 'Guardrails' to prevent generation of hate speech, misinformation, or adult content. Manual review of <1% of samples for QA.")

    # 2. Data Privacy
    doc.add_heading("2. DATA PRIVACY & SECURITY (GDPR/LGPD)", 1)
    doc.add_paragraph("Compliance with Brazil's LGPD and Mexico's LFPDPPP.")
    doc.add_paragraph("• Right to be Forgotten: Automated workflows for user data deletion.")
    doc.add_paragraph("• Data Sovereignty: Enterprise data stored in region-appropriate AWS availability zones where required.")

    # 3. Environmental Impact (ESG)
    doc.add_heading("3. ENVIRONMENTAL SUSTAINABILITY", 1)
    doc.add_paragraph("• Green Compute: We prioritize training on carbon-neutral cloud regions.")
    doc.add_paragraph("• Efficiency: Our proprietary 'Small Language Models' (SLMs) require 50% less energy for inference than generalist models like GPT-4.")

    doc.save(filename)
    print(f"ESG Policy created: {filename}")

if __name__ == "__main__":
    create_case_studies("Ventura_Capital_Case_Studies.docx")
    create_esg_policy("Ventura_Capital_ESG_Policy.docx")







