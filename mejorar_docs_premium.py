#!/usr/bin/env python3
"""
Script para MEJORAR documentos ya generados con:
- Gráficas más profesionales y detalladas
- Análisis más profundos
- Formato mejorado
- Dashboards interactivos
- Visualizaciones avanzadas
"""

import os
import re
import json
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any
import markdown
from collections import Counter
import numpy as np

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False
    print("⚠️  python-docx no disponible")

try:
    from openpyxl import Workbook, load_workbook
    from openpyxl.chart import BarChart, LineChart, PieChart, Reference, ScatterChart, AreaChart
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.drawing.image import Image as XLImage
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    print("⚠️  openpyxl no disponible")

try:
    import matplotlib.pyplot as plt
    import matplotlib
    matplotlib.use('Agg')
    import seaborn as sns
    sns.set_style("whitegrid")
    sns.set_palette("husl")
    MATPLOTLIB_AVAILABLE = True
except ImportError:
    MATPLOTLIB_AVAILABLE = False
    print("⚠️  matplotlib no disponible")

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image as RLImage
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("⚠️  reportlab no disponible")


class DocumentEnhancer:
    """Mejora documentos existentes con análisis avanzados y visualizaciones premium."""
    
    def __init__(self, output_dir="docs_premium_mejorados"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        self.charts_dir = self.output_dir / "charts_mejorados"
        self.charts_dir.mkdir(exist_ok=True)
        
    def read_markdown(self, file_path: Path) -> str:
        """Lee un archivo markdown."""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def analyze_content_deep(self, content: str) -> Dict[str, Any]:
        """Análisis profundo del contenido."""
        lines = content.split('\n')
        words = content.split()
        
        # Análisis de estructura
        headers = [line for line in lines if line.strip().startswith('#')]
        code_blocks = len(re.findall(r'```[\s\S]*?```', content))
        tables = len(re.findall(r'\|.*\|', content))
        links = len(re.findall(r'\[.*?\]\(.*?\)', content))
        images = len(re.findall(r'!\[.*?\]\(.*?\)', content))
        
        # Análisis de complejidad
        avg_words_per_line = len(words) / max(len(lines), 1)
        avg_chars_per_word = sum(len(w) for w in words) / max(len(words), 1)
        
        # Análisis de temas
        words_lower = [w.lower() for w in words if len(w) > 3]
        word_freq = Counter(words_lower)
        top_words = dict(word_freq.most_common(20))
        
        # Análisis de secciones
        section_levels = [len(h) - len(h.lstrip('#')) for h in headers]
        max_depth = max(section_levels) if section_levels else 0
        
        return {
            'total_lines': len(lines),
            'total_words': len(words),
            'total_chars': len(content),
            'headers_count': len(headers),
            'code_blocks': code_blocks,
            'tables': tables,
            'links': links,
            'images': images,
            'avg_words_per_line': round(avg_words_per_line, 2),
            'avg_chars_per_word': round(avg_chars_per_word, 2),
            'top_words': top_words,
            'max_depth': max_depth,
            'reading_time_minutes': round(len(words) / 200, 1)  # 200 palabras/minuto
        }
    
    def create_enhanced_charts(self, analysis: Dict[str, Any], doc_name: str) -> List[str]:
        """Crea gráficas mejoradas y profesionales."""
        if not MATPLOTLIB_AVAILABLE:
            return []
        
        chart_files = []
        
        # 1. Dashboard completo mejorado
        fig = plt.figure(figsize=(16, 10))
        gs = fig.add_gridspec(3, 3, hspace=0.3, wspace=0.3)
        
        # Gráfica 1: Métricas principales (arriba izquierda)
        ax1 = fig.add_subplot(gs[0, 0])
        metrics = ['Líneas', 'Palabras', 'Encabezados', 'Código']
        values = [
            analysis['total_lines'] / 100,
            analysis['total_words'] / 1000,
            analysis['headers_count'],
            analysis['code_blocks']
        ]
        bars = ax1.bar(metrics, values, color=['#2ecc71', '#3498db', '#9b59b6', '#e74c3c'])
        ax1.set_title('Métricas Principales', fontsize=12, fontweight='bold')
        ax1.set_ylabel('Valor Normalizado')
        for bar, val in zip(bars, values):
            height = bar.get_height()
            ax1.text(bar.get_x() + bar.get_width()/2., height,
                   f'{val:.1f}', ha='center', va='bottom', fontsize=9)
        
        # Gráfica 2: Distribución de contenido (arriba centro)
        ax2 = fig.add_subplot(gs[0, 1])
        content_types = ['Texto', 'Código', 'Tablas', 'Enlaces', 'Imágenes']
        content_values = [
            analysis['total_words'],
            analysis['code_blocks'] * 50,  # Estimado
            analysis['tables'] * 20,
            analysis['links'],
            analysis['images'] * 10
        ]
        colors_pie = ['#3498db', '#e74c3c', '#2ecc71', '#f39c12', '#9b59b6']
        wedges, texts, autotexts = ax2.pie(content_values, labels=content_types, 
                                          autopct='%1.1f%%', colors=colors_pie, startangle=90)
        ax2.set_title('Distribución de Contenido', fontsize=12, fontweight='bold')
        
        # Gráfica 3: Top palabras (arriba derecha)
        ax3 = fig.add_subplot(gs[0, 2])
        top_10_words = dict(list(analysis['top_words'].items())[:10])
        words_list = list(top_10_words.keys())
        counts = list(top_10_words.values())
        ax3.barh(words_list, counts, color='#3498db')
        ax3.set_title('Top 10 Palabras', fontsize=12, fontweight='bold')
        ax3.set_xlabel('Frecuencia')
        
        # Gráfica 4: Complejidad (centro izquierda)
        ax4 = fig.add_subplot(gs[1, 0])
        complexity_metrics = ['Profundidad', 'Prom. Palabras/Línea', 'Prom. Caracteres/Palabra']
        complexity_values = [
            analysis['max_depth'],
            analysis['avg_words_per_line'],
            analysis['avg_chars_per_word']
        ]
        ax4.bar(complexity_metrics, complexity_values, color=['#e74c3c', '#f39c12', '#2ecc71'])
        ax4.set_title('Métricas de Complejidad', fontsize=12, fontweight='bold')
        ax4.set_ylabel('Valor')
        plt.setp(ax4.xaxis.get_majorticklabels(), rotation=45, ha='right')
        
        # Gráfica 5: Tiempo de lectura (centro)
        ax5 = fig.add_subplot(gs[1, 1])
        reading_time = analysis['reading_time_minutes']
        categories = ['Lectura Rápida\n(200 wpm)', 'Lectura Normal\n(150 wpm)', 'Lectura Detallada\n(100 wpm)']
        times = [reading_time, reading_time * 1.33, reading_time * 2]
        bars = ax5.bar(categories, times, color=['#2ecc71', '#3498db', '#e74c3c'])
        ax5.set_title('Tiempo de Lectura Estimado', fontsize=12, fontweight='bold')
        ax5.set_ylabel('Minutos')
        for bar, time in zip(bars, times):
            height = bar.get_height()
            ax5.text(bar.get_x() + bar.get_width()/2., height,
                   f'{time:.1f}m', ha='center', va='bottom', fontsize=9)
        
        # Gráfica 6: Análisis de estructura (centro derecha)
        ax6 = fig.add_subplot(gs[1, 2])
        structure_data = {
            'Encabezados': analysis['headers_count'],
            'Bloques Código': analysis['code_blocks'],
            'Tablas': analysis['tables'],
            'Enlaces': analysis['links']
        }
        ax6.bar(structure_data.keys(), structure_data.values(), color='#9b59b6')
        ax6.set_title('Elementos Estructurales', fontsize=12, fontweight='bold')
        ax6.set_ylabel('Cantidad')
        plt.setp(ax6.xaxis.get_majorticklabels(), rotation=45, ha='right')
        
        # Gráfica 7: Resumen visual grande (abajo, ancho completo)
        ax7 = fig.add_subplot(gs[2, :])
        summary_labels = [
            f"📄 {analysis['total_lines']:,} líneas",
            f"📝 {analysis['total_words']:,} palabras",
            f"⏱️ {analysis['reading_time_minutes']:.1f} min lectura",
            f"📊 {analysis['headers_count']} secciones",
            f"💻 {analysis['code_blocks']} bloques código",
            f"📋 {analysis['tables']} tablas"
        ]
        summary_values = [1] * len(summary_labels)
        bars = ax7.barh(summary_labels, summary_values, 
                       color=['#2ecc71', '#3498db', '#f39c12', '#9b59b6', '#e74c3c', '#1abc9c'])
        ax7.set_title(f'Resumen Ejecutivo: {doc_name}', fontsize=14, fontweight='bold', pad=20)
        ax7.set_xlim(0, 1.2)
        ax7.axis('off')
        for i, (label, bar) in enumerate(zip(summary_labels, bars)):
            ax7.text(0.6, bar.get_y() + bar.get_height()/2, label,
                    ha='center', va='center', fontsize=11, fontweight='bold')
        
        plt.suptitle(f'Dashboard Mejorado: {doc_name}', fontsize=16, fontweight='bold', y=0.98)
        dashboard_path = self.charts_dir / f"{doc_name}_dashboard_mejorado.png"
        plt.savefig(dashboard_path, dpi=300, bbox_inches='tight', facecolor='white')
        plt.close()
        chart_files.append(str(dashboard_path))
        
        # 2. Gráfica de análisis de palabras mejorada
        if len(analysis['top_words']) > 0:
            fig, ax = plt.subplots(figsize=(12, 8))
            top_20 = dict(list(analysis['top_words'].items())[:20])
            words = list(top_20.keys())
            counts = list(top_20.values())
            
            bars = ax.barh(words, counts, color=plt.cm.viridis(np.linspace(0, 1, len(words))))
            ax.set_title('Análisis de Palabras Clave (Top 20)', fontsize=14, fontweight='bold')
            ax.set_xlabel('Frecuencia', fontsize=12)
            ax.set_ylabel('Palabra', fontsize=12)
            ax.grid(axis='x', alpha=0.3)
            
            for i, (word, count) in enumerate(zip(words, counts)):
                ax.text(count, i, f' {count}', va='center', fontsize=9)
            
            plt.tight_layout()
            words_path = self.charts_dir / f"{doc_name}_palabras_mejorado.png"
            plt.savefig(words_path, dpi=300, bbox_inches='tight', facecolor='white')
            plt.close()
            chart_files.append(str(words_path))
        
        return chart_files
    
    def enhance_word_document(self, original_path: Path, analysis: Dict[str, Any], charts: List[str]) -> Path:
        """Mejora un documento Word existente o crea uno nuevo mejorado."""
        if not WORD_AVAILABLE:
            return None
        
        doc = Document()
        
        # Portada mejorada
        title = doc.add_heading(original_path.stem.replace('_', ' ').title(), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(32)
        title_run.font.color.rgb = RGBColor(31, 78, 120)
        title_run.font.bold = True
        
        # Información del documento
        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info.add_run(f"Versión Mejorada Premium\n").font.size = Pt(14)
        info.add_run(f"Generado: {datetime.now().strftime('%d de %B de %Y a las %H:%M')}\n").font.size = Pt(11)
        info.add_run(f"Análisis Profundo con Visualizaciones Avanzadas").font.size = Pt(11)
        info.runs[1].font.color.rgb = RGBColor(100, 100, 100)
        
        doc.add_page_break()
        
        # Resumen ejecutivo
        doc.add_heading('📊 Resumen Ejecutivo', level=1)
        summary_para = doc.add_paragraph()
        summary_para.add_run(f"Este documento contiene ").font.size = Pt(11)
        summary_para.add_run(f"{analysis['total_words']:,} palabras").font.bold = True
        summary_para.add_run(f" distribuidas en ").font.size = Pt(11)
        summary_para.add_run(f"{analysis['total_lines']:,} líneas").font.bold = True
        summary_para.add_run(f". ").font.size = Pt(11)
        summary_para.add_run(f"Tiempo estimado de lectura: {analysis['reading_time_minutes']:.1f} minutos").font.bold = True
        summary_para.add_run(f".").font.size = Pt(11)
        
        # Métricas detalladas
        doc.add_heading('📈 Métricas Detalladas', level=1)
        metrics_table = doc.add_table(rows=1, cols=2)
        metrics_table.style = 'Light Grid Accent 1'
        hdr_cells = metrics_table.rows[0].cells
        hdr_cells[0].text = 'Métrica'
        hdr_cells[1].text = 'Valor'
        hdr_cells[0].paragraphs[0].runs[0].font.bold = True
        hdr_cells[1].paragraphs[0].runs[0].font.bold = True
        
        metrics_data = [
            ('Total de Líneas', f"{analysis['total_lines']:,}"),
            ('Total de Palabras', f"{analysis['total_words']:,}"),
            ('Total de Caracteres', f"{analysis['total_chars']:,}"),
            ('Encabezados', str(analysis['headers_count'])),
            ('Bloques de Código', str(analysis['code_blocks'])),
            ('Tablas', str(analysis['tables'])),
            ('Enlaces', str(analysis['links'])),
            ('Imágenes', str(analysis['images'])),
            ('Profundidad Máxima', str(analysis['max_depth'])),
            ('Promedio Palabras/Línea', f"{analysis['avg_words_per_line']:.2f}"),
            ('Promedio Caracteres/Palabra', f"{analysis['avg_chars_per_word']:.2f}"),
            ('Tiempo de Lectura (min)', f"{analysis['reading_time_minutes']:.1f}")
        ]
        
        for metric, value in metrics_data:
            row_cells = metrics_table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = value
        
        # Agregar gráficas
        if charts:
            doc.add_page_break()
            doc.add_heading('📊 Visualizaciones Mejoradas', level=1)
            for chart_path in charts:
                if Path(chart_path).exists():
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run()
                    run.add_picture(chart_path, width=Inches(6.5))
                    doc.add_paragraph()  # Espacio
        
        # Top palabras
        if analysis['top_words']:
            doc.add_page_break()
            doc.add_heading('🔤 Palabras Más Frecuentes', level=1)
            words_table = doc.add_table(rows=1, cols=2)
            words_table.style = 'Light Grid Accent 1'
            hdr_cells = words_table.rows[0].cells
            hdr_cells[0].text = 'Palabra'
            hdr_cells[1].text = 'Frecuencia'
            hdr_cells[0].paragraphs[0].runs[0].font.bold = True
            hdr_cells[1].paragraphs[0].runs[0].font.bold = True
            
            for word, count in list(analysis['top_words'].items())[:20]:
                row_cells = words_table.add_row().cells
                row_cells[0].text = word
                row_cells[1].text = str(count)
        
        # Guardar
        output_path = self.output_dir / f"{original_path.stem}_MEJORADO.docx"
        doc.save(str(output_path))
        return output_path
    
    def enhance_excel_document(self, original_path: Path, analysis: Dict[str, Any], charts: List[str]) -> Path:
        """Mejora un documento Excel existente o crea uno nuevo mejorado."""
        if not EXCEL_AVAILABLE:
            return None
        
        wb = Workbook()
        
        # Hoja 1: Dashboard
        ws1 = wb.active
        ws1.title = "Dashboard"
        
        # Título
        ws1['A1'] = f"Dashboard Mejorado: {original_path.stem}"
        ws1['A1'].font = Font(size=18, bold=True, color="1F4E78")
        ws1.merge_cells('A1:D1')
        
        # Métricas principales
        row = 3
        metrics = [
            ('Total Líneas', analysis['total_lines']),
            ('Total Palabras', analysis['total_words']),
            ('Total Caracteres', analysis['total_chars']),
            ('Encabezados', analysis['headers_count']),
            ('Bloques Código', analysis['code_blocks']),
            ('Tablas', analysis['tables']),
            ('Enlaces', analysis['links']),
            ('Imágenes', analysis['images']),
            ('Tiempo Lectura (min)', analysis['reading_time_minutes'])
        ]
        
        ws1['A3'] = 'Métrica'
        ws1['B3'] = 'Valor'
        ws1['A3'].font = Font(bold=True)
        ws1['B3'].font = Font(bold=True)
        
        for metric, value in metrics:
            row += 1
            ws1[f'A{row}'] = metric
            ws1[f'B{row}'] = value
        
        # Hoja 2: Análisis de Palabras
        ws2 = wb.create_sheet("Palabras Clave")
        ws2['A1'] = 'Palabra'
        ws2['B1'] = 'Frecuencia'
        ws2['A1'].font = Font(bold=True)
        ws2['B1'].font = Font(bold=True)
        
        row = 2
        for word, count in list(analysis['top_words'].items())[:50]:
            ws2[f'A{row}'] = word
            ws2[f'B{row}'] = count
            row += 1
        
        # Crear gráfica de barras
        chart = BarChart()
        chart.type = "col"
        chart.style = 10
        chart.title = "Top 10 Palabras Más Frecuentes"
        chart.y_axis.title = 'Frecuencia'
        chart.x_axis.title = 'Palabra'
        
        data = Reference(ws2, min_col=2, min_row=1, max_row=11)
        cats = Reference(ws2, min_col=1, min_row=2, max_row=11)
        chart.add_data(data, titles_from_data=True)
        chart.set_categories(cats)
        ws2.add_chart(chart, "D2")
        
        # Hoja 3: Estadísticas
        ws3 = wb.create_sheet("Estadísticas")
        ws3['A1'] = 'Estadística'
        ws3['B1'] = 'Valor'
        ws3['A1'].font = Font(bold=True)
        ws3['B1'].font = Font(bold=True)
        
        stats = [
            ('Promedio Palabras/Línea', analysis['avg_words_per_line']),
            ('Promedio Caracteres/Palabra', analysis['avg_chars_per_word']),
            ('Profundidad Máxima', analysis['max_depth']),
            ('Densidad de Código', f"{analysis['code_blocks'] / max(analysis['headers_count'], 1):.2f}"),
            ('Densidad de Tablas', f"{analysis['tables'] / max(analysis['headers_count'], 1):.2f}")
        ]
        
        row = 2
        for stat, value in stats:
            ws3[f'A{row}'] = stat
            ws3[f'B{row}'] = value
            row += 1
        
        # Guardar
        output_path = self.output_dir / f"{original_path.stem}_MEJORADO.xlsx"
        wb.save(str(output_path))
        return output_path
    
    def process_document(self, md_file: Path):
        """Procesa y mejora un documento."""
        print(f"\n🔄 Mejorando: {md_file.name}")
        print("-" * 70)
        
        try:
            # Leer y analizar
            content = self.read_markdown(md_file)
            analysis = self.analyze_content_deep(content)
            
            # Crear gráficas mejoradas
            print("  📊 Creando gráficas mejoradas...")
            charts = self.create_enhanced_charts(analysis, md_file.stem)
            print(f"  ✅ {len(charts)} gráficas creadas")
            
            # Mejorar Word
            if WORD_AVAILABLE:
                print("  📝 Mejorando documento Word...")
                word_path = self.enhance_word_document(md_file, analysis, charts)
                if word_path:
                    print(f"  ✅ Word mejorado: {word_path.name}")
            
            # Mejorar Excel
            if EXCEL_AVAILABLE:
                print("  📊 Mejorando libro Excel...")
                excel_path = self.enhance_excel_document(md_file, analysis, charts)
                if excel_path:
                    print(f"  ✅ Excel mejorado: {excel_path.name}")
            
            return True
            
        except Exception as e:
            print(f"  ❌ Error: {str(e)}")
            import traceback
            traceback.print_exc()
            return False


def main():
    """Función principal."""
    print("=" * 70)
    print("🚀 MEJORADOR DE DOCUMENTOS PREMIUM")
    print("=" * 70)
    print()
    
    base_dir = Path("/Users/adan/Documents/documentos_blatam")
    enhancer = DocumentEnhancer()
    
    # Documentos importantes a mejorar
    important_files = [
        base_dir / "airflow_automation_prompt.md",
        base_dir / "README.md",
        base_dir / "ARCHITECTURE.md",
        base_dir / "BEST_PRACTICES.md",
        base_dir / "truthgpt_collected/integration_code/production_code/ARCHITECTURE_IMPROVEMENTS.md",
        base_dir / "truthgpt_collected/integration_code/production_code/REFACTORING_PLAN.md",
    ]
    
    processed = 0
    for md_file in important_files:
        if md_file.exists():
            if enhancer.process_document(md_file):
                processed += 1
        else:
            print(f"⚠️  No encontrado: {md_file}")
    
    print("\n" + "=" * 70)
    print(f"✨ Procesados {processed} documentos mejorados")
    print(f"📁 Archivos guardados en: {enhancer.output_dir.absolute()}")
    print("=" * 70)


if __name__ == "__main__":
    main()

