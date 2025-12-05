#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script para generar documentos Word y PDF de respuestas para compañeros del foro
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT
from reportlab.lib.colors import HexColor, black

def create_word_responses():
    """Crea el documento Word con las respuestas"""
    doc = Document()
    
    # Configurar estilos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Título
    title = doc.add_heading('Respuestas para Compañeros en el Foro de Discusión', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('FASE: DESARROLLO - Comentarios a Compañeros')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.bold = True
    subtitle_format.italic = True
    doc.add_paragraph('')
    
    # Respuesta 1
    doc.add_heading('Respuesta 1: Comentario sobre la importancia de la inclusión financiera', level=1)
    
    saludo1 = doc.add_paragraph()
    saludo1.add_run('Hola [Nombre del Compañero 1],')
    saludo1_format = saludo1.runs[0]
    saludo1_format.bold = True
    
    p1 = doc.add_paragraph('Me pareció muy acertado tu análisis sobre la importancia de los mercados financieros, especialmente tu énfasis en la función de asignación de recursos. Coincido completamente contigo en que esta función es fundamental para el crecimiento económico.')
    
    p2 = doc.add_paragraph('Quisiera complementar tu aportación agregando una perspectiva adicional que considero relevante: la ')
    p2.add_run('inclusión financiera').bold = True
    p2.add_run(' como dimensión crítica de los mercados financieros. Aunque mencionaste la importancia de la canalización de recursos, creo que es importante destacar que esta función solo puede ser verdaderamente efectiva cuando los mercados financieros son accesibles para todos los segmentos de la población.')
    
    p3 = doc.add_paragraph('En el contexto mexicano, según datos de la Encuesta Nacional de Inclusión Financiera (ENIF) 2021, aproximadamente el 32% de la población adulta aún no tiene acceso a ningún producto financiero formal. Esta exclusión financiera limita significativamente el potencial de los mercados financieros para contribuir al desarrollo económico, ya que grandes segmentos de la población no pueden participar en el proceso de asignación de recursos.')
    
    p4 = doc.add_paragraph('Tu punto sobre la canalización de recursos hacia sectores estratégicos es muy válido, y creo que se complementa con la necesidad de ampliar el acceso a servicios financieros. Cuando más personas y empresas pueden acceder a servicios financieros, el proceso de asignación de recursos se vuelve más eficiente y equitativo, permitiendo que el capital fluya no solo hacia sectores estratégicos tradicionales, sino también hacia emprendimientos, pequeñas empresas y proyectos que tradicionalmente han tenido acceso limitado al financiamiento.')
    
    p5 = doc.add_paragraph('Además, la inclusión financiera puede tener efectos multiplicadores importantes. Cuando las personas tienen acceso a servicios financieros, pueden gestionar mejor sus recursos, invertir en educación y salud, y protegerse contra shocks económicos, lo que a su vez contribuye al crecimiento económico y la reducción de la desigualdad.')
    
    pregunta1 = doc.add_paragraph()
    pregunta1.add_run('¿Qué opinas sobre las políticas que podrían promover una mayor inclusión financiera sin comprometer la estabilidad del sistema? ').italic = True
    pregunta1.add_run('Me interesa conocer tu perspectiva sobre cómo equilibrar estos objetivos.')
    
    despedida1 = doc.add_paragraph('Saludos cordiales.')
    doc.add_paragraph('')
    
    # Respuesta 2
    doc.add_heading('Respuesta 2: Comentario sobre la efectividad de la regulación y desafíos futuros', level=1)
    
    saludo2 = doc.add_paragraph()
    saludo2.add_run('Hola [Nombre del Compañero 2],')
    saludo2_format = saludo2.runs[0]
    saludo2_format.bold = True
    
    p6 = doc.add_paragraph('Excelente análisis sobre el papel de las entidades reguladoras en México. Tu argumentación sobre las contribuciones positivas del sistema regulatorio es muy sólida y está bien fundamentada. Especialmente valioso fue tu reconocimiento de los avances logrados desde las crisis financieras del pasado.')
    
    p7 = doc.add_paragraph('Quisiera agregar una perspectiva complementaria sobre un desafío que considero particularmente relevante: la ')
    p7.add_run('velocidad de adaptación regulatoria').bold = True
    p7.add_run(' frente a la innovación financiera. Aunque mencionaste este punto, creo que vale la pena profundizar en sus implicaciones.')
    
    p8 = doc.add_paragraph('El ritmo de innovación en el sector financiero se ha acelerado significativamente en los últimos años, con el surgimiento de tecnologías como la inteligencia artificial, el blockchain, y nuevos modelos de negocio basados en plataformas digitales. La Ley Fintech de 2018, que mencionaste, es un ejemplo positivo de adaptación regulatoria, pero el desafío es mantener este ritmo de adaptación.')
    
    p9 = doc.add_paragraph('Un aspecto que me preocupa es el riesgo de que la regulación quede obsoleta o se convierta en un obstáculo para la innovación. Por ejemplo, mientras que la regulación de activos virtuales (criptomonedas) está en desarrollo, estas tecnologías continúan evolucionando rápidamente, y los reguladores pueden enfrentar dificultades para mantenerse al día.')
    
    p10 = doc.add_paragraph('Sin embargo, también veo oportunidades. El "sandbox regulatorio" establecido en la Ley Fintech es una herramienta innovadora que permite a las empresas probar nuevos productos en un entorno controlado, facilitando la innovación mientras se mantiene la protección al consumidor. Este tipo de enfoques flexibles y adaptativos pueden ser clave para el futuro de la regulación financiera.')
    
    p11 = doc.add_paragraph('Tu punto sobre la necesidad de recursos adecuados para la supervisión es muy relevante aquí. La supervisión de nuevas tecnologías y modelos de negocio requiere no solo recursos financieros, sino también capacidades técnicas y conocimiento especializado. ')
    pregunta2 = doc.add_paragraph()
    pregunta2.add_run('¿Crees que el sistema regulatorio mexicano tiene las capacidades necesarias para enfrentar estos desafíos, o se requieren inversiones adicionales en capacitación y tecnología?').italic = True
    
    p12 = doc.add_paragraph('También me interesa tu opinión sobre cómo otros países están enfrentando estos desafíos. ')
    pregunta3 = doc.add_paragraph()
    pregunta3.add_run('¿Hay ejemplos de mejores prácticas internacionales que México podría adoptar?').italic = True
    
    p13 = doc.add_paragraph('Gracias por tu excelente contribución, que ha enriquecido significativamente la discusión.')
    
    despedida2 = doc.add_paragraph('Saludos cordiales.')
    
    # Guardar
    filename = 'Respuestas_Companeros_Foro.docx'
    doc.save(filename)
    print(f"✓ Documento Word creado: {filename}")
    return filename

def create_pdf_responses():
    """Crea el documento PDF con las respuestas"""
    filename = 'Respuestas_Companeros_Foro.pdf'
    doc = SimpleDocTemplate(filename, pagesize=A4,
                            rightMargin=72, leftMargin=72,
                            topMargin=72, bottomMargin=72)
    
    Story = []
    styles = getSampleStyleSheet()
    
    # Estilos personalizados
    styles.add(ParagraphStyle(
        name='TitleStyle',
        parent=styles['Heading1'],
        fontSize=16,
        textColor=HexColor('#1a1a1a'),
        spaceAfter=12,
        alignment=1  # CENTER
    ))
    
    styles.add(ParagraphStyle(
        name='Justify',
        parent=styles['Normal'],
        alignment=TA_JUSTIFY,
        fontSize=11,
        leading=14
    ))
    
    # Título
    Story.append(Paragraph("Respuestas para Compañeros en el Foro de Discusión", styles['TitleStyle']))
    Story.append(Paragraph("<i>FASE: DESARROLLO - Comentarios a Compañeros</i>", styles['Normal']))
    Story.append(Spacer(1, 24))
    
    # Respuesta 1
    Story.append(Paragraph("<b>Respuesta 1: Comentario sobre la importancia de la inclusión financiera</b>", styles['Heading1']))
    Story.append(Paragraph("<b>Hola [Nombre del Compañero 1],</b>", styles['Normal']))
    Story.append(Spacer(1, 12))
    
    Story.append(Paragraph("Me pareció muy acertado tu análisis sobre la importancia de los mercados financieros, especialmente tu énfasis en la función de asignación de recursos. Coincido completamente contigo en que esta función es fundamental para el crecimiento económico.", styles['Justify']))
    Story.append(Spacer(1, 12))
    
    Story.append(Paragraph("Quisiera complementar tu aportación agregando una perspectiva adicional que considero relevante: la <b>inclusión financiera</b> como dimensión crítica de los mercados financieros. Aunque mencionaste la importancia de la canalización de recursos, creo que es importante destacar que esta función solo puede ser verdaderamente efectiva cuando los mercados financieros son accesibles para todos los segmentos de la población.", styles['Justify']))
    Story.append(Spacer(1, 12))
    
    Story.append(Paragraph("En el contexto mexicano, según datos de la Encuesta Nacional de Inclusión Financiera (ENIF) 2021, aproximadamente el 32% de la población adulta aún no tiene acceso a ningún producto financiero formal. Esta exclusión financiera limita significativamente el potencial de los mercados financieros para contribuir al desarrollo económico, ya que grandes segmentos de la población no pueden participar en el proceso de asignación de recursos.", styles['Justify']))
    Story.append(Spacer(1, 12))
    
    Story.append(Paragraph("Tu punto sobre la canalización de recursos hacia sectores estratégicos es muy válido, y creo que se complementa con la necesidad de ampliar el acceso a servicios financieros. Cuando más personas y empresas pueden acceder a servicios financieros, el proceso de asignación de recursos se vuelve más eficiente y equitativo, permitiendo que el capital fluya no solo hacia sectores estratégicos tradicionales, sino también hacia emprendimientos, pequeñas empresas y proyectos que tradicionalmente han tenido acceso limitado al financiamiento.", styles['Justify']))
    Story.append(Spacer(1, 12))
    
    Story.append(Paragraph("Además, la inclusión financiera puede tener efectos multiplicadores importantes. Cuando las personas tienen acceso a servicios financieros, pueden gestionar mejor sus recursos, invertir en educación y salud, y protegerse contra shocks económicos, lo que a su vez contribuye al crecimiento económico y la reducción de la desigualdad.", styles['Justify']))
    Story.append(Spacer(1, 12))
    
    Story.append(Paragraph("<i>¿Qué opinas sobre las políticas que podrían promover una mayor inclusión financiera sin comprometer la estabilidad del sistema? Me interesa conocer tu perspectiva sobre cómo equilibrar estos objetivos.</i>", styles['Justify']))
    Story.append(Spacer(1, 12))
    
    Story.append(Paragraph("Saludos cordiales.", styles['Normal']))
    Story.append(Spacer(1, 24))
    
    # Respuesta 2
    Story.append(Paragraph("<b>Respuesta 2: Comentario sobre la efectividad de la regulación y desafíos futuros</b>", styles['Heading1']))
    Story.append(Paragraph("<b>Hola [Nombre del Compañero 2],</b>", styles['Normal']))
    Story.append(Spacer(1, 12))
    
    Story.append(Paragraph("Excelente análisis sobre el papel de las entidades reguladoras en México. Tu argumentación sobre las contribuciones positivas del sistema regulatorio es muy sólida y está bien fundamentada. Especialmente valioso fue tu reconocimiento de los avances logrados desde las crisis financieras del pasado.", styles['Justify']))
    Story.append(Spacer(1, 12))
    
    Story.append(Paragraph("Quisiera agregar una perspectiva complementaria sobre un desafío que considero particularmente relevante: la <b>velocidad de adaptación regulatoria</b> frente a la innovación financiera. Aunque mencionaste este punto, creo que vale la pena profundizar en sus implicaciones.", styles['Justify']))
    Story.append(Spacer(1, 12))
    
    Story.append(Paragraph("El ritmo de innovación en el sector financiero se ha acelerado significativamente en los últimos años, con el surgimiento de tecnologías como la inteligencia artificial, el blockchain, y nuevos modelos de negocio basados en plataformas digitales. La Ley Fintech de 2018, que mencionaste, es un ejemplo positivo de adaptación regulatoria, pero el desafío es mantener este ritmo de adaptación.", styles['Justify']))
    Story.append(Spacer(1, 12))
    
    Story.append(Paragraph("Un aspecto que me preocupa es el riesgo de que la regulación quede obsoleta o se convierta en un obstáculo para la innovación. Por ejemplo, mientras que la regulación de activos virtuales (criptomonedas) está en desarrollo, estas tecnologías continúan evolucionando rápidamente, y los reguladores pueden enfrentar dificultades para mantenerse al día.", styles['Justify']))
    Story.append(Spacer(1, 12))
    
    Story.append(Paragraph("Sin embargo, también veo oportunidades. El \"sandbox regulatorio\" establecido en la Ley Fintech es una herramienta innovadora que permite a las empresas probar nuevos productos en un entorno controlado, facilitando la innovación mientras se mantiene la protección al consumidor. Este tipo de enfoques flexibles y adaptativos pueden ser clave para el futuro de la regulación financiera.", styles['Justify']))
    Story.append(Spacer(1, 12))
    
    Story.append(Paragraph("Tu punto sobre la necesidad de recursos adecuados para la supervisión es muy relevante aquí. La supervisión de nuevas tecnologías y modelos de negocio requiere no solo recursos financieros, sino también capacidades técnicas y conocimiento especializado. <i>¿Crees que el sistema regulatorio mexicano tiene las capacidades necesarias para enfrentar estos desafíos, o se requieren inversiones adicionales en capacitación y tecnología?</i>", styles['Justify']))
    Story.append(Spacer(1, 12))
    
    Story.append(Paragraph("También me interesa tu opinión sobre cómo otros países están enfrentando estos desafíos. <i>¿Hay ejemplos de mejores prácticas internacionales que México podría adoptar?</i>", styles['Justify']))
    Story.append(Spacer(1, 12))
    
    Story.append(Paragraph("Gracias por tu excelente contribución, que ha enriquecido significativamente la discusión.", styles['Justify']))
    Story.append(Spacer(1, 12))
    
    Story.append(Paragraph("Saludos cordiales.", styles['Normal']))
    
    doc.build(Story)
    print(f"✓ Documento PDF creado: {filename}")
    return filename

if __name__ == "__main__":
    print("Generando respuestas para compañeros del foro...")
    print("-" * 60)
    try:
        word_file = create_word_responses()
    except Exception as e:
        print(f"✗ Error al crear Word: {e}")
        word_file = None
    
    try:
        pdf_file = create_pdf_responses()
    except Exception as e:
        print(f"✗ Error al crear PDF: {e}")
        pdf_file = None
    
    print("-" * 60)
    if word_file:
        print(f"✓ Word: {word_file}")
    if pdf_file:
        print(f"✓ PDF: {pdf_file}")
    print("\n¡Documentos generados exitosamente!")







