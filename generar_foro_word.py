from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_forum_word(filename):
    doc = Document()
    
    # Title
    head = doc.add_heading('Participación en Foro: Protección de Equipo de Cómputo', 0)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph('Alumno: Adán Muñoz Pablo')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')

    # Main Contribution
    doc.add_heading('Aportación Principal', level=1)
    doc.add_paragraph('Asunto: 5 Estrategias Integrales para la Protección de Equipos de Cómputo', style='Intense Quote')
    
    doc.add_paragraph('Hola a todos,')
    doc.add_paragraph('Después de analizar las mejores prácticas de mantenimiento y seguridad, comparto las cinco medidas fundamentales que implemento para proteger mis equipos, abarcando tanto la integridad física (hardware) como la lógica (software):')

    points = [
        ("1. Uso de Sistema de Alimentación Ininterrumpida (No-Break/UPS)", 
         "Protege el hardware contra picos de voltaje y cortes repentinos de energía. Los apagones bruscos pueden dañar físicamente los discos duros mecánicos y corromper archivos del sistema operativo irremediablemente."),
        ("2. Limpieza Física Semestral (Aire Comprimido y Pasta Térmica)",
         "El polvo actúa como aislante térmico, provocando sobrecalentamiento que reduce la vida útil de los componentes (CPU, GPU). Mantener los ventiladores limpios asegura un flujo de aire óptimo."),
        ("3. Actualizaciones Automáticas del Sistema Operativo y Drivers",
         "Las actualizaciones no solo traen nuevas funciones, sino 'parches de seguridad' críticos. Los ciberdelincuentes explotan vulnerabilidades conocidas en versiones antiguas de software; mantener todo al día cierra esas puertas traseras."),
        ("4. Solución Antimalware con Protección en Tiempo Real",
         "Hoy en día no basta con escanear archivos manualmente. La protección en tiempo real bloquea scripts maliciosos al navegar por internet o descargar correos, actuando antes de que el virus infecte el sistema."),
        ("5. Estrategia de Respaldo 3-2-1",
         "Ningún sistema es infalible. Mantengo 3 copias de mis datos, en 2 medios diferentes (ej. disco duro local y disco externo), y 1 copia en la nube. Esto garantiza que, ante un fallo catastrófico de hardware o un ataque de ransomware, la información siempre sea recuperable.")
    ]

    for title, body in points:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        p.add_run(f"\n¿Por qué? {body}")

    doc.add_paragraph('Saludos,\nAdán Muñoz Pablo')
    doc.add_paragraph('')

    # Replies
    doc.add_heading('Aportaciones a Compañeros', level=1)
    
    doc.add_heading('Respuesta a Compañero 1 (Mantenimiento Físico)', level=2)
    doc.add_paragraph('Hola [Compañero],')
    doc.add_paragraph('Coincido totalmente contigo en la importancia de la ubicación física del equipo. A menudo subestimamos cómo factores ambientales como la humedad o la luz solar directa pueden degradar los plásticos y circuitos. Me pareció muy interesante tu punto, ya que complementa muy bien la necesidad de una limpieza interna periódica. ¡Buen aporte!')
    doc.add_paragraph('Saludos.')

    doc.add_heading('Respuesta a Compañero 2 (Seguridad Software)', level=2)
    doc.add_paragraph('Hola [Compañero],')
    doc.add_paragraph('Muy acertado tu comentario sobre el software. Quisiera agregar que, además del antivirus que mencionas, el uso de gestores de contraseñas y la autenticación en dos pasos (2FA) son barreras vitales hoy en día. Muchas veces la "ingeniería social" logra saltarse el antivirus, pero el 2FA detiene el acceso no autorizado. Gracias por compartir tus estrategias.')
    doc.add_paragraph('Saludos.')

    doc.save(filename)
    print(f"Word guardado como {filename}")

if __name__ == "__main__":
    create_forum_word("Adan_Munoz_Pablo_Foro_Discusion.docx")


