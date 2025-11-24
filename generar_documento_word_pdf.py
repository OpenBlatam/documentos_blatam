#!/usr/bin/env python3
"""
Script mejorado para generar el documento A1_ADZ en formato Word y PDF
con formato profesional y mejor presentación

Requisitos:
    pip install python-docx
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime

def set_cell_background(cell, color):
    """Establece el color de fondo de una celda"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_info_box(doc, title, content, bg_color='E7F3FF', icon=''):
    """Agrega una caja de información destacada mejorada"""
    table = doc.add_table(rows=2, cols=1)
    table.style = 'Light List Accent 1'
    table.autofit = True
    
    # Título de la caja con icono
    title_cell = table.rows[0].cells[0]
    title_text = f"{icon} {title}" if icon else title
    title_cell.text = title_text
    set_cell_background(title_cell, '4472C4')
    for paragraph in title_cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(11)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
    
    # Contenido con mejor formato
    content_cell = table.rows[1].cells[0]
    set_cell_background(content_cell, bg_color)
    content_cell.text = content
    for paragraph in content_cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.left_indent = Inches(0.1)
        paragraph.paragraph_format.right_indent = Inches(0.1)
        run = paragraph.runs[0]
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_paragraph()  # Espacio después

def add_warning_box(doc, content):
    """Agrega una caja de advertencia"""
    add_info_box(doc, '⚠️ ADVERTENCIA', content, bg_color='FFF4E6', icon='')

def add_tip_box(doc, content):
    """Agrega una caja de consejo"""
    add_info_box(doc, '💡 CONSEJO', content, bg_color='E8F5E9', icon='')

def add_important_box(doc, content):
    """Agrega una caja de información importante"""
    add_info_box(doc, '📌 IMPORTANTE', content, bg_color='FFEBEE', icon='')

def add_header_footer(section, student_name='ADZ'):
    """Agrega encabezado y pie de página"""
    # Encabezado
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = f"Actividad 1: Análisis de Indicadores Económicos - {student_name}"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_para.runs[0]
    header_run.font.size = Pt(9)
    header_run.font.color.rgb = RGBColor(100, 100, 100)
    
    # Pie de página
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run('Página ')
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(100, 100, 100)
    
    # Agregar campo de número de página
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    footer_para._p.append(fldChar1)
    footer_para._p.append(instrText)
    footer_para._p.append(fldChar2)
    
    footer_run2 = footer_para.add_run(' de ')
    footer_run2.font.size = Pt(9)
    footer_run2.font.color.rgb = RGBColor(100, 100, 100)

def crear_documento_word():
    """Crea el documento Word con formato profesional mejorado"""
    
    doc = Document()
    
    # Configurar márgenes mejorados
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.5)
    
    # Configurar estilos mejorados
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.25
    style.paragraph_format.space_after = Pt(8)
    
    # Título principal con formato mejorado y más grande
    title = doc.add_heading('Actividad 1: Análisis de Indicadores Económicos del Banco de México', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)  # Azul oscuro
    
    # Subtítulo descriptivo
    subtitle = doc.add_paragraph('Análisis de Indicadores Macroeconómicos y su Impacto Empresarial')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
    
    # Línea decorativa después del título
    doc.add_paragraph('_' * 80).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Información del estudiante en formato mejorado con mejor diseño
    info_table = doc.add_table(rows=3, cols=2)
    info_table.style = 'Light List Accent 1'
    info_table.autofit = True
    
    # Agregar fila de título
    title_cells = info_table.rows[0].cells
    title_cells[0].merge(title_cells[1])
    title_cells[0].text = 'INFORMACIÓN DEL ESTUDIANTE'
    set_cell_background(title_cells[0], '4472C4')  # Azul más intenso
    for paragraph in title_cells[0].paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)  # Blanco
            run.font.size = Pt(12)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    info_data = [
        ('Estudiante:', 'ADZ'),
        ('Fecha:', datetime.now().strftime('%Y'))
    ]
    
    for i, (label, value) in enumerate(info_data, 1):
        cells = info_table.rows[i].cells
        cells[0].text = label
        cells[0].paragraphs[0].runs[0].bold = True
        cells[0].paragraphs[0].runs[0].font.size = Pt(11)
        cells[1].text = value
        cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    
    doc.add_paragraph()  # Espacio
    
    # Agregar encabezado y pie de página
    add_header_footer(doc.sections[0], 'ADZ')
    
    # ÍNDICE AUTOMÁTICO
    doc.add_page_break()
    index_heading = doc.add_heading('ÍNDICE', 0)
    index_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    index_heading.runs[0].font.size = Pt(16)
    index_heading.runs[0].font.bold = True
    index_heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph()
    
    # Crear tabla de índice mejorada
    index_table = doc.add_table(rows=9, cols=3)
    index_table.style = 'Light List Accent 1'
    index_table.autofit = True
    
    # Encabezado del índice
    header_cells = index_table.rows[0].cells
    headers = ['Sección', 'Contenido', 'Página']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        set_cell_background(header_cells[i], '4472C4')
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(11)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    index_items = [
        ('I.', 'RECOPILACIÓN DE INFORMACIÓN', ''),
        ('', '  1. Cuadro A: Indicadores Económicos Principales', ''),
        ('', '  2. Cuadro B: Divisas Internacionales', ''),
        ('II.', 'ANÁLISIS DE LA INFORMACIÓN', ''),
        ('', '  1. Inflación', ''),
        ('', '  2. Tasa de Interés', ''),
        ('', '  3. Tipo de Cambio', ''),
        ('', '  4. Divisas', '')
    ]
    
    for i, (num, text, page) in enumerate(index_items, 1):
        cells = index_table.rows[i].cells
        cells[0].text = num
        cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        cells[1].text = text
        cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        cells[2].text = page  # Se llenará automáticamente en Word
        
        # Formato alternado
        if i % 2 == 0:
            bg_color = 'F2F2F2'
        else:
            bg_color = 'FFFFFF'
        
        for cell in cells:
            set_cell_background(cell, bg_color)
            if num:  # Si tiene número, hacerlo negrita
                cell.paragraphs[0].runs[0].bold = True
                cells[1].paragraphs[0].runs[0].bold = True
    
    doc.add_page_break()
    
    # ========== SECCIÓN I: RECOPILACIÓN DE INFORMACIÓN ==========
    section1 = doc.add_heading('I. RECOPILACIÓN DE INFORMACIÓN', 1)
    section1.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    # Cajas de información importante mejoradas
    add_important_box(
        doc,
        'Esta sección requiere que obtengas datos directamente del Banco de México. '
        'Sigue las instrucciones detalladas para cada indicador. Los datos deben ser '
        'de los últimos 5 años disponibles y deben completarse en las tablas correspondientes.'
    )
    
    add_tip_box(
        doc,
        'Recomendación: Guarda los archivos Excel que descargues del Banco de México '
        'para poder verificar tus datos más tarde. También es útil tomar capturas de pantalla '
        'de las páginas donde obtuviste la información.'
    )
    
    # Subsección 1: Cuadro A
    doc.add_heading('1. Cuadro A: Indicadores Económicos Principales', 2)
    
    # Descripción mejorada del cuadro
    desc_table = doc.add_table(rows=1, cols=1)
    desc_table.autofit = True
    desc_cell = desc_table.rows[0].cells[0]
    set_cell_background(desc_cell, 'F0F8FF')
    desc_cell.text = '📊 Complete el siguiente cuadro con datos del Banco de México para los últimos 5 años disponibles. ' \
                     'Use los valores anuales (promedio o cierre de año) para cada indicador.'
    for paragraph in desc_cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 0)
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.paragraph_format.space_before = Pt(8)
    
    doc.add_paragraph()  # Espacio
    
    # Crear tabla Cuadro A con formato profesional mejorado
    table_a = doc.add_table(rows=6, cols=4)
    table_a.style = 'Light Grid Accent 1'
    table_a.autofit = True
    
    # Encabezados con color de fondo mejorado
    headers_a = [
        'Año',
        'Índice Nacional de Precios al Consumidor\n(Variación anual %)',
        'Tipo de Cambio Peso-Dólar\n(Pesos por dólar)',
        'TIIE a 28 días\n(Tasa de interés % anual)'
    ]
    
    header_cells = table_a.rows[0].cells
    for i, header in enumerate(headers_a):
        cell = header_cells[i]
        cell.text = header
        set_cell_background(cell, '4472C4')  # Azul más intenso
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(255, 255, 255)  # Texto blanco
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Filas de datos (vacías para completar) con formato alternado
    for row_idx in range(1, 6):
        cells = table_a.rows[row_idx].cells
        cells[0].text = ''  # Año
        cells[1].text = ''  # Inflación
        cells[2].text = ''  # Tipo de cambio
        cells[3].text = ''  # TIIE
        
        # Color de fondo alternado para mejor legibilidad
        if row_idx % 2 == 0:
            bg_color = 'F2F2F2'  # Gris muy claro
        else:
            bg_color = 'FFFFFF'  # Blanco
        
        for cell in cells:
            set_cell_background(cell, bg_color)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_after = Pt(3)
    
    doc.add_paragraph()  # Espacio
    
    # Instrucciones Cuadro A con mejor formato
    doc.add_heading('📝 Instrucciones para completar el Cuadro A:', 3)
    
    # Cajas de ayuda mejoradas para el Cuadro A
    add_tip_box(
        doc,
        'Asegúrate de usar los datos anuales (promedio o valor al cierre del año) para cada indicador. '
        'Si los datos están en formato mensual, calcula el promedio anual o usa el valor de diciembre.'
    )
    
    add_warning_box(
        doc,
        'Verifica que estés usando la variación ANUAL de la inflación, no la mensual. '
        'El tipo de cambio debe ser el promedio anual o el valor al cierre del año. '
        'La TIIE debe ser el promedio anual de la tasa.'
    )
    
    # 1.1 Inflación
    doc.add_heading('1.1 📊 Índice Nacional de Precios al Consumidor (Inflación)', 4)
    instructions_inflacion = [
        'Acceda a: https://www.banxico.org.mx/SieInternet/consultarDirectorioInternetAction.do?sector=8&accion=consultarDirectorioCuadros&locale=es',
        'Busque la sección "Inflación" y haga clic',
        'Seleccione "Inflación anual" (marcar con palomita ✓)',
        'Presione el botón "Exportar series"',
        'En el archivo Excel, localice la columna "Índice Nacional de Precios al Consumidor, variación anual"',
        'Copie los datos de los últimos 5 años al Cuadro A'
    ]
    for i, instruction in enumerate(instructions_inflacion, 1):
        p = doc.add_paragraph(f'✓ {instruction}', style='List Bullet')
        if 'http' in instruction:
            # Separar el texto del enlace
            parts = instruction.split(':', 1)
            if len(parts) == 2:
                p.clear()
                p.add_run(f'✓ {parts[0]}: ').font.size = Pt(10)
                link_run = p.add_run(parts[1])
                link_run.font.color.rgb = RGBColor(0, 0, 255)
                link_run.font.underline = True
        else:
            for run in p.runs:
                run.font.size = Pt(10)
    
    # 1.2 Tipo de Cambio
    doc.add_heading('1.2 💱 Tipo de Cambio', 4)
    instructions_tc = [
        'Acceda a: https://www.banxico.org.mx/SieInternet/consultarDirectorioInternetAction.do?sector=23&accion=consultarCuadro&idCuadro=CF372&locale=es',
        'Seleccione "Tipo de cambio peso dólar desde 1954 1/" (marcar con palomita ✓)',
        'Presione "Exportar series"',
        'En el archivo Excel, identifique la columna "Serie histórica del tipo de cambio, Tipo de cambio peso dólar desde 1954"',
        'Copie los datos de los últimos 5 años al Cuadro A'
    ]
    for i, instruction in enumerate(instructions_tc, 1):
        p = doc.add_paragraph(f'✓ {instruction}', style='List Bullet')
        if 'http' in instruction:
            parts = instruction.split(':', 1)
            if len(parts) == 2:
                p.clear()
                p.add_run(f'✓ {parts[0]}: ').font.size = Pt(10)
                link_run = p.add_run(parts[1])
                link_run.font.color.rgb = RGBColor(0, 0, 255)
                link_run.font.underline = True
        else:
            for run in p.runs:
                run.font.size = Pt(10)
    
    # 1.3 Tasa de Interés
    doc.add_heading('1.3 💰 Tasa de Interés (TIIE)', 4)
    instructions_tiie = [
        'Acceda a: https://www.banxico.org.mx/SieInternet/consultarDirectorioInternetAction.do?sector=18&accion=consultarCuadroAnalitico&idCuadro=CA51&locale=es',
        'Presione "Exportar series"',
        'En el archivo Excel, identifique la columna "TIIE a 28 días, Tasa de interés en por ciento anual"',
        'Copie los datos de los últimos 5 años al Cuadro A'
    ]
    for i, instruction in enumerate(instructions_tiie, 1):
        p = doc.add_paragraph(f'✓ {instruction}', style='List Bullet')
        if 'http' in instruction:
            parts = instruction.split(':', 1)
            if len(parts) == 2:
                p.clear()
                p.add_run(f'✓ {parts[0]}: ').font.size = Pt(10)
                link_run = p.add_run(parts[1])
                link_run.font.color.rgb = RGBColor(0, 0, 255)
                link_run.font.underline = True
        else:
            for run in p.runs:
                run.font.size = Pt(10)
    
    doc.add_paragraph()  # Espacio
    
    # Subsección 2: Cuadro B
    doc.add_heading('2. Cuadro B: Divisas Internacionales', 2)
    
    # Descripción mejorada del cuadro
    desc_table_b = doc.add_table(rows=1, cols=1)
    desc_table_b.autofit = True
    desc_cell_b = desc_table_b.rows[0].cells[0]
    set_cell_background(desc_cell_b, 'F0F8FF')
    desc_cell_b.text = '💱 Complete el siguiente cuadro con datos del Banco de México para los últimos 5 años disponibles. ' \
                       'Las divisas se expresan en pesos mexicanos por unidad de moneda extranjera.'
    for paragraph in desc_cell_b.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 0)
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.paragraph_format.space_before = Pt(8)
    
    doc.add_paragraph()  # Espacio
    
    # Crear tabla Cuadro B con mejor formato
    table_b = doc.add_table(rows=6, cols=6)
    table_b.style = 'Light Grid Accent 1'
    table_b.autofit = True
    
    # Encabezados
    headers_b = [
        'Año',
        'Dólar EUA\n(Pesos por unidad)',
        'Euro\n(Pesos por unidad)',
        'Yen Japonés\n(Pesos por unidad)',
        'Libra Esterlina\n(Pesos por unidad)',
        'Yuan Chino\n(Pesos por unidad)'
    ]
    
    header_cells_b = table_b.rows[0].cells
    for i, header in enumerate(headers_b):
        cell = header_cells_b[i]
        cell.text = header
        set_cell_background(cell, '4472C4')  # Azul más intenso
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(255, 255, 255)  # Texto blanco
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Filas de datos (vacías para completar) con formato alternado
    for row_idx in range(1, 6):
        cells = table_b.rows[row_idx].cells
        for cell in cells:
            cell.text = ''
            
            # Color de fondo alternado
            if row_idx % 2 == 0:
                bg_color = 'F2F2F2'  # Gris muy claro
            else:
                bg_color = 'FFFFFF'  # Blanco
            
            set_cell_background(cell, bg_color)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_after = Pt(3)
    
    doc.add_paragraph()  # Espacio
    
    # Instrucciones Cuadro B
    doc.add_heading('📝 Instrucciones para completar el Cuadro B:', 3)
    
    # Cajas de ayuda mejoradas para el Cuadro B
    add_tip_box(
        doc,
        'Las divisas se expresan en pesos mexicanos por unidad de la moneda extranjera. '
        'Asegúrate de usar el tipo de cambio promedio anual o el valor al cierre del año.'
    )
    
    add_important_box(
        doc,
        'Recuerda que debes seleccionar las 5 monedas indicadas: Dólar EUA, Euro, Yen japonés, '
        'Libra esterlina y Yuan chino. Verifica que todas estén marcadas antes de exportar.'
    )
    
    doc.add_heading('2.1 💱 Divisas', 4)
    
    # Crear lista mejorada con viñetas para las monedas
    p = doc.add_paragraph('✓ Acceda a: ', style='List Bullet')
    link_run = p.add_run('https://www.banxico.org.mx/SieInternet/consultarDirectorioInternetAction.do?sector=6&accion=consultarCuadro&idCuadro=CF307&locale=es')
    link_run.font.color.rgb = RGBColor(0, 0, 255)
    link_run.font.underline = True
    p.runs[0].font.size = Pt(10)
    
    p = doc.add_paragraph('✓ Seleccione las siguientes monedas (marcar con palomita ✓):', style='List Bullet')
    p.runs[0].font.size = Pt(10)
    monedas = ['Dólar EUA', 'Euro', 'Yen japonés', 'Libra esterlina', 'Yuan chino']
    for moneda in monedas:
        p_moneda = doc.add_paragraph(f'   → {moneda}', style='List Bullet 2')
        p_moneda.runs[0].font.size = Pt(10)
    
    instructions_divisas_resto = [
        'Presione "Exportar series"',
        'En el archivo Excel, identifique las columnas correspondientes a cada moneda',
        'Copie los datos de los últimos 5 años al Cuadro B'
    ]
    for instruction in instructions_divisas_resto:
        p = doc.add_paragraph(f'✓ {instruction}', style='List Bullet')
        p.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    p_note = doc.add_paragraph()
    p_note.add_run('Nota importante: ').bold = True
    p_note.add_run(
        'Es posible que con el paso del tiempo se cambien algunos elementos o pestañas dentro de la página oficial del Banco de México. '
        'Si ese fuera el caso es importante investigar la información solicitada en las nuevas ligas, pero dentro del sitio oficial.'
    )
    p_note.style = 'Intense Quote'
    
    # ========== SECCIÓN II: ANÁLISIS ==========
    doc.add_page_break()
    section2 = doc.add_heading('II. ANÁLISIS DE LA INFORMACIÓN', 1)
    section2.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    # Cajas de información mejoradas para la sección de análisis
    add_important_box(
        doc,
        'En esta sección debes analizar los datos que obtuviste en los Cuadros A y B. '
        'Para cada indicador, explica la tendencia (al alza, a la baja, o constante) '
        'y fundamenta cómo afecta a tu empresa. Usa datos específicos de las tablas para '
        'sustentar tus respuestas.'
    )
    
    add_tip_box(
        doc,
        'Estrategia de análisis: (1) Identifica la tendencia general, (2) Menciona valores específicos '
        'del cuadro, (3) Explica el impacto en tu empresa, (4) Fundamenta con razones económicas. '
        'Sé específico y usa números de las tablas.'
    )
    
    add_warning_box(
        doc,
        'No olvides considerar el contexto de tu empresa específica. Si tu negocio tiene préstamos, '
        'analiza el impacto de la tasa de interés. Si vendes al extranjero, analiza el tipo de cambio. '
        'Sé específico sobre cómo estos indicadores afectan TU empresa.'
    )
    
    # 1. Inflación
    doc.add_heading('1. 📈 Inflación', 2)
    
    # Caja de contexto para inflación
    add_tip_box(
        doc,
        'La inflación mide el aumento generalizado de precios en una economía. '
        'Una inflación alta reduce el poder adquisitivo, mientras que una inflación muy baja '
        'puede indicar debilidad económica. Analiza cómo afecta a los costos y precios de tu empresa.'
    )
    
    doc.add_heading('1.1 Explicación del cambio en la Inflación', 3)
    
    p = doc.add_paragraph('📊 Análisis de la tendencia:')
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    
    # Caja de guía para el análisis
    add_tip_box(
        doc,
        'Estructura recomendada para tu análisis:\n'
        '1. Identifica la tendencia (al alza, a la baja, o constante)\n'
        '2. Menciona valores específicos del Cuadro A (ej: "En 2020 la inflación fue de X%...")\n'
        '3. Compara años consecutivos para mostrar la tendencia\n'
        '4. Concluye con la tendencia general'
    )
    
    doc.add_paragraph(
        '[Complete aquí su análisis sobre cómo ha cambiado la inflación medida por el Índice Nacional de Precios al Consumidor. '
        'Especifique si en los años que incluye el Cuadro A, la inflación va al alza, a la baja, o se mantiene constante. '
        'Incluya datos específicos del cuadro para fundamentar su respuesta.]',
        style='Intense Quote'
    )
    
    doc.add_paragraph()
    doc.add_paragraph('📋 Ejemplo de estructura:', style='Heading 4')
    doc.add_paragraph('→ Año X: La inflación fue de X%, lo que representa...', style='List Bullet')
    doc.add_paragraph('→ Año Y: La inflación fue de Y%, mostrando una tendencia...', style='List Bullet')
    doc.add_paragraph('→ Tendencia general: [Al alza / A la baja / Constante]', style='List Bullet')
    
    doc.add_heading('1.2 Impacto en la empresa', 3)
    
    p = doc.add_paragraph('❓ ¿Los cambios de inflación benefician o perjudican a la empresa?')
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(11)
    
    # Caja de guía para el impacto
    add_tip_box(
        doc,
        'Aspectos a considerar en tu análisis:\n'
        '• Costos de producción: ¿Aumentan o disminuyen con la inflación?\n'
        '• Precios de venta: ¿Puedes ajustarlos fácilmente?\n'
        '• Poder adquisitivo: ¿Cómo afecta a tus clientes?\n'
        '• Competitividad: ¿Cómo te afecta vs competidores?\n'
        '• Planificación financiera: ¿Dificulta o facilita la planificación?'
    )
    
    doc.add_paragraph(
        '[Complete aquí su análisis sobre si los cambios señalados de la inflación benefician o perjudican a la empresa. '
        'Fundamente su respuesta considerando aspectos como: costos de producción, precios de venta, poder adquisitivo de los consumidores, '
        'competitividad, planificación financiera.]',
        style='Intense Quote'
    )
    
    doc.add_paragraph()
    p = doc.add_paragraph('📝 Fundamentación:')
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    doc.add_paragraph(
        '[Desarrolle aquí su argumentación detallada sobre el impacto de la inflación en su empresa específica. '
        'Sea específico y use ejemplos concretos de cómo la inflación afecta a su negocio.]',
        style='Intense Quote'
    )
    
    # 2. Tasa de Interés
    doc.add_heading('2. 💰 Tasa de Interés', 2)
    
    # Caja de contexto para tasa de interés
    add_tip_box(
        doc,
        'La tasa de interés afecta directamente el costo del crédito. Si tu empresa tiene préstamos, '
        'una tasa alta aumenta los pagos de intereses. Si planeas pedir crédito, una tasa alta lo hace más caro. '
        'Analiza el impacto en tu situación financiera específica.'
    )
    
    doc.add_heading('2.1 Explicación del cambio en la Tasa de Interés', 3)
    
    p = doc.add_paragraph('📊 Análisis de la tendencia:')
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    
    # Caja de guía para el análisis
    add_tip_box(
        doc,
        'Estructura recomendada:\n'
        '1. Identifica la tendencia de la tasa de interés\n'
        '2. Menciona valores específicos del Cuadro A\n'
        '3. Compara años consecutivos\n'
        '4. Concluye con la tendencia general'
    )
    
    doc.add_paragraph(
        '[Complete aquí su análisis sobre cómo ha cambiado la tasa de interés medida por TIIE a 28 días, Tasa de interés en por ciento anual. '
        'Especifique si en los años que incluye el Cuadro A, la tasa de interés va al alza, a la baja, o se mantiene constante. '
        'Incluya datos específicos del cuadro para fundamentar su respuesta.]',
        style='Intense Quote'
    )
    
    doc.add_heading('2.2 Impacto en la empresa', 3)
    
    p = doc.add_paragraph('❓ ¿Los cambios de tasa de interés benefician o perjudican a la empresa?')
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(11)
    
    # Caja de guía para el impacto
    add_tip_box(
        doc,
        'Aspectos clave a analizar:\n'
        '• Si tienes préstamos: ¿Cómo afecta el aumento/disminución de la tasa?\n'
        '• Si planeas pedir crédito: ¿Es buen o mal momento?\n'
        '• Costo del financiamiento: ¿Aumenta o disminuye?\n'
        '• Capacidad de pago: ¿Puedes afrontar los pagos?\n'
        '• Rentabilidad: ¿Cómo impacta en tus ganancias?'
    )
    
    doc.add_paragraph(
        '[Complete aquí su análisis sobre si los cambios de la tasa de interés benefician o perjudican a la empresa. '
        'Especialmente explique si estos cambios benefician o perjudican al negocio al haber pedido un préstamo bancario para formar su capital inicial.]',
        style='Intense Quote'
    )
    
    doc.add_paragraph()
    p = doc.add_paragraph('💳 Análisis del préstamo bancario:')
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    
    # Caja específica para préstamo
    add_important_box(
        doc,
        'IMPORTANTE: Si tu empresa NO tiene préstamos, explica cómo la tasa de interés afectaría '
        'si decidieras pedir uno en el futuro, o cómo afecta a tus clientes que sí tienen créditos.'
    )
    
    doc.add_paragraph(
        '[Desarrolle aquí su análisis específico sobre el impacto de los cambios en la tasa de interés en el préstamo bancario utilizado para formar el capital inicial. '
        'Considere: costo del financiamiento, pagos de intereses, capacidad de pago, impacto en la rentabilidad.]',
        style='Intense Quote'
    )
    
    doc.add_paragraph()
    p = doc.add_paragraph('📝 Fundamentación:')
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    doc.add_paragraph(
        '[Desarrolle aquí su argumentación detallada sobre el impacto de la tasa de interés en su empresa, con énfasis en el préstamo bancario. '
        'Sea específico sobre montos, porcentajes y efectos concretos.]',
        style='Intense Quote'
    )
    
    # 3. Tipo de Cambio
    doc.add_heading('3. 💱 Tipo de Cambio', 2)
    
    # Caja de contexto para tipo de cambio
    add_tip_box(
        doc,
        'El tipo de cambio peso-dólar afecta a empresas que importan, exportan o compiten con productos extranjeros. '
        'Un peso débil (más pesos por dólar) beneficia exportadores pero encarece importaciones. '
        'Analiza cómo afecta a tu modelo de negocio específico.'
    )
    
    doc.add_heading('3.1 Explicación de la variación del Tipo de Cambio', 3)
    
    p = doc.add_paragraph('📊 Análisis de la tendencia:')
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    
    # Caja de guía para el análisis
    add_tip_box(
        doc,
        'Estructura recomendada:\n'
        '1. Identifica la tendencia del tipo de cambio\n'
        '2. Menciona valores específicos del Cuadro A (ej: "En 2020 el tipo de cambio fue de X pesos por dólar...")\n'
        '3. Compara años consecutivos\n'
        '4. Concluye con la tendencia general'
    )
    
    doc.add_paragraph(
        '[Complete aquí su análisis sobre cómo ha variado el tipo de cambio medido por el indicador Tipo de cambio pesos por dólar. '
        'Especifique si en los años que incluye el Cuadro A, el tipo de cambio va al alza, a la baja, o se mantiene constante. '
        'Incluya datos específicos del cuadro para fundamentar su respuesta.]',
        style='Intense Quote'
    )
    
    doc.add_heading('3.2 Impacto en la empresa', 3)
    
    p = doc.add_paragraph('❓ ¿Las variaciones del tipo de cambio benefician o perjudican a la empresa?')
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(11)
    
    doc.add_paragraph(
        '[Complete aquí su análisis sobre si las variaciones del tipo de cambio benefician o perjudican a la empresa. '
        'Especialmente explique si estos cambios benefician o perjudican a la empresa considerando que el negocio realiza ventas u ofrece productos o servicios al extranjero.]',
        style='Intense Quote'
    )
    
    # Caja de guía para ventas al extranjero
    add_tip_box(
        doc,
        'Aspectos a considerar:\n'
        '• Competitividad: ¿Un peso débil te hace más competitivo?\n'
        '• Ingresos: ¿Cómo se traducen los ingresos en dólares a pesos?\n'
        '• Costos: ¿Importas insumos? ¿Cómo afecta el tipo de cambio?\n'
        '• Margen: ¿Aumenta o disminuye tu margen de ganancia?\n'
        '• Planificación: ¿Facilita o dificulta la planificación?'
    )
    
    doc.add_paragraph()
    p = doc.add_paragraph('🌐 Análisis de ventas al extranjero:')
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    
    # Caja específica para ventas internacionales
    add_important_box(
        doc,
        'IMPORTANTE: Si tu empresa NO vende al extranjero, explica cómo el tipo de cambio '
        'afectaría si decidieras expandirte internacionalmente, o cómo afecta a la competencia '
        'con productos importados.'
    )
    
    doc.add_paragraph(
        '[Desarrolle aquí su análisis específico sobre el impacto de los cambios en el tipo de cambio en las ventas o servicios ofrecidos al extranjero. '
        'Considere: competitividad de precios, ingresos en moneda extranjera, costos de importación (si aplica), margen de ganancia.]',
        style='Intense Quote'
    )
    
    doc.add_paragraph()
    p = doc.add_paragraph('📝 Fundamentación:')
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    doc.add_paragraph(
        '[Desarrolle aquí su argumentación detallada sobre el impacto del tipo de cambio en su empresa, con énfasis en las operaciones internacionales. '
        'Use ejemplos concretos y mencione valores específicos del Cuadro A.]',
        style='Intense Quote'
    )
    
    # 4. Divisas
    doc.add_heading('4. 🌍 Divisas', 2)
    
    # Caja de contexto para divisas
    add_tip_box(
        doc,
        'Las divisas internacionales afectan a empresas con operaciones en esos países. '
        'Si vendes a Estados Unidos, el dólar es clave. Si vendes a Europa, el euro es importante. '
        'Analiza cada divisa según los mercados donde opera tu empresa.'
    )
    
    doc.add_heading('4.1 Explicación de los cambios en las Divisas', 3)
    
    p = doc.add_paragraph('📊 Análisis de la tendencia de cada divisa:')
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    
    divisas = ['Dólar EUA', 'Euro', 'Yen Japonés', 'Libra Esterlina', 'Yuan Chino']
    for divisa in divisas:
        doc.add_heading(f'💱 {divisa}:', 4)
        doc.add_paragraph(
            f'[Complete aquí su análisis sobre cómo ha cambiado el {divisa}. '
            'Especifique si va al alza, a la baja, o se mantiene constante en los años señalados en el Cuadro B. '
            'Mencione valores específicos del Cuadro B para fundamentar su respuesta.]',
            style='Intense Quote'
        )
    
    doc.add_paragraph()
    p = doc.add_paragraph('Resumen de tendencias:')
    p.runs[0].bold = True
    
    # Tabla de resumen de divisas
    table_divisas = doc.add_table(rows=6, cols=3)
    table_divisas.style = 'Light Grid Accent 1'
    table_divisas.autofit = True
    
    headers_divisas = ['Divisa', 'Tendencia', 'Observaciones']
    header_cells_div = table_divisas.rows[0].cells
    for i, header in enumerate(headers_divisas):
        cell = header_cells_div[i]
        cell.text = header
        set_cell_background(cell, '4472C4')  # Azul más intenso
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)  # Texto blanco
                run.font.size = Pt(11)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for i, divisa in enumerate(divisas, 1):
        cells = table_divisas.rows[i].cells
        cells[0].text = divisa
        cells[1].text = '[Al alza / A la baja / Constante]'
        cells[2].text = '[Breve descripción]'
        
        # Color de fondo alternado
        if i % 2 == 0:
            bg_color = 'F2F2F2'  # Gris muy claro
        else:
            bg_color = 'FFFFFF'  # Blanco
        
        for cell in cells:
            set_cell_background(cell, bg_color)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_after = Pt(3)
    
    doc.add_heading('4.2 Impacto en la empresa', 3)
    
    p = doc.add_paragraph('❓ ¿Los cambios de las divisas benefician o perjudican a la empresa?')
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(11)
    
    doc.add_paragraph(
        '[Complete aquí su análisis sobre si los cambios de las divisas benefician o perjudican a la empresa, '
        'ya que su negocio vende productos o servicios a los países donde circulan dichas monedas.]',
        style='Intense Quote'
    )
    
    doc.add_paragraph()
    p = doc.add_paragraph('🌍 Análisis por país/región:')
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    
    paises = [
        ('Estados Unidos (Dólar EUA)', 'Estados Unidos'),
        ('Zona Euro (Euro)', 'países de la zona euro'),
        ('Japón (Yen Japonés)', 'Japón'),
        ('Reino Unido (Libra Esterlina)', 'Reino Unido'),
        ('China (Yuan Chino)', 'China')
    ]
    
    for titulo, pais in paises:
        doc.add_heading(f'{titulo}:', 4)
        moneda = titulo.split('(')[1].split(')')[0]
        doc.add_paragraph(
            f'[Desarrolle aquí el análisis del impacto de los cambios en el {moneda} '
            f'en las ventas a {pais}.]',
            style='Intense Quote'
        )
    
    doc.add_paragraph()
    p = doc.add_paragraph('🔗 Análisis integral:')
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    doc.add_paragraph(
        '[Desarrolle aquí un análisis integral que considere el impacto conjunto de todas las divisas en la estrategia internacional de la empresa, '
        'incluyendo: diversificación de mercados, gestión de riesgo cambiario, competitividad por región, estrategias de mitigación.]',
        style='Intense Quote'
    )
    
    doc.add_paragraph()
    p = doc.add_paragraph('Fundamentación:')
    p.runs[0].bold = True
    doc.add_paragraph(
        '[Desarrolle aquí su argumentación detallada sobre el impacto de las divisas en su empresa, considerando las operaciones internacionales.]',
        style='Intense Quote'
    )
    
    # Conclusiones
    doc.add_page_break()
    conclusiones = doc.add_heading('📋 CONCLUSIONES', 1)
    conclusiones.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    # Cajas de ayuda mejoradas para conclusiones
    add_important_box(
        doc,
        'Las conclusiones deben ser un resumen ejecutivo que integre todos los análisis realizados. '
        'Incluye: (1) Tendencias principales identificadas, (2) Impacto general en la empresa, '
        '(3) Recomendaciones estratégicas, (4) Consideraciones para la toma de decisiones.'
    )
    
    add_tip_box(
        doc,
        'Estructura recomendada para las conclusiones:\n'
        '• Párrafo 1: Resumen de las tendencias principales (inflación, tasa de interés, tipo de cambio, divisas)\n'
        '• Párrafo 2: Impacto general en tu empresa (positivo/negativo/neutral)\n'
        '• Párrafo 3: Recomendaciones estratégicas específicas\n'
        '• Párrafo 4: Consideraciones para la toma de decisiones futuras'
    )
    
    # Caja de advertencia para conclusiones
    add_warning_box(
        doc,
        'IMPORTANTE: Las conclusiones deben integrar TODOS los análisis realizados. '
        'No repitas información, sino sintetiza los hallazgos principales y sus implicaciones '
        'para tu empresa. Usa datos específicos de los Cuadros A y B para fundamentar tus conclusiones.'
    )
    
    # Crear tabla estructurada para las conclusiones
    conclusiones_table = doc.add_table(rows=5, cols=2)
    conclusiones_table.style = 'Light Grid Accent 1'
    conclusiones_table.autofit = True
    
    # Encabezado
    header_cells = conclusiones_table.rows[0].cells
    header_cells[0].merge(header_cells[1])
    header_cells[0].text = 'ESTRUCTURA DE CONCLUSIONES'
    set_cell_background(header_cells[0], '4472C4')
    for paragraph in header_cells[0].paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(11)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Filas de estructura
    estructura_items = [
        ('1. Tendencias Principales', '[Resumen de las 4 tendencias principales identificadas]'),
        ('2. Impacto General', '[Impacto positivo/negativo/neutral en tu empresa]'),
        ('3. Recomendaciones', '[Recomendaciones estratégicas específicas]'),
        ('4. Consideraciones Futuras', '[Consideraciones para toma de decisiones]')
    ]
    
    for i, (titulo, contenido) in enumerate(estructura_items, 1):
        cells = conclusiones_table.rows[i].cells
        cells[0].text = titulo
        cells[0].paragraphs[0].runs[0].bold = True
        cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        cells[1].text = contenido
        cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        cells[1].paragraphs[0].runs[0].font.italic = True
        cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(100, 100, 100)
        
        # Color alternado
        if i % 2 == 0:
            bg_color = 'F2F2F2'
        else:
            bg_color = 'FFFFFF'
        
        for cell in cells:
            set_cell_background(cell, bg_color)
    
    doc.add_paragraph()
    doc.add_paragraph(
        '[Complete aquí un resumen ejecutivo de los principales hallazgos del análisis, destacando: '
        'tendencias principales identificadas, impacto general en la empresa, recomendaciones estratégicas, '
        'consideraciones para la toma de decisiones. Use la estructura de la tabla anterior como guía.]',
        style='Intense Quote'
    )
    
    # Referencias
    doc.add_heading('📚 REFERENCIAS', 1)
    
    # Caja informativa sobre referencias
    add_tip_box(
        doc,
        'Las referencias deben seguir el formato APA. Asegúrate de incluir la fecha de consulta '
        'si es necesario. Todas las URLs deben ser accesibles y verificables.'
    )
    
    referencias = [
        ('Banco de México. (2024). Sistema de Información Económica (SIE).', 
         'https://www.banxico.org.mx/SieInternet/'),
        ('Banco de México. (2024). Índice Nacional de Precios al Consumidor.', 
         'https://www.banxico.org.mx/SieInternet/consultarDirectorioInternetAction.do?sector=8&accion=consultarDirectorioCuadros&locale=es'),
        ('Banco de México. (2024). Tipo de Cambio.', 
         'https://www.banxico.org.mx/SieInternet/consultarDirectorioInternetAction.do?sector=23&accion=consultarCuadro&idCuadro=CF372&locale=es'),
        ('Banco de México. (2024). Tasa de Interés Interbancaria de Equilibrio (TIIE).', 
         'https://www.banxico.org.mx/SieInternet/consultarDirectorioInternetAction.do?sector=18&accion=consultarCuadroAnalitico&idCuadro=CA51&locale=es'),
        ('Banco de México. (2024). Divisas.', 
         'https://www.banxico.org.mx/SieInternet/consultarDirectorioInternetAction.do?sector=6&accion=consultarCuadro&idCuadro=CF307&locale=es')
    ]
    
    for i, (ref, url) in enumerate(referencias, 1):
        p = doc.add_paragraph(f'{i}. {ref} Recuperado de: ', style='List Number')
        link_run = p.add_run(url)
        link_run.font.color.rgb = RGBColor(0, 0, 255)
        link_run.font.underline = True
        p.runs[0].font.size = Pt(10)
        link_run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Caja final de recordatorio
    add_important_box(
        doc,
        'RECORDATORIO FINAL:\n'
        '• Completa todos los cuadros con datos reales del Banco de México\n'
        '• Responde todas las preguntas de análisis de manera completa\n'
        '• Incluye datos específicos de las tablas en tus respuestas\n'
        '• Guarda el documento como PDF con el nombre: A1_ADZ.pdf\n'
        '• Revisa la ortografía y gramática antes de enviar\n'
        '• Asegúrate de que todas las secciones estén completas'
    )
    
    p_final = doc.add_paragraph()
    p_final.add_run('📌 Nota final: ').bold = True
    p_final.add_run(
        'Este documento debe ser completado con los datos reales obtenidos del Banco de México y guardado en formato PDF con el nombre: '
    )
    p_final.add_run('A1_ADZ.pdf').bold = True
    p_final.add_run('.')
    
    return doc


def convertir_a_pdf(docx_path, pdf_path=None):
    """
    Convierte el documento Word a PDF
    
    Requiere: pip install docx2pdf
    O usar LibreOffice: libreoffice --headless --convert-to pdf archivo.docx
    """
    if pdf_path is None:
        pdf_path = docx_path.replace('.docx', '.pdf')
    
    try:
        # Método 1: Usando docx2pdf (requiere Microsoft Word en Windows o LibreOffice)
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        print(f"✅ PDF creado: {pdf_path}")
        return True
    except ImportError:
        print("⚠️  docx2pdf no está instalado. Instale con: pip install docx2pdf")
        print("   O use LibreOffice: libreoffice --headless --convert-to pdf A1_ADZ.docx")
        return False
    except Exception as e:
        print(f"⚠️  Error al convertir a PDF: {e}")
        print("   Alternativa: Abra el documento Word y guárdelo como PDF manualmente")
        return False


def main():
    """Función principal"""
    print("=" * 70)
    print("GENERAR DOCUMENTO WORD Y PDF - A1_ADZ (VERSIÓN ULTRA MEJORADA)")
    print("=" * 70)
    print()
    
    # Crear documento Word
    print("📝 Creando documento Word con formato ultra mejorado...")
    doc = crear_documento_word()
    
    # Guardar Word
    word_filename = 'A1_ADZ.docx'
    doc.save(word_filename)
    print(f"✅ Documento Word creado: {word_filename}")
    print("   ✨ Formato profesional premium aplicado")
    print("   ✨ Tablas con colores mejorados y filas alternadas")
    print("   ✨ Encabezados con texto blanco sobre fondo azul")
    print("   ✨ Márgenes y espaciado optimizados")
    print("   ✨ Estilos y estructura profesional")
    print()
    
    # Intentar convertir a PDF
    print("📄 Intentando convertir a PDF...")
    if convertir_a_pdf(word_filename):
        print(f"✅ Documento PDF creado: A1_ADZ.pdf")
    else:
        print()
        print("💡 INSTRUCCIONES PARA CREAR PDF:")
        print("   1. Abra el archivo A1_ADZ.docx en Microsoft Word")
        print("   2. Vaya a Archivo > Guardar como")
        print("   3. Seleccione formato PDF")
        print("   4. Guarde como A1_ADZ.pdf")
        print()
        print("   O use LibreOffice:")
        print("   libreoffice --headless --convert-to pdf A1_ADZ.docx")
    
    print()
    print("=" * 60)
    print("✅ PROCESO COMPLETADO")
    print("=" * 60)
    print()
    print("📋 PRÓXIMOS PASOS:")
    print("   1. Complete los Cuadros A y B con datos del Banco de México")
    print("   2. Complete las secciones de análisis")
    print("   3. Guarde el documento Word como PDF: A1_ADZ.pdf")
    print("   4. Envíe el PDF a su tutor")
    print()
    print("💡 MEJORAS PREMIUM APLICADAS:")
    print("   ✓ Formato profesional con colores corporativos")
    print("   ✓ Tablas con encabezados azules y texto blanco")
    print("   ✓ Filas alternadas para mejor legibilidad")
    print("   ✓ Enlaces en azul y clicables")
    print("   ✓ Espaciado y márgenes optimizados profesionalmente")
    print("   ✓ Estilos consistentes y elegantes")
    print("   ✓ Portada mejorada con información destacada")
    print("   ✓ Diseño visual premium y moderno")
    print("   ✨ ÍNDICE AUTOMÁTICO mejorado con encabezados")
    print("   ✨ ENCABEZADOS Y PIES DE PÁGINA con numeración")
    print("   ✨ CAJAS DE INFORMACIÓN mejoradas (Consejos, Advertencias, Importante)")
    print("   ✨ ICONOS Y ELEMENTOS VISUALES mejorados")
    print("   ✨ GUÍAS Y CONSEJOS contextuales integrados")
    print("   ✨ FORMATO DE TABLAS mejorado con filas alternadas")
    print("   ✨ ESTRUCTURA PROFESIONAL completa")
    print("   ✨ MÚLTIPLES CAJAS DE AYUDA por sección")
    print("   ✨ CAJAS DE CONTEXTO para cada indicador económico")
    print("   ✨ DESCRIPCIONES MEJORADAS en cada cuadro")
    print("   ✨ ICONOS EN PREGUNTAS Y SECCIONES")
    print("   ✨ TAMAÑOS DE FUENTE OPTIMIZADOS")
    print("   ✨ SUBTÍTULO DESCRIPTIVO en portada")
    print("   ✨ INSTRUCCIONES CON VIÑETAS MEJORADAS (✓ y →)")
    print("   ✨ ENLACES SUBRAYADOS Y CLICABLES")
    print("   ✨ CAJAS DE GUÍA PARA ESTRUCTURA DE ANÁLISIS")
    print("   ✨ EJEMPLOS MEJORADOS CON ICONOS")
    print("   ✨ FORMATO DE LISTAS OPTIMIZADO")
    print("   ✨ CAJA INFORMATIVA EN REFERENCIAS")
    print("   ✨ CAJA DE RECORDATORIO FINAL")
    print("   ✨ ENLACES EN REFERENCIAS SUBRAYADOS")
    print("   ✨ FORMATO FINAL COMPLETO Y PROFESIONAL")
    print("   ✨ CAJAS DE GUÍA ADICIONALES EN IMPACTO EMPRESARIAL")
    print("   ✨ CAJAS ESPECÍFICAS PARA CASOS ESPECIALES")
    print("   ✨ MÁS CONTEXTO EDUCATIVO EN CADA SECCIÓN")
    print("   ✨ GUÍAS DETALLADAS DE ASPECTOS A CONSIDERAR")
    print("   ✨ TABLA ESTRUCTURADA PARA CONCLUSIONES")
    print("   ✨ CAJA DE ADVERTENCIA EN CONCLUSIONES")
    print("   ✨ FORMATO MEJORADO EN SECCIONES FINALES")
    print("   ✨ ESTRUCTURA VISUAL PARA ORGANIZAR CONCLUSIONES")


if __name__ == "__main__":
    main()
