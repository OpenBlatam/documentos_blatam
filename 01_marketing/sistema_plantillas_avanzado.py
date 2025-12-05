#!/usr/bin/env python3
"""
Sistema de Plantillas Avanzado - Permite crear plantillas personalizables
para diferentes industrias, estilos y casos de uso específicos.
"""

import os
import json
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter

# Plantillas predefinidas por industria
PLANTILLAS_INDUSTRIA = {
    'tech': {
        'nombre': 'Tecnología',
        'colores': {
            'primary': '#2563EB',  # Azul tech
            'secondary': '#10B981',  # Verde éxito
            'accent': '#F59E0B',  # Naranja energía
        },
        'fuente': 'Calibri',
        'estilo': 'moderno'
    },
    'financiero': {
        'nombre': 'Financiero',
        'colores': {
            'primary': '#1F4E78',  # Azul corporativo
            'secondary': '#D4AF37',  # Dorado
            'accent': '#4CAF50',  # Verde éxito
        },
        'fuente': 'Times New Roman',
        'estilo': 'conservador'
    },
    'marketing': {
        'nombre': 'Marketing',
        'colores': {
            'primary': '#E1306C',  # Rosa Instagram
            'secondary': '#FF9800',  # Naranja
            'accent': '#2196F3',  # Azul
        },
        'fuente': 'Arial',
        'estilo': 'vibrante'
    },
    'salud': {
        'nombre': 'Salud',
        'colores': {
            'primary': '#2E7D32',  # Verde salud
            'secondary': '#1976D2',  # Azul confianza
            'accent': '#F44336',  # Rojo urgencia
        },
        'fuente': 'Arial',
        'estilo': 'profesional'
    },
    'educacion': {
        'nombre': 'Educación',
        'colores': {
            'primary': '#1976D2',  # Azul conocimiento
            'secondary': '#FFC107',  # Amarillo energía
            'accent': '#4CAF50',  # Verde crecimiento
        },
        'fuente': 'Calibri',
        'estilo': 'acogedor'
    }
}

def crear_plantilla_word_personalizada(industria, archivo_salida):
    """Crea plantilla Word personalizada según industria"""
    print(f"Creando plantilla Word para industria: {PLANTILLAS_INDUSTRIA[industria]['nombre']}...")
    
    plantilla = PLANTILLAS_INDUSTRIA[industria]
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)
    
    # Portada personalizada
    titulo_para = doc.add_paragraph()
    titulo_run = titulo_para.add_run("PLANTILLA PERSONALIZADA")
    titulo_run.font.name = plantilla['fuente']
    titulo_run.font.size = Pt(36)
    titulo_run.font.bold = True
    titulo_run.font.color.rgb = RGBColor.from_string(plantilla['colores']['primary'].replace('#', ''))
    titulo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitulo_para = doc.add_paragraph()
    subtitulo_run = subtitulo_para.add_run(f"Industria: {plantilla['nombre']}")
    subtitulo_run.font.size = Pt(18)
    subtitulo_run.font.color.rgb = RGBColor.from_string(plantilla['colores']['secondary'].replace('#', ''))
    subtitulo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Sección de ejemplo
    doc.add_heading('Sección de Ejemplo', 1)
    doc.add_paragraph('Este es un documento de plantilla personalizada según la industria seleccionada.')
    
    doc.save(archivo_salida)
    print(f"✓ Plantilla Word creada: {archivo_salida}")

def crear_plantilla_excel_personalizada(industria, archivo_salida):
    """Crea plantilla Excel personalizada según industria"""
    print(f"Creando plantilla Excel para industria: {PLANTILLAS_INDUSTRIA[industria]['nombre']}...")
    
    plantilla = PLANTILLAS_INDUSTRIA[industria]
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Dashboard"
    
    # Título
    ws['A1'] = f"Dashboard - {plantilla['nombre']}"
    ws['A1'].font = Font(bold=True, size=20, color="FFFFFF")
    ws['A1'].fill = PatternFill(start_color=plantilla['colores']['primary'].replace('#', ''), 
                               end_color=plantilla['colores']['primary'].replace('#', ''), 
                               fill_type="solid")
    ws.merge_cells('A1:D1')
    ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
    
    # KPIs de ejemplo
    kpis = [
        ['Métrica 1', '100', plantilla['colores']['primary']],
        ['Métrica 2', '200%', plantilla['colores']['secondary']],
        ['Métrica 3', '50', plantilla['colores']['accent']],
    ]
    
    for row, (kpi, valor, color) in enumerate(kpis, 3):
        ws.cell(row=row, column=1, value=kpi).font = Font(bold=True)
        cell_valor = ws.cell(row=row, column=2, value=valor)
        cell_valor.font = Font(bold=True, size=16, color="FFFFFF")
        cell_valor.fill = PatternFill(start_color=color.replace('#', ''), 
                                      end_color=color.replace('#', ''), 
                                      fill_type="solid")
        cell_valor.alignment = Alignment(horizontal='center')
    
    wb.save(archivo_salida)
    print(f"✓ Plantilla Excel creada: {archivo_salida}")

def crear_configuracion_plantilla(industria, archivo_config):
    """Crea archivo de configuración JSON para plantilla"""
    plantilla = PLANTILLAS_INDUSTRIA[industria]
    
    config = {
        'industria': industria,
        'nombre': plantilla['nombre'],
        'colores': plantilla['colores'],
        'fuente': plantilla['fuente'],
        'estilo': plantilla['estilo'],
        'fecha_creacion': datetime.now().isoformat(),
        'version': '1.0'
    }
    
    with open(archivo_config, 'w', encoding='utf-8') as f:
        json.dump(config, f, indent=2, ensure_ascii=False)
    
    print(f"✓ Configuración guardada: {archivo_config}")

def main():
    """Genera todas las plantillas disponibles"""
    directorio = '/Users/adan/Documents/documentos_blatam/01_marketing/plantillas'
    os.makedirs(directorio, exist_ok=True)
    
    print("🎨 Generando plantillas personalizadas por industria...\n")
    
    for industria in PLANTILLAS_INDUSTRIA.keys():
        print(f"\n📋 Procesando: {PLANTILLAS_INDUSTRIA[industria]['nombre']}")
        
        # Word
        archivo_word = os.path.join(directorio, f'plantilla_{industria}.docx')
        crear_plantilla_word_personalizada(industria, archivo_word)
        
        # Excel
        archivo_excel = os.path.join(directorio, f'plantilla_{industria}.xlsx')
        crear_plantilla_excel_personalizada(industria, archivo_excel)
        
        # Configuración
        archivo_config = os.path.join(directorio, f'config_{industria}.json')
        crear_configuracion_plantilla(industria, archivo_config)
    
    print("\n✅ Todas las plantillas generadas!")
    print(f"📁 Ubicación: {directorio}")

if __name__ == "__main__":
    main()








