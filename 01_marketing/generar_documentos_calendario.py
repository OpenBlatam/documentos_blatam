#!/usr/bin/env python3
"""
Script para generar documentos Word (.docx) y PDF a partir del 
Sistema de Calendario de Contenido de Redes Sociales en Markdown
"""

import os
import re
from datetime import datetime
from pathlib import Path

def markdown_to_docx(md_file, output_file):
    """Convierte Markdown a Word (.docx) con formato profesional"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        from docx.oxml.ns import qn
    except ImportError:
        print("❌ Error: Se requiere la librería python-docx")
        print("   Instala con: pip install python-docx")
        return False
    
    # Leer el archivo Markdown
    try:
        with open(md_file, 'r', encoding='utf-8') as f:
            content = f.read()
    except FileNotFoundError:
        print(f"❌ Error: No se encontró el archivo {md_file}")
        return False
    
    # Crear documento Word
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Configurar estilos personalizados
    styles = doc.styles
    
    # Color corporativo (azul/púrpura)
    primary_color = RGBColor(102, 126, 234)  # #667eea
    secondary_color = RGBColor(118, 75, 162)  # #764ba2
    
    # Procesar el contenido línea por línea
    lines = content.split('\n')
    i = 0
    in_table = False
    table_lines = []
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Saltar líneas vacías al inicio
        if not line and i == 0:
            i += 1
            continue
        
        try:
            # Título principal con emoji
            if line.startswith('# 📅'):
                title = line[2:].strip()
                # Remover emoji para el título
                title = re.sub(r'[📅📋🎯📐🚫📝🧠📤🎨📊🔧📈📌]', '', title).strip()
                p = doc.add_heading(title, level=1)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Aplicar color al título
                for run in p.runs:
                    run.font.color.rgb = primary_color
                    run.font.size = Pt(24)
                    run.font.bold = True
            
            # Títulos con emojis
            elif re.match(r'^##\s+[📋🎯📐🚫📝🧠📤🎨📊🔧📈📌]', line):
                title = re.sub(r'^##\s+[📋🎯📐🚫📝🧠📤🎨📊🔧📈📌]\s*', '', line).strip()
                p = doc.add_heading(title, level=2)
                for run in p.runs:
                    run.font.color.rgb = secondary_color
                    run.font.size = Pt(18)
            
            # Títulos de nivel 2
            elif line.startswith('## '):
                title = line[3:].strip()
                p = doc.add_heading(title, level=2)
                for run in p.runs:
                    run.font.color.rgb = secondary_color
                    run.font.size = Pt(18)
            
            # Títulos de nivel 3
            elif line.startswith('### '):
                title = line[4:].strip()
                doc.add_heading(title, level=3)
            
            # Títulos de nivel 4
            elif line.startswith('#### '):
                title = line[5:].strip()
                doc.add_heading(title, level=4)
            
            # Tablas
            elif line.startswith('|') and '|' in line:
                if not in_table:
                    in_table = True
                    table_lines = []
                table_lines.append(line)
            
            # Procesar tabla cuando termina
            elif in_table and not line.startswith('|'):
                in_table = False
                if table_lines:
                    _process_table(doc, table_lines)
                    table_lines = []
            
            # Listas con viñetas
            elif line.startswith('- ') or line.startswith('* '):
                items = []
                while i < len(lines) and (lines[i].strip().startswith('- ') or 
                                         lines[i].strip().startswith('* ')):
                    item = lines[i].strip()[2:].strip()
                    # Procesar texto en negrita y cursiva
                    item = _process_markdown_formatting(item)
                    items.append(item)
                    i += 1
                i -= 1
                
                for item in items:
                    p = doc.add_paragraph(item, style='List Bullet')
            
            # Listas numeradas
            elif re.match(r'^\d+\.\s+', line):
                items = []
                while i < len(lines) and re.match(r'^\d+\.\s+', lines[i].strip()):
                    item = lines[i].strip()
                    match = re.match(r'^\d+\.\s*(.*)', item)
                    if match:
                        item_text = _process_markdown_formatting(match.group(1))
                        items.append(item_text)
                    i += 1
                i -= 1
                
                for item in items:
                    p = doc.add_paragraph(item, style='List Number')
            
            # Citas (blockquotes)
            elif line.startswith('> '):
                quote_text = line[2:].strip()
                quote_text = _process_markdown_formatting(quote_text)
                p = doc.add_paragraph(quote_text, style='Intense Quote')
            
            # Separadores
            elif line.startswith('---'):
                # Agregar espacio en lugar de línea
                doc.add_paragraph()
            
            # Párrafos normales
            elif line and not line.startswith('```'):
                text = _process_markdown_formatting(line)
                p = doc.add_paragraph(text)
        
        except Exception as e:
            print(f"⚠️ Error procesando línea {i+1}: {e}")
        
        i += 1
    
    # Procesar última tabla si existe
    if in_table and table_lines:
        _process_table(doc, table_lines)
    
    # Guardar documento
    try:
        doc.save(output_file)
        print(f"✅ Documento Word creado exitosamente: {output_file}")
        return True
    except Exception as e:
        print(f"❌ Error al guardar documento Word: {e}")
        return False

def _process_table(doc, table_lines):
    """Procesa una tabla Markdown y la agrega al documento Word"""
    from docx import Document
    
    rows = []
    for table_line in table_lines:
        if '---' not in table_line:  # Saltar líneas de separación
            cells = [cell.strip() for cell in table_line.split('|')[1:-1]]
            if cells:
                rows.append(cells)
    
    if len(rows) < 2:  # Necesitamos al menos encabezado y una fila
        return
    
    # Crear tabla
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'
    
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_data in enumerate(row_data):
            if col_idx < len(table.rows[row_idx].cells):
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = _process_markdown_formatting(cell_data, remove_formatting=True)
                # Hacer la primera fila en negrita (encabezado)
                if row_idx == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

def _process_markdown_formatting(text, remove_formatting=False):
    """Procesa formato Markdown (negrita, cursiva, código)"""
    if remove_formatting:
        # Remover todo el formato
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'\*(.*?)\*', r'\1', text)
        text = re.sub(r'`(.*?)`', r'\1', text)
        text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)  # Links
    else:
        # Mantener el texto pero limpiar formato básico
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Negrita
        text = re.sub(r'\*(.*?)\*', r'\1', text)  # Cursiva
        text = re.sub(r'`(.*?)`', r'\1', text)  # Código
        text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)  # Links
    
    return text

def markdown_to_pdf(md_file, output_file):
    """Convierte Markdown a PDF usando markdown2 y weasyprint"""
    try:
        import markdown2
        from weasyprint import HTML, CSS
        from weasyprint.text.fonts import FontConfiguration
    except ImportError:
        print("❌ Error: Se requieren las librerías markdown2 y weasyprint")
        print("   Instala con: pip install markdown2 weasyprint")
        return False
    
    # Leer el archivo Markdown
    try:
        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
    except FileNotFoundError:
        print(f"❌ Error: No se encontró el archivo {md_file}")
        return False
    
    # Convertir Markdown a HTML
    html_content = markdown2.markdown(
        md_content,
        extras=['tables', 'fenced-code-blocks', 'toc']
    )
    
    # HTML completo con estilos profesionales
    full_html = f"""
    <!DOCTYPE html>
    <html lang="es">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Sistema de Calendario de Contenido de Redes Sociales</title>
        <style>
            @page {{
                size: A4;
                margin: 2cm;
            }}
            body {{
                font-family: 'Segoe UI', 'Helvetica Neue', Arial, sans-serif;
                line-height: 1.6;
                color: #333;
                font-size: 11pt;
            }}
            h1 {{
                color: #667eea;
                border-bottom: 3px solid #667eea;
                padding-bottom: 10px;
                font-size: 24pt;
                text-align: center;
                margin-top: 20px;
                margin-bottom: 30px;
            }}
            h2 {{
                color: #764ba2;
                margin-top: 30px;
                margin-bottom: 15px;
                font-size: 18pt;
                border-left: 4px solid #764ba2;
                padding-left: 10px;
            }}
            h3 {{
                color: #555;
                margin-top: 25px;
                margin-bottom: 12px;
                font-size: 14pt;
            }}
            h4 {{
                color: #666;
                margin-top: 20px;
                margin-bottom: 10px;
                font-size: 12pt;
            }}
            p {{
                margin: 10px 0;
                text-align: justify;
            }}
            ul, ol {{
                padding-left: 25px;
                margin: 15px 0;
            }}
            li {{
                margin: 8px 0;
            }}
            table {{
                width: 100%;
                border-collapse: collapse;
                margin: 20px 0;
                font-size: 10pt;
            }}
            th {{
                background-color: #667eea;
                color: white;
                padding: 10px;
                text-align: left;
                font-weight: bold;
            }}
            td {{
                padding: 8px;
                border: 1px solid #ddd;
            }}
            tr:nth-child(even) {{
                background-color: #f8f9fa;
            }}
            blockquote {{
                border-left: 4px solid #667eea;
                padding-left: 15px;
                margin: 20px 0;
                font-style: italic;
                color: #555;
            }}
            code {{
                background-color: #f4f4f4;
                padding: 2px 6px;
                border-radius: 3px;
                font-family: 'Courier New', monospace;
                font-size: 10pt;
            }}
            pre {{
                background-color: #f4f4f4;
                padding: 15px;
                border-radius: 5px;
                overflow-x: auto;
            }}
            hr {{
                border: none;
                border-top: 2px solid #ddd;
                margin: 30px 0;
            }}
            .toc {{
                background-color: #f8f9fa;
                padding: 20px;
                border-radius: 5px;
                margin: 20px 0;
            }}
        </style>
    </head>
    <body>
        {html_content}
    </body>
    </html>
    """
    
    # Generar PDF
    try:
        HTML(string=full_html).write_pdf(output_file)
        print(f"✅ Documento PDF creado exitosamente: {output_file}")
        return True
    except Exception as e:
        print(f"❌ Error al generar PDF: {e}")
        print("💡 Alternativa: Usa el archivo HTML generado y conviértelo manualmente")
        # Guardar HTML como alternativa
        html_file = output_file.replace('.pdf', '.html')
        with open(html_file, 'w', encoding='utf-8') as f:
            f.write(full_html)
        print(f"✅ Archivo HTML generado como alternativa: {html_file}")
        return False

def main():
    """Función principal"""
    print("🚀 Generando documentos del Sistema de Calendario de Contenido...")
    print("=" * 60)
    
    # Archivos
    base_dir = Path(__file__).parent
    md_file = base_dir / "SISTEMA_CALENDARIO_CONTENIDO_REDES_SOCIALES.md"
    
    if not md_file.exists():
        print(f"❌ Error: No se encontró el archivo {md_file}")
        return
    
    # Generar Word
    print("\n📄 Generando documento Word (.docx)...")
    docx_file = base_dir / "SISTEMA_CALENDARIO_CONTENIDO_REDES_SOCIALES.docx"
    markdown_to_docx(str(md_file), str(docx_file))
    
    # Generar PDF
    print("\n📄 Generando documento PDF...")
    pdf_file = base_dir / "SISTEMA_CALENDARIO_CONTENIDO_REDES_SOCIALES.pdf"
    markdown_to_pdf(str(md_file), str(pdf_file))
    
    print("\n" + "=" * 60)
    print("✅ Proceso completado!")
    print(f"\n📁 Archivos generados en: {base_dir}")
    print(f"   • Markdown: {md_file.name}")
    if docx_file.exists():
        print(f"   • Word: {docx_file.name}")
    if pdf_file.exists():
        print(f"   • PDF: {pdf_file.name}")
    
    print("\n💡 Notas:")
    print("   • Si el PDF no se generó, revisa que tengas weasyprint instalado")
    print("   • Puedes abrir el HTML generado en un navegador y guardar como PDF")
    print("   • El documento Word puede requerir ajustes manuales de formato")

if __name__ == "__main__":
    main()




