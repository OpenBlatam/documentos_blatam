"""
WordProcessor - Procesador profesional para documentos Word
"""

import os
from pathlib import Path
from typing import Dict, Any, List
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from ..config import ConfigManager
from ..logger import LoggerManager
from ..utils.image_handler import ImageHandler

class WordProcessor:
    """Procesador avanzado para documentos Word"""
    
    def __init__(self, config: ConfigManager, logger: LoggerManager, 
                 image_handler: ImageHandler):
        self.config = config
        self.logger = logger.get_logger()
        self.image_handler = image_handler
        self.word_config = config.config.word
    
    def convert(self, parsed_content: Dict[str, Any], 
                base_name: str, output_dir: str) -> str:
        """Convierte contenido parseado a Word"""
        doc = Document()
        
        # Configurar documento
        self._setup_document(doc)
        
        # Crear estilos personalizados
        self._create_custom_styles(doc)
        
        # Agregar portada
        if self.word_config.include_header:
            self._add_cover_page(doc, parsed_content.get('title', base_name))
        
        # Agregar tabla de contenidos
        if self.word_config.include_toc:
            self._add_table_of_contents(doc)
        
        # Procesar contenido
        self._process_content(doc, parsed_content)
        
        # Agregar pie de página
        if self.word_config.include_footer:
            self._add_footer(doc)
        
        # Guardar
        output_file = os.path.join(output_dir, f"{base_name}.docx")
        doc.save(output_file)
        
        self.logger.info(f"Documento Word creado: {output_file}")
        return output_file
    
    def _setup_document(self, doc: Document):
        """Configura el documento"""
        # Configurar márgenes
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(self.word_config.page_margin_top)
            section.bottom_margin = Cm(self.word_config.page_margin_bottom)
            section.left_margin = Cm(self.word_config.page_margin_left)
            section.right_margin = Cm(self.word_config.page_margin_right)
    
    def _create_custom_styles(self, doc: Document):
        """Crea estilos personalizados"""
        styles = doc.styles
        
        # Estilo para títulos
        if 'CustomTitle' not in [s.name for s in styles]:
            title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = self.word_config.title_font
            title_font.size = Pt(self.word_config.title_font_size)
            title_font.bold = True
            title_font.color.rgb = RGBColor(31, 78, 120)  # #1F4E78
        
        # Estilo para encabezados
        if 'CustomHeading' not in [s.name for s in styles]:
            heading_style = styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
            heading_font = heading_style.font
            heading_font.name = self.word_config.title_font
            heading_font.size = Pt(self.word_config.heading_font_size)
            heading_font.bold = True
            heading_font.color.rgb = RGBColor(31, 78, 120)
    
    def _add_cover_page(self, doc: Document, title: str):
        """Agrega portada"""
        # Título principal
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.name = self.word_config.title_font
        title_run.font.size = Pt(self.word_config.title_font_size * 1.5)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(31, 78, 120)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Espaciado
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Fecha
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(datetime.now().strftime('%d de %B de %Y'))
        date_run.font.size = Pt(12)
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Salto de página
        doc.add_page_break()
    
    def _add_table_of_contents(self, doc: Document):
        """Agrega tabla de contenidos"""
        toc_para = doc.add_paragraph()
        toc_run = toc_para.add_run("Tabla de Contenidos")
        toc_run.font.bold = True
        toc_run.font.size = Pt(14)
        
        doc.add_paragraph("(La tabla de contenidos se generará automáticamente)")
        doc.add_paragraph()
    
    def _process_content(self, doc: Document, parsed_content: Dict[str, Any]):
        """Procesa el contenido"""
        # Procesar secciones
        for section in parsed_content.get('sections', []):
            # Título de sección
            if 'title' in section:
                heading = doc.add_heading(section['title'], level=1)
                heading.style = 'CustomHeading'
            
            # Contenido
            if 'content' in section:
                self._add_paragraphs(doc, section['content'])
            
            # Imágenes
            if 'images' in section:
                for image_path in section['images']:
                    self._add_image(doc, image_path)
            
            # Tablas
            if 'tables' in section:
                for table_data in section['tables']:
                    self._add_table(doc, table_data)
    
    def _add_paragraphs(self, doc: Document, content: str):
        """Agrega párrafos"""
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph(para_text.strip())
                para.style.font.name = self.word_config.default_font
                para.style.font.size = Pt(self.word_config.default_font_size)
    
    def _add_image(self, doc: Document, image_path: str):
        """Agrega imagen"""
        try:
            if os.path.exists(image_path):
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(image_path, width=Inches(5))
        except Exception as e:
            self.logger.warning(f"Error agregando imagen {image_path}: {e}")
    
    def _add_table(self, doc: Document, table_data: List[List[str]]):
        """Agrega tabla"""
        if not table_data:
            return
        
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        table.style = 'Light Grid Accent 1'
        
        for i, row_data in enumerate(table_data):
            for j, cell_data in enumerate(row_data):
                cell = table.rows[i].cells[j]
                cell.text = str(cell_data)
                
                # Formato de encabezado
                if i == 0:
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                    shading = cell._element.get_or_add_tcPr().get_or_add_shd()
                    shading.set(qn('w:fill'), '1F4E78')
    
    def _add_footer(self, doc: Document):
        """Agrega pie de página"""
        for section in doc.sections:
            footer = section.footer
            footer_para = footer.paragraphs[0]
            
            if self.word_config.footer_text:
                footer_para.text = self.word_config.footer_text
            else:
                footer_para.text = f"Generado el {datetime.now().strftime('%d/%m/%Y %H:%M')}"
            
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_para.runs[0].font.size = Pt(9)








