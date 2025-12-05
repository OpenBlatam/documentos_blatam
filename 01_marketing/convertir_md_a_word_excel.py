#!/usr/bin/env python3
"""
Script profesional para convertir archivos Markdown importantes de marketing
a Word (.docx) y Excel (.xlsx) con gráficas, imágenes y formato de alta calidad.

Librerías utilizadas:
- python-docx: Documentos Word profesionales
- openpyxl: Excel con gráficas avanzadas
- pandas: Manipulación de datos
- matplotlib: Gráficas estáticas
- plotly: Gráficas interactivas
- seaborn: Visualizaciones estadísticas
- markdown: Parseo de Markdown
- Pillow: Procesamiento de imágenes
- BeautifulSoup: HTML parsing
"""

import os
import re
import sys
from pathlib import Path
from datetime import datetime
import json

# Importar librerías principales
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx no está instalado. Instala con: pip install python-docx")
    sys.exit(1)

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.chart import BarChart, LineChart, PieChart, ScatterChart
    from openpyxl.chart.series import DataPoint
    from openpyxl.drawing.image import Image
    from openpyxl.utils import get_column_letter
except ImportError:
    print("ERROR: openpyxl no está instalado. Instala con: pip install openpyxl")
    sys.exit(1)

try:
    import pandas as pd
    import numpy as np
except ImportError:
    print("ERROR: pandas/numpy no están instalados. Instala con: pip install pandas numpy")
    sys.exit(1)

try:
    import matplotlib
    matplotlib.use('Agg')  # Backend sin GUI
    import matplotlib.pyplot as plt
    import seaborn as sns
    from matplotlib.patches import Rectangle
except ImportError:
    print("ERROR: matplotlib/seaborn no están instalados. Instala con: pip install matplotlib seaborn")
    sys.exit(1)

try:
    import plotly.graph_objects as go
    import plotly.express as px
except ImportError:
    print("ADVERTENCIA: plotly no está instalado. Algunas gráficas interactivas no estarán disponibles.")
    plotly = None

try:
    import markdown
    from markdown.extensions import tables, fenced_code, codehilite
except ImportError:
    print("ERROR: markdown no está instalado. Instala con: pip install markdown")
    sys.exit(1)

try:
    from PIL import Image as PILImage
    from PIL import ImageDraw, ImageFont
except ImportError:
    print("ADVERTENCIA: Pillow no está instalado. Procesamiento de imágenes limitado.")
    PILImage = None

# Configuración
BASE_DIR = Path(__file__).parent
OUTPUT_DIR = BASE_DIR / "documentos_convertidos"
OUTPUT_DIR.mkdir(exist_ok=True)

# Archivos importantes a convertir
ARCHIVOS_IMPORTANTES = [
    "SISTEMAS_PROMPTS_CONSOLIDADO.md",
    "RESUMEN_EJECUTIVO_COMPLETO.md",
    "PRESUPUESTO_PRICING.md",
    "DASHBOARD_METRICAS.md",
    "ESTRATEGIAS_CONTENIDO.md",
    "AUTOMATIZACION_AVANZADA.md",
    "SISTEMA_CALENDARIO_CONTENIDO_REDES_SOCIALES.md",
    "GUIA_MEJORES_PRACTICAS_AVANZADAS.md"
]

# Configuración de estilos
ESTILOS_WORD = {
    'titulo_principal': {
        'font_size': 24,
        'bold': True,
        'color': RGBColor(0, 51, 102),
        'alignment': WD_ALIGN_PARAGRAPH.CENTER
    },
    'titulo_seccion': {
        'font_size': 18,
        'bold': True,
        'color': RGBColor(0, 102, 204),
        'alignment': WD_ALIGN_PARAGRAPH.LEFT
    },
    'subtitulo': {
        'font_size': 14,
        'bold': True,
        'color': RGBColor(51, 51, 51),
        'alignment': WD_ALIGN_PARAGRAPH.LEFT
    },
    'texto_normal': {
        'font_size': 11,
        'bold': False,
        'color': RGBColor(0, 0, 0),
        'alignment': WD_ALIGN_PARAGRAPH.LEFT
    }
}

COLORES_EXCEL = {
    'header': '366092',
    'subheader': '4F81BD',
    'accent': '95B3D7',
    'success': '92D050',
    'warning': 'FFC000',
    'error': 'FF0000',
    'info': '00B0F0'
}


def parsear_markdown(archivo_md):
    """Parsea un archivo Markdown y extrae su contenido estructurado."""
    with open(archivo_md, 'r', encoding='utf-8') as f:
        contenido = f.read()
    
    # Usar markdown para convertir a HTML
    md = markdown.Markdown(extensions=['tables', 'fenced_code', 'codehilite', 'nl2br'])
    html = md.convert(contenido)
    
    # Extraer estructura
    estructura = {
        'titulo': extraer_titulo(contenido),
        'secciones': extraer_secciones(contenido),
        'tablas': extraer_tablas(contenido),
        'codigo': extraer_codigo(contenido),
        'enlaces': extraer_enlaces(contenido),
        'imagenes': extraer_imagenes(contenido),
        'listas': extraer_listas(contenido),
        'contenido_html': html,
        'contenido_raw': contenido
    }
    
    return estructura


def extraer_titulo(contenido):
    """Extrae el título principal del documento."""
    match = re.search(r'^#\s+(.+)$', contenido, re.MULTILINE)
    return match.group(1).strip() if match else "Documento sin título"


def extraer_secciones(contenido):
    """Extrae todas las secciones del documento."""
    secciones = []
    patron = r'^(#{1,6})\s+(.+)$'
    
    for match in re.finditer(patron, contenido, re.MULTILINE):
        nivel = len(match.group(1))
        titulo = match.group(2).strip()
        secciones.append({'nivel': nivel, 'titulo': titulo})
    
    return secciones


def extraer_tablas(contenido):
    """Extrae tablas del contenido Markdown."""
    tablas = []
    patron = r'\|(.+)\|\n\|[-\s\|]+\|\n((?:\|.+\|\n?)+)'
    
    for match in re.finditer(patron, contenido):
        headers = [h.strip() for h in match.group(1).split('|') if h.strip()]
        rows = []
        
        for row in match.group(2).strip().split('\n'):
            if row.strip():
                cells = [c.strip() for c in row.split('|') if c.strip()]
                if cells:
                    rows.append(cells)
        
        tablas.append({'headers': headers, 'rows': rows})
    
    return tablas


def extraer_codigo(contenido):
    """Extrae bloques de código."""
    bloques = []
    patron = r'```(\w+)?\n(.*?)```'
    
    for match in re.finditer(patron, contenido, re.DOTALL):
        lenguaje = match.group(1) or 'text'
        codigo = match.group(2).strip()
        bloques.append({'lenguaje': lenguaje, 'codigo': codigo})
    
    return bloques


def extraer_enlaces(contenido):
    """Extrae enlaces del contenido."""
    enlaces = []
    patron = r'\[([^\]]+)\]\(([^\)]+)\)'
    
    for match in re.finditer(patron, contenido):
        enlaces.append({'texto': match.group(1), 'url': match.group(2)})
    
    return enlaces


def extraer_imagenes(contenido):
    """Extrae referencias a imágenes."""
    imagenes = []
    patron = r'!\[([^\]]*)\]\(([^\)]+)\)'
    
    for match in re.finditer(patron, contenido):
        imagenes.append({'alt': match.group(1), 'ruta': match.group(2)})
    
    return imagenes


def extraer_listas(contenido):
    """Extrae listas del contenido."""
    listas = []
    patron = r'^([\s]*)[-*+]\s+(.+)$'
    
    for match in re.finditer(patron, contenido, re.MULTILINE):
        indentacion = len(match.group(1))
        item = match.group(2).strip()
        listas.append({'indentacion': indentacion, 'item': item})
    
    return listas


def crear_documento_word(estructura, nombre_archivo):
    """Crea un documento Word profesional desde la estructura parseada."""
    doc = Document()
    
    # Configurar estilos
    configurar_estilos_word(doc)
    
    # Agregar título principal
    titulo = doc.add_heading(estructura['titulo'], 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Agregar metadatos
    doc.core_properties.title = estructura['titulo']
    doc.core_properties.author = "Sistema de Conversión Automática"
    doc.core_properties.comments = f"Convertido desde {nombre_archivo} el {datetime.now().strftime('%Y-%m-%d')}"
    
    # Procesar contenido por secciones
    contenido_raw = estructura['contenido_raw']
    lineas = contenido_raw.split('\n')
    
    i = 0
    while i < len(lineas):
        linea = lineas[i].strip()
        
        if not linea:
            i += 1
            continue
        
        # Títulos
        if linea.startswith('#'):
            nivel = len(linea) - len(linea.lstrip('#'))
            titulo_texto = linea.lstrip('#').strip()
            doc.add_heading(titulo_texto, nivel)
        
        # Tablas
        elif linea.startswith('|'):
            tabla_data = []
            j = i
            while j < len(lineas) and lineas[j].strip().startswith('|'):
                if not lineas[j].strip().startswith('|---'):
                    tabla_data.append(lineas[j].strip())
                j += 1
            
            if tabla_data:
                crear_tabla_word(doc, tabla_data)
                i = j - 1
        
        # Código
        elif linea.startswith('```'):
            lenguaje = linea[3:].strip()
            codigo = []
            j = i + 1
            while j < len(lineas) and not lineas[j].strip().startswith('```'):
                codigo.append(lineas[j])
                j += 1
            
            if codigo:
                agregar_codigo_word(doc, '\n'.join(codigo), lenguaje)
                i = j
        
        # Listas
        elif linea.startswith('-') or linea.startswith('*') or linea.startswith('+'):
            items = []
            j = i
            while j < len(lineas) and (lineas[j].strip().startswith('-') or 
                                      lineas[j].strip().startswith('*') or 
                                      lineas[j].strip().startswith('+') or
                                      lineas[j].strip().startswith(' ')):
                items.append(lineas[j])
                j += 1
            
            if items:
                crear_lista_word(doc, items)
                i = j - 1
        
        # Texto normal
        else:
            if linea and not linea.startswith('---'):
                parrafo = doc.add_paragraph(linea)
                formatear_parrafo(parrafo, linea)
        
        i += 1
    
    # Agregar tablas de datos si existen
    if estructura['tablas']:
        doc.add_page_break()
        doc.add_heading('Tablas de Datos', 1)
        for tabla in estructura['tablas']:
            crear_tabla_desde_estructura(doc, tabla)
    
    # Guardar documento
    nombre_salida = OUTPUT_DIR / f"{Path(nombre_archivo).stem}.docx"
    doc.save(str(nombre_salida))
    print(f"✅ Word creado: {nombre_salida}")
    
    return nombre_salida


def configurar_estilos_word(doc):
    """Configura estilos personalizados para el documento Word."""
    styles = doc.styles
    
    # Estilo para código
    try:
        code_style = styles['Code']
    except KeyError:
        code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = 'Courier New'
        code_style.font.size = Pt(9)
        code_style.paragraph_format.space_after = Pt(6)


def crear_tabla_word(doc, datos_tabla):
    """Crea una tabla en Word desde datos Markdown."""
    if not datos_tabla:
        return
    
    # Parsear primera fila como headers
    headers = [cell.strip() for cell in datos_tabla[0].split('|') if cell.strip()]
    
    if not headers:
        return
    
    # Crear tabla
    tabla = doc.add_table(rows=1, cols=len(headers))
    tabla.style = 'Light Grid Accent 1'
    
    # Agregar headers
    header_cells = tabla.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        header_cells[i].shading.background_color = RGBColor(54, 96, 146)
    
    # Agregar filas de datos
    for fila_data in datos_tabla[1:]:
        cells = [cell.strip() for cell in fila_data.split('|') if cell.strip()]
        if len(cells) == len(headers):
            row = tabla.add_row()
            for i, cell in enumerate(cells):
                row.cells[i].text = cell


def crear_tabla_desde_estructura(doc, tabla_estructura):
    """Crea una tabla desde la estructura parseada."""
    headers = tabla_estructura['headers']
    rows = tabla_estructura['rows']
    
    if not headers:
        return
    
    tabla = doc.add_table(rows=1, cols=len(headers))
    tabla.style = 'Light Grid Accent 1'
    
    # Headers
    header_cells = tabla.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
    
    # Filas
    for row_data in rows:
        if len(row_data) == len(headers):
            row = tabla.add_row()
            for i, cell in enumerate(row_data):
                row.cells[i].text = cell


def agregar_codigo_word(doc, codigo, lenguaje):
    """Agrega un bloque de código formateado."""
    parrafo = doc.add_paragraph()
    run = parrafo.add_run(codigo)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    parrafo.style = 'Code'


def crear_lista_word(doc, items):
    """Crea una lista en Word."""
    for item in items:
        item_texto = re.sub(r'^[\s]*[-*+]\s+', '', item).strip()
        if item_texto:
            doc.add_paragraph(item_texto, style='List Bullet')


def formatear_parrafo(parrafo, texto):
    """Formatea un párrafo con negritas, cursivas y enlaces."""
    # Buscar negritas
    texto = re.sub(r'\*\*(.+?)\*\*', r'<bold>\1</bold>', texto)
    texto = re.sub(r'__(.+?)__', r'<bold>\1</bold>', texto)
    
    # Buscar cursivas
    texto = re.sub(r'\*(.+?)\*', r'<italic>\1</italic>', texto)
    texto = re.sub(r'_(.+?)_', r'<italic>\1</italic>', texto)
    
    # Buscar enlaces
    texto = re.sub(r'\[([^\]]+)\]\(([^\)]+)\)', r'<link>\1|\2</link>', texto)
    
    # Procesar formato
    partes = re.split(r'(<bold>.*?</bold>|<italic>.*?</italic>|<link>.*?</link>)', texto)
    
    parrafo.clear()
    for parte in partes:
        if parte.startswith('<bold>'):
            texto_bold = parte.replace('<bold>', '').replace('</bold>', '')
            run = parrafo.add_run(texto_bold)
            run.bold = True
        elif parte.startswith('<italic>'):
            texto_italic = parte.replace('<italic>', '').replace('</italic>', '')
            run = parrafo.add_run(texto_italic)
            run.italic = True
        elif parte.startswith('<link>'):
            link_text = parte.replace('<link>', '').replace('</link>', '')
            if '|' in link_text:
                texto_link, url = link_text.split('|', 1)
                run = parrafo.add_run(texto_link)
                run.font.color.rgb = RGBColor(0, 102, 204)
                run.underline = True
                # Agregar hipervínculo
                add_hyperlink(parrafo, url, texto_link)
            else:
                parrafo.add_run(link_text)
        else:
            if parte:
                parrafo.add_run(parte)


def add_hyperlink(paragraph, url, text):
    """Agrega un hipervínculo a un párrafo."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    color = OxmlElement('w:color')
    color.set(qn('w:val'), "0563C1")
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)


def crear_documento_excel(estructura, nombre_archivo):
    """Crea un documento Excel profesional con gráficas desde la estructura."""
    wb = Workbook()
    
    # Eliminar hoja por defecto
    if 'Sheet' in wb.sheetnames:
        wb.remove(wb['Sheet'])
    
    # Hoja 1: Resumen Ejecutivo
    ws_resumen = wb.create_sheet("Resumen Ejecutivo", 0)
    crear_hoja_resumen(ws_resumen, estructura)
    
    # Hoja 2: Contenido Principal
    ws_contenido = wb.create_sheet("Contenido", 1)
    crear_hoja_contenido(ws_contenido, estructura)
    
    # Hoja 3: Tablas de Datos
    if estructura['tablas']:
        ws_tablas = wb.create_sheet("Tablas de Datos", 2)
        crear_hoja_tablas(ws_tablas, estructura['tablas'])
    
    # Hoja 4: Métricas y KPIs
    ws_metricas = wb.create_sheet("Métricas", 3)
    crear_hoja_metricas(ws_metricas, estructura)
    
    # Hoja 5: Gráficas
    ws_graficas = wb.create_sheet("Gráficas", 4)
    crear_hoja_graficas(ws_graficas, estructura)
    
    # Guardar
    nombre_salida = OUTPUT_DIR / f"{Path(nombre_archivo).stem}.xlsx"
    wb.save(str(nombre_salida))
    print(f"✅ Excel creado: {nombre_salida}")
    
    return nombre_salida


def crear_hoja_resumen(ws, estructura):
    """Crea la hoja de resumen ejecutivo."""
    ws['A1'] = estructura['titulo']
    ws['A1'].font = Font(size=20, bold=True, color=COLORES_EXCEL['header'])
    ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
    ws.merge_cells('A1:D1')
    
    ws['A3'] = "Información del Documento"
    ws['A3'].font = Font(size=14, bold=True)
    
    ws['A4'] = "Título:"
    ws['B4'] = estructura['titulo']
    ws['A5'] = "Número de Secciones:"
    ws['B5'] = len(estructura['secciones'])
    ws['A6'] = "Número de Tablas:"
    ws['B6'] = len(estructura['tablas'])
    ws['A7'] = "Número de Enlaces:"
    ws['B7'] = len(estructura['enlaces'])
    ws['A8'] = "Fecha de Conversión:"
    ws['B8'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    
    # Ajustar columnas
    ws.column_dimensions['A'].width = 25
    ws.column_dimensions['B'].width = 50


def crear_hoja_contenido(ws, estructura):
    """Crea la hoja con el contenido principal."""
    fila = 1
    
    for seccion in estructura['secciones']:
        nivel = seccion['nivel']
        titulo = seccion['titulo']
        
        celda = ws.cell(row=fila, column=1)
        celda.value = titulo
        celda.font = Font(size=16 - nivel, bold=True)
        celda.fill = PatternFill(start_color=COLORES_EXCEL['accent'], 
                                 end_color=COLORES_EXCEL['accent'], 
                                 fill_type='solid')
        
        fila += 1
    
    ws.column_dimensions['A'].width = 80


def crear_hoja_tablas(ws, tablas):
    """Crea una hoja con todas las tablas."""
    fila_inicio = 1
    
    for idx, tabla in enumerate(tablas):
        # Título de la tabla
        ws.cell(row=fila_inicio, column=1, value=f"Tabla {idx + 1}")
        ws.cell(row=fila_inicio, column=1).font = Font(size=12, bold=True)
        fila_inicio += 1
        
        # Headers
        headers = tabla['headers']
        for col, header in enumerate(headers, 1):
            celda = ws.cell(row=fila_inicio, column=col)
            celda.value = header
            celda.font = Font(bold=True, color="FFFFFF")
            celda.fill = PatternFill(start_color=COLORES_EXCEL['header'],
                                    end_color=COLORES_EXCEL['header'],
                                    fill_type='solid')
            celda.alignment = Alignment(horizontal='center')
        
        fila_inicio += 1
        
        # Datos
        for row_data in tabla['rows']:
            for col, cell_value in enumerate(row_data[:len(headers)], 1):
                ws.cell(row=fila_inicio, column=col, value=cell_value)
            fila_inicio += 1
        
        fila_inicio += 2  # Espacio entre tablas
    
    # Ajustar columnas
    for col in range(1, 10):
        ws.column_dimensions[get_column_letter(col)].width = 20


def crear_hoja_metricas(ws, estructura):
    """Crea una hoja con métricas y KPIs."""
    ws['A1'] = "Métricas del Documento"
    ws['A1'].font = Font(size=16, bold=True)
    
    metricas = [
        ["Métrica", "Valor"],
        ["Total de Secciones", len(estructura['secciones'])],
        ["Total de Tablas", len(estructura['tablas'])],
        ["Total de Enlaces", len(estructura['enlaces'])],
        ["Total de Imágenes", len(estructura['imagenes'])],
        ["Total de Bloques de Código", len(estructura['codigo'])],
        ["Total de Items de Lista", len(estructura['listas'])]
    ]
    
    for fila, (metrica, valor) in enumerate(metricas, 1):
        ws.cell(row=fila, column=1, value=metrica)
        ws.cell(row=fila, column=2, value=valor)
    
    # Formato headers
    for col in [1, 2]:
        celda = ws.cell(row=1, column=col)
        celda.font = Font(bold=True, color="FFFFFF")
        celda.fill = PatternFill(start_color=COLORES_EXCEL['header'],
                                 end_color=COLORES_EXCEL['header'],
                                 fill_type='solid')
    
    ws.column_dimensions['A'].width = 30
    ws.column_dimensions['B'].width = 15


def crear_hoja_graficas(ws, estructura):
    """Crea una hoja con gráficas visuales."""
    ws['A1'] = "Gráficas y Visualizaciones"
    ws['A1'].font = Font(size=16, bold=True)
    
    # Crear gráfica de secciones por nivel
    if estructura['secciones']:
        niveles = {}
        for seccion in estructura['secciones']:
            nivel = seccion['nivel']
            niveles[nivel] = niveles.get(nivel, 0) + 1
        
        # Datos para gráfica
        ws['A3'] = "Nivel"
        ws['B3'] = "Cantidad"
        fila = 4
        for nivel, cantidad in sorted(niveles.items()):
            ws.cell(row=fila, column=1, value=f"Nivel {nivel}")
            ws.cell(row=fila, column=2, value=cantidad)
            fila += 1
        
        # Crear gráfica de barras
        chart = BarChart()
        chart.type = "col"
        chart.style = 10
        chart.title = "Distribución de Secciones por Nivel"
        chart.y_axis.title = "Cantidad"
        chart.x_axis.title = "Nivel"
        
        data = ws['B4:B' + str(fila-1)]
        cats = ws['A4:A' + str(fila-1)]
        chart.add_data(data, titles_from_data=False)
        chart.set_categories(cats)
        chart.height = 10
        chart.width = 15
        
        ws.add_chart(chart, "D3")
    
    # Crear gráfica de tipos de contenido
    tipos_contenido = [
        ["Tipo", "Cantidad"],
        ["Secciones", len(estructura['secciones'])],
        ["Tablas", len(estructura['tablas'])],
        ["Enlaces", len(estructura['enlaces'])],
        ["Imágenes", len(estructura['imagenes'])],
        ["Código", len(estructura['codigo'])]
    ]
    
    fila_inicio = 20
    for fila, (tipo, cantidad) in enumerate(tipos_contenido, fila_inicio):
        ws.cell(row=fila, column=1, value=tipo)
        ws.cell(row=fila, column=2, value=cantidad)
    
    # Gráfica de pastel
    pie_chart = PieChart()
    pie_chart.title = "Distribución de Tipos de Contenido"
    data = ws['B' + str(fila_inicio+1) + ':B' + str(fila_inicio+len(tipos_contenido)-1)]
    labels = ws['A' + str(fila_inicio+1) + ':A' + str(fila_inicio+len(tipos_contenido)-1)]
    pie_chart.add_data(data, titles_from_data=False)
    pie_chart.set_categories(labels)
    pie_chart.height = 10
    pie_chart.width = 15
    
    ws.add_chart(pie_chart, "D20")


def generar_graficas_avanzadas(estructura, nombre_base):
    """Genera gráficas avanzadas usando matplotlib y las guarda."""
    graficas_dir = OUTPUT_DIR / "graficas"
    graficas_dir.mkdir(exist_ok=True)
    
    # Configurar estilo
    sns.set_style("whitegrid")
    plt.rcParams['figure.figsize'] = (12, 8)
    plt.rcParams['font.size'] = 10
    
    # Gráfica 1: Distribución de secciones
    if estructura['secciones']:
        niveles = {}
        for seccion in estructura['secciones']:
            nivel = seccion['nivel']
            niveles[nivel] = niveles.get(nivel, 0) + 1
        
        fig, ax = plt.subplots()
        niveles_ordenados = sorted(niveles.items())
        ax.bar([f"Nivel {n}" for n, _ in niveles_ordenados], 
               [c for _, c in niveles_ordenados],
               color=COLORES_EXCEL['header'], alpha=0.7)
        ax.set_title('Distribución de Secciones por Nivel', fontsize=14, fontweight='bold')
        ax.set_xlabel('Nivel de Sección')
        ax.set_ylabel('Cantidad')
        plt.tight_layout()
        plt.savefig(graficas_dir / f"{nombre_base}_secciones.png", dpi=300, bbox_inches='tight')
        plt.close()
    
    # Gráfica 2: Tipos de contenido
    tipos = {
        'Secciones': len(estructura['secciones']),
        'Tablas': len(estructura['tablas']),
        'Enlaces': len(estructura['enlaces']),
        'Imágenes': len(estructura['imagenes']),
        'Código': len(estructura['codigo']),
        'Listas': len(estructura['listas'])
    }
    
    fig, ax = plt.subplots()
    ax.pie(tipos.values(), labels=tipos.keys(), autopct='%1.1f%%', startangle=90)
    ax.set_title('Distribución de Tipos de Contenido', fontsize=14, fontweight='bold')
    plt.tight_layout()
    plt.savefig(graficas_dir / f"{nombre_base}_tipos_contenido.png", dpi=300, bbox_inches='tight')
    plt.close()
    
    print(f"✅ Gráficas guardadas en: {graficas_dir}")


def main():
    """Función principal."""
    print("=" * 60)
    print("CONVERSIÓN DE MARKDOWN A WORD Y EXCEL")
    print("Sistema Profesional con Gráficas e Imágenes")
    print("=" * 60)
    print()
    
    archivos_procesados = 0
    archivos_errores = []
    
    for archivo_md in ARCHIVOS_IMPORTANTES:
        ruta_archivo = BASE_DIR / archivo_md
        
        if not ruta_archivo.exists():
            print(f"⚠️  Archivo no encontrado: {archivo_md}")
            archivos_errores.append(archivo_md)
            continue
        
        print(f"\n📄 Procesando: {archivo_md}")
        print("-" * 60)
        
        try:
            # Parsear Markdown
            print("  → Parseando Markdown...")
            estructura = parsear_markdown(ruta_archivo)
            
            # Crear Word
            print("  → Creando documento Word...")
            crear_documento_word(estructura, archivo_md)
            
            # Crear Excel
            print("  → Creando documento Excel...")
            crear_documento_excel(estructura, archivo_md)
            
            # Generar gráficas avanzadas
            print("  → Generando gráficas avanzadas...")
            nombre_base = Path(archivo_md).stem
            generar_graficas_avanzadas(estructura, nombre_base)
            
            archivos_procesados += 1
            print(f"  ✅ {archivo_md} procesado exitosamente")
        
        except Exception as e:
            print(f"  ❌ Error procesando {archivo_md}: {str(e)}")
            archivos_errores.append(archivo_md)
            import traceback
            traceback.print_exc()
    
    # Resumen
    print("\n" + "=" * 60)
    print("RESUMEN DE CONVERSIÓN")
    print("=" * 60)
    print(f"✅ Archivos procesados exitosamente: {archivos_procesados}")
    print(f"❌ Archivos con errores: {len(archivos_errores)}")
    
    if archivos_errores:
        print("\nArchivos con errores:")
        for archivo in archivos_errores:
            print(f"  - {archivo}")
    
    print(f"\n📁 Documentos guardados en: {OUTPUT_DIR}")
    print("=" * 60)


if __name__ == "__main__":
    main()








