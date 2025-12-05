#!/usr/bin/env python3
"""
Script para convertir archivos Markdown importantes a documentos Word y Excel profesionales
con gráficos, imágenes y formato de alta calidad.
"""

import re
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import openpyxl
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.chart import BarChart, LineChart, PieChart, Reference
from openpyxl.drawing.image import Image
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')  # Para generar imágenes sin GUI
import pandas as pd
from PIL import Image as PILImage
import io

# Configuración de colores corporativos
COLORS = {
    'primary': RGBColor(25, 118, 210),      # Azul
    'secondary': RGBColor(255, 152, 0),      # Naranja
    'success': RGBColor(76, 175, 80),        # Verde
    'danger': RGBColor(244, 67, 54),         # Rojo
    'dark': RGBColor(33, 33, 33),            # Negro
    'light': RGBColor(245, 245, 245)         # Gris claro
}

def crear_estilos_word(doc):
    """Crea estilos personalizados para Word"""
    styles = doc.styles
    
    # Estilo para títulos principales
    try:
        heading_style = styles['Heading 1']
    except:
        heading_style = styles.add_style('Custom Heading 1', WD_STYLE_TYPE.PARAGRAPH)
    
    heading_style.font.name = 'Calibri'
    heading_style.font.size = Pt(24)
    heading_style.font.bold = True
    heading_style.font.color.rgb = COLORS['primary']
    
    # Estilo para subtítulos
    try:
        subheading_style = styles['Heading 2']
    except:
        subheading_style = styles.add_style('Custom Heading 2', WD_STYLE_TYPE.PARAGRAPH)
    
    subheading_style.font.name = 'Calibri'
    subheading_style.font.size = Pt(18)
    subheading_style.font.bold = True
    subheading_style.font.color.rgb = COLORS['secondary']

def parsear_markdown(texto):
    """Parsea markdown básico a estructura"""
    lineas = texto.split('\n')
    elementos = []
    
    for linea in lineas:
        # Títulos
        if linea.startswith('# '):
            elementos.append(('h1', linea[2:].strip()))
        elif linea.startswith('## '):
            elementos.append(('h2', linea[3:].strip()))
        elif linea.startswith('### '):
            elementos.append(('h3', linea[4:].strip()))
        # Listas
        elif linea.startswith('- ') or linea.startswith('* '):
            elementos.append(('li', linea[2:].strip()))
        # Código
        elif linea.startswith('```'):
            elementos.append(('code', ''))
        # Tablas (básico)
        elif '|' in linea and linea.count('|') >= 2:
            elementos.append(('table_row', linea))
        # Párrafo normal
        elif linea.strip():
            elementos.append(('p', linea.strip()))
        else:
            elementos.append(('br', ''))
    
    return elementos

def crear_grafico_presupuesto():
    """Crea gráfico de presupuesto para Excel"""
    fig, ax = plt.subplots(figsize=(10, 6))
    
    # Datos de ejemplo
    categorias = ['Nano\n(1K-10K)', 'Micro\n(10K-100K)', 'Macro\n(100K-1M)']
    instagram = [50, 500, 5000]
    tiktok = [30, 300, 3000]
    youtube = [100, 1000, 10000]
    
    x = range(len(categorias))
    width = 0.25
    
    ax.bar([i - width for i in x], instagram, width, label='Instagram', color='#E1306C')
    ax.bar(x, tiktok, width, label='TikTok', color='#000000')
    ax.bar([i + width for i in x], youtube, width, label='YouTube', color='#FF0000')
    
    ax.set_xlabel('Tipo de Influencer', fontsize=12, fontweight='bold')
    ax.set_ylabel('Precio Promedio (USD)', fontsize=12, fontweight='bold')
    ax.set_title('Tarifas de Influencers por Plataforma', fontsize=14, fontweight='bold')
    ax.set_xticks(x)
    ax.set_xticklabels(categorias)
    ax.legend()
    ax.grid(axis='y', alpha=0.3)
    
    plt.tight_layout()
    buffer = io.BytesIO()
    plt.savefig(buffer, format='png', dpi=300, bbox_inches='tight')
    buffer.seek(0)
    plt.close()
    return buffer

def crear_grafico_metricas():
    """Crea gráfico de métricas para Excel"""
    fig, ((ax1, ax2), (ax3, ax4)) = plt.subplots(2, 2, figsize=(12, 10))
    
    # Gráfico 1: Tasa de respuesta
    labels = ['Respuestas', 'Sin Respuesta', 'Rechazaron', 'Pendientes']
    sizes = [25, 58, 8, 8]
    colors = ['#4CAF50', '#FF9800', '#F44336', '#9E9E9E']
    ax1.pie(sizes, labels=labels, autopct='%1.1f%%', colors=colors, startangle=90)
    ax1.set_title('Distribución de Respuestas', fontweight='bold')
    
    # Gráfico 2: Colaboraciones
    estados = ['Negociadas', 'Activas', 'Completadas']
    valores = [12, 8, 4]
    ax2.bar(estados, valores, color=['#2196F3', '#4CAF50', '#FF9800'])
    ax2.set_title('Estado de Colaboraciones', fontweight='bold')
    ax2.set_ylabel('Cantidad')
    ax2.grid(axis='y', alpha=0.3)
    
    # Gráfico 3: ROI
    meses = ['Ene', 'Feb', 'Mar', 'Abr']
    inversion = [500, 600, 700, 700]
    ventas = [1200, 1500, 2000, 2800]
    ax3.plot(meses, inversion, marker='o', label='Inversión', color='#F44336', linewidth=2)
    ax3.plot(meses, ventas, marker='s', label='Ventas', color='#4CAF50', linewidth=2)
    ax3.set_title('Evolución ROI', fontweight='bold')
    ax3.set_ylabel('USD')
    ax3.legend()
    ax3.grid(alpha=0.3)
    
    # Gráfico 4: Por plataforma
    plataformas = ['Instagram', 'TikTok', 'Twitter', 'YouTube']
    candidatos = [80, 40, 20, 10]
    ax4.barh(plataformas, candidatos, color='#2196F3')
    ax4.set_title('Candidatos por Plataforma', fontweight='bold')
    ax4.set_xlabel('Cantidad')
    ax4.grid(axis='x', alpha=0.3)
    
    plt.tight_layout()
    buffer = io.BytesIO()
    plt.savefig(buffer, format='png', dpi=300, bbox_inches='tight')
    buffer.seek(0)
    plt.close()
    return buffer

def convertir_md_a_word(archivo_md, archivo_word):
    """Convierte archivo Markdown a Word profesional"""
    print(f"Convirtiendo {archivo_md} a Word...")
    
    # Leer markdown
    with open(archivo_md, 'r', encoding='utf-8') as f:
        contenido = f.read()
    
    # Crear documento Word
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)
    
    # Crear estilos
    crear_estilos_word(doc)
    
    # Agregar portada
    titulo = doc.add_heading(archivo_md.replace('.md', '').replace('_', ' ').title(), 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo.runs[0].font.color.rgb = COLORS['primary']
    
    fecha = doc.add_paragraph(f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fecha.runs[0].font.size = Pt(10)
    fecha.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_page_break()
    
    # Parsear y agregar contenido
    elementos = parsear_markdown(contenido)
    
    for tipo, contenido_elem in elementos:
        if tipo == 'h1':
            doc.add_heading(contenido_elem, 1)
        elif tipo == 'h2':
            doc.add_heading(contenido_elem, 2)
        elif tipo == 'h3':
            doc.add_heading(contenido_elem, 3)
        elif tipo == 'li':
            doc.add_paragraph(contenido_elem, style='List Bullet')
        elif tipo == 'p':
            p = doc.add_paragraph(contenido_elem)
            p.runs[0].font.size = Pt(11)
        elif tipo == 'br':
            doc.add_paragraph()
    
    # Guardar
    doc.save(archivo_word)
    print(f"✓ Word creado: {archivo_word}")

def convertir_md_a_excel(archivo_md, archivo_excel):
    """Convierte archivo Markdown a Excel profesional con gráficos"""
    print(f"Convirtiendo {archivo_md} a Excel...")
    
    # Leer markdown
    with open(archivo_md, 'r', encoding='utf-8') as f:
        contenido = f.read()
    
    # Crear workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Resumen"
    
    # Estilos
    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=12)
    title_font = Font(bold=True, size=16, color="1F4E78")
    
    # Título
    ws['A1'] = archivo_md.replace('.md', '').replace('_', ' ').title()
    ws['A1'].font = title_font
    ws.merge_cells('A1:D1')
    ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
    
    # Fecha
    ws['A2'] = f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    ws['A2'].font = Font(size=10, italic=True)
    ws.merge_cells('A2:D2')
    
    # Agregar datos según el tipo de archivo
    if 'PRESUPUESTO' in archivo_md.upper() or 'PRICING' in archivo_md.upper():
        crear_hoja_presupuesto(wb, contenido)
    elif 'DASHBOARD' in archivo_md.upper() or 'METRICAS' in archivo_md.upper():
        crear_hoja_dashboard(wb, contenido)
    elif 'ANALISIS' in archivo_md.upper() or 'COMPETITIVO' in archivo_md.upper():
        crear_hoja_analisis(wb, contenido)
    else:
        crear_hoja_generica(wb, contenido)
    
    # Guardar
    wb.save(archivo_excel)
    print(f"✓ Excel creado: {archivo_excel}")

def crear_hoja_presupuesto(wb, contenido):
    """Crea hoja de presupuesto con datos y gráficos"""
    ws = wb.active
    
    # Encabezados
    headers = ['Tipo Influencer', 'Plataforma', 'Tipo Contenido', 'Precio Min', 'Precio Max', 'Promedio']
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=4, column=col, value=header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
        cell.alignment = Alignment(horizontal='center', vertical='center')
    
    # Datos de ejemplo
    datos = [
        ['Nano (1K-10K)', 'Instagram', 'Post Feed', 10, 100, 55],
        ['Nano (1K-10K)', 'Instagram', 'Reel/Video', 20, 150, 85],
        ['Nano (1K-10K)', 'TikTok', 'Video', 5, 50, 27],
        ['Micro (10K-100K)', 'Instagram', 'Post Feed', 100, 1000, 550],
        ['Micro (10K-100K)', 'Instagram', 'Reel/Video', 200, 1500, 850],
        ['Micro (10K-100K)', 'TikTok', 'Video', 50, 500, 275],
        ['Micro (10K-100K)', 'YouTube', 'Video Dedicado', 500, 3000, 1750],
    ]
    
    for row, dato in enumerate(datos, 5):
        for col, valor in enumerate(dato, 1):
            ws.cell(row=row, column=col, value=valor)
    
    # Ajustar columnas
    for col in range(1, 7):
        ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = 18
    
    # Crear gráfico
    chart = BarChart()
    chart.type = "col"
    chart.style = 10
    chart.title = "Tarifas por Tipo de Influencer"
    chart.y_axis.title = 'Precio (USD)'
    chart.x_axis.title = 'Tipo'
    
    data = Reference(ws, min_col=6, min_row=4, max_row=11)
    cats = Reference(ws, min_col=1, min_row=5, max_row=11)
    chart.add_data(data, titles_from_data=False)
    chart.set_categories(cats)
    chart.height = 10
    chart.width = 15
    
    ws.add_chart(chart, "H4")
    
    # Agregar gráfico de imagen
    grafico_img = crear_grafico_presupuesto()
    img = Image(grafico_img)
    img.width = 600
    img.height = 360
    ws.add_image(img, "H20")

def crear_hoja_dashboard(wb, contenido):
    """Crea hoja de dashboard con métricas y gráficos"""
    ws = wb.active
    
    # KPIs principales
    ws['A4'] = "MÉTRICAS PRINCIPALES"
    ws['A4'].font = Font(bold=True, size=14, color="1F4E78")
    
    kpis = [
        ['Candidatos Identificados', 150],
        ['Verificados', 120],
        ['Contactados', 60],
        ['Respuestas', 15],
        ['Tasa de Respuesta', '25%'],
        ['Colaboraciones Activas', 8],
        ['ROI', '200%'],
    ]
    
    for row, (kpi, valor) in enumerate(kpis, 5):
        ws.cell(row=row, column=1, value=kpi).font = Font(bold=True)
        ws.cell(row=row, column=2, value=valor).font = Font(size=12, bold=True, color="1F4E78")
    
    # Tabla de métricas por plataforma
    ws['D4'] = "POR PLATAFORMA"
    ws['D4'].font = Font(bold=True, size=14, color="1F4E78")
    
    headers = ['Plataforma', 'Candidatos', 'Contactados', 'Respuestas']
    for col, header in enumerate(headers, 4):
        cell = ws.cell(row=5, column=col, value=header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    
    datos_plataforma = [
        ['Instagram', 80, 35, 10],
        ['TikTok', 40, 15, 3],
        ['Twitter', 20, 8, 2],
        ['YouTube', 10, 2, 0],
    ]
    
    for row, dato in enumerate(datos_plataforma, 6):
        for col, valor in enumerate(dato, 4):
            ws.cell(row=row, column=col, value=valor)
    
    # Gráfico de barras
    chart = BarChart()
    chart.type = "col"
    chart.style = 10
    chart.title = "Candidatos por Plataforma"
    chart.y_axis.title = 'Cantidad'
    
    data = Reference(ws, min_col=5, min_row=5, max_row=9)
    cats = Reference(ws, min_col=4, min_row=6, max_row=9)
    chart.add_data(data, titles_from_data=True)
    chart.set_categories(cats)
    chart.height = 8
    chart.width = 12
    
    ws.add_chart(chart, "A15")
    
    # Agregar gráfico de imagen
    grafico_img = crear_grafico_metricas()
    img = Image(grafico_img)
    img.width = 800
    img.height = 600
    ws.add_image(img, "H15")

def crear_hoja_analisis(wb, contenido):
    """Crea hoja de análisis competitivo"""
    ws = wb.active
    
    ws['A4'] = "ANÁLISIS COMPETITIVO"
    ws['A4'].font = Font(bold=True, size=14, color="1F4E78")
    
    headers = ['Competidor', 'Influencers', 'Plataformas', 'Engagement Avg', 'Tipo Contenido']
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=5, column=col, value=header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    
    datos = [
        ['Competidor A', 15, 'Instagram, TikTok', '4.5%', 'Reviews, Tutoriales'],
        ['Competidor B', 8, 'YouTube, LinkedIn', '6.2%', 'Casos de Uso'],
        ['Competidor C', 12, 'Instagram, Twitter', '3.8%', 'Comparativas'],
    ]
    
    for row, dato in enumerate(datos, 6):
        for col, valor in enumerate(dato, 1):
            ws.cell(row=row, column=col, value=valor)
    
    for col in range(1, 6):
        ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = 20

def crear_hoja_generica(wb, contenido):
    """Crea hoja genérica para otros tipos de documentos"""
    ws = wb.active
    ws['A4'] = "CONTENIDO"
    ws['A4'].font = Font(bold=True, size=14, color="1F4E78")
    
    # Agregar contenido básico
    lineas = contenido.split('\n')[:50]  # Primeras 50 líneas
    for row, linea in enumerate(lineas, 5):
        if linea.strip():
            ws.cell(row=row, column=1, value=linea.strip()[:100])

def main():
    """Función principal"""
    archivos_importantes = [
        'PRESUPUESTO_PRICING.md',
        'DASHBOARD_METRICAS.md',
        'ANALISIS_COMPETITIVO.md',
    ]
    
    directorio = '/Users/adan/Documents/documentos_blatam/01_marketing'
    
    for archivo_md in archivos_importantes:
        ruta_md = os.path.join(directorio, archivo_md)
        
        if os.path.exists(ruta_md):
            # Crear Word
            archivo_word = ruta_md.replace('.md', '_PROFESIONAL.docx')
            convertir_md_a_word(ruta_md, archivo_word)
            
            # Crear Excel
            archivo_excel = ruta_md.replace('.md', '_PROFESIONAL.xlsx')
            convertir_md_a_excel(ruta_md, archivo_excel)
        else:
            print(f"⚠ Archivo no encontrado: {ruta_md}")
    
    print("\n✅ Conversión completada!")

if __name__ == "__main__":
    main()








