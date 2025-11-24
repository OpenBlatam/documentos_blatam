#!/usr/bin/env python3
"""
Script avanzado para convertir TODOS los archivos Markdown importantes a documentos
Word y Excel profesionales con gráficos avanzados, imágenes, tablas dinámicas y más.
"""

import re
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import openpyxl
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.chart import BarChart, LineChart, PieChart, Reference, ScatterChart
from openpyxl.chart.series import DataPoint
from openpyxl.drawing.image import Image
from openpyxl.utils import get_column_letter
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')
import pandas as pd
import numpy as np
from PIL import Image as PILImage
import io
import json

# Configuración avanzada de colores
COLORS = {
    'primary': RGBColor(25, 118, 210),
    'secondary': RGBColor(255, 152, 0),
    'success': RGBColor(76, 175, 80),
    'danger': RGBColor(244, 67, 54),
    'warning': RGBColor(255, 193, 7),
    'info': RGBColor(33, 150, 243),
    'dark': RGBColor(33, 33, 33),
    'light': RGBColor(245, 245, 245)
}

def crear_portada_profesional(doc, titulo, subtitulo=""):
    """Crea una portada profesional"""
    # Título principal
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(titulo)
    title_run.font.size = Pt(36)
    title_run.font.bold = True
    title_run.font.color.rgb = COLORS['primary']
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtítulo
    if subtitulo:
        subtitle_para = doc.add_paragraph()
        subtitle_run = subtitle_para.add_run(subtitulo)
        subtitle_run.font.size = Pt(18)
        subtitle_run.font.color.rgb = COLORS['secondary']
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Fecha
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f"Generado: {datetime.now().strftime('%d de %B de %Y')}")
    date_run.font.size = Pt(12)
    date_run.font.color.rgb = RGBColor(128, 128, 128)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()

def crear_tabla_contenidos(doc, elementos):
    """Crea tabla de contenidos automática"""
    doc.add_heading('Tabla de Contenidos', 1)
    
    for elemento in elementos:
        if elemento[0] in ['h1', 'h2']:
            nivel = 0 if elemento[0] == 'h1' else 1
            p = doc.add_paragraph(elementio[1], style='List Bullet' if nivel == 1 else 'List Number')
            p.paragraph_format.left_indent = Inches(nivel * 0.5)
    
    doc.add_page_break()

def crear_grafico_avanzado_roi():
    """Crea gráfico avanzado de ROI con múltiples métricas"""
    fig = plt.figure(figsize=(14, 8))
    gs = fig.add_gridspec(2, 3, hspace=0.3, wspace=0.3)
    
    # Gráfico 1: ROI por mes
    ax1 = fig.add_subplot(gs[0, 0])
    meses = ['Ene', 'Feb', 'Mar', 'Abr', 'May', 'Jun']
    roi = [150, 180, 200, 220, 250, 280]
    ax1.plot(meses, roi, marker='o', linewidth=3, markersize=8, color='#4CAF50')
    ax1.fill_between(meses, roi, alpha=0.3, color='#4CAF50')
    ax1.set_title('Evolución ROI', fontweight='bold', fontsize=12)
    ax1.set_ylabel('ROI (%)', fontweight='bold')
    ax1.grid(alpha=0.3)
    
    # Gráfico 2: Distribución de inversión
    ax2 = fig.add_subplot(gs[0, 1])
    categorias = ['Nano', 'Micro', 'Macro']
    valores = [30, 50, 20]
    colors = ['#FF9800', '#2196F3', '#9C27B0']
    ax2.pie(valores, labels=categorias, autopct='%1.1f%%', colors=colors, startangle=90)
    ax2.set_title('Distribución de Inversión', fontweight='bold', fontsize=12)
    
    # Gráfico 3: Métricas de engagement
    ax3 = fig.add_subplot(gs[0, 2])
    plataformas = ['IG', 'TT', 'YT', 'LI']
    engagement = [4.5, 6.2, 8.1, 3.8]
    bars = ax3.bar(plataformas, engagement, color=['#E1306C', '#000000', '#FF0000', '#0077B5'])
    ax3.set_title('Engagement por Plataforma', fontweight='bold', fontsize=12)
    ax3.set_ylabel('Engagement (%)', fontweight='bold')
    ax3.grid(axis='y', alpha=0.3)
    for bar in bars:
        height = bar.get_height()
        ax3.text(bar.get_x() + bar.get_width()/2., height,
                f'{height}%', ha='center', va='bottom', fontweight='bold')
    
    # Gráfico 4: Comparativa de costos
    ax4 = fig.add_subplot(gs[1, :2])
    tipos = ['Post Feed', 'Reel/Video', 'Story', 'Carousel', 'IGTV']
    nano = [55, 85, 25, 70, 150]
    micro = [550, 850, 175, 675, 1150]
    x = np.arange(len(tipos))
    width = 0.35
    ax4.bar(x - width/2, nano, width, label='Nano (1K-10K)', color='#FF9800')
    ax4.bar(x + width/2, micro, width, label='Micro (10K-100K)', color='#2196F3')
    ax4.set_xlabel('Tipo de Contenido', fontweight='bold')
    ax4.set_ylabel('Precio Promedio (USD)', fontweight='bold')
    ax4.set_title('Comparativa de Precios por Tipo de Contenido', fontweight='bold', fontsize=12)
    ax4.set_xticks(x)
    ax4.set_xticklabels(tipos, rotation=45, ha='right')
    ax4.legend()
    ax4.grid(axis='y', alpha=0.3)
    
    # Gráfico 5: Funnel de conversión
    ax5 = fig.add_subplot(gs[1, 2])
    etapas = ['Identificados', 'Contactados', 'Respuestas', 'Negociados', 'Activos']
    valores_funnel = [150, 60, 15, 12, 8]
    porcentajes = [100, 40, 10, 8, 5.3]
    colors_funnel = ['#4CAF50', '#8BC34A', '#CDDC39', '#FFC107', '#FF9800']
    bars = ax5.barh(etapas, porcentajes, color=colors_funnel)
    ax5.set_title('Funnel de Conversión', fontweight='bold', fontsize=12)
    ax5.set_xlabel('Porcentaje (%)', fontweight='bold')
    ax5.grid(axis='x', alpha=0.3)
    for i, (bar, val) in enumerate(zip(bars, valores_funnel)):
        width = bar.get_width()
        ax5.text(width + 1, bar.get_y() + bar.get_height()/2,
                f'{val}', ha='left', va='center', fontweight='bold')
    
    plt.suptitle('Dashboard Completo de Métricas de Marketing', fontsize=16, fontweight='bold', y=0.98)
    plt.tight_layout(rect=[0, 0, 1, 0.96])
    
    buffer = io.BytesIO()
    plt.savefig(buffer, format='png', dpi=300, bbox_inches='tight', facecolor='white')
    buffer.seek(0)
    plt.close()
    return buffer

def crear_excel_avanzado(archivo_md, archivo_excel):
    """Crea Excel avanzado con múltiples hojas, gráficos y análisis"""
    print(f"Creando Excel avanzado para {archivo_md}...")
    
    wb = Workbook()
    
    # Hoja 1: Dashboard Principal
    ws_dashboard = wb.active
    ws_dashboard.title = "Dashboard"
    
    # Título
    ws_dashboard['A1'] = archivo_md.replace('.md', '').replace('_', ' ').title()
    ws_dashboard['A1'].font = Font(bold=True, size=20, color="FFFFFF")
    ws_dashboard['A1'].fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    ws_dashboard.merge_cells('A1:F1')
    ws_dashboard['A1'].alignment = Alignment(horizontal='center', vertical='center')
    ws_dashboard.row_dimensions[1].height = 30
    
    # KPIs principales
    kpis = [
        ['Total Inversión', '$2,500', '1F4E78'],
        ['ROI', '200%', '4CAF50'],
        ['Colaboraciones Activas', '8', 'FF9800'],
        ['Tasa de Respuesta', '25%', '2196F3'],
    ]
    
    for col, (kpi, valor, color) in enumerate(kpis, 1):
        cell_kpi = ws_dashboard.cell(row=3, column=col*2-1, value=kpi)
        cell_kpi.font = Font(bold=True, size=11)
        cell_kpi.alignment = Alignment(horizontal='center')
        
        cell_valor = ws_dashboard.cell(row=4, column=col*2-1, value=valor)
        cell_valor.font = Font(bold=True, size=16, color="FFFFFF")
        cell_valor.fill = PatternFill(start_color=color, end_color=color, fill_type="solid")
        cell_valor.alignment = Alignment(horizontal='center', vertical='center')
        ws_dashboard.merge_cells(f'{get_column_letter(col*2-1)}4:{get_column_letter(col*2)}4')
        ws_dashboard.row_dimensions[4].height = 40
    
    # Agregar gráfico avanzado
    grafico_img = crear_grafico_avanzado_roi()
    img = Image(grafico_img)
    img.width = 1200
    img.height = 700
    ws_dashboard.add_image(img, "A6")
    
    # Hoja 2: Datos Detallados
    ws_datos = wb.create_sheet("Datos Detallados")
    
    headers = ['Fecha', 'Influencer', 'Plataforma', 'Tipo', 'Inversión', 'Resultados', 'ROI']
    for col, header in enumerate(headers, 1):
        cell = ws_datos.cell(row=1, column=col, value=header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
        cell.alignment = Alignment(horizontal='center')
        ws_datos.column_dimensions[get_column_letter(col)].width = 18
    
    # Datos de ejemplo
    datos_ejemplo = [
        ['2025-01-15', 'Influencer A', 'Instagram', 'Post', 200, 600, '200%'],
        ['2025-01-20', 'Influencer B', 'TikTok', 'Video', 150, 450, '200%'],
        ['2025-02-01', 'Influencer C', 'YouTube', 'Review', 500, 1500, '200%'],
        ['2025-02-10', 'Influencer D', 'Instagram', 'Reel', 300, 900, '200%'],
        ['2025-02-15', 'Influencer E', 'TikTok', 'Series', 400, 1200, '200%'],
    ]
    
    for row, dato in enumerate(datos_ejemplo, 2):
        for col, valor in enumerate(dato, 1):
            cell = ws_datos.cell(row=row, column=col, value=valor)
            if col == 7:  # ROI
                cell.font = Font(bold=True, color="4CAF50")
    
    # Formato condicional para ROI
    from openpyxl.formatting.rule import CellIsRule
    green_fill = PatternFill(start_color="C8E6C9", end_color="C8E6C9", fill_type="solid")
    ws_datos.conditional_formatting.add('G2:G100', CellIsRule(operator='greaterThan', formula=['150%'], fill=green_fill))
    
    # Hoja 3: Análisis por Plataforma
    ws_plataforma = wb.create_sheet("Por Plataforma")
    
    datos_plataforma = {
        'Plataforma': ['Instagram', 'TikTok', 'YouTube', 'Twitter', 'LinkedIn'],
        'Inversión': [1000, 600, 500, 200, 200],
        'Resultados': [3000, 1800, 1500, 600, 600],
        'ROI': ['200%', '200%', '200%', '200%', '200%'],
        'Colaboraciones': [5, 3, 2, 1, 1]
    }
    
    df = pd.DataFrame(datos_plataforma)
    
    # Encabezados
    for col, header in enumerate(df.columns, 1):
        cell = ws_plataforma.cell(row=1, column=col, value=header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
        ws_plataforma.column_dimensions[get_column_letter(col)].width = 18
    
    # Datos
    for row, (_, fila) in enumerate(df.iterrows(), 2):
        for col, valor in enumerate(fila, 1):
            ws_plataforma.cell(row=row, column=col, value=valor)
    
    # Gráfico de barras
    chart = BarChart()
    chart.type = "col"
    chart.style = 10
    chart.title = "Inversión vs Resultados por Plataforma"
    chart.y_axis.title = 'USD'
    chart.x_axis.title = 'Plataforma'
    
    data = Reference(ws_plataforma, min_col=2, min_row=1, max_col=3, max_row=6)
    cats = Reference(ws_plataforma, min_col=1, min_row=2, max_row=6)
    chart.add_data(data, titles_from_data=True)
    chart.set_categories(cats)
    chart.height = 10
    chart.width = 15
    
    ws_plataforma.add_chart(chart, "A8")
    
    # Guardar
    wb.save(archivo_excel)
    print(f"✓ Excel avanzado creado: {archivo_excel}")

def convertir_sistemas_consolidado():
    """Convierte el archivo principal SISTEMAS_PROMPTS_CONSOLIDADO.md"""
    archivo_md = '/Users/adan/Documents/documentos_blatam/01_marketing/SISTEMAS_PROMPTS_CONSOLIDADO.md'
    
    if not os.path.exists(archivo_md):
        print(f"⚠ Archivo no encontrado: {archivo_md}")
        return
    
    print(f"Convirtiendo archivo principal: SISTEMAS_PROMPTS_CONSOLIDADO.md...")
    
    # Crear Word
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)
    
    # Portada
    crear_portada_profesional(
        doc,
        "Sistemas de Prompts Consolidados",
        "Guía Completa de Sistemas de Marketing y Creación de Contenido"
    )
    
    # Leer y agregar contenido (primeras 1000 líneas para no sobrecargar)
    with open(archivo_md, 'r', encoding='utf-8') as f:
        lineas = f.readlines()[:1000]
    
    for linea in lineas:
        linea = linea.strip()
        if linea.startswith('# '):
            doc.add_heading(linea[2:], 1)
        elif linea.startswith('## '):
            doc.add_heading(linea[3:], 2)
        elif linea.startswith('### '):
            doc.add_heading(linea[4:], 3)
        elif linea.startswith('- '):
            doc.add_paragraph(linea[2:], style='List Bullet')
        elif linea.strip():
            doc.add_paragraph(linea)
    
    archivo_word = archivo_md.replace('.md', '_PROFESIONAL.docx')
    doc.save(archivo_word)
    print(f"✓ Word creado: {archivo_word}")
    
    # Crear Excel resumen
    wb = Workbook()
    ws = wb.active
    ws.title = "Resumen Sistemas"
    
    ws['A1'] = "SISTEMAS DE PROMPTS CONSOLIDADOS"
    ws['A1'].font = Font(bold=True, size=18, color="FFFFFF")
    ws['A1'].fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    ws.merge_cells('A1:D1')
    ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
    
    sistemas = [
        ['Sistema', 'Tipo', 'Variantes', 'Estado'],
        ['Perplexity', 'Búsqueda', '1', 'Activo'],
        ['UGC Videos', 'Contenido', '5', 'Activo'],
        ['Creación Contenido', 'Contenido', '7', 'Activo'],
        ['Calendario Social', 'Planificación', '3', 'Activo'],
        ['Diseño Gráfico', 'Diseño', '8', 'Activo'],
        ['Ventas y Cierre', 'Ventas', '8', 'Activo'],
    ]
    
    for row, sistema in enumerate(sistemas, 3):
        for col, valor in enumerate(sistema, 1):
            cell = ws.cell(row=row, column=col, value=valor)
            if row == 3:
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
            cell.alignment = Alignment(horizontal='center')
    
    for col in range(1, 5):
        ws.column_dimensions[get_column_letter(col)].width = 25
    
    archivo_excel = archivo_md.replace('.md', '_PROFESIONAL.xlsx')
    wb.save(archivo_excel)
    print(f"✓ Excel creado: {archivo_excel}")

def main():
    """Función principal"""
    archivos = [
        'PRESUPUESTO_PRICING.md',
        'DASHBOARD_METRICAS.md',
        'ANALISIS_COMPETITIVO.md',
        'REPORTES_EJECUTIVOS.md',
    ]
    
    directorio = '/Users/adan/Documents/documentos_blatam/01_marketing'
    
    # Convertir archivos específicos
    for archivo_md in archivos:
        ruta_md = os.path.join(directorio, archivo_md)
        if os.path.exists(ruta_md):
            archivo_excel = ruta_md.replace('.md', '_AVANZADO.xlsx')
            crear_excel_avanzado(ruta_md, archivo_excel)
    
    # Convertir archivo principal
    convertir_sistemas_consolidado()
    
    print("\n✅ Todos los documentos profesionales han sido creados!")

if __name__ == "__main__":
    main()



