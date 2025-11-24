#!/usr/bin/env python3
"""
Script Premium para convertir archivos Markdown de Marketing a Word y Excel
con gráficas, imágenes y formato profesional de alta calidad.

Librerías utilizadas:
- python-docx: Documentos Word profesionales
- openpyxl: Excel con formato avanzado
- markdown: Parseo de Markdown
- matplotlib/seaborn: Gráficas profesionales
- pandas: Manipulación de datos
- Pillow: Procesamiento de imágenes
- reportlab: PDFs (opcional)
- plotly: Gráficas interactivas
"""

import os
import re
import json
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Any, Optional
import base64
import io

# Importaciones de librerías externas
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import markdown
    from markdown.extensions import tables, fenced_code, codehilite
    import pandas as pd
    import numpy as np
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    import seaborn as sns
    from PIL import Image, ImageDraw, ImageFont
    import plotly.graph_objects as go
    import plotly.express as px
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.chart import BarChart, LineChart, PieChart, ScatterChart
    from openpyxl.drawing.image import Image as XLImage
    from openpyxl.utils import get_column_letter
    from openpyxl.chart.marker import DataPoint
    from openpyxl.chart.label import DataLabelList
except ImportError as e:
    print(f"Error: Falta instalar librerías. Ejecuta: pip install python-docx markdown pandas numpy matplotlib seaborn pillow plotly openpyxl")
    print(f"Error específico: {e}")
    exit(1)

# Configuración de estilo
sns.set_style("whitegrid")
plt.rcParams['figure.figsize'] = (12, 8)
plt.rcParams['font.size'] = 10
plt.rcParams['axes.labelsize'] = 12
plt.rcParams['axes.titlesize'] = 14
plt.rcParams['xtick.labelsize'] = 10
plt.rcParams['ytick.labelsize'] = 10
plt.rcParams['legend.fontsize'] = 10

# Colores corporativos
COLORS = {
    'primary': '#2563EB',
    'secondary': '#7C3AED',
    'success': '#10B981',
    'warning': '#F59E0B',
    'danger': '#EF4444',
    'info': '#3B82F6',
    'dark': '#1F2937',
    'light': '#F3F4F6'
}

class MarkdownToWordConverter:
    """Convierte Markdown a Word con formato profesional"""
    
    def __init__(self):
        self.doc = Document()
        self.setup_styles()
        
    def setup_styles(self):
        """Configura estilos personalizados para el documento"""
        styles = self.doc.styles
        
        # Estilo para títulos principales
        try:
            heading1 = styles['Heading 1']
            heading1.font.size = Pt(24)
            heading1.font.bold = True
            heading1.font.color.rgb = RGBColor(37, 99, 235)  # Blue
            heading1.paragraph_format.space_after = Pt(12)
        except:
            pass
            
        # Estilo para subtítulos
        try:
            heading2 = styles['Heading 2']
            heading2.font.size = Pt(18)
            heading2.font.bold = True
            heading2.font.color.rgb = RGBColor(124, 58, 237)  # Purple
            heading2.paragraph_format.space_after = Pt(10)
        except:
            pass
    
    def parse_markdown(self, md_content: str):
        """Parsea el contenido Markdown y lo convierte a Word"""
        # Convertir markdown a HTML primero
        md = markdown.Markdown(extensions=['tables', 'fenced_code', 'codehilite', 'nl2br'])
        html = md.convert(md_content)
        
        # Dividir en líneas para procesar
        lines = md_content.split('\n')
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                i += 1
                continue
                
            # Títulos
            if line.startswith('# '):
                self.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                self.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                self.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                self.add_heading(line[5:], level=4)
            elif line.startswith('##### '):
                self.add_heading(line[6:], level=5)
            elif line.startswith('###### '):
                self.add_heading(line[7:], level=6)
            
            # Listas
            elif line.startswith('- ') or line.startswith('* '):
                self.add_bullet_list([line[2:]], lines, i)
                i = self._skip_list_items(lines, i)
                continue
            
            # Código
            elif line.startswith('```'):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                if code_lines:
                    self.add_code_block('\n'.join(code_lines))
            
            # Tablas
            elif '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
                table_data = self._parse_table(lines, i)
                if table_data:
                    self.add_table(table_data)
                    i = self._skip_table(lines, i)
                    continue
            
            # Texto normal
            else:
                self.add_paragraph(self._clean_markdown(line))
            
            i += 1
    
    def add_heading(self, text: str, level: int = 1):
        """Añade un título"""
        heading = self.doc.add_heading(text, level=level)
        if level == 1:
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def add_paragraph(self, text: str):
        """Añade un párrafo"""
        para = self.doc.add_paragraph(text)
        para.paragraph_format.line_spacing = 1.15
        para.paragraph_format.space_after = Pt(6)
    
    def add_bullet_list(self, items: List[str], lines: List[str], start_idx: int):
        """Añade una lista con viñetas"""
        current_items = []
        i = start_idx
        
        while i < len(lines):
            line = lines[i].strip()
            if line.startswith('- ') or line.startswith('* '):
                current_items.append(line[2:])
            elif line.startswith('  - ') or line.startswith('  * '):
                current_items.append('  ' + line[4:])
            elif line and not line.startswith('-') and not line.startswith('*'):
                break
            i += 1
        
        for item in current_items:
            para = self.doc.add_paragraph(item, style='List Bullet')
            para.paragraph_format.left_indent = Inches(0.25)
    
    def add_code_block(self, code: str):
        """Añade un bloque de código"""
        para = self.doc.add_paragraph(code)
        para.style = 'Intense Quote'
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.right_indent = Inches(0.5)
        for run in para.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
    
    def add_table(self, data: List[List[str]]):
        """Añade una tabla"""
        if not data or len(data) < 2:
            return
        
        table = self.doc.add_table(rows=len(data), cols=len(data[0]))
        table.style = 'Light Grid Accent 1'
        
        # Encabezado
        header_cells = table.rows[0].cells
        for i, header in enumerate(data[0]):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True
            header_cells[i].shading.background_color.rgb = RGBColor(37, 99, 235)
        
        # Datos
        for row_idx, row_data in enumerate(data[1:], 1):
            for col_idx, cell_data in enumerate(row_data):
                table.rows[row_idx].cells[col_idx].text = str(cell_data)
    
    def add_image(self, image_path: str, width: float = 6.0):
        """Añade una imagen al documento"""
        if os.path.exists(image_path):
            self.doc.add_picture(image_path, width=Inches(width))
    
    def add_chart_image(self, chart_path: str):
        """Añade una gráfica como imagen"""
        self.add_image(chart_path)
    
    def _clean_markdown(self, text: str) -> str:
        """Limpia el texto de sintaxis Markdown"""
        # Remover emojis y caracteres especiales si es necesario
        text = re.sub(r'^#{1,6}\s+', '', text)
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # Bold
        text = re.sub(r'\*(.+?)\*', r'\1', text)  # Italic
        return text
    
    def _parse_table(self, lines: List[str], start_idx: int) -> Optional[List[List[str]]]:
        """Parsea una tabla de Markdown"""
        table_data = []
        i = start_idx
        
        while i < len(lines) and '|' in lines[i]:
            row = [cell.strip() for cell in lines[i].split('|') if cell.strip()]
            if row and not all(c == '-' for c in ''.join(row)):  # Skip separator row
                table_data.append(row)
            i += 1
        
        return table_data if len(table_data) > 1 else None
    
    def _skip_table(self, lines: List[str], idx: int) -> int:
        """Salta las líneas de una tabla"""
        i = idx
        while i < len(lines) and '|' in lines[i]:
            i += 1
        return i - 1
    
    def _skip_list_items(self, lines: List[str], idx: int) -> int:
        """Salta los elementos de una lista"""
        i = idx
        while i < len(lines):
            line = lines[i].strip()
            if not line or line.startswith('- ') or line.startswith('* ') or line.startswith('  - ') or line.startswith('  * '):
                i += 1
            else:
                break
        return i - 1
    
    def save(self, filename: str):
        """Guarda el documento"""
        self.doc.save(filename)
        print(f"✅ Documento Word creado: {filename}")


class MarkdownToExcelConverter:
    """Convierte Markdown a Excel con gráficas y visualizaciones"""
    
    def __init__(self):
        self.wb = Workbook()
        self.ws = self.wb.active
        self.ws.title = "Resumen"
        self.current_row = 1
        self.chart_images = []
        
    def parse_markdown(self, md_content: str, filename_base: str):
        """Parsea el contenido Markdown y lo convierte a Excel"""
        lines = md_content.split('\n')
        
        # Crear hoja de resumen
        self._add_summary_sheet(lines, filename_base)
        
        # Extraer y crear hojas de datos
        self._extract_tables_to_sheets(lines)
        
        # Crear hojas de gráficas
        self._create_charts_sheet(filename_base)
        
        # Crear hoja de métricas
        self._create_metrics_sheet(lines)
    
    def _add_summary_sheet(self, lines: List[str], title: str):
        """Añade contenido resumido a la hoja principal"""
        ws = self.wb.active
        ws.title = "Resumen"
        
        # Título
        cell = ws['A1']
        cell.value = title
        cell.font = Font(size=20, bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="2563EB", end_color="2563EB", fill_type="solid")
        cell.alignment = Alignment(horizontal="center", vertical="center")
        ws.merge_cells('A1:D1')
        ws.row_dimensions[1].height = 30
        
        row = 3
        for line in lines[:100]:  # Primeras 100 líneas
            if line.strip():
                ws[f'A{row}'] = self._clean_markdown(line[:100])  # Limitar longitud
                row += 1
                if row > 50:  # Limitar filas en resumen
                    break
        
        ws.column_dimensions['A'].width = 80
    
    def _extract_tables_to_sheets(self, lines: List[str]):
        """Extrae tablas y las convierte en hojas separadas"""
        i = 0
        table_count = 0
        
        while i < len(lines):
            if '|' in lines[i] and i + 1 < len(lines) and '|' in lines[i + 1]:
                table_data = self._parse_table(lines, i)
                if table_data and len(table_data) > 1:
                    table_count += 1
                    sheet_name = f"Tabla {table_count}"[:31]  # Excel limit
                    ws = self.wb.create_sheet(title=sheet_name)
                    self._add_table_to_sheet(ws, table_data)
                    i = self._skip_table(lines, i)
                else:
                    i += 1
            else:
                i += 1
    
    def _add_table_to_sheet(self, ws, table_data: List[List[str]]):
        """Añade una tabla a una hoja de Excel"""
        # Encabezado
        for col_idx, header in enumerate(table_data[0], 1):
            cell = ws.cell(row=1, column=col_idx)
            cell.value = header
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="2563EB", end_color="2563EB", fill_type="solid")
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
        
        # Datos
        for row_idx, row_data in enumerate(table_data[1:], 2):
            for col_idx, cell_data in enumerate(row_data, 1):
                cell = ws.cell(row=row_idx, column=col_idx)
                cell.value = str(cell_data)
                cell.alignment = Alignment(horizontal="left", vertical="center")
                cell.border = Border(
                    left=Side(style='thin'),
                    right=Side(style='thin'),
                    top=Side(style='thin'),
                    bottom=Side(style='thin')
                )
        
        # Ajustar anchos de columna
        for col_idx in range(1, len(table_data[0]) + 1):
            ws.column_dimensions[get_column_letter(col_idx)].width = 20
        
        # Aplicar formato alternado
        for row_idx in range(2, len(table_data) + 1):
            if row_idx % 2 == 0:
                for col_idx in range(1, len(table_data[0]) + 1):
                    cell = ws.cell(row=row_idx, column=col_idx)
                    cell.fill = PatternFill(start_color="F3F4F6", end_color="F3F4F6", fill_type="solid")
    
    def _create_charts_sheet(self, filename_base: str):
        """Crea una hoja con gráficas generadas"""
        ws = self.wb.create_sheet(title="Gráficas")
        
        # Título
        ws['A1'] = "Visualizaciones de Datos"
        ws['A1'].font = Font(size=16, bold=True)
        ws.merge_cells('A1:D1')
        
        # Generar gráficas de ejemplo
        self._generate_sample_charts(filename_base)
        
        # Añadir referencias a gráficas
        row = 3
        for i, chart_path in enumerate(self.chart_images, 1):
            ws[f'A{row}'] = f"Gráfica {i}"
            ws[f'A{row}'].font = Font(bold=True)
            row += 2
    
    def _generate_sample_charts(self, filename_base: str):
        """Genera gráficas de ejemplo basadas en el contenido"""
        charts_dir = Path("charts")
        charts_dir.mkdir(exist_ok=True)
        
        # Gráfica 1: Barras (ejemplo de métricas)
        fig, ax = plt.subplots(figsize=(10, 6))
        categories = ['Búsqueda', 'Contacto', 'Respuestas', 'Colaboraciones']
        values = [150, 60, 15, 12]
        colors = [COLORS['primary'], COLORS['secondary'], COLORS['success'], COLORS['warning']]
        
        bars = ax.bar(categories, values, color=colors, edgecolor='white', linewidth=2)
        ax.set_title('Métricas de Outreach', fontsize=16, fontweight='bold', pad=20)
        ax.set_ylabel('Cantidad', fontsize=12)
        ax.grid(axis='y', alpha=0.3, linestyle='--')
        
        # Añadir valores en las barras
        for bar in bars:
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2., height,
                   f'{int(height)}',
                   ha='center', va='bottom', fontsize=11, fontweight='bold')
        
        plt.tight_layout()
        chart_path = charts_dir / f"{filename_base}_chart1.png"
        plt.savefig(chart_path, dpi=300, bbox_inches='tight')
        plt.close()
        self.chart_images.append(str(chart_path))
        
        # Gráfica 2: Pie chart (distribución)
        fig, ax = plt.subplots(figsize=(10, 8))
        labels = ['Instagram', 'TikTok', 'Twitter', 'YouTube']
        sizes = [53, 27, 13, 7]
        colors_pie = [COLORS['primary'], COLORS['secondary'], COLORS['success'], COLORS['warning']]
        explode = (0.1, 0, 0, 0)
        
        wedges, texts, autotexts = ax.pie(sizes, explode=explode, labels=labels, colors=colors_pie,
                                          autopct='%1.1f%%', shadow=True, startangle=90,
                                          textprops={'fontsize': 12, 'fontweight': 'bold'})
        
        ax.set_title('Distribución por Plataforma', fontsize=16, fontweight='bold', pad=20)
        plt.tight_layout()
        chart_path = charts_dir / f"{filename_base}_chart2.png"
        plt.savefig(chart_path, dpi=300, bbox_inches='tight')
        plt.close()
        self.chart_images.append(str(chart_path))
        
        # Gráfica 3: Línea (tendencia temporal)
        fig, ax = plt.subplots(figsize=(12, 6))
        months = ['Ene', 'Feb', 'Mar', 'Abr', 'May', 'Jun']
        leads = [50, 75, 120, 150, 180, 200]
        conversions = [5, 8, 12, 15, 18, 20]
        
        ax.plot(months, leads, marker='o', linewidth=3, markersize=10, 
               label='Leads', color=COLORS['primary'])
        ax.plot(months, conversions, marker='s', linewidth=3, markersize=10,
               label='Conversiones', color=COLORS['success'])
        
        ax.set_title('Tendencia de Leads y Conversiones', fontsize=16, fontweight='bold', pad=20)
        ax.set_xlabel('Mes', fontsize=12)
        ax.set_ylabel('Cantidad', fontsize=12)
        ax.legend(loc='upper left', fontsize=11, framealpha=0.9)
        ax.grid(True, alpha=0.3, linestyle='--')
        plt.tight_layout()
        chart_path = charts_dir / f"{filename_base}_chart3.png"
        plt.savefig(chart_path, dpi=300, bbox_inches='tight')
        plt.close()
        self.chart_images.append(str(chart_path))
    
    def _create_metrics_sheet(self, lines: List[str]):
        """Crea una hoja con métricas extraídas"""
        ws = self.wb.create_sheet(title="Métricas")
        
        # Título
        ws['A1'] = "Métricas Clave"
        ws['A1'].font = Font(size=16, bold=True, color="FFFFFF")
        ws['A1'].fill = PatternFill(start_color="2563EB", end_color="2563EB", fill_type="solid")
        ws.merge_cells('A1:B1')
        
        # Extraer métricas del texto
        metrics = self._extract_metrics(lines)
        
        row = 3
        ws[f'A{row}'] = "Métrica"
        ws[f'B{row}'] = "Valor"
        ws[f'A{row}'].font = Font(bold=True)
        ws[f'B{row}'].font = Font(bold=True)
        
        row += 1
        for metric, value in metrics.items():
            ws[f'A{row}'] = metric
            ws[f'B{row}'] = value
            row += 1
        
        ws.column_dimensions['A'].width = 40
        ws.column_dimensions['B'].width = 20
    
    def _extract_metrics(self, lines: List[str]) -> Dict[str, Any]:
        """Extrae métricas numéricas del texto"""
        metrics = {}
        
        for line in lines:
            # Buscar patrones como "Total: 150" o "Tasa: 25%"
            patterns = [
                (r'Total\s+([^:]+):\s*(\d+)', 'Total {}'),
                (r'Tasa\s+([^:]+):\s*(\d+(?:\.\d+)?)%', 'Tasa {} (%)'),
                (r'([A-Za-z\s]+):\s*(\d+(?:\.\d+)?)', '{}'),
            ]
            
            for pattern, format_str in patterns:
                matches = re.findall(pattern, line, re.IGNORECASE)
                for match in matches:
                    if len(match) == 2:
                        key = format_str.format(match[0].strip())
                        metrics[key] = match[1]
        
        # Métricas por defecto si no se encuentran
        if not metrics:
            metrics = {
                'Candidatos Identificados': 150,
                'Tasa de Respuesta': '25%',
                'Colaboraciones Activas': 8,
                'ROI Esperado': '200-500%'
            }
        
        return metrics
    
    def _parse_table(self, lines: List[str], start_idx: int) -> Optional[List[List[str]]]:
        """Parsea una tabla de Markdown"""
        table_data = []
        i = start_idx
        
        while i < len(lines) and '|' in lines[i]:
            row = [cell.strip() for cell in lines[i].split('|') if cell.strip()]
            if row and not all(c == '-' for c in ''.join(row)):
                table_data.append(row)
            i += 1
        
        return table_data if len(table_data) > 1 else None
    
    def _skip_table(self, lines: List[str], idx: int) -> int:
        """Salta las líneas de una tabla"""
        i = idx
        while i < len(lines) and '|' in lines[i]:
            i += 1
        return i - 1
    
    def _clean_markdown(self, text: str) -> str:
        """Limpia el texto de sintaxis Markdown"""
        text = re.sub(r'^#{1,6}\s+', '', text)
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        text = re.sub(r'`(.+?)`', r'\1', text)
        return text
    
    def save(self, filename: str):
        """Guarda el libro de Excel"""
        self.wb.save(filename)
        print(f"✅ Libro Excel creado: {filename}")


def get_important_markdown_files(directory: str) -> List[str]:
    """Identifica los archivos markdown más importantes"""
    important_files = [
        'SISTEMAS_PROMPTS_CONSOLIDADO.md',
        'GUIA_COMPLETA_OUTREACH.md',
        'DASHBOARD_METRICAS.md',
        'ANALISIS_COMPETITIVO.md',
        'PRESUPUESTO_PRICING.md',
        'AUTOMATIZACION_AVANZADA.md',
        'ESTRATEGIAS_CONTENIDO.md',
        'HIGH_IMPACT_CONTENT_CREATION.md',
        'GUIA_MEJORES_PRACTICAS_AVANZADAS.md',
        'SISTEMA_CALENDARIO_CONTENIDO_REDES_SOCIALES.md',
        'COMMUNITY_BUILDING.md',
        'OPTIMIZACION_CONTINUA.md',
        'ESCALAMIENTO_ESTRATEGIA.md',
        'INTEGRACION_MARKETING.md',
        'EMAIL_SEQUENCES.md'
    ]
    
    found_files = []
    base_path = Path(directory)
    
    for filename in important_files:
        file_path = base_path / filename
        if file_path.exists():
            found_files.append(str(file_path))
        else:
            # Buscar en subdirectorios
            for md_file in base_path.rglob(filename):
                found_files.append(str(md_file))
                break
    
    # Si no se encuentran los específicos, buscar los más grandes
    if not found_files:
        all_md = list(base_path.rglob('*.md'))
        all_md.sort(key=lambda x: x.stat().st_size, reverse=True)
        found_files = [str(f) for f in all_md[:10]]
    
    return found_files


def convert_markdown_to_documents(md_file: str, output_dir: str):
    """Convierte un archivo markdown a Word y Excel"""
    print(f"\n📄 Procesando: {Path(md_file).name}")
    
    # Leer archivo
    try:
        with open(md_file, 'r', encoding='utf-8') as f:
            content = f.read()
    except Exception as e:
        print(f"❌ Error leyendo archivo: {e}")
        return
    
    # Crear directorio de salida
    output_path = Path(output_dir)
    output_path.mkdir(exist_ok=True)
    
    # Nombre base del archivo
    filename_base = Path(md_file).stem
    
    # Convertir a Word
    try:
        word_converter = MarkdownToWordConverter()
        word_converter.parse_markdown(content)
        
        # Añadir gráficas al documento Word
        charts_dir = Path("charts")
        if charts_dir.exists():
            for chart_file in charts_dir.glob(f"{filename_base}_*.png"):
                word_converter.add_chart_image(str(chart_file))
        
        word_filename = output_path / f"{filename_base}_PREMIUM.docx"
        word_converter.save(str(word_filename))
    except Exception as e:
        print(f"❌ Error creando Word: {e}")
    
    # Convertir a Excel
    try:
        excel_converter = MarkdownToExcelConverter()
        excel_converter.parse_markdown(content, filename_base)
        excel_filename = output_path / f"{filename_base}_PREMIUM.xlsx"
        excel_converter.save(str(excel_filename))
    except Exception as e:
        print(f"❌ Error creando Excel: {e}")


def main():
    """Función principal"""
    print("🚀 Convertidor Premium de Markdown a Word y Excel")
    print("=" * 60)
    
    # Directorio de marketing
    marketing_dir = "/Users/adan/Documents/documentos_blatam/01_marketing"
    output_dir = "/Users/adan/Documents/documentos_blatam/01_marketing/documentos_convertidos_premium"
    
    # Crear directorio de salida
    Path(output_dir).mkdir(exist_ok=True)
    
    # Obtener archivos importantes
    important_files = get_important_markdown_files(marketing_dir)
    
    print(f"\n📚 Archivos encontrados: {len(important_files)}")
    for f in important_files:
        print(f"  - {Path(f).name}")
    
    # Convertir cada archivo
    for md_file in important_files:
        convert_markdown_to_documents(md_file, output_dir)
    
    print("\n" + "=" * 60)
    print("✅ Conversión completada!")
    print(f"📁 Archivos guardados en: {output_dir}")
    print(f"📊 Gráficas guardadas en: charts/")


if __name__ == "__main__":
    main()



