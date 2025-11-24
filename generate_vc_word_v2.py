import docx
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_enhanced_word(filename):
    doc = docx.Document()
    
    # --- Document Setup ---
    section = doc.sections[0]
    section.page_height = Cm(29.7) # A4
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # --- Styles ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Headings
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Calibri Light'
    h1.font.size = Pt(24)
    h1.font.color.rgb = RGBColor(32, 55, 100) # Navy
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)

    h2 = doc.styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(16)
    h2.font.color.rgb = RGBColor(47, 117, 181) # Blue
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(18)

    # --- 1. Cover Page ---
    for _ in range(5): doc.add_paragraph()
    
    title_para = doc.add_paragraph('VENTURA CAPITAL')
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.runs[0]
    title_run.bold = True
    title_run.font.size = Pt(36)
    title_run.font.color.rgb = RGBColor(32, 55, 100)

    subtitle_para = doc.add_paragraph('KIT ESTRATÉGICO DE LEVANTAMIENTO DE CAPITAL')
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.runs[0]
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.color.rgb = RGBColor(128, 128, 128)

    for _ in range(3): doc.add_paragraph()

    details_para = doc.add_paragraph('SaaS IA Marketing/Copywriting\nLATAM Market Entry\n\nPreparado para: Inversionistas Serie A\nFecha: Diciembre 2024')
    details_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()

    # --- 2. Table of Contents (Simulation) ---
    doc.add_heading('Índice de Contenidos', 0)
    toc_items = [
        ("1. Executive Summary", "3"),
        ("2. Due Diligence & Risk Analysis", "4"),
        ("3. Term Sheet (Propuesta)", "6"),
        ("4. Guiones de Negociación", "8"),
        ("5. Anexos Legales", "10")
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Cm(14)
    table.columns[1].width = Cm(2)
    
    # Remove borders for TOC look - simplified approach
    # table.style = 'Table Normal' # Removing style assignment to use default
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Sección"
    hdr_cells[1].text = "Pág."
    
    for section, page in toc_items:
        row_cells = table.add_row().cells
        row_cells[0].text = section
        row_cells[1].text = page
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_page_break()

    # --- 3. Content Sections ---
    
    # Executive Summary
    doc.add_heading('1. Executive Summary', 1)
    doc.add_paragraph('Este documento presenta la estrategia integral de levantamiento de capital para nuestra plataforma SaaS de IA. Incluye análisis de riesgos, estructuras legales propuestas y herramientas de negociación.')
    
    # Risk Analysis (Table Format)
    doc.add_heading('2. Due Diligence & Risk Analysis', 1)
    doc.add_paragraph('Matriz de riesgos identificados y estrategias de mitigación:')

    risk_table = doc.add_table(rows=1, cols=3)
    risk_table.style = 'Light Shading Accent 1'
    hdr_cells = risk_table.rows[0].cells
    hdr_cells[0].text = 'Categoría'
    hdr_cells[1].text = 'Riesgo Identificado'
    hdr_cells[2].text = 'Estrategia de Mitigación'

    risks = [
        ('Financiero', 'Falta de liquidez corto plazo', 'Línea de crédito puente $500k'),
        ('Técnico', 'Dependencia de APIs externas', 'Desarrollo de modelos propios (Roadmap Q3)'),
        ('Legal', 'Regulación de IA en LATAM', 'Compliance proactivo y auditoría legal'),
        ('Mercado', 'Entrada de competidores US', 'Foco en localización cultural profunda')
    ]

    for cat, risk, mit in risks:
        row_cells = risk_table.add_row().cells
        row_cells[0].text = cat
        row_cells[1].text = risk
        row_cells[2].text = mit

    doc.add_page_break()

    # Term Sheet
    doc.add_heading('3. Term Sheet (Propuesta)', 1)
    doc.add_paragraph('Términos clave para la ronda de inversión Series A.')

    ts_table = doc.add_table(rows=1, cols=2)
    ts_table.style = 'Light List Accent 1'
    
    terms = [
        ('Investment Amount', '$2,000,000 USD'),
        ('Pre-money Valuation', '$8,000,000 USD'),
        ('Instrument', 'Preferred Stock (Series A)'),
        ('Liquidation Pref.', '1x Non-participating'),
        ('Board Seats', '2 Founders, 1 Investor, 1 Independent'),
        ('ESOP Pool', '10% Post-money')
    ]

    for term, val in terms:
        row_cells = ts_table.add_row().cells
        row_cells[0].text = term
        row_cells[0].paragraphs[0].runs[0].bold = True
        row_cells[1].text = val

    # Negotiation Scripts
    doc.add_heading('4. Guiones de Negociación', 1)
    
    doc.add_heading('Escenario: Valuación', 2)
    script_para = doc.add_paragraph()
    script_para.add_run('Inversionista: ').bold = True
    script_para.add_run('"La valuación de $10M nos parece alta dado el MRR actual."')
    
    response_para = doc.add_paragraph()
    response_para.paragraph_format.left_indent = Cm(1)
    response_para.add_run('Respuesta Sugerida: ').bold = True
    response_para.add_run('"Entendemos el punto. Sin embargo, si miramos los comparables de SaaS vertical en LATAM y nuestra tasa de crecimiento del 20% MoM, el múltiplo está justificado. Estamos dispuestos a ajustar si estructuramos un earn-out basado en hitos de Q4."').italic = True

    doc.add_heading('Escenario: Control del Board', 2)
    script_para2 = doc.add_paragraph()
    script_para2.add_run('Inversionista: ').bold = True
    script_para2.add_run('"Queremos 2 asientos en el board para tener mayor control."')
    
    response_para2 = doc.add_paragraph()
    response_para2.paragraph_format.left_indent = Cm(1)
    response_para2.add_run('Respuesta Sugerida: ').bold = True
    response_para2.add_run('"Valoramos su experiencia operativa. Nuestra propuesta es mantener el board ágil con 1 asiento de inversionista y 1 independiente de mutuo acuerdo, lo cual alinea incentivos sin burocratizar la ejecución."').italic = True

    # Footer
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Ventura Capital - Confidencial - 2024"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.style = doc.styles['Footer']

    doc.save(filename)
    print(f"Enhanced Word created: {filename}")

if __name__ == "__main__":
    create_enhanced_word("Ventura_Capital_Documentacion_V2.docx")

