#!/usr/bin/env python3
"""
Script para generar documentos premium (PDF, Word, Excel) con gráficas
de los documentos más importantes del proyecto.
"""

import os
import re
import json
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Any
import markdown
from collections import Counter

SPANISH_STOPWORDS = {
    "de", "la", "que", "el", "en", "y", "a", "los", "del", "se", "las", "por",
    "un", "para", "con", "no", "una", "su", "al", "lo", "como", "más", "pero",
    "sus", "le", "ya", "o", "este", "sí", "porque", "esta", "entre", "cuando",
    "muy", "sin", "sobre", "también", "me", "hasta", "hay", "donde", "quien",
    "desde", "todo", "nos", "durante", "todos", "uno", "les", "ni", "contra",
    "otros", "ese", "eso", "ante", "ellos", "e", "esto", "mí", "antes", "algunos",
    "qué", "unos", "yo", "otro", "otras", "otra", "él", "tanto", "esa", "estos",
    "mucho", "quienes", "nada", "muchos", "cual", "poco", "ella", "estar",
    "estas", "algunas", "algo", "nosotros", "mi", "mis", "tú", "te", "ti",
    "tu", "tus", "ellas", "nosotras", "vosotros", "vosotras", "os", "mío",
    "mí", "mía", "míos", "mías", "tuyo", "tuya", "tuyos", "tuyas", "suyo",
    "suya", "suyos", "suyas", "nuestro", "nuestra", "nuestros", "nuestras",
    "vuestro", "vuestra", "vuestros", "vuestras", "esos", "esas", "estoy",
    "estás", "está", "estamos", "estáis", "están", "esté", "estés", "estemos",
    "estéis", "estén", "estaré", "estarás", "estará", "estaremos", "estaréis",
    "estarán", "estaría", "estarías", "estaríamos", "estaríais", "estarían",
    "estaba", "estabas", "estábamos", "estabais", "estaban", "estuve",
    "estuviste", "estuvo", "estuvimos", "estuvisteis", "estuvieron"
}

# PDF
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
    from reportlab.pdfgen import canvas
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("⚠️  reportlab no disponible. Instala con: pip install reportlab")

# Word
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False
    print("⚠️  python-docx no disponible. Instala con: pip install python-docx")

# Excel
try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.chart import BarChart, LineChart, PieChart, Reference
    from openpyxl.drawing.image import Image as XLImage
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    print("⚠️  openpyxl no disponible. Instala con: pip install openpyxl")

# Gráficas
try:
    import matplotlib
    matplotlib.use('Agg')  # Backend sin GUI
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    import numpy as np
    MATPLOTLIB_AVAILABLE = True
except ImportError:
    MATPLOTLIB_AVAILABLE = False
    print("⚠️  matplotlib no disponible. Instala con: pip install matplotlib numpy")

# Markdown
try:
    import markdown
    from markdown.extensions import codehilite, tables, fenced_code
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False
    print("⚠️  markdown no disponible. Instala con: pip install markdown")


class DocumentConverter:
    """Convertidor de documentos Markdown a PDF, Word y Excel con gráficas."""
    
    def __init__(self, output_dir: str = "docs_premium"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        self.graphs_dir = self.output_dir / "graphs"
        self.graphs_dir.mkdir(exist_ok=True)
        self.stopwords = SPANISH_STOPWORDS
        
    def identify_important_docs(self, base_dir: str) -> List[Dict[str, Any]]:
        """Identifica los documentos más importantes."""
        base_path = Path(base_dir)
        important_patterns = [
            "ARCHITECTURE",
            "REFACTORING",
            "README",
            "CHANGELOG",
            "BEST_PRACTICES",
            "airflow_automation",
            "MEJORAS",
            "IMPROVEMENTS",
            "PLAN",
            "STRATEGY"
        ]
        
        important_docs = []
        
        # Documentos específicos prioritarios
        priority_files = [
            "airflow_automation_prompt.md",
            "truthgpt_collected/integration_code/production_code/ARCHITECTURE_IMPROVEMENTS.md",
            "truthgpt_collected/integration_code/production_code/REFACTORING_PLAN.md",
            "truthgpt_collected/integration_code/production_code/README.md",
            "truthgpt_collected/integration_code/production_code/ARCHITECTURE.md"
        ]
        
        for file_path in priority_files:
            full_path = base_path / file_path
            if full_path.exists():
                important_docs.append({
                    "path": full_path,
                    "name": full_path.stem,
                    "priority": 1
                })
        
        # Buscar otros documentos importantes
        for pattern in important_patterns:
            for md_file in base_path.rglob(f"*{pattern}*.md"):
                if md_file not in [doc["path"] for doc in important_docs]:
                    # Calcular prioridad basada en tamaño y nombre
                    priority = 2
                    if "ARCHITECTURE" in md_file.name.upper() or "README" in md_file.name.upper():
                        priority = 1
                    
                    important_docs.append({
                        "path": md_file,
                        "name": md_file.stem,
                        "priority": priority
                    })
        
        # Ordenar por prioridad
        important_docs.sort(key=lambda x: (x["priority"], x["path"].stat().st_size), reverse=True)
        
        # Limitar a los 10 más importantes
        return important_docs[:10]
    
    def parse_markdown(self, file_path: Path) -> Dict[str, Any]:
        """Parsea un archivo Markdown y extrae información estructurada."""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Convertir markdown a HTML para análisis
        md = markdown.Markdown(extensions=['codehilite', 'tables', 'fenced_code'])
        html_content = md.convert(content)
        
        # Extraer secciones
        sections = self._extract_sections(content)
        
        # Extraer métricas y datos para gráficas
        metrics = self._extract_metrics(content)
        
        # Extraer tablas
        tables_data = self._extract_tables(content)
        keywords = self._extract_keywords(content)
        summary_points = self._build_summary(sections, keywords)
        action_items = self._extract_action_items(content)
        timeline_dates = self._extract_dates(content)
        
        return {
            "raw_content": content,
            "html_content": html_content,
            "sections": sections,
            "metrics": metrics,
            "tables": tables_data,
            "keywords": keywords,
            "summary_points": summary_points,
            "action_items": action_items,
            "timeline_dates": timeline_dates,
            "title": self._extract_title(content),
            "metadata": self._extract_metadata(content)
        }
    
    def _extract_title(self, content: str) -> str:
        """Extrae el título del documento."""
        lines = content.split('\n')
        for line in lines[:10]:
            if line.startswith('# '):
                return line[2:].strip()
        return "Documento sin título"
    
    def _extract_metadata(self, content: str) -> Dict[str, Any]:
        """Extrae metadatos del documento."""
        metadata = {
            "date": datetime.now().strftime("%Y-%m-%d"),
            "author": "Sistema de Conversión",
            "version": "1.0"
        }
        
        # Buscar metadatos en formato YAML frontmatter
        if content.startswith('---'):
            try:
                parts = content.split('---', 2)
                if len(parts) >= 3:
                    import yaml
                    frontmatter = yaml.safe_load(parts[1])
                    metadata.update(frontmatter)
            except:
                pass
        
        # Buscar metadatos en el contenido
        date_match = re.search(r'Date[:\s]+(\d{4}-\d{2}-\d{2})', content, re.IGNORECASE)
        if date_match:
            metadata["date"] = date_match.group(1)
        
        return metadata
    
    def _extract_sections(self, content: str) -> List[Dict[str, Any]]:
        """Extrae secciones del documento."""
        sections = []
        lines = content.split('\n')
        current_section = None
        current_content = []
        
        for line in lines:
            if line.startswith('#'):
                if current_section:
                    sections.append({
                        "title": current_section,
                        "content": '\n'.join(current_content),
                        "level": len(current_section) - len(current_section.lstrip('#'))
                    })
                current_section = line.strip()
                current_content = []
            else:
                current_content.append(line)
        
        if current_section:
            sections.append({
                "title": current_section,
                "content": '\n'.join(current_content),
                "level": len(current_section) - len(current_section.lstrip('#'))
            })
        
        return sections
    
    def _extract_metrics(self, content: str) -> Dict[str, Any]:
        """Extrae métricas y datos numéricos para gráficas."""
        metrics = {
            "numbers": [],
            "percentages": [],
            "counts": {},
            "status": {"completed": 0, "pending": 0, "in_progress": 0},
            "word_count": 0,
            "char_count": 0,
            "estimated_minutes": 0.0
        }
        
        # Buscar números
        numbers = re.findall(r'\b\d+\.?\d*\b', content)
        metrics["numbers"] = [float(n) for n in numbers[:50]]  # Limitar a 50
        
        # Buscar porcentajes
        percentages = re.findall(r'(\d+\.?\d*)%', content)
        metrics["percentages"] = [float(p) for p in percentages[:20]]
        
        # Contar estados
        metrics["status"]["completed"] = len(re.findall(r'✅|\[x\]|completed|completado', content, re.IGNORECASE))
        metrics["status"]["pending"] = len(re.findall(r'⏳|\[ \]|pending|pendiente', content, re.IGNORECASE))
        metrics["status"]["in_progress"] = len(re.findall(r'🔄|in progress|en progreso', content, re.IGNORECASE))
        
        # Contar secciones por nivel
        h1_count = len(re.findall(r'^# [^#]', content, re.MULTILINE))
        h2_count = len(re.findall(r'^## [^#]', content, re.MULTILINE))
        h3_count = len(re.findall(r'^### [^#]', content, re.MULTILINE))
        metrics["counts"]["h1"] = h1_count
        metrics["counts"]["h2"] = h2_count
        metrics["counts"]["h3"] = h3_count

        # Métricas de longitud y tiempo de lectura
        words = re.findall(r'\w+', content)
        metrics["word_count"] = len(words)
        metrics["char_count"] = len(content)
        metrics["estimated_minutes"] = round(len(words) / 200.0, 1) if words else 0.0
        
        return metrics
    
    def _extract_tables(self, content: str) -> List[List[List[str]]]:
        """Extrae tablas del contenido Markdown."""
        tables = []
        lines = content.split('\n')
        current_table = []
        in_table = False
        
        for line in lines:
            if '|' in line and not line.strip().startswith('```'):
                if not in_table:
                    in_table = True
                    current_table = []
                current_table.append([cell.strip() for cell in line.split('|') if cell.strip()])
            else:
                if in_table and current_table:
                    if len(current_table) > 1:  # Al menos header + 1 fila
                        tables.append(current_table)
                    current_table = []
                    in_table = False
        
        if in_table and current_table:
            tables.append(current_table)
        
        return tables

    def _extract_keywords(self, content: str, top_n: int = 8) -> List[Dict[str, Any]]:
        """Identifica palabras clave frecuentes excluyendo stopwords."""
        words = re.findall(r'[A-Za-zÁÉÍÓÚÜÑáéíóúüñ]{4,}', content.lower())
        filtered = [word for word in words if word not in self.stopwords]
        frequency = Counter(filtered)
        return [{"word": word, "count": count} for word, count in frequency.most_common(top_n)]

    def _build_summary(self, sections: List[Dict[str, Any]], keywords: List[Dict[str, Any]]) -> List[str]:
        """Construye un resumen ejecutivo con bullets clave."""
        summary_points: List[str] = []

        # Priorizar bullets existentes en el contenido
        for section in sections:
            for line in section["content"].split('\n'):
                stripped = line.strip()
                if stripped.startswith(('-', '*', '•')) and len(stripped) > 5:
                    summary_points.append(stripped.lstrip('-*• ').strip())
                if len(summary_points) >= 5:
                    break
            if len(summary_points) >= 5:
                break

        # Complementar con keywords si faltan bullets
        if len(summary_points) < 5 and keywords:
            remaining = 5 - len(summary_points)
            for keyword in keywords[:remaining]:
                summary_points.append(
                    f"Profundizar en iniciativas relacionadas con '{keyword['word']}' (menciones: {keyword['count']})."
                )

        if not summary_points:
            summary_points.append("Revisar el documento para extraer hallazgos y próximos pasos.")

        return summary_points[:5]

    def _extract_action_items(self, content: str, max_items: int = 8) -> List[Dict[str, str]]:
        """Extrae acciones y tareas del contenido."""
        action_items: List[Dict[str, str]] = []
        lines = content.split('\n')

        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            status = "pendiente"
            if any(token in stripped.lower() for token in ["✅", "[x]", "[X]", "completado", "done"]):
                status = "completado"
            elif any(token in stripped.lower() for token in ["⏳", "en progreso", "in progress"]):
                status = "en progreso"

            if stripped.startswith(('-', '*')) or stripped.startswith('•'):
                clean_text = stripped.lstrip('-*• ').strip()
                if clean_text:
                    action_items.append({"descripcion": clean_text, "estado": status})
            elif stripped.startswith('1.') or stripped[:2].isdigit():
                action_items.append({"descripcion": stripped, "estado": status})

            if len(action_items) >= max_items:
                break

        return action_items

    def _extract_dates(self, content: str, max_dates: int = 10) -> List[str]:
        """Identifica fechas mencionadas para generar una cronología simple."""
        date_patterns = [
            r'\b\d{4}-\d{2}-\d{2}\b',          # 2025-11-30
            r'\b\d{2}/\d{2}/\d{4}\b',          # 30/11/2025
            r'\b\d{2}-\d{2}-\d{4}\b',          # 30-11-2025
            r'\b\d{1,2} de [A-Za-z]+ de \d{4}\b'  # 30 de noviembre de 2025
        ]

        dates_found: List[str] = []
        for pattern in date_patterns:
            matches = re.findall(pattern, content, flags=re.IGNORECASE)
            for match in matches:
                if match not in dates_found:
                    dates_found.append(match)
                if len(dates_found) >= max_dates:
                    return dates_found

        return dates_found
    
    def create_graphs(self, doc_name: str, metrics: Dict[str, Any]) -> List[str]:
        """Crea gráficas basadas en las métricas extraídas."""
        graph_files = []
        
        if not MATPLOTLIB_AVAILABLE:
            return graph_files
        
        # Gráfica 1: Estados del proyecto (Bar Chart)
        if sum(metrics["status"].values()) > 0:
            fig, ax = plt.subplots(figsize=(8, 6))
            statuses = list(metrics["status"].keys())
            values = list(metrics["status"].values())
            colors_list = ['#2ecc71', '#f39c12', '#3498db']
            
            bars = ax.bar(statuses, values, color=colors_list)
            ax.set_title('Estados del Proyecto', fontsize=14, fontweight='bold')
            ax.set_ylabel('Cantidad', fontsize=12)
            ax.set_xlabel('Estado', fontsize=12)
            
            # Agregar valores en las barras
            for bar in bars:
                height = bar.get_height()
                ax.text(bar.get_x() + bar.get_width()/2., height,
                       f'{int(height)}',
                       ha='center', va='bottom', fontsize=10)
            
            plt.tight_layout()
            status_graph = self.graphs_dir / f"{doc_name}_status.png"
            plt.savefig(status_graph, dpi=300, bbox_inches='tight')
            plt.close()
            graph_files.append(str(status_graph))
            
            # Gráfica 1b: Estados del proyecto (Pie Chart)
            if sum(values) > 0:
                fig, ax = plt.subplots(figsize=(8, 8))
                labels = [s.replace('_', ' ').title() for s in statuses]
                explode = [0.05 if v == max(values) else 0 for v in values]
                wedges, texts, autotexts = ax.pie(values, labels=labels, colors=colors_list,
                                                   autopct='%1.1f%%', startangle=90, explode=explode,
                                                   shadow=True, textprops={'fontsize': 11})
                ax.set_title('Distribución de Estados (%)', fontsize=14, fontweight='bold', pad=20)
                for autotext in autotexts:
                    autotext.set_color('white')
                    autotext.set_fontweight('bold')
                plt.tight_layout()
                status_pie_graph = self.graphs_dir / f"{doc_name}_status_pie.png"
                plt.savefig(status_pie_graph, dpi=300, bbox_inches='tight')
                plt.close()
                graph_files.append(str(status_pie_graph))
        
        # Gráfica 2: Distribución de secciones
        if sum(metrics["counts"].values()) > 0:
            fig, ax = plt.subplots(figsize=(8, 6))
            levels = ['H1', 'H2', 'H3']
            counts = [metrics["counts"]["h1"], metrics["counts"]["h2"], metrics["counts"]["h3"]]
            
            bars = ax.bar(levels, counts, color=['#e74c3c', '#3498db', '#9b59b6'])
            ax.set_title('Distribución de Secciones', fontsize=14, fontweight='bold')
            ax.set_ylabel('Cantidad de Secciones', fontsize=12)
            ax.set_xlabel('Nivel de Encabezado', fontsize=12)
            
            for bar in bars:
                height = bar.get_height()
                ax.text(bar.get_x() + bar.get_width()/2., height,
                       f'{int(height)}',
                       ha='center', va='bottom', fontsize=10)
            
            plt.tight_layout()
            sections_graph = self.graphs_dir / f"{doc_name}_sections.png"
            plt.savefig(sections_graph, dpi=300, bbox_inches='tight')
            plt.close()
            graph_files.append(str(sections_graph))
        
        # Gráfica 3: Distribución de porcentajes (si hay)
        if metrics["percentages"]:
            fig, ax = plt.subplots(figsize=(10, 6))
            percentages = metrics["percentages"][:15]  # Limitar a 15
            ax.hist(percentages, bins=10, color='#3498db', edgecolor='black', alpha=0.7)
            ax.set_title('Distribución de Porcentajes', fontsize=14, fontweight='bold')
            ax.set_ylabel('Frecuencia', fontsize=12)
            ax.set_xlabel('Porcentaje (%)', fontsize=12)
            ax.grid(True, alpha=0.3)
            
            plt.tight_layout()
            percent_graph = self.graphs_dir / f"{doc_name}_percentages.png"
            plt.savefig(percent_graph, dpi=300, bbox_inches='tight')
            plt.close()
            graph_files.append(str(percent_graph))

        # Gráfica 4: Palabras clave
        keywords_data = metrics.get("keywords_graph_data", [])
        if keywords_data:
            fig, ax = plt.subplots(figsize=(10, 6))
            words = [item["word"] for item in keywords_data]
            counts = [item["count"] for item in keywords_data]
            bars = ax.bar(words, counts, color='#16a085')
            ax.set_title('Palabras Clave Más Frecuentes', fontsize=14, fontweight='bold')
            ax.set_ylabel('Frecuencia', fontsize=12)
            ax.set_xlabel('Keyword', fontsize=12)
            ax.tick_params(axis='x', rotation=45)

            for bar in bars:
                height = bar.get_height()
                ax.text(bar.get_x() + bar.get_width()/2., height,
                        f'{int(height)}',
                        ha='center', va='bottom', fontsize=9)

            plt.tight_layout()
            keywords_graph = self.graphs_dir / f"{doc_name}_keywords.png"
            plt.savefig(keywords_graph, dpi=300, bbox_inches='tight')
            plt.close()
            graph_files.append(str(keywords_graph))
        
        return graph_files
    
    def generate_pdf(self, doc_data: Dict[str, Any], output_path: Path):
        """Genera un documento PDF con gráficas."""
        if not PDF_AVAILABLE:
            print(f"⚠️  PDF no disponible para {output_path.name}")
            return
        
        doc = SimpleDocTemplate(str(output_path), pagesize=A4,
                              rightMargin=72, leftMargin=72,
                              topMargin=72, bottomMargin=18)
        
        story = []
        styles = getSampleStyleSheet()
        
        # Portada mejorada
        cover_style = ParagraphStyle(
            'CoverTitle',
            parent=styles['Heading1'],
            fontSize=28,
            textColor=colors.HexColor('#1a1a1a'),
            spaceAfter=20,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        story.append(Spacer(1, 1.5*inch))
        story.append(Paragraph(doc_data["title"], cover_style))
        story.append(Spacer(1, 0.3*inch))
        
        # Título estilo normal para contenido
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#2c3e50'),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        # Metadatos
        meta_style = ParagraphStyle(
            'Meta',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.HexColor('#7f8c8d'),
            alignment=TA_CENTER
        )
        metadata = doc_data["metadata"]
        meta_text = f"Fecha: {metadata.get('date', 'N/A')} | Versión: {metadata.get('version', '1.0')}"
        story.append(Paragraph(meta_text, meta_style))
        story.append(Spacer(1, 0.3*inch))

        # Métricas clave
        metrics_table_data = [
            ['Palabras', doc_data["metrics"]["word_count"],
             'Secciones H1/H2/H3', f"{doc_data['metrics']['counts']['h1']}/"
                                   f"{doc_data['metrics']['counts']['h2']}/"
                                   f"{doc_data['metrics']['counts']['h3']}"],
            ['Tiempo estimado (min)', doc_data["metrics"]["estimated_minutes"],
             'Actividades completadas', doc_data["metrics"]["status"]["completed"]]
        ]
        metrics_table = Table(metrics_table_data, colWidths=[2.0*inch, 1.5*inch, 2.0*inch, 1.5*inch])
        metrics_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#ecf0f1')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor('#2c3e50')),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#bdc3c7'))
        ]))
        story.append(metrics_table)
        story.append(Spacer(1, 0.3*inch))

        # Resumen ejecutivo
        summary_style = ParagraphStyle(
            'SummaryHeading',
            parent=styles['Heading2'],
            textColor=colors.HexColor('#1abc9c')
        )
        story.append(Paragraph("Resumen Ejecutivo", summary_style))
        story.append(Spacer(1, 0.1*inch))
        for point in doc_data["summary_points"]:
            story.append(Paragraph(f"• {point}", styles['Normal']))
            story.append(Spacer(1, 0.05*inch))
        story.append(Spacer(1, 0.2*inch))

        # Acciones prioritarias
        if doc_data["action_items"]:
            actions_style = ParagraphStyle(
                'ActionsHeading',
                parent=styles['Heading2'],
                textColor=colors.HexColor('#2980b9')
            )
            story.append(Paragraph("Acciones Prioritarias", actions_style))
            story.append(Spacer(1, 0.1*inch))

            table_data = [['Descripción', 'Estado']]
            for action in doc_data["action_items"]:
                table_data.append([action["descripcion"][:120], action["estado"].title()])

            actions_table = Table(table_data, colWidths=[4.5*inch, 1.5*inch])
            actions_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#d6eaf8')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#1b4f72')),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#aed6f1')),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f2f4f4')])
            ]))
            story.append(actions_table)
            story.append(Spacer(1, 0.2*inch))

        # Cronología simple
        if doc_data["timeline_dates"]:
            timeline_style = ParagraphStyle(
                'TimelineHeading',
                parent=styles['Heading2'],
                textColor=colors.HexColor('#8e44ad')
            )
            story.append(Paragraph("Hitos y Fechas Identificadas", timeline_style))
            story.append(Spacer(1, 0.1*inch))
            for date in doc_data["timeline_dates"]:
                story.append(Paragraph(f"• {date}", styles['Normal']))
            story.append(Spacer(1, 0.2*inch))
        
        # Índice de contenido
        toc_style = ParagraphStyle(
            'TOCHeading',
            parent=styles['Heading2'],
            textColor=colors.HexColor('#34495e')
        )
        story.append(PageBreak())
        story.append(Paragraph("Índice de Contenido", toc_style))
        story.append(Spacer(1, 0.2*inch))
        
        for i, section in enumerate(doc_data["sections"][:30], 1):  # Limitar a 30 secciones
            level = section["level"]
            title_clean = section["title"].replace('#', '').strip()
            if level <= 3 and title_clean:
                indent = "  " * (level - 1)
                toc_text = f"{indent}{i}. {title_clean}"
                story.append(Paragraph(toc_text, styles['Normal']))
        story.append(Spacer(1, 0.3*inch))
        story.append(PageBreak())
        
        # Agregar gráficas
        graph_files = self.create_graphs(output_path.stem, doc_data["metrics"])
        for graph_file in graph_files:
            if os.path.exists(graph_file):
                img = Image(graph_file, width=6*inch, height=4.5*inch)
                story.append(img)
                story.append(Spacer(1, 0.2*inch))
        
        # Contenido por secciones
        heading_styles = {
            1: styles['Heading1'],
            2: styles['Heading2'],
            3: styles['Heading3']
        }
        
        for section in doc_data["sections"]:
            level = section["level"]
            if level <= 3:
                style = heading_styles.get(level, styles['Heading2'])
                story.append(Paragraph(section["title"], style))
                story.append(Spacer(1, 0.1*inch))
            
            # Procesar contenido
            content = section["content"]
            # Limpiar HTML problemático
            content = re.sub(r'<br\s*/?>', '\n', content, flags=re.IGNORECASE)
            content = re.sub(r'<para>.*?</para>', '', content, flags=re.DOTALL)
            # Limpiar markdown básico
            content = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', content)
            content = re.sub(r'\*(.*?)\*', r'<i>\1</i>', content)
            # Escapar caracteres especiales de reportlab
            content = content.replace('&', '&amp;')
            content = content.replace('<', '&lt;').replace('>', '&gt;')
            # Restaurar tags permitidos
            content = re.sub(r'&lt;b&gt;(.*?)&lt;/b&gt;', r'<b>\1</b>', content)
            content = re.sub(r'&lt;i&gt;(.*?)&lt;/i&gt;', r'<i>\1</i>', content)
            # Limpiar código inline
            content = re.sub(r'`([^`]+)`', r'<font face="Courier" size="9">\1</font>', content)
            
            # Procesar tablas primero
            lines = content.split('\n')
            current_table = []
            in_table = False
            processed_lines = []
            
            for line in lines:
                if '|' in line and line.strip().startswith('|'):
                    if not in_table:
                        in_table = True
                        current_table = []
                    current_table.append([cell.strip() for cell in line.split('|') if cell.strip()])
                else:
                    if in_table and len(current_table) > 1:
                        # Procesar tabla
                        try:
                            table = Table(current_table)
                            table.setStyle(TableStyle([
                                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#34495e')),
                                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                                ('FONTSIZE', (0, 0), (-1, 0), 10),
                                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                                ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                                ('FONTSIZE', (0, 1), (-1, -1), 9),
                                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8f9fa')])
                            ]))
                            story.append(table)
                            story.append(Spacer(1, 0.2*inch))
                        except:
                            pass
                        current_table = []
                        in_table = False
                    if line.strip() and not line.strip().startswith('|'):
                        processed_lines.append(line)
            
            # Procesar líneas restantes como párrafos
            content_clean = '\n'.join(processed_lines)
            paragraphs = content_clean.split('\n\n')
            for para in paragraphs:
                para = para.strip()
                if para and len(para) > 0:
                    try:
                        # Limitar longitud para evitar problemas
                        if len(para) > 2000:
                            para = para[:2000] + "..."
                        story.append(Paragraph(para, styles['Normal']))
                        story.append(Spacer(1, 0.1*inch))
                    except Exception as e:
                        # Si falla, agregar como texto plano
                        try:
                            para_clean = para.replace('<', '').replace('>', '').replace('&', 'and')
                            story.append(Paragraph(para_clean[:500], styles['Normal']))
                            story.append(Spacer(1, 0.1*inch))
                        except:
                            pass  # Saltar párrafos problemáticos
            
            story.append(Spacer(1, 0.2*inch))
        
        # Construir PDF
        doc.build(story)
        print(f"✅ PDF generado: {output_path}")
    
    def generate_word(self, doc_data: Dict[str, Any], output_path: Path):
        """Genera un documento Word con gráficas."""
        if not WORD_AVAILABLE:
            print(f"⚠️  Word no disponible para {output_path.name}")
            return
        
        doc = Document()
        
        # Configurar estilos
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        # Título
        title = doc.add_heading(doc_data["title"], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadatos
        metadata = doc_data["metadata"]
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta_para.add_run(f"Fecha: {metadata.get('date', 'N/A')} | Versión: {metadata.get('version', '1.0')}")
        meta_run.font.size = Pt(9)
        meta_run.font.color.rgb = RGBColor(127, 127, 127)
        
        doc.add_paragraph()  # Espacio

        # Resumen ejecutivo
        doc.add_heading("Resumen Ejecutivo", level=1)
        for point in doc_data["summary_points"]:
            doc.add_paragraph(point, style='List Bullet')
        doc.add_paragraph()

        # Acciones prioritarias
        if doc_data["action_items"]:
            doc.add_heading("Acciones Prioritarias", level=2)
            table = doc.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Descripción'
            hdr_cells[1].text = 'Estado'
            for action in doc_data["action_items"]:
                row_cells = table.add_row().cells
                row_cells[0].text = action["descripcion"]
                row_cells[1].text = action["estado"].title()
            doc.add_paragraph()

        # Métricas clave
        doc.add_heading("Métricas Clave", level=2)
        metrics_para = doc.add_paragraph()
        metrics_para.add_run(f"Palabras totales: {doc_data['metrics']['word_count']}\n")
        metrics_para.add_run(f"Lectura estimada (min): {doc_data['metrics']['estimated_minutes']}\n")
        metrics_para.add_run("Secciones H1/H2/H3: "
                             f"{doc_data['metrics']['counts']['h1']}/"
                             f"{doc_data['metrics']['counts']['h2']}/"
                             f"{doc_data['metrics']['counts']['h3']}")
        doc.add_paragraph()

        # Palabras clave
        doc.add_heading("Palabras Clave Prioritarias", level=2)
        if doc_data["keywords"]:
            table = doc.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Palabra'
            hdr_cells[1].text = 'Frecuencia'
            for item in doc_data["keywords"]:
                row_cells = table.add_row().cells
                row_cells[0].text = item["word"]
                row_cells[1].text = str(item["count"])
        else:
            doc.add_paragraph("No se identificaron palabras clave dominantes.")
        doc.add_paragraph()

        if doc_data["timeline_dates"]:
            doc.add_heading("Cronología Identificada", level=2)
            for date in doc_data["timeline_dates"]:
                doc.add_paragraph(date, style='List Bullet')
            doc.add_paragraph()

        doc.add_page_break()
        
        # Agregar gráficas
        graph_files = self.create_graphs(output_path.stem, doc_data["metrics"])
        for graph_file in graph_files:
            if os.path.exists(graph_file):
                doc.add_picture(graph_file, width=Inches(6))
                doc.add_paragraph()  # Espacio después de la gráfica
        
        # Contenido por secciones
        for section in doc_data["sections"]:
            level = min(section["level"], 3)  # Word soporta hasta nivel 3 bien
            doc.add_heading(section["title"].replace('#', '').strip(), level)
            
            # Procesar contenido
            content = section["content"]
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                para = para.strip()
                if para and not para.startswith('|'):
                    # Limpiar markdown básico
                    para = re.sub(r'\*\*(.*?)\*\*', r'\1', para)  # Negrita
                    para = re.sub(r'\*(.*?)\*', r'\1', para)  # Cursiva
                    para = re.sub(r'`(.*?)`', r'\1', para)  # Código
                    doc.add_paragraph(para)
        
        doc.save(str(output_path))
        print(f"✅ Word generado: {output_path}")
    
    def generate_excel(self, doc_data: Dict[str, Any], output_path: Path):
        """Genera un documento Excel con datos y gráficas."""
        if not EXCEL_AVAILABLE:
            print(f"⚠️  Excel no disponible para {output_path.name}")
            return
        
        wb = Workbook()
        ws = wb.active
        ws.title = "Resumen"
        
        # Estilos
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF", size=12)
        title_font = Font(bold=True, size=16)
        
        # Título
        ws['A1'] = doc_data["title"]
        ws['A1'].font = title_font
        ws.merge_cells('A1:D1')
        ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
        
        # Metadatos
        metadata = doc_data["metadata"]
        ws['A3'] = f"Fecha: {metadata.get('date', 'N/A')}"
        ws['B3'] = f"Versión: {metadata.get('version', '1.0')}"
        
        # Hoja de Métricas
        ws_metrics = wb.create_sheet("Métricas")
        
        # Estados
        ws_metrics['A1'] = "Estados del Proyecto"
        ws_metrics['A1'].font = header_font
        ws_metrics['A1'].fill = header_fill
        
        ws_metrics['A2'] = "Estado"
        ws_metrics['B2'] = "Cantidad"
        ws_metrics['A2'].font = header_font
        ws_metrics['A2'].fill = header_fill
        ws_metrics['B2'].font = header_font
        ws_metrics['B2'].fill = header_fill
        
        row = 3
        for status, count in doc_data["metrics"]["status"].items():
            ws_metrics[f'A{row}'] = status.replace('_', ' ').title()
            ws_metrics[f'B{row}'] = count
            row += 1
        
        # Gráfica de estados
        chart1 = BarChart()
        chart1.type = "col"
        chart1.style = 10
        chart1.title = "Estados del Proyecto"
        chart1.y_axis.title = "Cantidad"
        chart1.x_axis.title = "Estado"
        
        data = Reference(ws_metrics, min_col=2, min_row=2, max_row=row-1)
        cats = Reference(ws_metrics, min_col=1, min_row=3, max_row=row-1)
        chart1.add_data(data, titles_from_data=False)
        chart1.set_categories(cats)
        chart1.height = 10
        chart1.width = 15
        ws_metrics.add_chart(chart1, "D2")
        
        # Secciones
        ws_metrics['A10'] = "Distribución de Secciones"
        ws_metrics['A10'].font = header_font
        ws_metrics['A10'].fill = header_fill
        
        ws_metrics['A11'] = "Nivel"
        ws_metrics['B11'] = "Cantidad"
        ws_metrics['A11'].font = header_font
        ws_metrics['A11'].fill = header_fill
        ws_metrics['B11'].font = header_font
        ws_metrics['B11'].fill = header_fill
        
        row = 12
        for level, count in doc_data["metrics"]["counts"].items():
            ws_metrics[f'A{row}'] = f"H{level[-1]}"
            ws_metrics[f'B{row}'] = count
            row += 1
        
        # Gráfica de secciones
        chart2 = BarChart()
        chart2.type = "col"
        chart2.style = 10
        chart2.title = "Distribución de Secciones"
        chart2.y_axis.title = "Cantidad"
        chart2.x_axis.title = "Nivel"
        
        data = Reference(ws_metrics, min_col=2, min_row=11, max_row=row-1)
        cats = Reference(ws_metrics, min_col=1, min_row=12, max_row=row-1)
        chart2.add_data(data, titles_from_data=False)
        chart2.set_categories(cats)
        chart2.height = 10
        chart2.width = 15
        ws_metrics.add_chart(chart2, "D10")

        # Palabras clave
        ws_metrics['A18'] = "Palabras Clave"
        ws_metrics['A18'].font = header_font
        ws_metrics['A18'].fill = header_fill
        ws_metrics['A19'] = "Palabra"
        ws_metrics['B19'] = "Frecuencia"
        ws_metrics['A19'].font = header_font
        ws_metrics['A19'].fill = header_fill
        ws_metrics['B19'].font = header_font
        ws_metrics['B19'].fill = header_fill
        row = 20
        for keyword in doc_data["keywords"]:
            ws_metrics[f'A{row}'] = keyword["word"]
            ws_metrics[f'B{row}'] = keyword["count"]
            row += 1
        
        # Hoja de Contenido
        ws_content = wb.create_sheet("Contenido")
        ws_content['A1'] = "Sección"
        ws_content['B1'] = "Contenido (Resumen)"
        ws_content['A1'].font = header_font
        ws_content['A1'].fill = header_fill
        ws_content['B1'].font = header_font
        ws_content['B1'].fill = header_fill
        
        row = 2
        for section in doc_data["sections"][:20]:  # Limitar a 20 secciones
            ws_content[f'A{row}'] = section["title"].replace('#', '').strip()[:50]
            content_preview = section["content"][:200].replace('\n', ' ')
            ws_content[f'B{row}'] = content_preview
            row += 1

        # Hoja de Insights
        ws_insights = wb.create_sheet("Insights")
        ws_insights['A1'] = "Resumen Ejecutivo"
        ws_insights['A1'].font = header_font
        ws_insights['A1'].fill = header_fill
        row = 2
        for point in doc_data["summary_points"]:
            ws_insights[f'A{row}'] = f"• {point}"
            row += 1
        row += 1
        ws_insights[f'A{row}'] = "KPIs Clave"
        ws_insights[f'A{row}'].font = header_font
        ws_insights[f'A{row}'].fill = header_fill
        ws_insights[f'A{row+1}'] = f"Palabras totales: {doc_data['metrics']['word_count']}"
        ws_insights[f'A{row+2}'] = f"Lectura estimada (min): {doc_data['metrics']['estimated_minutes']}"
        ws_insights[f'A{row+3}'] = ("Secciones H1/H2/H3: "
                                    f"{doc_data['metrics']['counts']['h1']}/"
                                    f"{doc_data['metrics']['counts']['h2']}/"
                                    f"{doc_data['metrics']['counts']['h3']}")
        row = row + 5
        if doc_data["timeline_dates"]:
            ws_insights[f'A{row}'] = "Cronología"
            ws_insights[f'A{row}'].font = header_font
            ws_insights[f'A{row}'].fill = header_fill
            row += 1
            for date in doc_data["timeline_dates"]:
                ws_insights[f'A{row}'] = f"• {date}"
                row += 1

        # Hoja de Acciones
        if doc_data["action_items"]:
            ws_actions = wb.create_sheet("Acciones")
            ws_actions['A1'] = "Acciones Prioritarias"
            ws_actions['A1'].font = header_font
            ws_actions['A1'].fill = header_fill
            ws_actions['A2'] = "Descripción"
            ws_actions['B2'] = "Estado"
            ws_actions['A2'].font = header_font
            ws_actions['A2'].fill = header_fill
            ws_actions['B2'].font = header_font
            ws_actions['B2'].fill = header_fill

            row = 3
            for action in doc_data["action_items"]:
                ws_actions[f'A{row}'] = action["descripcion"]
                ws_actions[f'B{row}'] = action["estado"].title()
                row += 1

            ws_actions.column_dimensions['A'].width = 80
            ws_actions.column_dimensions['B'].width = 20
        
        # Ajustar ancho de columnas
        ws_metrics.column_dimensions['A'].width = 25
        ws_metrics.column_dimensions['B'].width = 15
        ws_content.column_dimensions['A'].width = 30
        ws_content.column_dimensions['B'].width = 80
        
        wb.save(str(output_path))
        print(f"✅ Excel generado: {output_path}")
    
    def convert_document(self, doc_info: Dict[str, Any]):
        """Convierte un documento a los tres formatos."""
        print(f"\n📄 Procesando: {doc_info['name']}")
        
        try:
            # Parsear documento
            doc_data = self.parse_markdown(doc_info["path"])
            
            base_name = doc_info["name"]
            doc_data["metrics"]["keywords_graph_data"] = doc_data["keywords"]
            
            # Generar PDF
            if PDF_AVAILABLE:
                pdf_path = self.output_dir / f"{base_name}.pdf"
                self.generate_pdf(doc_data, pdf_path)
            
            # Generar Word
            if WORD_AVAILABLE:
                word_path = self.output_dir / f"{base_name}.docx"
                self.generate_word(doc_data, word_path)
            
            # Generar Excel
            if EXCEL_AVAILABLE:
                excel_path = self.output_dir / f"{base_name}.xlsx"
                self.generate_excel(doc_data, excel_path)
            
            print(f"✅ Completado: {base_name}\n")
            
        except Exception as e:
            print(f"❌ Error procesando {doc_info['name']}: {str(e)}")
            import traceback
            traceback.print_exc()
    
    def convert_all(self, base_dir: str):
        """Convierte todos los documentos importantes."""
        print("🔍 Identificando documentos importantes...")
        important_docs = self.identify_important_docs(base_dir)
        
        print(f"📚 Encontrados {len(important_docs)} documentos importantes")
        print("\n" + "="*60)
        
        for doc_info in important_docs:
            self.convert_document(doc_info)
        
        print("="*60)
        print(f"\n✨ Conversión completada!")
        print(f"📁 Archivos guardados en: {self.output_dir.absolute()}")


def main():
    """Función principal."""
    import sys
    
    base_dir = sys.argv[1] if len(sys.argv) > 1 else "."
    output_dir = sys.argv[2] if len(sys.argv) > 2 else "docs_premium"
    
    converter = DocumentConverter(output_dir)
    converter.convert_all(base_dir)


if __name__ == "__main__":
    main()

