#!/usr/bin/env python3
"""
Script mejorado para convertir documentos importantes a PDF, Word y Excel con gráficas profesionales
"""

import os
import sys
import re
from pathlib import Path
import markdown
from datetime import datetime
import json
from collections import Counter
import numpy as np

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Instalando python-docx...")
    os.system("pip install python-docx")
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.chart import BarChart, PieChart, LineChart, Reference, ScatterChart
    from openpyxl.chart.label import DataLabelList
except ImportError:
    print("Instalando openpyxl...")
    os.system("pip install openpyxl")
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.chart import BarChart, PieChart, LineChart, Reference

try:
    import matplotlib.pyplot as plt
    import matplotlib
    matplotlib.use('Agg')
    import seaborn as sns
    sns.set_style("whitegrid")
    plt.rcParams['figure.figsize'] = (12, 8)
    plt.rcParams['font.size'] = 10
except ImportError:
    print("Instalando matplotlib y seaborn...")
    os.system("pip install matplotlib seaborn")
    import matplotlib.pyplot as plt
    import matplotlib
    matplotlib.use('Agg')
    import seaborn as sns

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak, KeepTogether
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
    from reportlab.pdfgen import canvas
except ImportError:
    print("Instalando reportlab...")
    os.system("pip install reportlab")
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

class AdvancedDocumentConverter:
    def __init__(self, output_dir="converted_docs"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        self.temp_dir = self.output_dir / "temp"
        self.temp_dir.mkdir(exist_ok=True)
        
    def read_markdown(self, file_path):
        """Lee un archivo markdown"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def parse_markdown_sections(self, content):
        """Parsea el markdown en secciones con más detalle"""
        sections = []
        current_section = {
            "title": "", 
            "content": "", 
            "level": 0, 
            "subsections": [],
            "code_blocks": 0,
            "links": 0,
            "words": 0
        }
        
        lines = content.split('\n')
        in_code_block = False
        code_language = ""
        
        for line in lines:
            # Detectar bloques de código
            if line.strip().startswith('```'):
                in_code_block = not in_code_block
                if in_code_block:
                    code_language = line.strip()[3:].strip()
                continue
            
            if in_code_block:
                current_section["content"] += line + "\n"
                continue
            
            if line.startswith('#'):
                if current_section["content"] or current_section["title"]:
                    sections.append(current_section)
                level = len(line) - len(line.lstrip('#'))
                title = line.lstrip('#').strip()
                current_section = {
                    "title": title, 
                    "content": "", 
                    "level": level,
                    "subsections": [],
                    "code_blocks": 0,
                    "links": 0,
                    "words": 0
                }
            else:
                current_section["content"] += line + "\n"
                # Contar elementos
                if '`' in line:
                    current_section["code_blocks"] += line.count('`') // 2
                if '[' in line and ']' in line:
                    current_section["links"] += min(line.count('['), line.count(']'))
                current_section["words"] += len(line.split())
        
        if current_section["content"] or current_section["title"]:
            sections.append(current_section)
        
        return sections
    
    def create_advanced_statistics(self, content, sections):
        """Crea estadísticas avanzadas del documento"""
        lines = content.split('\n')
        code_blocks = content.count('```')
        links = len(re.findall(r'\[([^\]]+)\]\(([^\)]+)\)', content))
        images = len(re.findall(r'!\[([^\]]*)\]\(([^\)]+)\)', content))
        headers = len([l for l in lines if l.strip().startswith('#')])
        
        # Análisis de palabras clave
        words = re.findall(r'\b[a-zA-Z]{4,}\b', content.lower())
        word_freq = Counter(words)
        top_words = dict(word_freq.most_common(20))
        
        # Análisis de secciones
        section_stats = {
            "total": len(sections),
            "by_level": Counter([s["level"] for s in sections]),
            "avg_words_per_section": sum(s["words"] for s in sections) / max(len(sections), 1),
            "sections_with_code": sum(1 for s in sections if s["code_blocks"] > 0),
            "sections_with_links": sum(1 for s in sections if s["links"] > 0),
        }
        
        stats = {
            "total_lines": len(lines),
            "total_words": len(content.split()),
            "total_chars": len(content),
            "total_chars_no_spaces": len(content.replace(' ', '')),
            "sections": headers,
            "code_blocks": code_blocks // 2,
            "links": links,
            "images": images,
            "top_keywords": top_words,
            "section_analysis": section_stats,
            "readability_score": self.calculate_readability(content),
            "complexity_score": self.calculate_complexity(content, sections)
        }
        return stats
    
    def calculate_readability(self, content):
        """Calcula un score de legibilidad"""
        sentences = re.split(r'[.!?]+', content)
        words = content.split()
        avg_sentence_length = len(words) / max(len(sentences), 1)
        avg_word_length = sum(len(w) for w in words) / max(len(words), 1)
        
        # Score simple (0-100, más alto = más legible)
        score = 100 - (avg_sentence_length * 2) - (avg_word_length * 3)
        return max(0, min(100, int(score)))
    
    def calculate_complexity(self, content, sections):
        """Calcula un score de complejidad"""
        code_ratio = content.count('`') / max(len(content), 1) * 100
        section_depth = max([s["level"] for s in sections] + [0])
        avg_section_size = sum(len(s["content"]) for s in sections) / max(len(sections), 1)
        
        complexity = (code_ratio * 0.3) + (section_depth * 5) + (avg_section_size / 100)
        return min(100, int(complexity))
    
    def create_professional_charts(self, stats, sections, doc_name):
        """Crea gráficas profesionales con múltiples visualizaciones"""
        charts = []
        
        # 1. Gráfica de barras mejorada - Estadísticas principales
        fig, ax = plt.subplots(figsize=(12, 7))
        categories = ['Líneas\n(x100)', 'Palabras\n(x100)', 'Secciones', 'Bloques\nCódigo', 'Enlaces']
        values = [
            stats['total_lines'] / 100,
            stats['total_words'] / 100,
            stats['sections'],
            stats['code_blocks'],
            stats['links']
        ]
        colors_bar = ['#3498db', '#2ecc71', '#e74c3c', '#f39c12', '#9b59b6']
        bars = ax.bar(categories, values, color=colors_bar, edgecolor='black', linewidth=1.5)
        ax.set_title(f'Análisis Estadístico: {doc_name}', fontsize=16, fontweight='bold', pad=20)
        ax.set_ylabel('Valor Normalizado', fontsize=12, fontweight='bold')
        ax.grid(axis='y', alpha=0.3, linestyle='--')
        
        # Agregar valores en las barras
        for bar, val in zip(bars, values):
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2., height,
                   f'{val:.1f}', ha='center', va='bottom', fontweight='bold')
        
        plt.tight_layout()
        chart_path = self.temp_dir / f"{doc_name}_stats_bar.png"
        plt.savefig(chart_path, dpi=300, bbox_inches='tight', facecolor='white')
        plt.close()
        charts.append(("Estadísticas Principales", chart_path))
        
        # 2. Gráfica de pie mejorada - Distribución de contenido
        fig, ax = plt.subplots(figsize=(10, 10))
        labels = ['Texto', 'Código', 'Enlaces', 'Imágenes', 'Estructura']
        sizes = [
            stats['total_words'] * 0.5,
            stats['code_blocks'] * 100,
            stats['links'] * 50,
            stats['images'] * 200,
            stats['sections'] * 30
        ]
        colors_pie = ['#3498db', '#2ecc71', '#e74c3c', '#f39c12', '#9b59b6']
        explode = (0.05, 0.1, 0.05, 0.05, 0.05)
        
        wedges, texts, autotexts = ax.pie(sizes, labels=labels, autopct='%1.1f%%', 
                                          colors=colors_pie, startangle=90, explode=explode,
                                          shadow=True, textprops={'fontsize': 12, 'fontweight': 'bold'})
        ax.set_title('Distribución de Contenido', fontsize=16, fontweight='bold', pad=20)
        plt.tight_layout()
        chart_path = self.temp_dir / f"{doc_name}_content_pie.png"
        plt.savefig(chart_path, dpi=300, bbox_inches='tight', facecolor='white')
        plt.close()
        charts.append(("Distribución de Contenido", chart_path))
        
        # 3. Gráfica de palabras clave más frecuentes
        if stats['top_keywords']:
            fig, ax = plt.subplots(figsize=(12, 8))
            top_10 = dict(list(stats['top_keywords'].items())[:10])
            words = list(top_10.keys())
            counts = list(top_10.values())
            
            bars = ax.barh(words, counts, color='#3498db', edgecolor='black', linewidth=1)
            ax.set_xlabel('Frecuencia', fontsize=12, fontweight='bold')
            ax.set_title('Top 10 Palabras Clave', fontsize=16, fontweight='bold', pad=20)
            ax.grid(axis='x', alpha=0.3, linestyle='--')
            
            # Agregar valores
            for i, (bar, count) in enumerate(zip(bars, counts)):
                ax.text(count, bar.get_y() + bar.get_height()/2,
                       f' {count}', va='center', fontweight='bold')
            
            plt.tight_layout()
            chart_path = self.temp_dir / f"{doc_name}_keywords.png"
            plt.savefig(chart_path, dpi=300, bbox_inches='tight', facecolor='white')
            plt.close()
            charts.append(("Palabras Clave", chart_path))
        
        # 4. Gráfica de análisis de secciones por nivel
        if sections:
            fig, ax = plt.subplots(figsize=(10, 6))
            section_levels = [s["level"] for s in sections if s["level"] > 0]
            if section_levels:
                level_counts = Counter(section_levels)
                levels = sorted(level_counts.keys())
                counts = [level_counts[l] for l in levels]
                
                bars = ax.bar([f'Nivel {l}' for l in levels], counts, 
                            color='#2ecc71', edgecolor='black', linewidth=1.5)
                ax.set_ylabel('Cantidad de Secciones', fontsize=12, fontweight='bold')
                ax.set_title('Distribución de Secciones por Nivel', fontsize=16, fontweight='bold', pad=20)
                ax.grid(axis='y', alpha=0.3, linestyle='--')
                
                for bar, count in zip(bars, counts):
                    height = bar.get_height()
                    ax.text(bar.get_x() + bar.get_width()/2., height,
                           f'{count}', ha='center', va='bottom', fontweight='bold')
                
                plt.tight_layout()
                chart_path = self.temp_dir / f"{doc_name}_sections.png"
                plt.savefig(chart_path, dpi=300, bbox_inches='tight', facecolor='white')
                plt.close()
                charts.append(("Análisis de Secciones", chart_path))
        
        # 5. Gráfica de métricas de calidad
        fig, ax = plt.subplots(figsize=(10, 6))
        metrics = ['Legibilidad', 'Complejidad']
        scores = [stats['readability_score'], stats['complexity_score']]
        colors_metrics = ['#2ecc71' if s > 50 else '#e74c3c' for s in scores]
        
        bars = ax.bar(metrics, scores, color=colors_metrics, edgecolor='black', linewidth=2)
        ax.set_ylabel('Score (0-100)', fontsize=12, fontweight='bold')
        ax.set_title('Métricas de Calidad del Documento', fontsize=16, fontweight='bold', pad=20)
        ax.set_ylim(0, 100)
        ax.grid(axis='y', alpha=0.3, linestyle='--')
        
        for bar, score in zip(bars, scores):
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2., height,
                   f'{score}', ha='center', va='bottom', fontweight='bold', fontsize=14)
        
        plt.tight_layout()
        chart_path = self.temp_dir / f"{doc_name}_quality.png"
        plt.savefig(chart_path, dpi=300, bbox_inches='tight', facecolor='white')
        plt.close()
        charts.append(("Métricas de Calidad", chart_path))
        
        return charts
    
    def convert_to_word_improved(self, file_path, output_name):
        """Convierte markdown a Word con formato profesional mejorado"""
        print(f"Convirtiendo {file_path} a Word (versión mejorada)...")
        
        content = self.read_markdown(file_path)
        sections = self.parse_markdown_sections(content)
        stats = self.create_advanced_statistics(content, sections)
        charts = self.create_professional_charts(stats, sections, output_name)
        
        doc = Document()
        
        # Configurar estilos personalizados
        styles = doc.styles
        
        # Título principal
        title = doc.add_heading(output_name.replace('_', ' ').title(), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(28)
        title_run.font.color.rgb = RGBColor(54, 96, 146)
        
        # Información del documento
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_run1 = info_para.add_run(f"Generado el: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        info_run1.font.size = Pt(11)
        info_run1.font.color.rgb = RGBColor(100, 100, 100)
        info_run2 = info_para.add_run(f"Archivo original: {Path(file_path).name}")
        info_run2.font.size = Pt(11)
        info_run2.font.color.rgb = RGBColor(100, 100, 100)
        
        doc.add_page_break()
        
        # Tabla de contenido
        doc.add_heading('Tabla de Contenido', level=1)
        toc_para = doc.add_paragraph()
        toc_para.add_run("1. Resumen Ejecutivo\n")
        toc_para.add_run("2. Estadísticas del Documento\n")
        toc_para.add_run("3. Análisis de Contenido\n")
        toc_para.add_run("4. Contenido Completo\n")
        doc.add_page_break()
        
        # Resumen ejecutivo
        doc.add_heading('Resumen Ejecutivo', level=1)
        summary_para = doc.add_paragraph()
        summary_para.add_run(f"Este documento contiene {stats['sections']} secciones principales, ")
        summary_para.add_run(f"{stats['total_words']:,} palabras, y ")
        summary_para.add_run(f"{stats['code_blocks']} bloques de código. ")
        summary_para.add_run(f"El score de legibilidad es {stats['readability_score']}/100 y ")
        summary_para.add_run(f"el score de complejidad es {stats['complexity_score']}/100.")
        
        doc.add_heading('Estadísticas del Documento', level=1)
        
        # Agregar todas las gráficas
        for chart_title, chart_path in charts:
            doc.add_heading(chart_title, level=2)
            doc.add_picture(str(chart_path), width=Inches(6.5))
            doc.add_paragraph()  # Espacio
        
        # Tabla de estadísticas detalladas
        doc.add_heading('Estadísticas Detalladas', level=2)
        stats_table = doc.add_table(rows=1, cols=2)
        stats_table.style = 'Light Grid Accent 1'
        hdr_cells = stats_table.rows[0].cells
        hdr_cells[0].text = 'Métrica'
        hdr_cells[1].text = 'Valor'
        hdr_cells[0].paragraphs[0].runs[0].font.bold = True
        hdr_cells[1].paragraphs[0].runs[0].font.bold = True
        
        for key, value in stats.items():
            if key not in ['top_keywords', 'section_analysis']:
                row_cells = stats_table.add_row().cells
                row_cells[0].text = key.replace('_', ' ').title()
                if isinstance(value, (int, float)):
                    row_cells[1].text = f"{value:,}" if isinstance(value, int) else f"{value:.2f}"
                else:
                    row_cells[1].text = str(value)
        
        doc.add_page_break()
        
        # Análisis de palabras clave
        if stats['top_keywords']:
            doc.add_heading('Top Palabras Clave', level=2)
            keywords_table = doc.add_table(rows=1, cols=2)
            keywords_table.style = 'Light List Accent 1'
            hdr_cells = keywords_table.rows[0].cells
            hdr_cells[0].text = 'Palabra'
            hdr_cells[1].text = 'Frecuencia'
            hdr_cells[0].paragraphs[0].runs[0].font.bold = True
            hdr_cells[1].paragraphs[0].runs[0].font.bold = True
            
            for word, count in list(stats['top_keywords'].items())[:15]:
                row_cells = keywords_table.add_row().cells
                row_cells[0].text = word
                row_cells[1].text = str(count)
        
        doc.add_page_break()
        
        # Contenido completo
        doc.add_heading('Contenido Completo del Documento', level=1)
        
        for i, section in enumerate(sections, 1):
            if section['title']:
                level = min(section['level'], 9)
                doc.add_heading(section['title'], level=level)
            
            # Procesar contenido
            content_lines = section['content'].split('\n')
            for line in content_lines:
                if line.strip():
                    if line.startswith('- ') or line.startswith('* '):
                        doc.add_paragraph(line[2:], style='List Bullet')
                    elif line.startswith('```'):
                        continue
                    elif '`' in line:
                        para = doc.add_paragraph()
                        parts = line.split('`')
                        for j, part in enumerate(parts):
                            if j % 2 == 0:
                                para.add_run(part)
                            else:
                                run = para.add_run(part)
                                run.font.name = 'Courier New'
                                run.font.size = Pt(9)
                                run.font.color.rgb = RGBColor(200, 0, 0)
                    else:
                        doc.add_paragraph(line)
            
            doc.add_paragraph()  # Espacio entre secciones
        
        # Guardar
        output_path = self.output_dir / f"{output_name}_IMPROVED.docx"
        doc.save(str(output_path))
        print(f"✓ Word mejorado guardado: {output_path}")
        return output_path
    
    def convert_to_excel_improved(self, file_path, output_name):
        """Convierte markdown a Excel con análisis avanzado"""
        print(f"Convirtiendo {file_path} a Excel (versión mejorada)...")
        
        content = self.read_markdown(file_path)
        sections = self.parse_markdown_sections(content)
        stats = self.create_advanced_statistics(content, sections)
        
        wb = Workbook()
        
        # Hoja 1: Resumen Ejecutivo
        ws = wb.active
        ws.title = "Resumen Ejecutivo"
        
        # Estilos
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF", size=14)
        title_font = Font(bold=True, size=18, color="366092")
        subheader_font = Font(bold=True, size=12)
        
        # Título
        ws['A1'] = output_name.replace('_', ' ').title()
        ws['A1'].font = title_font
        ws.merge_cells('A1:D1')
        ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
        ws.row_dimensions[1].height = 30
        
        # Información
        ws['A3'] = f"Generado el: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        ws['A4'] = f"Archivo original: {Path(file_path).name}"
        
        # Resumen
        row = 6
        ws[f'A{row}'] = "Resumen Ejecutivo"
        ws[f'A{row}'].font = subheader_font
        ws[f'A{row}'].fill = header_fill
        ws[f'A{row}'].font = Font(bold=True, color="FFFFFF", size=12)
        ws.merge_cells(f'A{row}:D{row}')
        
        row += 1
        ws[f'A{row}'] = f"Total de secciones: {stats['sections']}"
        ws[f'B{row}'] = f"Total de palabras: {stats['total_words']:,}"
        row += 1
        ws[f'A{row}'] = f"Score de legibilidad: {stats['readability_score']}/100"
        ws[f'B{row}'] = f"Score de complejidad: {stats['complexity_score']}/100"
        
        # Estadísticas detalladas
        row += 2
        ws[f'A{row}'] = "Estadísticas Detalladas"
        ws[f'A{row}'].font = subheader_font
        ws[f'A{row}'].fill = header_fill
        ws[f'A{row}'].font = Font(bold=True, color="FFFFFF", size=12)
        ws.merge_cells(f'A{row}:B{row}')
        
        row += 1
        ws[f'A{row}'] = "Métrica"
        ws[f'B{row}'] = "Valor"
        ws[f'A{row}'].font = Font(bold=True)
        ws[f'B{row}'].font = Font(bold=True)
        
        for key, value in stats.items():
            if key not in ['top_keywords', 'section_analysis']:
                row += 1
                ws[f'A{row}'] = key.replace('_', ' ').title()
                if isinstance(value, (int, float)):
                    ws[f'B{row}'] = f"{value:,}" if isinstance(value, int) else f"{value:.2f}"
                else:
                    ws[f'B{row}'] = str(value)
        
        # Gráfica
        chart = BarChart()
        chart.type = "col"
        chart.style = 10
        chart.title = "Estadísticas Principales"
        chart.y_axis.title = "Valor"
        chart.x_axis.title = "Métrica"
        
        data = Reference(ws, min_col=2, min_row=row - len([k for k in stats.keys() if k not in ['top_keywords', 'section_analysis']]) + 1, 
                        max_row=row)
        cats = Reference(ws, min_col=1, min_row=row - len([k for k in stats.keys() if k not in ['top_keywords', 'section_analysis']]) + 2, 
                        max_row=row)
        chart.add_data(data, titles_from_data=False)
        chart.set_categories(cats)
        chart.height = 12
        chart.width = 18
        ws.add_chart(chart, "D6")
        
        # Hoja 2: Palabras Clave
        ws2 = wb.create_sheet("Palabras Clave")
        ws2['A1'] = "Palabra"
        ws2['B1'] = "Frecuencia"
        ws2['A1'].font = header_font
        ws2['B1'].font = header_font
        ws2['A1'].fill = header_fill
        ws2['B1'].fill = header_fill
        
        row = 2
        for word, count in list(stats['top_keywords'].items())[:50]:
            ws2[f'A{row}'] = word
            ws2[f'B{row}'] = count
            row += 1
        
        # Gráfica de palabras clave
        chart2 = BarChart()
        chart2.type = "bar"
        chart2.style = 10
        chart2.title = "Top 10 Palabras Clave"
        chart2.y_axis.title = "Palabra"
        chart2.x_axis.title = "Frecuencia"
        
        data2 = Reference(ws2, min_col=2, min_row=1, max_row=11)
        cats2 = Reference(ws2, min_col=1, min_row=2, max_row=11)
        chart2.add_data(data2, titles_from_data=False)
        chart2.set_categories(cats2)
        chart2.height = 10
        chart2.width = 15
        ws2.add_chart(chart2, "D2")
        
        # Hoja 3: Análisis de Secciones
        ws3 = wb.create_sheet("Análisis Secciones")
        ws3['A1'] = "Sección"
        ws3['B1'] = "Nivel"
        ws3['C1'] = "Palabras"
        ws3['D1'] = "Bloques Código"
        ws3['E1'] = "Enlaces"
        
        for col in ['A', 'B', 'C', 'D', 'E']:
            ws3[f'{col}1'].font = header_font
            ws3[f'{col}1'].fill = header_fill
        
        row = 2
        for section in sections[:100]:  # Limitar a 100 secciones
            if section['title']:
                ws3[f'A{row}'] = section['title'][:50]
                ws3[f'B{row}'] = section['level']
                ws3[f'C{row}'] = section['words']
                ws3[f'D{row}'] = section['code_blocks']
                ws3[f'E{row}'] = section['links']
                row += 1
        
        # Hoja 4: Contenido
        ws4 = wb.create_sheet("Contenido")
        ws4['A1'] = "Sección"
        ws4['B1'] = "Contenido"
        ws4['A1'].font = header_font
        ws4['B1'].font = header_font
        ws4['A1'].fill = header_fill
        ws4['B1'].fill = header_fill
        
        row = 2
        for section in sections[:200]:  # Limitar contenido
            if section['title']:
                ws4[f'A{row}'] = section['title']
                content_preview = section['content'][:500].replace('\n', ' ')
                ws4[f'B{row}'] = content_preview
                row += 1
                if row > 500:
                    break
        
        # Ajustar ancho de columnas
        for ws in wb.worksheets:
            ws.column_dimensions['A'].width = 30
            ws.column_dimensions['B'].width = 40
            if len(ws.column_dimensions) > 2:
                ws.column_dimensions['C'].width = 15
                ws.column_dimensions['D'].width = 15
                ws.column_dimensions['E'].width = 15
        
        # Guardar
        output_path = self.output_dir / f"{output_name}_IMPROVED.xlsx"
        wb.save(str(output_path))
        print(f"✓ Excel mejorado guardado: {output_path}")
        return output_path
    
    def convert_to_pdf_improved(self, file_path, output_name):
        """Convierte markdown a PDF con formato profesional mejorado"""
        print(f"Convirtiendo {file_path} a PDF (versión mejorada)...")
        
        content = self.read_markdown(file_path)
        sections = self.parse_markdown_sections(content)
        stats = self.create_advanced_statistics(content, sections)
        charts = self.create_professional_charts(stats, sections, output_name)
        
        output_path = self.output_dir / f"{output_name}_IMPROVED.pdf"
        doc = SimpleDocTemplate(str(output_path), pagesize=A4,
                              rightMargin=72, leftMargin=72,
                              topMargin=72, bottomMargin=18)
        
        styles = getSampleStyleSheet()
        
        # Estilos personalizados
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1a1a1a'),
            spaceAfter=30,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        story = []
        
        # Portada
        title = Paragraph(output_name.replace('_', ' ').title(), title_style)
        story.append(title)
        story.append(Spacer(1, 20))
        
        info_style = ParagraphStyle(
            'Info',
            parent=styles['Normal'],
            fontSize=11,
            textColor=colors.HexColor('#666666'),
            alignment=TA_CENTER
        )
        story.append(Paragraph(f"Generado el: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", info_style))
        story.append(Paragraph(f"Archivo original: {Path(file_path).name}", info_style))
        story.append(PageBreak())
        
        # Resumen ejecutivo
        story.append(Paragraph("Resumen Ejecutivo", styles['Heading1']))
        story.append(Spacer(1, 12))
        
        summary_text = (f"Este documento contiene {stats['sections']} secciones principales, "
                       f"{stats['total_words']:,} palabras, y {stats['code_blocks']} bloques de código. "
                       f"El score de legibilidad es {stats['readability_score']}/100 y "
                       f"el score de complejidad es {stats['complexity_score']}/100.")
        story.append(Paragraph(summary_text, styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Agregar gráficas
        story.append(Paragraph("Análisis Visual del Documento", styles['Heading1']))
        story.append(Spacer(1, 12))
        
        for chart_title, chart_path in charts:
            story.append(Paragraph(chart_title, styles['Heading2']))
            story.append(Spacer(1, 6))
            img = Image(str(chart_path), width=6*inch, height=4*inch)
            story.append(img)
            story.append(Spacer(1, 12))
        
        story.append(PageBreak())
        
        # Estadísticas detalladas
        story.append(Paragraph("Estadísticas Detalladas", styles['Heading1']))
        story.append(Spacer(1, 12))
        
        stats_data = [['Métrica', 'Valor']]
        for key, value in stats.items():
            if key not in ['top_keywords', 'section_analysis']:
                if isinstance(value, (int, float)):
                    display_value = f"{value:,}" if isinstance(value, int) else f"{value:.2f}"
                else:
                    display_value = str(value)
                stats_data.append([key.replace('_', ' ').title(), display_value])
        
        stats_table = Table(stats_data, colWidths=[3.5*inch, 2.5*inch])
        stats_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#366092')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f8f9fa')),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#dee2e6')),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8f9fa')]),
        ]))
        story.append(stats_table)
        story.append(PageBreak())
        
        # Palabras clave
        if stats['top_keywords']:
            story.append(Paragraph("Top Palabras Clave", styles['Heading1']))
            story.append(Spacer(1, 12))
            
            keywords_data = [['Palabra', 'Frecuencia']]
            for word, count in list(stats['top_keywords'].items())[:20]:
                keywords_data.append([word, str(count)])
            
            keywords_table = Table(keywords_data, colWidths=[3*inch, 2*inch])
            keywords_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#366092')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f8f9fa')),
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#dee2e6')),
            ]))
            story.append(keywords_table)
            story.append(PageBreak())
        
        # Contenido
        story.append(Paragraph("Contenido Completo del Documento", styles['Heading1']))
        story.append(Spacer(1, 12))
        
        for section in sections:
            if section['title']:
                level = min(section['level'], 6)
                heading_style = styles[f'Heading{level}']
                story.append(Paragraph(section['title'], heading_style))
                story.append(Spacer(1, 6))
            
            content_lines = section['content'].split('\n')
            for line in content_lines:
                if line.strip() and not line.startswith('```'):
                    # Limpiar y escapar texto
                    line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    line = ''.join(char for char in line if ord(char) < 128 or char.isprintable())
                    if len(line) > 120:
                        # Dividir líneas largas
                        words = line.split()
                        current_line = ""
                        for word in words:
                            if len(current_line + word) < 120:
                                current_line += word + " "
                            else:
                                if current_line:
                                    story.append(Paragraph(current_line.strip(), styles['Normal']))
                                current_line = word + " "
                        if current_line:
                            story.append(Paragraph(current_line.strip(), styles['Normal']))
                    else:
                        try:
                            story.append(Paragraph(line, styles['Normal']))
                        except:
                            pass
                    story.append(Spacer(1, 2))
            
            story.append(Spacer(1, 8))
        
        doc.build(story)
        print(f"✓ PDF mejorado guardado: {output_path}")
        return output_path

def main():
    """Función principal"""
    base_dir = Path("/Users/adan/Documents/documentos_blatam")
    production_dir = base_dir / "truthgpt_collected/integration_code/production_code"
    
    documents = [
        {
            "path": base_dir / "airflow_automation_prompt.md",
            "name": "Automation_Expert_Prompt"
        },
        {
            "path": production_dir / "ARCHITECTURE_IMPROVEMENTS.md",
            "name": "Architecture_Improvements"
        },
        {
            "path": production_dir / "REFACTORING_PLAN.md",
            "name": "Refactoring_Plan"
        }
    ]
    
    converter = AdvancedDocumentConverter()
    
    print("=" * 70)
    print("CONVERSIÓN MEJORADA DE DOCUMENTOS A PDF, WORD Y EXCEL")
    print("=" * 70)
    print()
    
    results = []
    
    for doc in documents:
        if not doc["path"].exists():
            print(f"⚠ Archivo no encontrado: {doc['path']}")
            continue
        
        print(f"\n📄 Procesando: {doc['name']}")
        print("-" * 70)
        
        try:
            word_path = converter.convert_to_word_improved(doc["path"], doc["name"])
            excel_path = converter.convert_to_excel_improved(doc["path"], doc["name"])
            pdf_path = converter.convert_to_pdf_improved(doc["path"], doc["name"])
            
            results.append({
                "document": doc["name"],
                "word": str(word_path),
                "excel": str(excel_path),
                "pdf": str(pdf_path)
            })
            
        except Exception as e:
            print(f"❌ Error procesando {doc['name']}: {str(e)}")
            import traceback
            traceback.print_exc()
    
    print("\n" + "=" * 70)
    print("RESUMEN DE CONVERSIÓN MEJORADA")
    print("=" * 70)
    print(f"\nDocumentos procesados: {len(results)}")
    print(f"Directorio de salida: {converter.output_dir}")
    print("\nArchivos generados (versión mejorada):")
    for result in results:
        print(f"\n  📄 {result['document']}:")
        print(f"     - Word: {result['word']}")
        print(f"     - Excel: {result['excel']}")
        print(f"     - PDF: {result['pdf']}")
    
    print("\n✅ Conversión mejorada completada!")

if __name__ == "__main__":
    main()

