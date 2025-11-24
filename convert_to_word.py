#!/usr/bin/env python3
"""
Script para convertir SISTEMAS_PROMPTS_CONSOLIDADO.md a formato Word (.docx)
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, url, text):
    """Agrega un hipervínculo a un párrafo"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)
    
    new_run.append(rPr)
    new_run.text = text
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    
    return hyperlink

def parse_markdown_to_docx(md_file, docx_file):
    """Convierte un archivo markdown a formato Word"""
    
    # Leer el archivo markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Crear documento Word
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_block_lines = []
    
    while i < len(lines):
        line = lines[i]
        original_line = line
        line_stripped = line.strip()
        
        # Bloque de código
        if line_stripped.startswith('```'):
            if in_code_block:
                # Fin del bloque de código
                if code_block_lines:
                    p = doc.add_paragraph()
                    p.style = 'Intense Quote'
                    run = p.add_run('\n'.join(code_block_lines))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0, 100, 0)  # Verde oscuro
                code_block_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue
        
        # Título nivel 1 (#)
        if line_stripped.startswith('# ') and not line_stripped.startswith('##'):
            text = line_stripped[2:].strip()
            heading = doc.add_heading(text, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Título nivel 2 (##)
        elif line_stripped.startswith('## ') and not line_stripped.startswith('###'):
            text = line_stripped[3:].strip()
            heading = doc.add_heading(text, level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Título nivel 3 (###)
        elif line_stripped.startswith('### '):
            text = line_stripped[4:].strip()
            heading = doc.add_heading(text, level=3)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Línea horizontal (---)
        elif line_stripped == '---' or (line_stripped.startswith('---') and len(line_stripped) <= 5):
            p = doc.add_paragraph()
            p.add_run('_' * 80)
        
        # Lista ordenada
        elif re.match(r'^\d+\.\s+', line_stripped):
            text = re.sub(r'^\d+\.\s+', '', line_stripped)
            p = doc.add_paragraph(text, style='List Number')
            process_inline_formatting(p, text)
        
        # Lista con viñetas (- o *)
        elif line_stripped.startswith('- ') or line_stripped.startswith('* '):
            text = line_stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            process_inline_formatting(p, text)
        
        # Etiquetas XML (goal, format_rules, etc.)
        elif line_stripped.startswith('<') and line_stripped.endswith('>'):
            p = doc.add_paragraph()
            run = p.add_run(line_stripped)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 128)  # Azul oscuro
        
        # Párrafo normal
        elif line_stripped:
            p = doc.add_paragraph()
            process_inline_formatting(p, line_stripped)
        
        # Línea vacía
        else:
            # Solo agregar espacio si el último párrafo no está vacío
            if doc.paragraphs and doc.paragraphs[-1].text.strip():
                doc.add_paragraph()
        
        i += 1
    
    # Guardar documento
    doc.save(docx_file)
    print(f"✅ Documento Word creado exitosamente: {docx_file}")
    print(f"📄 Ubicación: {docx_file}")

def process_inline_formatting(paragraph, text):
    """Procesa formato inline: negrita, cursiva, enlaces, código"""
    
    # Patrones
    link_pattern = r'\[([^\]]+)\]\(([^\)]+)\)'
    bold_pattern = r'\*\*([^\*]+)\*\*'
    italic_pattern = r'(?<!\*)\*([^\*]+)\*(?!\*)'
    code_pattern = r'`([^`]+)`'
    
    # Procesar en orden: enlaces, negrita, cursiva, código
    parts = [text]
    
    # Procesar enlaces
    new_parts = []
    for part in parts:
        if re.search(link_pattern, part):
            link_parts = re.split(link_pattern, part)
            for j in range(0, len(link_parts), 3):
                if j < len(link_parts):
                    if link_parts[j]:
                        new_parts.append(('text', link_parts[j]))
                if j + 1 < len(link_parts):
                    new_parts.append(('link', link_parts[j+1], link_parts[j+2]))
        else:
            new_parts.append(('text', part))
    
    parts = new_parts
    
    # Procesar cada parte
    for part_type, *part_data in parts:
        if part_type == 'link':
            link_text, link_url = part_data
            add_hyperlink(paragraph, link_url, link_text)
        else:
            text = part_data[0]
            # Procesar negrita
            bold_parts = re.split(bold_pattern, text)
            for j in range(0, len(bold_parts), 2):
                if j < len(bold_parts) and bold_parts[j]:
                    # Procesar cursiva en el texto
                    italic_parts = re.split(italic_pattern, bold_parts[j])
                    for k in range(0, len(italic_parts), 2):
                        if k < len(italic_parts) and italic_parts[k]:
                            # Procesar código inline
                            code_parts = re.split(code_pattern, italic_parts[k])
                            for m in range(0, len(code_parts), 2):
                                if m < len(code_parts) and code_parts[m]:
                                    run = paragraph.add_run(code_parts[m])
                                    run.font.name = 'Calibri'
                                    run.font.size = Pt(11)
                                if m + 1 < len(code_parts):
                                    run = paragraph.add_run(code_parts[m+1])
                                    run.font.name = 'Courier New'
                                    run.font.size = Pt(10)
                                    run.font.color.rgb = RGBColor(0, 100, 0)
                        if k + 1 < len(italic_parts):
                            run = paragraph.add_run(italic_parts[k+1])
                            run.italic = True
                            run.font.name = 'Calibri'
                            run.font.size = Pt(11)
                if j + 1 < len(bold_parts):
                    run = paragraph.add_run(bold_parts[j+1])
                    run.bold = True
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)

if __name__ == '__main__':
    input_file = '01_marketing/SISTEMAS_PROMPTS_CONSOLIDADO.md'
    output_file = '01_marketing/SISTEMAS_PROMPTS_CONSOLIDADO.docx'
    
    try:
        parse_markdown_to_docx(input_file, output_file)
    except FileNotFoundError:
        print(f"❌ Error: No se encontró el archivo {input_file}")
    except Exception as e:
        print(f"❌ Error al convertir: {e}")
        import traceback
        traceback.print_exc()
