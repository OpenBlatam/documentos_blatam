#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script para convertir el documento ultra completo final a HTML, PDF y Word
Versión: 2024 - Ultra Completa Final
"""

import markdown
from docx import Document
from docx.shared import Inches
import os
import subprocess
import sys
import platform
from bs4 import BeautifulSoup
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.units import inch
from reportlab.lib.colors import black

def check_dependencies():
    """Verifica que todas las dependencias estén instaladas"""
    print("🔍 Verificando dependencias...")
    
    required_packages = [
        'markdown',
        'python-docx', 
        'beautifulsoup4',
        'reportlab'
    ]
    
    missing_packages = []
    
    for package in required_packages:
        try:
            __import__(package.replace('-', '_'))
        except ImportError:
            missing_packages.append(package)
    
    if missing_packages:
        print(f"❌ Faltan dependencias: {', '.join(missing_packages)}")
        print("📦 Instalando dependencias faltantes...")
        
        for package in missing_packages:
            try:
                subprocess.check_call([sys.executable, '-m', 'pip', 'install', package])
                print(f"✅ {package} instalado correctamente")
            except subprocess.CalledProcessError:
                print(f"❌ Error instalando {package}")
                return False
    else:
        print("✅ Todas las dependencias ya están instaladas")
    
    return True

def markdown_to_html(md_file_path, html_file_path):
    """Convierte Markdown a HTML con estilos mejorados"""
    print("📄 Convirtiendo Markdown a HTML...")
    
    with open(md_file_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Configuración de Markdown con extensiones
    md = markdown.Markdown(extensions=['tables', 'fenced_code', 'toc'])
    html_content = md.convert(md_content)
    
    # HTML completo con estilos
    full_html = f"""
    <!DOCTYPE html>
    <html lang="es">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Frases Populares de Búsqueda Google 2024 - Ultra Completo</title>
        <style>
            body {{
                font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                line-height: 1.6;
                color: #333;
                max-width: 1200px;
                margin: 0 auto;
                padding: 20px;
                background-color: #f8f9fa;
            }}
            .container {{
                background: white;
                padding: 40px;
                border-radius: 10px;
                box-shadow: 0 4px 6px rgba(0,0,0,0.1);
            }}
            h1 {{
                color: #2c3e50;
                border-bottom: 3px solid #3498db;
                padding-bottom: 10px;
                text-align: center;
                font-size: 2.5em;
                margin-bottom: 30px;
            }}
            h2 {{
                color: #34495e;
                border-left: 4px solid #3498db;
                padding-left: 15px;
                margin-top: 40px;
                font-size: 1.8em;
            }}
            h3 {{
                color: #2c3e50;
                margin-top: 30px;
                font-size: 1.4em;
            }}
            h4 {{
                color: #34495e;
                margin-top: 25px;
                font-size: 1.2em;
            }}
            p {{
                margin-bottom: 15px;
                text-align: justify;
            }}
            ul, ol {{
                margin-bottom: 20px;
                padding-left: 30px;
            }}
            li {{
                margin-bottom: 8px;
            }}
            strong {{
                color: #2c3e50;
                font-weight: 600;
            }}
            em {{
                color: #7f8c8d;
                font-style: italic;
            }}
            code {{
                background-color: #f1f2f6;
                padding: 2px 6px;
                border-radius: 3px;
                font-family: 'Courier New', monospace;
                color: #e74c3c;
            }}
            pre {{
                background-color: #2c3e50;
                color: #ecf0f1;
                padding: 20px;
                border-radius: 5px;
                overflow-x: auto;
                margin: 20px 0;
            }}
            pre code {{
                background: none;
                color: inherit;
                padding: 0;
            }}
            table {{
                width: 100%;
                border-collapse: collapse;
                margin: 20px 0;
                background: white;
                border-radius: 5px;
                overflow: hidden;
                box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            }}
            th, td {{
                padding: 12px 15px;
                text-align: left;
                border-bottom: 1px solid #ddd;
            }}
            th {{
                background-color: #3498db;
                color: white;
                font-weight: 600;
            }}
            tr:hover {{
                background-color: #f5f5f5;
            }}
            blockquote {{
                border-left: 4px solid #3498db;
                margin: 20px 0;
                padding: 15px 20px;
                background-color: #f8f9fa;
                font-style: italic;
            }}
            .toc {{
                background-color: #f8f9fa;
                padding: 20px;
                border-radius: 5px;
                margin: 20px 0;
            }}
            .toc h2 {{
                margin-top: 0;
                color: #2c3e50;
            }}
            .toc ul {{
                list-style-type: none;
                padding-left: 0;
            }}
            .toc ul ul {{
                padding-left: 20px;
            }}
            .toc a {{
                color: #3498db;
                text-decoration: none;
            }}
            .toc a:hover {{
                text-decoration: underline;
            }}
            .highlight {{
                background-color: #fff3cd;
                padding: 15px;
                border-radius: 5px;
                border-left: 4px solid #ffc107;
                margin: 20px 0;
            }}
            .success {{
                background-color: #d4edda;
                padding: 15px;
                border-radius: 5px;
                border-left: 4px solid #28a745;
                margin: 20px 0;
            }}
            .info {{
                background-color: #d1ecf1;
                padding: 15px;
                border-radius: 5px;
                border-left: 4px solid #17a2b8;
                margin: 20px 0;
            }}
            .warning {{
                background-color: #f8d7da;
                padding: 15px;
                border-radius: 5px;
                border-left: 4px solid #dc3545;
                margin: 20px 0;
            }}
            @media (max-width: 768px) {{
                body {{
                    padding: 10px;
                }}
                .container {{
                    padding: 20px;
                }}
                h1 {{
                    font-size: 2em;
                }}
                h2 {{
                    font-size: 1.5em;
                }}
                table {{
                    font-size: 0.9em;
                }}
            }}
        </style>
    </head>
    <body>
        <div class="container">
            {html_content}
        </div>
    </body>
    </html>
    """
    
    with open(html_file_path, 'w', encoding='utf-8') as f:
        f.write(full_html)
    
    print(f"✅ HTML creado: {html_file_path}")

def html_to_pdf_reportlab(html_content, output_path):
    """Convierte HTML a PDF usando ReportLab"""
    print("📄 Convirtiendo HTML a PDF usando ReportLab...")
    
    # Crear documento PDF
    doc = SimpleDocTemplate(output_path, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
    
    # Estilos
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=16, spaceAfter=30, alignment=TA_CENTER)
    heading1_style = ParagraphStyle('CustomHeading1', parent=styles['Heading1'], fontSize=14, spaceAfter=12, spaceBefore=12)
    heading2_style = ParagraphStyle('CustomHeading2', parent=styles['Heading2'], fontSize=12, spaceAfter=10, spaceBefore=10)
    heading3_style = ParagraphStyle('CustomHeading3', parent=styles['Heading3'], fontSize=11, spaceAfter=8, spaceBefore=8)
    normal_style = ParagraphStyle('CustomNormal', parent=styles['Normal'], fontSize=10, spaceAfter=6)
    
    # Parsear HTML
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Remover estilos y scripts
    for style in soup(["style", "script"]):
        style.decompose()
    
    # Construir contenido PDF
    story = []
    
    # Procesar elementos
    for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'ul', 'ol', 'li', 'strong', 'em']):
        if element.name == 'h1':
            story.append(Paragraph(element.get_text(), title_style))
            story.append(Spacer(1, 12))
        elif element.name == 'h2':
            story.append(Paragraph(element.get_text(), heading1_style))
            story.append(Spacer(1, 8))
        elif element.name == 'h3':
            story.append(Paragraph(element.get_text(), heading2_style))
            story.append(Spacer(1, 6))
        elif element.name == 'h4':
            story.append(Paragraph(element.get_text(), heading3_style))
            story.append(Spacer(1, 4))
        elif element.name == 'p':
            if element.get_text().strip():
                story.append(Paragraph(element.get_text(), normal_style))
                story.append(Spacer(1, 6))
        elif element.name in ['ul', 'ol']:
            for li in element.find_all('li'):
                if li.get_text().strip():
                    story.append(Paragraph(f"• {li.get_text()}", normal_style))
                    story.append(Spacer(1, 3))
            story.append(Spacer(1, 6))
        elif element.name == 'li':
            if element.get_text().strip():
                story.append(Paragraph(f"• {element.get_text()}", normal_style))
                story.append(Spacer(1, 3))
        elif element.name in ['strong', 'em']:
            if element.get_text().strip():
                story.append(Paragraph(element.get_text(), normal_style))
                story.append(Spacer(1, 3))
    
    # Construir PDF
    try:
        doc.build(story)
        print(f"✅ PDF creado exitosamente: {output_path}")
    except Exception as e:
        print(f"❌ Error creando PDF: {e}")
        raise

def markdown_to_docx(md_file_path, docx_file_path):
    """Convierte Markdown a Word con formato mejorado"""
    print("📄 Convirtiendo Markdown a Word...")
    
    document = Document()
    styles = getSampleStyleSheet()
    
    # Usar estilos básicos
    title_style = 'Title'
    heading1_style = 'Heading1'
    heading2_style = 'Heading2'
    heading3_style = 'Heading3'
    normal_style = 'Normal'
    
    # Leer archivo Markdown
    with open(md_file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    in_list = False
    list_level = 0
    list_counter = [0] * 10  # Support for nested lists up to 10 levels
    
    for line in lines:
        line = line.strip()
        if not line:  # Handle empty lines gracefully
            if in_list:
                in_list = False
                list_level = 0
                list_counter = [0] * 10
            continue
        
        # Headers
        if line.startswith('# '):
            if in_list:
                in_list = False
                list_level = 0
                list_counter = [0] * 10
            p = document.add_paragraph(line[2:], style=title_style)
        elif line.startswith('## '):
            if in_list:
                in_list = False
                list_level = 0
                list_counter = [0] * 10
            p = document.add_paragraph(line[3:], style=heading1_style)
        elif line.startswith('### '):
            if in_list:
                in_list = False
                list_level = 0
                list_counter = [0] * 10
            p = document.add_paragraph(line[4:], style=heading2_style)
        elif line.startswith('#### '):
            if in_list:
                in_list = False
                list_level = 0
                list_counter = [0] * 10
            p = document.add_paragraph(line[5:], style=heading3_style)
        
        # Lists
        elif line.startswith('- ') or line.startswith('* '):
            if not in_list:
                in_list = True
                list_level = 0
            list_counter[list_level] += 1
            indent = list_level * 0.5
            p = document.add_paragraph(f"• {line[2:]}", style=normal_style)
            p.paragraph_format.left_indent = Inches(indent)
        
        elif line[0].isdigit() and '. ' in line:
            if not in_list:
                in_list = True
                list_level = 0
            list_counter[list_level] += 1
            indent = list_level * 0.5
            p = document.add_paragraph(line, style=normal_style)
            p.paragraph_format.left_indent = Inches(indent)
        
        # Regular paragraphs
        else:
            if in_list:
                in_list = False
                list_level = 0
                list_counter = [0] * 10
            
            # Handle bold and italic
            text = line
            if '**' in text:
                text = text.replace('**', '')
            if '*' in text:
                text = text.replace('*', '')
            
            p = document.add_paragraph(text, style=normal_style)
    
    # Guardar documento
    document.save(docx_file_path)
    print(f"✅ Documento Word creado exitosamente: {docx_file_path}")

def main():
    """Función principal"""
    print("🚀 Iniciando conversión del documento ultra completo final...")
    
    # Verificar dependencias
    if not check_dependencies():
        print("❌ Error: No se pudieron instalar las dependencias necesarias")
        return
    
    # Archivos
    md_file = "frases-populares-busqueda-google-2024-ultra-completo-final.md"
    html_file = "frases-populares-busqueda-google-2024-ultra-completo-final.html"
    pdf_file = "frases-populares-busqueda-google-2024-ultra-completo-final.pdf"
    docx_file = "frases-populares-busqueda-google-2024-ultra-completo-final.docx"
    
    # Verificar que el archivo Markdown existe
    if not os.path.exists(md_file):
        print(f"❌ Error: No se encontró el archivo {md_file}")
        return
    
    try:
        # Convertir a HTML
        markdown_to_html(md_file, html_file)
        
        # Leer HTML para PDF
        with open(html_file, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        # Convertir a PDF
        html_to_pdf_reportlab(html_content, pdf_file)
        
        # Convertir a Word
        markdown_to_docx(md_file, docx_file)
        
        print("\n🎉 ¡Conversión completada exitosamente!")
        print(f"\n📁 Archivos generados:")
        print(f"   • HTML: {os.path.abspath(html_file)}")
        print(f"   • PDF: {os.path.abspath(pdf_file)}")
        print(f"   • Word: {os.path.abspath(docx_file)}")
        
        # Mostrar tamaños de archivos
        print(f"\n📊 Resumen de archivos creados:")
        for file_path in [html_file, pdf_file, docx_file]:
            if os.path.exists(file_path):
                size = os.path.getsize(file_path)
                print(f"   ✅ {os.path.basename(file_path)}: {size} bytes")
        
    except Exception as e:
        print(f"❌ Error durante la conversión: {e}")
        return

if __name__ == "__main__":
    main()
