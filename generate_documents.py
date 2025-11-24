from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

# Content
TITLE = "ANÁLISIS ESTRATÉGICO DE ARQUITECTURA DE RED: MARCOS OSI VS. TCP/IP"
DATE = "23 de Noviembre, 2025"
AUTHOR = "Autor: Departamento de Infraestructura TI"
CONFIDENTIAL = "Clasificación: Confidencial & Propietario"

SECTIONS = [
    {
        "title": "1. Resumen Ejecutivo",
        "body": "La infraestructura de comunicaciones global opera sobre dos modelos de referencia críticos: el modelo teórico OSI (7 capas) y el modelo práctico TCP/IP (4 capas). Mientras que TCP/IP se ha consolidado como el estándar de facto para la transmisión de datos en Internet y redes corporativas modernas, el modelo OSI retiene un valor insustituible para el diagnóstico granular de ingeniería y el diseño de interoperabilidad.",
        "bullets": [
            "Dicotomía Operativa: TCP/IP prioriza la conectividad y velocidad de implementación; OSI prioriza la estandarización y modularidad.",
            "Gestión de Crisis: El modelo OSI reduce el tiempo medio de reparación (MTTR) al permitir una segmentación precisa de fallos (Física vs. Enlace).",
            "Imperativo de Modernización: La adopción de tecnologías Cloud y IoT exige migrar de topologías jerárquicas tradicionales a arquitecturas Spine-Leaf para mitigar latencia."
        ]
    },
    {
        "title": "2. Contexto y Antecedentes",
        "body": "Las organizaciones modernas enfrentan el desafío de escalar sus redes para soportar cargas de trabajo distribuidas. La elección y comprensión del modelo de referencia adecuado no es meramente académica; define la capacidad de la organización para gestionar la seguridad, escalar la infraestructura y resolver incidencias críticas. Este documento evalúa la aplicabilidad de ambos estándares en el entorno empresarial actual.",
        "diagram": """
   MODELO OSI (7 Capas)             MODELO TCP/IP (4 Capas)
+-----------------------+          +-----------------------+
| 7. Aplicación         |          |                       |
| 6. Presentación       | <------> | 4. Aplicación         |
| 5. Sesión             |          |                       |
+-----------------------+          +-----------------------+
| 4. Transporte         | <------> | 3. Transporte         |
+-----------------------+          +-----------------------+
| 3. Red                | <------> | 2. Internet           |
+-----------------------+          +-----------------------+
| 2. Enlace de Datos    |          |                       |
| 1. Física             | <------> | 1. Acceso a la Red    |
+-----------------------+          +-----------------------+
        """,
        "bullets": []
    },
    {
        "title": "3. Análisis Central: Evaluación Comparativa de Marcos",
        "body": "3.1 Desglose Estructural y Funcional\n\nLa siguiente matriz contrasta las capacidades operativas de ambos modelos, destacando sus dominios de aplicación óptimos.",
        "table": [
            ["Dimensión", "Modelo OSI (Referencia)", "Modelo TCP/IP (Implementación)"],
            ["Arquitectura", "7 Capas (Modularidad Estricta)", "4 Capas (Integración Práctica)"],
            ["Enfoque de Diseño", "Prescriptivo: Define qué debe hacerse.", "Descriptivo: Define cómo hacerlo."],
            ["Gestión de Protocolos", "Independiente de la tecnología subyacente.", "Estrechamente acoplado a la suite de protocolos de Internet."],
            ["Utilidad Primaria", "Diagnóstico de ingeniería, formación, estandarización.", "Conectividad global, transmisión de datos en tiempo real."],
            ["Seguridad", "Modularizada por capa específica.", "Aditiva (e.g., TLS sobre Transporte)."]
        ],
        "bullets": []
    },
    {
        "title": "3.2 Escenarios de Aplicación Específica",
        "body": "El análisis indica que el modelo OSI es superior en escenarios de Alta Complejidad Técnica:",
        "bullets": [
            "Diagnóstico de Fallos (Troubleshooting): Permite aislar incidencias. Ejemplo: Diferenciar una pérdida de señal (Capa 1) de un error de protocolo ARP (Capa 2).",
            "Interoperabilidad: Crítico para integrar sistemas legacy o propietarios que no utilizan TCP/IP nativo."
        ],
        "extra_body": "Por el contrario, TCP/IP es mandatorio para:\n\n• Despliegue Operativo: Configuración de enrutamiento global y servicios web."
    },
    {
        "title": "4. Recomendaciones Estratégicas",
        "body": "4.1 Optimización de Infraestructura para Nuevas Tecnologías (Cloud & IoT)\n\nLa transición hacia ecosistemas digitales avanzados presenta desafíos de arquitectura que requieren intervención inmediata.\n\nDesafíos Identificados:\n1. Cuellos de Botella en Topologías Legacy\n2. Vulnerabilidad Perimetral (IoT)",
        "bullets": [
            "Plan de Acción:",
            "Migración a Spine-Leaf: Implementar topologías aplanadas para garantizar ancho de banda consistente.",
            "Segmentación de Red Zero-Trust: Aislar tráfico IoT en VLANs dedicadas.",
            "Actualización del Edge: Incrementar la capacidad de enlaces troncales (10G/40G)."
        ]
    }
]

# --- GENERATE DOCX ---
doc = Document()

# Title Block
title = doc.add_heading(TITLE, 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(DATE).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(AUTHOR).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(CONFIDENTIAL).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("---").alignment = WD_ALIGN_PARAGRAPH.CENTER

for section in SECTIONS:
    doc.add_heading(section["title"], level=1)
    if section.get("body"):
        doc.add_paragraph(section["body"])

    if section.get("diagram"):
        p = doc.add_paragraph()
        run = p.add_run(section["diagram"])
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    
    if section.get("table"):
        table_data = section["table"]
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        table.style = 'Table Grid'
        for i, row in enumerate(table_data):
            for j, cell_text in enumerate(row):
                cell = table.cell(i, j)
                cell.text = cell_text
                if i == 0: # Header bold
                    run = cell.paragraphs[0].runs[0]
                    run.font.bold = True

    if section.get("bullets"):
        for bullet in section["bullets"]:
            doc.add_paragraph(bullet, style='List Bullet')

    if section.get("extra_body"):
        doc.add_paragraph(section["extra_body"])

doc.save("Analisis_Estrategico_Redes.docx")

# --- GENERATE PDF ---
class PDF(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 15)
        # Calculate width of title and position
        w = self.get_string_width(TITLE) + 6
        self.set_x((210 - w) / 2)
        self.cell(w, 9, TITLE, 0, 1, 'C')
        self.ln(10)

    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

pdf = FPDF()
pdf.add_page()
pdf.set_font("Arial", size=12)

# Meta info
pdf.set_font("Arial", 'I', 10)
pdf.cell(0, 5, DATE, 0, 1, 'C')
pdf.cell(0, 5, AUTHOR, 0, 1, 'C')
pdf.cell(0, 5, CONFIDENTIAL, 0, 1, 'C')
pdf.ln(10)

pdf.set_font("Arial", size=11)

for section in SECTIONS:
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(0, 10, section["title"].encode('latin-1', 'replace').decode('latin-1'), 0, 1)
    pdf.set_font("Arial", size=11)
    
    if section.get("body"):
        # Multi_cell for text wrapping
        pdf.multi_cell(0, 6, section["body"].encode('latin-1', 'replace').decode('latin-1'))
        pdf.ln(3)

    if section.get("diagram"):
        pdf.set_font("Courier", size=8)
        pdf.multi_cell(0, 4, section["diagram"].encode('latin-1', 'replace').decode('latin-1'))
        pdf.set_font("Arial", size=11)
        pdf.ln(3)

    if section.get("table"):
        pdf.ln(2)
        col_width = pdf.w / 3.5
        row_height = 8
        table_data = section["table"]
        for row in table_data:
            for item in row:
                # Simple table handling
                pdf.cell(col_width, row_height, str(item)[:25].encode('latin-1', 'replace').decode('latin-1'), border=1)
            pdf.ln(row_height)
        pdf.ln(5)

    if section.get("bullets"):
        for bullet in section["bullets"]:
            pdf.cell(5) # indent
            pdf.multi_cell(0, 6, f"- {bullet}".encode('latin-1', 'replace').decode('latin-1'))
        pdf.ln(3)

    if section.get("extra_body"):
        pdf.multi_cell(0, 6, section["extra_body"].encode('latin-1', 'replace').decode('latin-1'))
        pdf.ln(3)
    
    pdf.ln(5)

pdf.output("Analisis_Estrategico_Redes.pdf")

