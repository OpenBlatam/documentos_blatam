#!/usr/bin/env python3
"""
Script para convertir documentos importantes a PDF, Word y Excel con gráficas de alta calidad
"""

import os
import re
import json
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Tuple, Optional
import markdown
from markdown.extensions import tables, codehilite, fenced_code

# PDF
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle, Image
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    print("⚠️  reportlab no disponible. Instalando...")

# Word
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  python-docx no disponible. Instalando...")

# Excel
try:
    from openpyxl import Workbook
    from openpyxl.chart import BarChart, LineChart, PieChart, ScatterChart
    from openpyxl.chart.series import DataPoint
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    print("⚠️  openpyxl no disponible. Instalando...")

# Gráficas
try:
    import matplotlib
    matplotlib.use('Agg')  # Backend sin GUI
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    import seaborn as sns
    import numpy as np
    import pandas as pd
    MATPLOTLIB_AVAILABLE = True
except ImportError:
    MATPLOTLIB_AVAILABLE = False
    print("⚠️  matplotlib/seaborn no disponible. Instalando...")

# Configuración
sns.set_style("whitegrid")
plt.rcParams['figure.figsize'] = (12, 8)
plt.rcParams['font.size'] = 10
plt.rcParams['axes.labelsize'] = 12
plt.rcParams['axes.titlesize'] = 14
plt.rcParams['xtick.labelsize'] = 10
plt.rcParams['ytick.labelsize'] = 10

# Documentos importantes a convertir
DOCUMENTOS_IMPORTANTES = [
    {
        'path': 'airflow_automation_prompt.md',
        'nombre': 'Airflow Automation Prompt',
        'categoria': 'Automatización'
    },
    {
        'path': 'truthgpt_collected/integration_code/production_code/ARCHITECTURE_IMPROVEMENTS.md',
        'nombre': 'Architecture Improvements',
        'categoria': 'Arquitectura'
    },
    {
        'path': 'truthgpt_collected/integration_code/production_code/REFACTORING_PLAN.md',
        'nombre': 'Refactoring Plan',
        'categoria': 'Refactorización'
    },
    {
        'path': 'truthgpt_collected/integration_code/production_code/RESUMEN_FINAL_MEJORAS.md',
        'nombre': 'Resumen Final Mejoras',
        'categoria': 'Resumen Ejecutivo'
    },
    {
        'path': 'truthgpt_collected/integration_code/production_code/MEJORAS_ARQUITECTURA_COMPLETAS.md',
        'nombre': 'Mejoras Arquitectura Completas',
        'categoria': 'Arquitectura'
    },
    {
        'path': 'README.md',
        'nombre': 'README Principal',
        'categoria': 'Documentación Principal'
    },
    {
        'path': 'ARCHITECTURE.md',
        'nombre': 'Architecture',
        'categoria': 'Arquitectura'
    },
    {
        'path': 'truthgpt_collected/integration_code/production_code/MEJORAS_ADICIONALES_RECOMENDADAS.md',
        'nombre': 'Mejoras Adicionales Recomendadas',
        'categoria': 'Mejoras'
    },
    {
        'path': 'truthgpt_collected/integration_code/production_code/ARCHITECTURE.md',
        'nombre': 'Architecture Production Code',
        'categoria': 'Arquitectura'
    },
    {
        'path': '06_documentation/Playbooks/ai_playbook.md',
        'nombre': 'AI Playbook',
        'categoria': 'IA'
    },
    {
        'path': '05_technology/Tech_stack_docs/technical_specifications.md',
        'nombre': 'Technical Specifications',
        'categoria': 'Especificaciones Técnicas'
    },
    {
        'path': 'ROADMAP.md',
        'nombre': 'Roadmap',
        'categoria': 'Roadmap'
    },
    {
        'path': '04_business_strategy/Business_plans/comprehensive_business_plan.md',
        'nombre': 'Comprehensive Business Plan',
        'categoria': 'Plan de Negocio'
    },
    {
        'path': '06_strategy/Strategic_plans/business_plan_executive.md',
        'nombre': 'Business Plan Executive',
        'categoria': 'Plan de Negocio'
    },
    {
        'path': '04_business_strategy/Strategic_plans/advanced_strategic_planning_guide.md',
        'nombre': 'Advanced Strategic Planning Guide',
        'categoria': 'Planificación Estratégica'
    },
    {
        'path': '06_documentation/Resumes/resumen_ejecutivo_completo.md',
        'nombre': 'Resumen Ejecutivo Completo',
        'categoria': 'Resumen Ejecutivo'
    },
    {
        'path': '06_documentation/resumen_final_completo.md',
        'nombre': 'Resumen Final Completo',
        'categoria': 'Resumen'
    },
    {
        'path': 'truthgpt_collected/integration_code/production_code/INDICE_DOCUMENTACION.md',
        'nombre': 'Indice Documentacion Production Code',
        'categoria': 'Documentación'
    },
    {
        'path': '06_documentation/Master_documents/master_implementation_guide.md',
        'nombre': 'Master Implementation Guide',
        'categoria': 'Guía de Implementación'
    },
    {
        'path': '11_system_architecture/Documentation/master_implementation_roadmap.md',
        'nombre': 'Master Implementation Roadmap',
        'categoria': 'Roadmap'
    }
]


class DocumentConverter:
    """Clase principal para convertir documentos a PDF, Word y Excel"""
    
    def __init__(self, base_path: str):
        self.base_path = Path(base_path)
        self.output_dir = self.base_path / 'documentos_convertidos'
        self.output_dir.mkdir(exist_ok=True)
        self.graphs_dir = self.output_dir / 'graficas'
        self.graphs_dir.mkdir(exist_ok=True)
        
    def leer_markdown(self, file_path: Path) -> Tuple[str, Dict]:
        """Lee un archivo markdown y extrae contenido y metadatos"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                contenido = f.read()
            
            # Extraer metadatos si existen
            metadata = {}
            if contenido.startswith('---'):
                parts = contenido.split('---', 2)
                if len(parts) >= 3:
                    try:
                        import yaml
                        metadata = yaml.safe_load(parts[1]) or {}
                    except:
                        pass
                    contenido = parts[2].strip()
            
            return contenido, metadata
        except Exception as e:
            print(f"❌ Error leyendo {file_path}: {e}")
            return "", {}
    
    def extraer_tablas(self, contenido: str) -> List[Dict]:
        """Extrae tablas del contenido markdown con mejor detección"""
        tablas = []
        lines = contenido.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            # Detectar inicio de tabla
            if '|' in line and not line.startswith('```'):
                # Verificar que no sea código
                if i > 0 and '```' in lines[i-1]:
                    i += 1
                    continue
                
                # Leer headers
                headers = [h.strip() for h in line.split('|') if h.strip()]
                
                # Leer separador
                if i + 1 < len(lines):
                    separator = lines[i + 1].strip()
                    if '---' in separator or '===' in separator or all(c in '-|: ' for c in separator):
                        i += 2
                        rows = []
                        
                        # Leer filas de datos
                        while i < len(lines):
                            row_line = lines[i].strip()
                            
                            # Fin de tabla
                            if not row_line or ('|' not in row_line) or row_line.startswith('```'):
                                break
                            
                            # Verificar que no sea otra tabla
                            if i + 1 < len(lines):
                                next_line = lines[i + 1].strip()
                                if '---' in next_line or '===' in next_line:
                                    break
                            
                            cells = [c.strip() for c in row_line.split('|') if c.strip()]
                            if cells and len(cells) >= len(headers) - 1:  # Permitir variación
                                # Ajustar número de celdas
                                while len(cells) < len(headers):
                                    cells.append('')
                                rows.append(cells[:len(headers)])
                            
                            i += 1
                        
                        if headers and rows:
                            tablas.append({
                                'headers': headers,
                                'rows': rows
                            })
                        continue
            
            i += 1
        
        return tablas
    
    def extraer_metricas(self, contenido: str) -> Dict:
        """Extrae métricas y números del contenido"""
        metricas = {}
        
        # Buscar patrones comunes de métricas
        patterns = {
            'fases_completadas': r'(\d+)\s*/\s*(\d+)\s*fases?',
            'archivos_modificados': r'(\d+)\+?\s*archivos?\s*modificados',
            'lineas_eliminadas': r'(\d+)\+?\s*líneas?\s*eliminadas',
            'imports_estandarizados': r'(\d+)\+?\s*imports?\s*estandarizados',
            'porcentaje': r'(\d+)%',
        }
        
        for key, pattern in patterns.items():
            matches = re.findall(pattern, contenido, re.IGNORECASE)
            if matches:
                metricas[key] = matches
        
        # Buscar estados de fases
        fases = re.findall(r'Phase\s+(\d+):\s*([^✅❌\n]+)', contenido, re.IGNORECASE)
        if fases:
            metricas['fases'] = fases
        
        return metricas
    
    def generar_graficas(self, contenido: str, nombre_doc: str) -> List[str]:
        """Genera gráficas relevantes basadas en el contenido"""
        graficas_generadas = []
        metricas = self.extraer_metricas(contenido)
        tablas = self.extraer_tablas(contenido)
        
        # Gráfica 1: Fases completadas (si hay información de fases)
        if 'fases' in metricas or any('fase' in contenido.lower() for _ in [1]):
            try:
                fig, ax = plt.subplots(figsize=(10, 6))
                
                # Extraer datos de fases de tablas
                fases_data = []
                for tabla in tablas:
                    if any('fase' in str(h).lower() or 'phase' in str(h).lower() for h in tabla['headers']):
                        estado_col = None
                        nombre_col = None
                        for i, header in enumerate(tabla['headers']):
                            if 'estado' in header.lower() or 'status' in header.lower():
                                estado_col = i
                            if 'fase' in header.lower() or 'phase' in header.lower():
                                nombre_col = i
                        
                        if estado_col is not None:
                            for row in tabla['rows']:
                                if len(row) > max(estado_col, nombre_col or 0):
                                    estado = row[estado_col] if estado_col < len(row) else ''
                                    nombre = row[nombre_col] if nombre_col and nombre_col < len(row) else ''
                                    if '✅' in estado or 'complet' in estado.lower():
                                        fases_data.append(nombre or 'Fase')
                
                if not fases_data:
                    # Crear datos de ejemplo basados en el contenido
                    fases_data = ['Phase 1', 'Phase 2', 'Phase 3', 'Phase 4', 'Phase 5', 'Phase 6']
                    estados = [1] * len(fases_data)  # Todas completadas
                else:
                    estados = [1] * len(fases_data)
                
                    if fases_data:
                        colors_list = plt.cm.viridis(np.linspace(0, 1, len(fases_data)))
                        bars = ax.barh(fases_data, estados, color=colors_list)
                        ax.set_xlabel('Estado (Completada)', fontsize=12, fontweight='bold')
                        ax.set_ylabel('Fases', fontsize=12, fontweight='bold')
                        ax.set_title(f'Estado de Fases - {nombre_doc}', fontsize=14, fontweight='bold', pad=20)
                        ax.set_xlim(0, 1.2)
                        ax.grid(axis='x', alpha=0.3)
                        
                        # Añadir valores (usar texto simple en lugar de emoji)
                        for i, (bar, estado) in enumerate(zip(bars, estados)):
                            ax.text(estado + 0.05, bar.get_y() + bar.get_height()/2, 
                                   'Completada', va='center', fontweight='bold')
                    
                    plt.tight_layout()
                    graph_path = self.graphs_dir / f'{nombre_doc}_fases.png'
                    plt.savefig(graph_path, dpi=300, bbox_inches='tight')
                    plt.close()
                    graficas_generadas.append(str(graph_path))
            except Exception as e:
                print(f"⚠️  Error generando gráfica de fases: {e}")
        
        # Gráfica 2: Métricas de mejora (mejorada con múltiples tipos)
        for tabla_idx, tabla in enumerate(tablas):
            if any('métrica' in str(h).lower() or 'metric' in str(h).lower() or 
                   'archivo' in str(h).lower() or 'línea' in str(h).lower() or
                   'valor' in str(h).lower() or 'cantidad' in str(h).lower() or
                   'total' in str(h).lower() or 'número' in str(h).lower()
                   for h in tabla['headers']):
                try:
                    # Intentar extraer datos numéricos mejorado
                    metricas_data = {}
                    metricas_labels = []
                    
                    for row in tabla['rows']:
                        if len(row) >= 2:
                            nombre = row[0]
                            # Buscar números en las celdas (mejorado)
                            valores = []
                            for cell in row[1:]:
                                # Extraer números incluyendo decimales y porcentajes
                                cell_str = str(cell).replace(',', '').replace('+', '').replace('%', '')
                                nums = re.findall(r'\d+\.?\d*', cell_str)
                                if nums:
                                    try:
                                        valores.extend([float(n) for n in nums])
                                    except:
                                        pass
                            
                            if valores:
                                # Usar el valor más representativo
                                metricas_data[nombre[:40]] = max(valores) if valores else 0
                                metricas_labels.append(nombre[:40])
                    
                    if metricas_data and len(metricas_data) > 0:
                        # Determinar tipo de gráfica según cantidad de datos
                        num_items = len(metricas_data)
                        
                        if num_items <= 8:
                            # Gráfica de barras horizontales
                            fig, ax = plt.subplots(figsize=(12, max(6, num_items * 0.6)))
                            items = sorted(metricas_data.items(), key=lambda x: x[1], reverse=True)[:10]
                            nombres = [item[0] for item in items]
                            valores = [item[1] for item in items]
                            
                            colors_list = plt.cm.viridis(np.linspace(0.2, 0.8, len(nombres)))
                            bars = ax.barh(nombres, valores, color=colors_list, edgecolor='black', linewidth=0.5)
                            
                            ax.set_xlabel('Valor', fontsize=12, fontweight='bold')
                            ax.set_ylabel('Métrica', fontsize=12, fontweight='bold')
                            ax.set_title(f'Métricas de Mejora - {nombre_doc}', fontsize=14, fontweight='bold', pad=20)
                            ax.grid(axis='x', alpha=0.3, linestyle='--')
                            
                            # Añadir valores en las barras
                            max_val = max(valores) if valores else 1
                            for bar, valor in zip(bars, valores):
                                ax.text(valor + max_val*0.02, bar.get_y() + bar.get_height()/2, 
                                       f'{valor:.1f}', va='center', fontweight='bold', fontsize=9)
                            
                            plt.tight_layout()
                            graph_path = self.graphs_dir / f'{nombre_doc}_metricas_{tabla_idx}.png'
                            plt.savefig(graph_path, dpi=300, bbox_inches='tight', facecolor='white')
                            plt.close()
                            graficas_generadas.append(str(graph_path))
                            
                        elif num_items <= 15:
                            # Gráfica de pastel para muchos elementos
                            fig, ax = plt.subplots(figsize=(10, 10))
                            items = sorted(metricas_data.items(), key=lambda x: x[1], reverse=True)[:15]
                            nombres = [item[0][:25] for item in items]
                            valores = [item[1] for item in items]
                            
                            colors_list = plt.cm.Set3(np.linspace(0, 1, len(nombres)))
                            wedges, texts, autotexts = ax.pie(valores, labels=nombres, autopct='%1.1f%%',
                                                             colors=colors_list, startangle=90,
                                                             textprops={'fontsize': 9})
                            
                            ax.set_title(f'Distribución de Métricas - {nombre_doc}', 
                                        fontsize=14, fontweight='bold', pad=20)
                            
                            plt.tight_layout()
                            graph_path = self.graphs_dir / f'{nombre_doc}_metricas_pie_{tabla_idx}.png'
                            plt.savefig(graph_path, dpi=300, bbox_inches='tight', facecolor='white')
                            plt.close()
                            graficas_generadas.append(str(graph_path))
                        
                        if tabla_idx >= 2:  # Limitar a 3 gráficas de métricas
                            break
                            
                except Exception as e:
                    print(f"⚠️  Error generando gráfica de métricas: {e}")
                    import traceback
                    traceback.print_exc()
        
        # Gráfica 3: Timeline de mejoras (mejorada)
        fechas = re.findall(r'\d{4}-\d{2}-\d{2}', contenido)
        if fechas and len(set(fechas)) >= 2:
            try:
                fig, ax = plt.subplots(figsize=(14, 8))
                
                fechas_unicas = sorted(set(fechas))
                num_fechas = len(fechas_unicas)
                
                # Crear timeline mejorado
                y_pos = np.arange(num_fechas)
                colors_timeline = plt.cm.plasma(np.linspace(0.2, 0.8, num_fechas))
                
                # Barras horizontales
                bars = ax.barh(y_pos, [1]*num_fechas, color=colors_timeline, 
                              edgecolor='black', linewidth=1.5, alpha=0.8)
                
                # Etiquetas mejoradas
                labels = []
                for i, fecha in enumerate(fechas_unicas):
                    # Contar eventos en esa fecha
                    count = fechas.count(fecha)
                    label = f'{fecha}'
                    if count > 1:
                        label += f' ({count} eventos)'
                    labels.append(label)
                
                ax.set_yticks(y_pos)
                ax.set_yticklabels(labels, fontsize=10)
                ax.set_xlabel('Timeline', fontsize=12, fontweight='bold')
                ax.set_title(f'Timeline de Mejoras - {nombre_doc}', fontsize=14, fontweight='bold', pad=20)
                ax.set_xlim(0, 1.2)
                ax.grid(axis='x', alpha=0.3, linestyle='--')
                ax.invert_yaxis()  # Fechas más recientes arriba
                
                plt.tight_layout()
                graph_path = self.graphs_dir / f'{nombre_doc}_timeline.png'
                plt.savefig(graph_path, dpi=300, bbox_inches='tight', facecolor='white')
                plt.close()
                graficas_generadas.append(str(graph_path))
            except Exception as e:
                print(f"⚠️  Error generando timeline: {e}")
        
        # Gráfica 4: Comparación de fases (si hay múltiples fases)
        if len(tablas) >= 2:
            try:
                # Buscar tablas con información de fases
                fases_tablas = [t for t in tablas if any('fase' in str(h).lower() or 'phase' in str(h).lower() 
                                                        for h in t['headers'])]
                
                if len(fases_tablas) >= 1:
                    fig, ax = plt.subplots(figsize=(14, 8))
                    
                    # Extraer datos de comparación
                    comparacion_data = {}
                    for tabla in fases_tablas[:3]:  # Máximo 3 tablas
                        for row in tabla['rows'][:10]:  # Máximo 10 filas
                            if len(row) >= 2:
                                nombre = row[0][:30]
                                # Buscar valores numéricos
                                for cell in row[1:]:
                                    nums = re.findall(r'\d+', str(cell))
                                    if nums:
                                        try:
                                            valor = int(nums[0])
                                            if nombre not in comparacion_data:
                                                comparacion_data[nombre] = []
                                            comparacion_data[nombre].append(valor)
                                            break
                                        except:
                                            pass
                    
                    if comparacion_data:
                        nombres = list(comparacion_data.keys())[:8]
                        valores = [max(comparacion_data[n]) if comparacion_data[n] else 0 for n in nombres]
                        
                        bars = ax.bar(range(len(nombres)), valores, 
                                     color=plt.cm.coolwarm(np.linspace(0, 1, len(nombres))),
                                     edgecolor='black', linewidth=1)
                        
                        ax.set_xticks(range(len(nombres)))
                        ax.set_xticklabels(nombres, rotation=45, ha='right', fontsize=9)
                        ax.set_ylabel('Valor', fontsize=12, fontweight='bold')
                        ax.set_title(f'Comparación de Métricas - {nombre_doc}', 
                                   fontsize=14, fontweight='bold', pad=20)
                        ax.grid(axis='y', alpha=0.3, linestyle='--')
                        
                        # Añadir valores
                        for bar, valor in zip(bars, valores):
                            ax.text(bar.get_x() + bar.get_width()/2, valor + max(valores)*0.02,
                                   f'{valor}', ha='center', va='bottom', fontweight='bold', fontsize=9)
                        
                        plt.tight_layout()
                        graph_path = self.graphs_dir / f'{nombre_doc}_comparacion.png'
                        plt.savefig(graph_path, dpi=300, bbox_inches='tight', facecolor='white')
                        plt.close()
                        graficas_generadas.append(str(graph_path))
            except Exception as e:
                print(f"⚠️  Error generando gráfica de comparación: {e}")
        
        return graficas_generadas
    
    def convertir_a_pdf(self, contenido: str, nombre_doc: str, graficas: List[str], metadata: Dict = None):
        """Convierte contenido a PDF con alta calidad"""
        if not REPORTLAB_AVAILABLE:
            print("❌ reportlab no disponible. Instalando dependencias...")
            return None
        
        try:
            output_path = self.output_dir / f'{nombre_doc}.pdf'
            
            doc = SimpleDocTemplate(str(output_path), pagesize=A4,
                                   rightMargin=72, leftMargin=72,
                                   topMargin=72, bottomMargin=18)
            
            # Estilos
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor=colors.HexColor('#1a1a1a'),
                spaceAfter=30,
                alignment=TA_CENTER
            )
            
            heading1_style = ParagraphStyle(
                'CustomHeading1',
                parent=styles['Heading1'],
                fontSize=18,
                textColor=colors.HexColor('#2c3e50'),
                spaceAfter=12,
                spaceBefore=12
            )
            
            heading2_style = ParagraphStyle(
                'CustomHeading2',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=colors.HexColor('#34495e'),
                spaceAfter=10,
                spaceBefore=10
            )
            
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#2c3e50'),
                spaceAfter=6,
                leading=14
            )
            
            # Construir contenido
            story = []
            
            # Título
            story.append(Paragraph(nombre_doc, title_style))
            story.append(Spacer(1, 0.2*inch))
            
            if metadata:
                fecha = metadata.get('date', metadata.get('created', datetime.now().strftime('%Y-%m-%d')))
                story.append(Paragraph(f"<b>Fecha:</b> {fecha}", normal_style))
                story.append(Spacer(1, 0.1*inch))
            
            story.append(Spacer(1, 0.3*inch))
            
            # Procesar contenido markdown - usar tablas extraídas
            tablas = self.extraer_tablas(contenido)
            tabla_idx = 0
            
            lines = contenido.split('\n')
            in_code_block = False
            code_lines = []
            in_table = False
            table_lines = []
            
            i = 0
            while i < len(lines):
                line = lines[i]
                
                # Títulos
                if line.startswith('# '):
                    story.append(Spacer(1, 0.2*inch))
                    title_text = line[2:].strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    story.append(Paragraph(title_text, heading1_style))
                    story.append(Spacer(1, 0.1*inch))
                elif line.startswith('## '):
                    story.append(Spacer(1, 0.15*inch))
                    title_text = line[3:].strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    story.append(Paragraph(title_text, heading2_style))
                    story.append(Spacer(1, 0.08*inch))
                elif line.startswith('### '):
                    story.append(Spacer(1, 0.1*inch))
                    title_text = line[4:].strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    story.append(Paragraph(f"<b>{title_text}</b>", normal_style))
                    story.append(Spacer(1, 0.05*inch))
                # Tablas - procesar usando las tablas extraídas
                elif '|' in line and not in_code_block and not line.strip().startswith('```'):
                    if not in_table:
                        in_table = True
                        table_lines = [line]
                    else:
                        table_lines.append(line)
                    
                    # Verificar si es el final de la tabla (siguiente línea no tiene |)
                    if i + 1 < len(lines):
                        next_line = lines[i + 1]
                        if '|' not in next_line and next_line.strip() and not next_line.strip().startswith('|'):
                            # Procesar tabla acumulada
                            if tabla_idx < len(tablas):
                                tabla = tablas[tabla_idx]
                                try:
                                    # Crear tabla PDF
                                    table_data = [tabla['headers']] + tabla['rows'][:20]  # Limitar a 20 filas
                                    pdf_table = Table(table_data, repeatRows=1)
                                    pdf_table.setStyle(TableStyle([
                                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#366092')),
                                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                                        ('FONTSIZE', (0, 0), (-1, 0), 10),
                                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                                        ('FONTSIZE', (0, 1), (-1, -1), 9),
                                        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f5f5f5')]),
                                    ]))
                                    story.append(pdf_table)
                                    story.append(Spacer(1, 0.2*inch))
                                except Exception as e:
                                    print(f"⚠️  Error procesando tabla: {e}")
                                tabla_idx += 1
                            in_table = False
                            table_lines = []
                # Código
                elif line.strip().startswith('```'):
                    if in_code_block:
                        # Finalizar bloque de código
                        code_text = '\n'.join(code_lines)
                        code_text = code_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        story.append(Paragraph(f"<font face='Courier'>{code_text}</font>", 
                                             ParagraphStyle('Code', parent=normal_style, 
                                                          fontName='Courier', fontSize=8,
                                                          backColor=colors.HexColor('#f5f5f5'),
                                                          leftIndent=20)))
                        code_lines = []
                        in_code_block = False
                    else:
                        in_code_block = True
                elif in_code_block:
                    code_lines.append(line)
                # Listas
                elif line.strip().startswith('- ') or line.strip().startswith('* '):
                    list_text = line.strip()[2:].replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    story.append(Paragraph(f"• {list_text}", normal_style))
                # Párrafos normales
                elif line.strip() and not in_table:
                    # Limpiar markdown básico y escapar HTML problemático
                    text = line.strip()
                    # Escapar caracteres especiales primero
                    text = text.replace('&', '&amp;')
                    text = text.replace('<', '&lt;')
                    text = text.replace('>', '&gt;')
                    # Luego aplicar formato (solo si no hay problemas)
                    text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
                    text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
                    text = re.sub(r'`(.+?)`', r'<font face="Courier">\1</font>', text)
                    # Limpiar tags HTML no balanceados
                    try:
                        story.append(Paragraph(text, normal_style))
                    except Exception as e:
                        # Si falla, usar texto plano
                        text_plain = text.replace('<b>', '').replace('</b>', '').replace('<i>', '').replace('</i>', '')
                        text_plain = text_plain.replace('<font face="Courier">', '').replace('</font>', '')
                        story.append(Paragraph(text_plain, normal_style))
                elif not line.strip() and not in_table:
                    story.append(Spacer(1, 0.05*inch))
                
                i += 1
            
            # Añadir gráficas
            if graficas:
                story.append(PageBreak())
                story.append(Paragraph("Gráficas y Visualizaciones", heading1_style))
                story.append(Spacer(1, 0.2*inch))
                
                for graph_path in graficas:
                    if os.path.exists(graph_path):
                        try:
                            img = Image(graph_path, width=6*inch, height=4*inch)
                            story.append(img)
                            story.append(Spacer(1, 0.2*inch))
                        except Exception as e:
                            print(f"⚠️  Error añadiendo gráfica {graph_path}: {e}")
            
            # Construir PDF
            doc.build(story)
            print(f"✅ PDF generado: {output_path}")
            return str(output_path)
            
        except Exception as e:
            print(f"❌ Error generando PDF: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def convertir_a_word(self, contenido: str, nombre_doc: str, graficas: List[str], metadata: Dict = None):
        """Convierte contenido a Word con formato profesional"""
        if not DOCX_AVAILABLE:
            print("❌ python-docx no disponible. Instalando dependencias...")
            return None
        
        try:
            doc = Document()
            
            # Configurar estilos
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(11)
            
            # Título
            title = doc.add_heading(nombre_doc, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadatos
            if metadata:
                fecha = metadata.get('date', metadata.get('created', datetime.now().strftime('%Y-%m-%d')))
                para = doc.add_paragraph()
                para.add_run(f"Fecha: {fecha}").bold = True
                doc.add_paragraph()
            
            # Procesar contenido
            lines = contenido.split('\n')
            in_code_block = False
            code_lines = []
            
            for line in lines:
                if line.startswith('# '):
                    doc.add_heading(line[2:].strip(), level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:].strip(), level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:].strip(), level=3)
                elif line.strip().startswith('```'):
                    if in_code_block:
                        # Finalizar código
                        para = doc.add_paragraph()
                        run = para.add_run('\n'.join(code_lines))
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                        para.style = 'No Spacing'
                        code_lines = []
                        in_code_block = False
                    else:
                        in_code_block = True
                elif in_code_block:
                    code_lines.append(line)
                elif '|' in line and not in_code_block:
                    # Procesar tabla
                    cells = [c.strip() for c in line.split('|') if c.strip()]
                    if cells and not all('-' in c or '=' in c for c in cells):
                        # Crear tabla (simplificado - solo primera fila)
                        continue
                elif line.strip().startswith('- ') or line.strip().startswith('* '):
                    para = doc.add_paragraph(line.strip()[2:], style='List Bullet')
                elif line.strip():
                    para = doc.add_paragraph(line.strip())
                else:
                    doc.add_paragraph()
            
            # Añadir gráficas
            if graficas:
                doc.add_page_break()
                doc.add_heading('Gráficas y Visualizaciones', level=1)
                
                for graph_path in graficas:
                    if os.path.exists(graph_path):
                        try:
                            para = doc.add_paragraph()
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = para.add_run()
                            run.add_picture(graph_path, width=Inches(6))
                            doc.add_paragraph()
                        except Exception as e:
                            print(f"⚠️  Error añadiendo gráfica {graph_path}: {e}")
            
            # Guardar
            output_path = self.output_dir / f'{nombre_doc}.docx'
            doc.save(str(output_path))
            print(f"✅ Word generado: {output_path}")
            return str(output_path)
            
        except Exception as e:
            print(f"❌ Error generando Word: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def convertir_a_excel(self, contenido: str, nombre_doc: str, graficas: List[str], metadata: Dict = None):
        """Convierte contenido a Excel con tablas y gráficas"""
        if not OPENPYXL_AVAILABLE:
            print("❌ openpyxl no disponible. Instalando dependencias...")
            return None
        
        try:
            wb = Workbook()
            ws = wb.active
            ws.title = "Resumen"
            
            # Estilos
            header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            header_font = Font(bold=True, color="FFFFFF", size=12)
            title_font = Font(bold=True, size=16)
            
            # Título
            ws['A1'] = nombre_doc
            ws['A1'].font = title_font
            ws.merge_cells('A1:D1')
            ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
            
            row = 3
            
            # Metadatos
            if metadata:
                ws[f'A{row}'] = f"Fecha: {metadata.get('date', metadata.get('created', datetime.now().strftime('%Y-%m-%d')))}"
                row += 2
            
            # Extraer y añadir tablas
            tablas = self.extraer_tablas(contenido)
            
            for tabla_idx, tabla in enumerate(tablas[:5]):  # Máximo 5 tablas
                # Crear nueva hoja para cada tabla grande
                if tabla_idx > 0:
                    ws = wb.create_sheet(title=f"Tabla {tabla_idx + 1}")
                    row = 1
                
                # Headers
                for col_idx, header in enumerate(tabla['headers'], 1):
                    cell = ws.cell(row=row, column=col_idx, value=header)
                    cell.fill = header_fill
                    cell.font = header_font
                    cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                    ws.column_dimensions[get_column_letter(col_idx)].width = 20
                
                row += 1
                
                # Datos
                for tabla_row in tabla['rows']:
                    for col_idx, value in enumerate(tabla_row[:len(tabla['headers'])], 1):
                        cell = ws.cell(row=row, column=col_idx, value=value)
                        cell.alignment = Alignment(vertical='top', wrap_text=True)
                    row += 1
                
                row += 2
                
                # Crear gráfica si hay datos numéricos
                try:
                    # Intentar encontrar columnas numéricas
                    numeric_cols = []
                    for col_idx in range(1, len(tabla['headers']) + 1):
                        try:
                            # Verificar si la columna tiene números
                            has_numbers = False
                            for r in range(row - len(tabla['rows']) - 1, row - 1):
                                cell_value = ws.cell(row=r, column=col_idx).value
                                if cell_value:
                                    cell_str = str(cell_value)
                                    # Verificar si es numérico (incluye decimales y negativos)
                                    try:
                                        float(cell_str.replace(',', '').replace('+', '').replace('%', ''))
                                        has_numbers = True
                                        break
                                    except:
                                        pass
                            if has_numbers:
                                numeric_cols.append(col_idx)
                        except:
                            pass
                    
                    if numeric_cols and len(tabla['rows']) > 0:
                        try:
                            # Determinar tipo de gráfica según datos
                            data_start_row = row - len(tabla['rows']) - 1
                            data_end_row = row - 1
                            
                            if len(numeric_cols) == 1 and len(tabla['rows']) <= 10:
                                # Gráfica de pastel para pocos elementos
                                chart = PieChart()
                                chart.title = f"Distribución - {tabla['headers'][0]}"
                                data_col = get_column_letter(numeric_cols[0])
                                cat_col = get_column_letter(1)
                                
                                chart.add_data(ws[f'{data_col}{data_start_row}:{data_col}{data_end_row}'], 
                                             titles_from_data=True)
                                chart.set_categories(ws[f'{cat_col}{data_start_row}:{cat_col}{data_end_row}'])
                                chart.height = 10
                                chart.width = 15
                                
                                ws.add_chart(chart, f'A{row + 2}')
                                row += 20
                            else:
                                # Gráfica de barras mejorada
                                chart = BarChart()
                                chart.type = "col"
                                chart.style = 10
                                chart.title = f"Métricas - {tabla['headers'][0]}"
                                chart.y_axis.title = 'Valor'
                                chart.x_axis.title = tabla['headers'][0]
                                chart.height = 10
                                chart.width = 15
                                
                                # Añadir todas las columnas numéricas
                                for col_idx in numeric_cols[:3]:  # Máximo 3 series
                                    col_letter = get_column_letter(col_idx)
                                    chart.add_data(ws[f'{col_letter}{data_start_row}:{col_letter}{data_end_row}'], 
                                                 titles_from_data=True)
                                
                                # Categorías (primera columna)
                                cat_col = get_column_letter(1)
                                chart.set_categories(ws[f'{cat_col}{data_start_row}:{cat_col}{data_end_row}'])
                                
                                # Añadir gráfica
                                ws.add_chart(chart, f'A{row + 2}')
                                row += 20
                        except Exception as e:
                            print(f"⚠️  Error creando gráfica avanzada en Excel: {e}")
                except Exception as e:
                    print(f"⚠️  Error creando gráfica en Excel: {e}")
            
            # Hoja de métricas
            if tablas:
                ws_metrics = wb.create_sheet(title="Métricas", index=1)
                metricas = self.extraer_metricas(contenido)
                
                ws_metrics['A1'] = "Métrica"
                ws_metrics['B1'] = "Valor"
                ws_metrics['A1'].fill = header_fill
                ws_metrics['B1'].fill = header_fill
                ws_metrics['A1'].font = header_font
                ws_metrics['B1'].font = header_font
                
                row = 2
                for key, values in metricas.items():
                    if values:
                        ws_metrics[f'A{row}'] = key.replace('_', ' ').title()
                        ws_metrics[f'B{row}'] = str(values[0]) if isinstance(values[0], (int, str)) else len(values)
                        row += 1
                
                ws_metrics.column_dimensions['A'].width = 30
                ws_metrics.column_dimensions['B'].width = 20
            
            # Guardar
            output_path = self.output_dir / f'{nombre_doc}.xlsx'
            wb.save(str(output_path))
            print(f"✅ Excel generado: {output_path}")
            return str(output_path)
            
        except Exception as e:
            print(f"❌ Error generando Excel: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def procesar_documento(self, doc_info: Dict):
        """Procesa un documento completo"""
        doc_path = self.base_path / doc_info['path']
        
        if not doc_path.exists():
            print(f"⚠️  Documento no encontrado: {doc_path}")
            return None
        
        print(f"\n📄 Procesando: {doc_info['nombre']}")
        print(f"   Ruta: {doc_path}")
        
        # Leer contenido
        contenido, metadata = self.leer_markdown(doc_path)
        
        if not contenido:
            print(f"❌ No se pudo leer el contenido")
            return None
        
        # Generar gráficas
        print("   📊 Generando gráficas...")
        graficas = self.generar_graficas(contenido, doc_info['nombre'])
        print(f"   ✅ {len(graficas)} gráficas generadas")
        
        # Convertir a PDF
        print("   📄 Generando PDF...")
        pdf_path = self.convertir_a_pdf(contenido, doc_info['nombre'], graficas, metadata)
        
        # Convertir a Word
        print("   📝 Generando Word...")
        word_path = self.convertir_a_word(contenido, doc_info['nombre'], graficas, metadata)
        
        # Convertir a Excel
        print("   📊 Generando Excel...")
        excel_path = self.convertir_a_excel(contenido, doc_info['nombre'], graficas, metadata)
        
        return {
            'nombre': doc_info['nombre'],
            'pdf': pdf_path,
            'word': word_path,
            'excel': excel_path,
            'graficas': graficas
        }
    
    def procesar_todos(self):
        """Procesa todos los documentos importantes"""
        resultados = []
        
        print("🚀 Iniciando conversión de documentos importantes...")
        print(f"📁 Directorio de salida: {self.output_dir}")
        print(f"📊 Directorio de gráficas: {self.graphs_dir}\n")
        
        for doc_info in DOCUMENTOS_IMPORTANTES:
            resultado = self.procesar_documento(doc_info)
            if resultado:
                resultados.append(resultado)
        
        # Resumen
        print("\n" + "="*60)
        print("📋 RESUMEN DE CONVERSIÓN")
        print("="*60)
        
        for resultado in resultados:
            print(f"\n📄 {resultado['nombre']}")
            if resultado['pdf']:
                print(f"   ✅ PDF: {resultado['pdf']}")
            if resultado['word']:
                print(f"   ✅ Word: {resultado['word']}")
            if resultado['excel']:
                print(f"   ✅ Excel: {resultado['excel']}")
            print(f"   📊 Gráficas: {len(resultado['graficas'])}")
        
        print(f"\n✅ Proceso completado. Archivos guardados en: {self.output_dir}")
        
        return resultados


def main():
    """Función principal"""
    import sys
    
    # Verificar dependencias
    missing_deps = []
    if not REPORTLAB_AVAILABLE:
        missing_deps.append('reportlab')
    if not DOCX_AVAILABLE:
        missing_deps.append('python-docx')
    if not OPENPYXL_AVAILABLE:
        missing_deps.append('openpyxl')
    if not MATPLOTLIB_AVAILABLE:
        missing_deps.extend(['matplotlib', 'seaborn', 'pandas', 'numpy'])
    
    if missing_deps:
        print("⚠️  Faltan dependencias. Instalando...")
        import subprocess
        subprocess.check_call([sys.executable, '-m', 'pip', 'install'] + missing_deps)
        print("✅ Dependencias instaladas. Por favor ejecuta el script nuevamente.")
        return
    
    # Obtener directorio base
    base_path = Path(__file__).parent
    
    # Crear convertidor y procesar
    converter = DocumentConverter(str(base_path))
    converter.procesar_todos()


if __name__ == '__main__':
    main()

