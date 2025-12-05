#!/usr/bin/env python3
"""
Script MASTER con todas las funcionalidades: similitud, índices, recomendaciones, LaTeX, Markdown mejorado
"""

import os
import sys
import re
from pathlib import Path
from datetime import datetime
import json
from collections import Counter, defaultdict
import math
from difflib import SequenceMatcher

try:
    import numpy as np
except ImportError:
    os.system("pip install numpy")
    import numpy as np

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    os.system("pip install python-docx")
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor

try:
    import matplotlib.pyplot as plt
    import matplotlib
    matplotlib.use('Agg')
    import seaborn as sns
    sns.set_style("whitegrid")
except ImportError:
    os.system("pip install matplotlib seaborn")
    import matplotlib.pyplot as plt
    import matplotlib
    matplotlib.use('Agg')
    import seaborn as sns

# Importar funcionalidades anteriores
sys.path.insert(0, str(Path(__file__).parent))
try:
    from convert_docs_advanced_analysis import AdvancedAnalysisConverter
except ImportError:
    try:
        from convert_docs_ultimate_plus import UltimatePlusConverter as AdvancedAnalysisConverter
    except ImportError:
        from convert_docs_to_formats_improved import AdvancedDocumentConverter as AdvancedAnalysisConverter

class MasterDocumentConverter(AdvancedAnalysisConverter):
    def __init__(self, output_dir="converted_docs"):
        super().__init__(output_dir)
        self.master_dir = self.output_dir / "master_analysis"
        self.master_dir.mkdir(exist_ok=True)
        
    def calculate_similarity(self, text1, text2):
        """Calcula similitud entre dos textos"""
        return SequenceMatcher(None, text1.lower(), text2.lower()).ratio()
    
    def analyze_document_similarity(self, all_docs_data):
        """Analiza similitud entre documentos"""
        similarities = []
        doc_names = [d["name"] for d in all_docs_data]
        
        for i, doc1 in enumerate(all_docs_data):
            for j, doc2 in enumerate(all_docs_data):
                if i < j:
                    # Leer contenidos
                    content1 = self.read_markdown(doc1["path"]) if isinstance(doc1["path"], (str, Path)) else ""
                    content2 = self.read_markdown(doc2["path"]) if isinstance(doc2["path"], (str, Path)) else ""
                    
                    similarity = self.calculate_similarity(content1, content2)
                    
                    # Similitud de temas
                    themes1 = set(self.detect_main_themes(content1, doc1.get("sections", []))["main_themes"].keys())
                    themes2 = set(self.detect_main_themes(content2, doc2.get("sections", []))["main_themes"].keys())
                    theme_similarity = len(themes1 & themes2) / max(len(themes1 | themes2), 1)
                    
                    similarities.append({
                        "doc1": doc1["name"],
                        "doc2": doc2["name"],
                        "text_similarity": similarity,
                        "theme_similarity": theme_similarity,
                        "combined_similarity": (similarity + theme_similarity) / 2
                    })
        
        return similarities
    
    def generate_auto_index(self, sections):
        """Genera índice automático del documento"""
        index = []
        for i, section in enumerate(sections, 1):
            if section["title"]:
                indent = "  " * (section["level"] - 1)
                index.append(f"{indent}{i}. {section['title']}")
        return "\n".join(index)
    
    def analyze_writing_quality(self, content, sections):
        """Analiza calidad de escritura"""
        issues = []
        score = 100
        
        # Verificar longitud de párrafos
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        long_paragraphs = [p for p in paragraphs if len(p) > 500]
        if long_paragraphs:
            issues.append(f"{len(long_paragraphs)} párrafos muy largos (>500 caracteres)")
            score -= len(long_paragraphs) * 2
        
        # Verificar repetición de palabras
        words = re.findall(r'\b[a-z]{4,}\b', content.lower())
        word_freq = Counter(words)
        repeated_words = [(w, c) for w, c in word_freq.items() if c > 20 and len(w) > 4]
        if repeated_words:
            issues.append(f"Palabras repetidas excesivamente: {', '.join([w for w, c in repeated_words[:5]])}")
            score -= 5
        
        # Verificar uso de listas
        list_items = content.count('- ') + content.count('* ')
        if len(paragraphs) > 0 and list_items / len(paragraphs) < 0.1:
            issues.append("Bajo uso de listas para organizar información")
            score -= 5
        
        # Verificar títulos
        headers = [l for l in content.split('\n') if l.strip().startswith('#')]
        if len(headers) < len(sections) * 0.8:
            issues.append("Algunas secciones podrían necesitar títulos más claros")
            score -= 5
        
        return {
            "quality_score": max(0, score),
            "issues": issues,
            "recommendations": self.generate_recommendations(issues, content, sections)
        }
    
    def generate_recommendations(self, issues, content, sections):
        """Genera recomendaciones automáticas"""
        recommendations = []
        
        if any("párrafos muy largos" in issue for issue in issues):
            recommendations.append("Dividir párrafos largos en párrafos más cortos para mejorar la legibilidad")
        
        if any("Palabras repetidas" in issue for issue in issues):
            recommendations.append("Usar sinónimos para evitar repetición excesiva de palabras")
        
        if any("Bajo uso de listas" in issue for issue in issues):
            recommendations.append("Considerar usar más listas para organizar información compleja")
        
        # Recomendaciones basadas en métricas
        word_count = len(content.split())
        if word_count > 10000:
            recommendations.append("Documento muy extenso - considerar dividirlo en múltiples documentos")
        
        section_count = len(sections)
        if section_count > 50:
            recommendations.append("Muchas secciones - considerar agrupar secciones relacionadas")
        
        code_blocks = content.count('```')
        if code_blocks > 20:
            recommendations.append("Muchos bloques de código - considerar crear un documento separado de ejemplos")
        
        return recommendations
    
    def detect_inconsistencies(self, content, sections):
        """Detecta inconsistencias en el documento"""
        inconsistencies = []
        
        # Verificar formato de títulos
        header_levels = []
        for line in content.split('\n'):
            if line.strip().startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                header_levels.append(level)
        
        # Verificar saltos de nivel
        for i in range(len(header_levels) - 1):
            if header_levels[i+1] - header_levels[i] > 1:
                inconsistencies.append(f"Salto de nivel de título: {header_levels[i]} -> {header_levels[i+1]}")
        
        # Verificar formato de código
        code_blocks = re.findall(r'```[\s\S]*?```', content)
        languages = []
        for block in code_blocks:
            first_line = block.split('\n')[0].strip()
            if first_line.startswith('```'):
                lang = first_line[3:].strip()
                if lang:
                    languages.append(lang)
        
        # Verificar consistencia de lenguajes
        lang_counter = Counter(languages)
        if len(lang_counter) > 5:
            inconsistencies.append(f"Muchos lenguajes de programación diferentes ({len(lang_counter)}) - considerar estandarizar")
        
        # Verificar enlaces rotos (básico)
        links = re.findall(r'\[([^\]]+)\]\(([^\)]+)\)', content)
        broken_links = [link for link in links if link[1].startswith('#') and link[1][1:] not in [s["title"].lower().replace(' ', '-') for s in sections if s["title"]]]
        if broken_links:
            inconsistencies.append(f"Posibles enlaces internos rotos: {len(broken_links)}")
        
        return {
            "total_inconsistencies": len(inconsistencies),
            "inconsistencies": inconsistencies,
            "consistency_score": max(0, 100 - len(inconsistencies) * 10)
        }
    
    def analyze_accessibility(self, content, sections):
        """Analiza accesibilidad del documento"""
        accessibility_issues = []
        score = 100
        
        # Verificar uso de encabezados
        headers = [l for l in content.split('\n') if l.strip().startswith('#')]
        if len(headers) < 3:
            accessibility_issues.append("Pocos encabezados - dificulta la navegación")
            score -= 10
        
        # Verificar texto alternativo en imágenes
        images = re.findall(r'!\[([^\]]*)\]\(([^\)]+)\)', content)
        images_without_alt = [img for img in images if not img[0].strip()]
        if images_without_alt:
            accessibility_issues.append(f"{len(images_without_alt)} imágenes sin texto alternativo")
            score -= len(images_without_alt) * 5
        
        # Verificar uso de listas
        list_count = content.count('- ') + content.count('* ') + content.count('1. ')
        if list_count < 5:
            accessibility_issues.append("Pocas listas - considerar más estructura")
            score -= 5
        
        # Verificar longitud de líneas (en código)
        code_blocks = re.findall(r'```[\s\S]*?```', content)
        long_lines = 0
        for block in code_blocks:
            lines = block.split('\n')
            long_lines += sum(1 for line in lines if len(line) > 100)
        
        if long_lines > 10:
            accessibility_issues.append(f"{long_lines} líneas de código muy largas (>100 caracteres)")
            score -= min(10, long_lines // 5)
        
        return {
            "accessibility_score": max(0, score),
            "issues": accessibility_issues,
            "level": "Excelente" if score > 80 else "Buena" if score > 60 else "Necesita mejoras"
        }
    
    def export_to_latex(self, doc_data, output_name):
        """Exporta a LaTeX"""
        content = self.read_markdown(doc_data["path"]) if isinstance(doc_data["path"], (str, Path)) else ""
        sections = doc_data["sections"]
        
        latex_content = f"""\\documentclass[11pt,a4paper]{{article}}
\\usepackage[utf8]{{inputenc}}
\\usepackage[spanish,english]{{babel}}
\\usepackage{{graphicx}}
\\usepackage{{hyperref}}
\\usepackage{{listings}}
\\usepackage{{xcolor}}

\\title{{{output_name.replace('_', ' ').title()}}}
\\author{{Generado automáticamente}}
\\date{{\\today}}

\\begin{{document}}

\\maketitle

\\tableofcontents
\\newpage

"""
        
        for section in sections:
            if section["title"]:
                level = min(section["level"], 4)
                section_cmd = ["", "\\section", "\\subsection", "\\subsubsection", "\\paragraph"][level]
                latex_content += f"{section_cmd}{{{section['title']}}}\n\n"
            
            # Procesar contenido
            content_lines = section["content"].split('\n')
            for line in content_lines:
                if line.strip():
                    if line.startswith('- ') or line.startswith('* '):
                        latex_content += f"\\begin{{itemize}}\n\\item {line[2:]}\n\\end{{itemize}}\n\n"
                    elif line.startswith('```'):
                        latex_content += "\\begin{lstlisting}\n"
                    elif line.startswith('`') and line.endswith('`'):
                        latex_content += f"\\texttt{{{line[1:-1]}}}\n"
                    else:
                        # Escapar caracteres especiales de LaTeX
                        line = line.replace('\\', '\\textbackslash{}')
                        line = line.replace('&', '\\&')
                        line = line.replace('%', '\\%')
                        line = line.replace('$', '\\$')
                        line = line.replace('#', '\\#')
                        line = line.replace('^', '\\textasciicircum{}')
                        line = line.replace('_', '\\_')
                        line = line.replace('{', '\\{')
                        line = line.replace('}', '\\}')
                        latex_content += f"{line}\n\n"
        
        latex_content += "\\end{document}\n"
        
        output_path = self.master_dir / f"{output_name}.tex"
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(latex_content)
        
        print(f"✓ LaTeX exportado: {output_path}")
        return output_path
    
    def export_to_enhanced_markdown(self, doc_data, output_name):
        """Exporta a Markdown mejorado con TOC y metadatos"""
        content = self.read_markdown(doc_data["path"]) if isinstance(doc_data["path"], (str, Path)) else ""
        sections = doc_data["sections"]
        stats = doc_data.get("stats", {})
        
        # Generar TOC
        toc = "## Tabla de Contenidos\n\n"
        for i, section in enumerate(sections, 1):
            if section["title"]:
                indent = "  " * (section["level"] - 1)
                anchor = section["title"].lower().replace(' ', '-').replace('_', '-')
                toc += f"{indent}- [{section['title']}](#{anchor})\n"
        
        # Metadatos YAML front matter
        yaml_front = f"""---
title: {output_name.replace('_', ' ').title()}
author: Sistema de Análisis Automático
date: {datetime.now().strftime('%Y-%m-%d')}
generated: {datetime.now().isoformat()}
statistics:
  words: {stats.get('total_words', 0)}
  sections: {stats.get('sections', 0)}
  readability: {stats.get('readability_score', 0)}/100
  complexity: {stats.get('complexity_score', 0)}/100
---

"""
        
        enhanced_md = yaml_front + toc + "\n---\n\n" + content
        
        output_path = self.master_dir / f"{output_name}_ENHANCED.md"
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(enhanced_md)
        
        print(f"✓ Markdown mejorado exportado: {output_path}")
        return output_path
    
    def create_similarity_matrix(self, all_docs_data):
        """Crea matriz de similitud entre documentos"""
        similarities = self.analyze_document_similarity(all_docs_data)
        
        if not similarities:
            return None
        
        doc_names = list(set([s["doc1"] for s in similarities] + [s["doc2"] for s in similarities]))
        n = len(doc_names)
        matrix = np.zeros((n, n))
        
        # Crear diccionario de similitudes
        sim_dict = {}
        for sim in similarities:
            key = (sim["doc1"], sim["doc2"])
            sim_dict[key] = sim["combined_similarity"]
        
        # Llenar matriz
        for i, doc1 in enumerate(doc_names):
            for j, doc2 in enumerate(doc_names):
                if i == j:
                    matrix[i][j] = 1.0
                elif (doc1, doc2) in sim_dict:
                    matrix[i][j] = sim_dict[(doc1, doc2)]
                elif (doc2, doc1) in sim_dict:
                    matrix[i][j] = sim_dict[(doc2, doc1)]
        
        # Crear gráfica
        fig, ax = plt.subplots(figsize=(10, 8))
        im = ax.imshow(matrix, cmap='YlOrRd', aspect='auto', vmin=0, vmax=1)
        
        ax.set_xticks(np.arange(n))
        ax.set_yticks(np.arange(n))
        ax.set_xticklabels(doc_names, rotation=45, ha='right')
        ax.set_yticklabels(doc_names)
        
        # Agregar valores
        for i in range(n):
            for j in range(n):
                text = ax.text(j, i, f'{matrix[i, j]:.2f}',
                             ha="center", va="center", color="black", fontweight='bold')
        
        ax.set_title('Matriz de Similitud entre Documentos', fontsize=16, fontweight='bold', pad=20)
        plt.colorbar(im, ax=ax)
        plt.tight_layout()
        
        chart_path = self.temp_dir / "similarity_matrix.png"
        plt.savefig(chart_path, dpi=300, bbox_inches='tight', facecolor='white')
        plt.close()
        
        print(f"✓ Matriz de similitud guardada: {chart_path}")
        return chart_path
    
    def create_master_report(self, all_docs_data, all_analyses):
        """Crea reporte maestro con todos los análisis"""
        doc = Document()
        
        # Portada
        title = doc.add_heading('Reporte Maestro - Análisis Completo de Documentos', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run(f"Generado el {datetime.now().strftime('%d de %B de %Y a las %H:%M')}")
        
        doc.add_page_break()
        
        # Resumen Ejecutivo
        doc.add_heading('Resumen Ejecutivo', level=1)
        doc.add_paragraph(f"Se analizaron {len(all_docs_data)} documentos con análisis completo incluyendo:")
        doc.add_paragraph("• Análisis de sentimiento y tono", style='List Bullet')
        doc.add_paragraph("• Detección de temas principales", style='List Bullet')
        doc.add_paragraph("• Análisis de coherencia", style='List Bullet')
        doc.add_paragraph("• Análisis de calidad de escritura", style='List Bullet')
        doc.add_paragraph("• Análisis de accesibilidad", style='List Bullet')
        doc.add_paragraph("• Análisis de similitud entre documentos", style='List Bullet')
        doc.add_paragraph("• Recomendaciones automáticas", style='List Bullet')
        
        doc.add_page_break()
        
        # Análisis de Similitud
        if len(all_docs_data) >= 2:
            doc.add_heading('Análisis de Similitud entre Documentos', level=1)
            similarities = self.analyze_document_similarity(all_docs_data)
            
            sim_table = doc.add_table(rows=1, cols=4)
            sim_table.style = 'Light Grid Accent 1'
            hdr_cells = sim_table.rows[0].cells
            hdr_cells[0].text = 'Documento 1'
            hdr_cells[1].text = 'Documento 2'
            hdr_cells[2].text = 'Similitud Texto'
            hdr_cells[3].text = 'Similitud Temas'
            
            for sim in similarities:
                row_cells = sim_table.add_row().cells
                row_cells[0].text = sim["doc1"]
                row_cells[1].text = sim["doc2"]
                row_cells[2].text = f"{sim['text_similarity']:.2%}"
                row_cells[3].text = f"{sim['theme_similarity']:.2%}"
        
        doc.add_page_break()
        
        # Análisis Individual Mejorado
        for i, (doc_data, analysis) in enumerate(zip(all_docs_data, all_analyses), 1):
            doc.add_heading(f'Análisis Detallado: {doc_data["name"]}', level=1)
            
            # Calidad de escritura
            content = self.read_markdown(doc_data["path"]) if isinstance(doc_data["path"], (str, Path)) else ""
            writing_quality = self.analyze_writing_quality(content, doc_data["sections"])
            
            doc.add_heading('Calidad de Escritura', level=2)
            para = doc.add_paragraph()
            para.add_run(f"Score: {writing_quality['quality_score']}/100").bold = True
            
            if writing_quality["issues"]:
                doc.add_paragraph("Problemas detectados:")
                for issue in writing_quality["issues"]:
                    doc.add_paragraph(issue, style='List Bullet')
            
            if writing_quality["recommendations"]:
                doc.add_paragraph("Recomendaciones:")
                for rec in writing_quality["recommendations"]:
                    doc.add_paragraph(rec, style='List Bullet')
            
            # Inconsistencias
            inconsistencies = self.detect_inconsistencies(content, doc_data["sections"])
            doc.add_heading('Análisis de Consistencia', level=2)
            para = doc.add_paragraph()
            para.add_run(f"Score de Consistencia: {inconsistencies['consistency_score']}/100").bold = True
            
            if inconsistencies["inconsistencies"]:
                doc.add_paragraph("Inconsistencias detectadas:")
                for inc in inconsistencies["inconsistencies"][:5]:
                    doc.add_paragraph(inc, style='List Bullet')
            
            # Accesibilidad
            accessibility = self.analyze_accessibility(content, doc_data["sections"])
            doc.add_heading('Análisis de Accesibilidad', level=2)
            para = doc.add_paragraph()
            para.add_run(f"Score: {accessibility['accessibility_score']}/100 - Nivel: {accessibility['level']}").bold = True
            
            if accessibility["issues"]:
                doc.add_paragraph("Problemas de accesibilidad:")
                for issue in accessibility["issues"]:
                    doc.add_paragraph(issue, style='List Bullet')
            
            # Índice automático
            auto_index = self.generate_auto_index(doc_data["sections"])
            doc.add_heading('Índice Automático', level=2)
            doc.add_paragraph(auto_index)
            
            if i < len(all_docs_data):
                doc.add_page_break()
        
        output_path = self.master_dir / "MASTER_REPORT.docx"
        doc.save(str(output_path))
        print(f"✓ Reporte maestro guardado: {output_path}")
        return output_path

def main():
    """Función principal"""
    base_dir = Path("/Users/adan/Documents/documentos_blatam")
    production_dir = base_dir / "truthgpt_collected/integration_code/production_code"
    
    documents = [
        {
            "path": base_dir / "airflow_automation_prompt.md",
            "name": "Automation_Expert_Prompt"
        },
        {
            "path": production_dir / "ARCHITECTURE_IMPROVEMENTS.md",
            "name": "Architecture_Improvements"
        },
        {
            "path": production_dir / "REFACTORING_PLAN.md",
            "name": "Refactoring_Plan"
        }
    ]
    
    converter = MasterDocumentConverter()
    
    print("=" * 70)
    print("ANÁLISIS MASTER COMPLETO")
    print("=" * 70)
    print()
    
    all_docs_data = []
    all_analyses = []
    
    # Procesar documentos
    for doc in documents:
        if not doc["path"].exists():
            continue
        
        print(f"\n📄 Procesando: {doc['name']}")
        print("-" * 70)
        
        try:
            content = converter.read_markdown(doc["path"])
            sections = converter.parse_markdown_sections(content)
            stats = converter.create_advanced_statistics(content, sections)
            
            doc_data = {
                "name": doc["name"],
                "path": doc["path"],
                "sections": sections,
                "stats": stats
            }
            
            all_docs_data.append(doc_data)
            
            # Análisis avanzado
            sentiment = converter.analyze_sentiment_tone(content)
            themes = converter.detect_main_themes(content, sections)
            coherence = converter.analyze_coherence(sections)
            summary = converter.generate_auto_summary(content, sections)
            code_analysis = converter.analyze_code_complexity(content)
            dependencies = converter.analyze_dependencies(sections)
            
            all_analyses.append({
                "name": doc["name"],
                "sentiment": sentiment,
                "themes": themes,
                "coherence": coherence,
                "summary": summary,
                "code": code_analysis,
                "dependencies": dependencies
            })
            
            # Exportar a LaTeX y Markdown mejorado
            converter.export_to_latex(doc_data, doc["name"])
            converter.export_to_enhanced_markdown(doc_data, doc["name"])
            
        except Exception as e:
            print(f"❌ Error: {str(e)}")
            import traceback
            traceback.print_exc()
    
    # Crear análisis de similitud
    if len(all_docs_data) >= 2:
        print("\n" + "=" * 70)
        print("ANÁLISIS DE SIMILITUD")
        print("=" * 70)
        converter.create_similarity_matrix(all_docs_data)
    
    # Crear reporte maestro
    print("\n" + "=" * 70)
    print("CREANDO REPORTE MAESTRO")
    print("=" * 70)
    converter.create_master_report(all_docs_data, all_analyses)
    
    print("\n" + "=" * 70)
    print("ANÁLISIS MASTER COMPLETADO")
    print("=" * 70)
    print(f"\nArchivos generados en: {converter.master_dir}")
    print("\n✅ Análisis master completado!")

if __name__ == "__main__":
    main()



