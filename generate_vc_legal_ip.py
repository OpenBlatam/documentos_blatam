import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_ip_agreement(filename):
    doc = docx.Document()
    
    doc.add_paragraph("PROPRIETARY INFORMATION AND INVENTIONS ASSIGNMENT AGREEMENT", style='Title').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("(Standard Template for Employees & Contractors)", style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("This Agreement is made between Ventura Capital Inc. (the 'Company') and the undersigned individual (the 'Assignor').")
    
    doc.add_heading("1. Assignment of Inventions", level=1)
    doc.add_paragraph("Assignor hereby assigns to the Company all right, title, and interest in and to any and all inventions, original works of authorship, developments, concepts, improvements, designs, discoveries, ideas, trademarks, or trade secrets, whether or not patentable or registrable under copyright or similar laws, which Assignor may solely or jointly conceive or develop or reduce to practice, or cause to be conceived or developed or reduced to practice, during the period of time Assignor is in the service of the Company.")

    doc.add_heading("2. Proprietary Information", level=1)
    doc.add_paragraph("Assignor understands that his/her/their work for the Company creates a relationship of trust and confidence between Assignor and the Company. Assignor agrees that all Proprietary Information (including but not limited to: AI models, datasets, training weights, customer lists, and source code) is the sole property of the Company.")

    doc.add_heading("3. Moral Rights Waiver", level=1)
    doc.add_paragraph("To the extent allowed by law, Assignor hereby waives all moral rights relating to the Inventions, including the right to attribution of authorship.")
    
    doc.add_heading("4. Governing Law", level=1)
    doc.add_paragraph("This Agreement will be governed by the laws of [Jurisdiction].")
    
    doc.add_paragraph("_" * 40).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Signatures").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(filename)
    print(f"IP Agreement created: {filename}")

if __name__ == "__main__":
    create_ip_agreement("Ventura_Capital_IP_Assignment_Template.docx")







