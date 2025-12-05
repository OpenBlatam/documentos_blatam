import docx
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_msa(filename):
    doc = docx.Document()
    
    # Styles
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.space_before = Pt(12)

    # Title
    doc.add_paragraph("MASTER SERVICE AGREEMENT (SAAS)", style='Title').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Ventura Capital Inc.", style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("This Master Service Agreement ('Agreement') is entered into by and between Ventura Capital Inc. ('Provider') and the Customer identified in the Order Form.")

    # 1. Definitions
    doc.add_heading("1. DEFINITIONS", 1)
    doc.add_paragraph("'Service' means the AI-powered marketing content generation platform provided by Provider.")
    doc.add_paragraph("'Customer Data' means all data and materials uploaded by Customer to the Service.")

    # 2. License & Use
    doc.add_heading("2. ACCESS AND USE", 1)
    doc.add_paragraph("2.1 Provision of Access. Subject to payment of Fees, Provider grants Customer a non-exclusive, non-transferable right to access and use the Service.")
    doc.add_paragraph("2.2 Usage Limits. Services are subject to usage limits (e.g., API calls, generated words) specified in the Order Form.")

    # 3. SLA
    doc.add_heading("3. SERVICE LEVEL AGREEMENT (SLA)", 1)
    doc.add_paragraph("Provider commits to 99.9% Uptime during business hours. Service Credits apply for downtime exceeding 1% per month.")

    # 4. Data Protection
    doc.add_heading("4. DATA PROTECTION & AI TRAINING", 1)
    doc.add_paragraph("4.1 Confidentiality. Provider treats all Customer Data as confidential.")
    doc.add_paragraph("4.2 AI Training Opt-Out. Customer may opt-out of having their data used to fine-tune Provider's global models. Enterprise plans default to 'Private Mode' (no training on customer data).")

    # 5. Fees
    doc.add_heading("5. FEES AND PAYMENT", 1)
    doc.add_paragraph("Fees are invoiced annually in advance. Late payments accrue interest at 1.5% per month.")

    # 6. Term
    doc.add_heading("6. TERM AND TERMINATION", 1)
    doc.add_paragraph("This Agreement commences on the Effective Date and continues for the Initial Term (typically 12 months), auto-renewing unless cancelled with 30 days notice.")

    doc.save(filename)
    print(f"MSA created: {filename}")

if __name__ == "__main__":
    create_msa("Ventura_Capital_SaaS_MSA.docx")







