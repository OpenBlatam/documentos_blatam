import docx
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_talent_docs(filename):
    doc = docx.Document()
    
    # Styles
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(16)
    h1.font.color.rgb = RGBColor(32, 55, 100)
    h1.font.bold = True
    
    doc.add_paragraph("VENTURA CAPITAL", style='Title').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("KEY HIRES - JOB DESCRIPTIONS (SERIES A)", style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("_" * 80).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Role 1: Head of Sales
    doc.add_heading("ROLE 1: HEAD OF SALES (LATAM)", 1)
    
    doc.add_paragraph("The Mission:", style='Heading 3')
    doc.add_paragraph("Scale our revenue from $150k ARR to $3.6M ARR in 18 months by building a high-velocity outbound sales engine targeting Enterprise Marketing Agencies.")
    
    doc.add_paragraph("Key Responsibilities:", style='Heading 3')
    doc.add_paragraph("• Build and lead a team of 5 SDRs and 3 Account Executives.", style='List Bullet')
    doc.add_paragraph("• Define and optimize the outbound sales playbook (cold email, LinkedIn, demos).", style='List Bullet')
    doc.add_paragraph("• Personally close the first 10 'Whale' accounts ($50k+ ACV).", style='List Bullet')
    
    doc.add_paragraph("Requirements:", style='Heading 3')
    doc.add_paragraph("• 5+ years in B2B SaaS Sales (MarTech preferred).", style='List Bullet')
    doc.add_paragraph("• Experience scaling from Seed to Series B.", style='List Bullet')

    doc.add_paragraph("_" * 40).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Role 2: VP of Engineering
    doc.add_heading("ROLE 2: VP OF ENGINEERING", 1)
    
    doc.add_paragraph("The Mission:", style='Heading 3')
    doc.add_paragraph("Lead the technical strategy and scaling of our proprietary AI infrastructure, ensuring 99.9% uptime while shipping dialect-specific models weekly.")
    
    doc.add_paragraph("Key Responsibilities:", style='Heading 3')
    doc.add_paragraph("• Manage a team of 12 engineers (ML Ops, Backend, Frontend).", style='List Bullet')
    doc.add_paragraph("• Oversee the transition from prototype architecture to SOC2-compliant enterprise grade.", style='List Bullet')
    doc.add_paragraph("• attract top AI talent from across the region.", style='List Bullet')
    
    doc.add_paragraph("Requirements:", style='Heading 3')
    doc.add_paragraph("• Ex-CTO or VP Eng at a scaling startup.", style='List Bullet')
    doc.add_paragraph("• Deep understanding of LLM inference at scale (Kubernetes, Ray).", style='List Bullet')

    doc.save(filename)
    print(f"Talent Docs created: {filename}")

def create_demo_script(filename):
    doc = docx.Document()
    
    doc.add_paragraph("VENTURA CAPITAL", style='Title').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("PRODUCT DEMO VIDEO SCRIPT (2 MIN)", style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = "Visual"
    hdr[1].text = "Audio / Narration"
    
    scenes = [
        ("Intro: Logo animation -> Split screen of chaotic agency office vs Calm Ventura user.", 
         "Marketing in LATAM is broken. Agencies are drowning in generic content. But what if you could generate weeks of culturally-calibrated campaigns in seconds?"),
        
        ("Screen Record: Dashboard. User selects 'Campaign Mode', chooses 'Mexico City' dialect.", 
         "Meet Ventura. Unlike generic AI, Ventura speaks your language. Literally. Select your target dialect - say, 'Chilango Moderno' for a Mexico City launch."),
         
        ("Screen Record: Input brand URL. AI scrapes brand assets. Generates 10 post variations.", 
         "Just drop your brand URL. Ventura analyzes your visual identity and tone. In 30 seconds, you get 10 platform-native assets, ready to publish."),
         
        ("Screen Record: 'Cultural Check' feature highlighting a slang term correction.", 
         "See this? Our 'Cultural Guardrails' flagged a term that works in Spain but means something... else... in Colombia. Ventura fixes it automatically."),
         
        ("Outro: Happy user publishing. Metrics chart going up. Ventura Logo.", 
         "Ventura Capital. The operating system for Latin American growth. Book your demo today.")
    ]
    
    for vis, aud in scenes:
        row = table.add_row().cells
        row[0].text = vis
        row[1].text = aud

    doc.save(filename)
    print(f"Demo Script created: {filename}")

if __name__ == "__main__":
    create_talent_docs("Ventura_Capital_Key_Hires_JDs.docx")
    create_demo_script("Ventura_Capital_Demo_Script.docx")







