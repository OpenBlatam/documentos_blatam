#!/usr/bin/env python3
"""
Script con análisis avanzado: sentimiento, temas, coherencia, resúmenes automáticos
"""

import os
import sys
import re
from pathlib import Path
from datetime import datetime
import json
from collections import Counter, defaultdict
import math

try:
    import numpy as np
except ImportError:
    os.system("pip install numpy")
    import numpy as np

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    os.system("pip install python-docx")
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor

try:
    import matplotlib.pyplot as plt
    import matplotlib
    matplotlib.use('Agg')
    import seaborn as sns
    sns.set_style("whitegrid")
except ImportError:
    os.system("pip install matplotlib seaborn")
    import matplotlib.pyplot as plt
    import matplotlib
    matplotlib.use('Agg')
    import seaborn as sns

# Importar funcionalidades base
sys.path.insert(0, str(Path(__file__).parent))
try:
    from convert_docs_ultimate_plus import UltimatePlusConverter
except ImportError:
    try:
        from convert_docs_ultimate import UltimateDocumentConverter as UltimatePlusConverter
    except ImportError:
        from convert_docs_to_formats_improved import AdvancedDocumentConverter as UltimatePlusConverter

class AdvancedAnalysisConverter(UltimatePlusConverter):
    def __init__(self, output_dir="converted_docs"):
        super().__init__(output_dir)
        self.analysis_dir = self.output_dir / "advanced_analysis"
        self.analysis_dir.mkdir(exist_ok=True)
        
    def analyze_sentiment_tone(self, content):
        """Analiza el tono y sentimiento del documento"""
        # Palabras positivas
        positive_words = ['excelente', 'bueno', 'mejor', 'éxito', 'eficiente', 'optimizado', 
                         'mejorado', 'avanzado', 'profesional', 'completo', 'robusto']
        # Palabras negativas
        negative_words = ['error', 'fallo', 'problema', 'dificultad', 'complejo', 'lento',
                         'ineficiente', 'limitado', 'incompleto']
        # Palabras técnicas
        technical_words = ['implementar', 'configurar', 'desplegar', 'optimizar', 'integrar',
                          'automatizar', 'orquestar', 'escalar', 'monitorear']
        
        content_lower = content.lower()
        positive_count = sum(1 for word in positive_words if word in content_lower)
        negative_count = sum(1 for word in negative_words if word in content_lower)
        technical_count = sum(1 for word in technical_words if word in content_lower)
        
        total_sentiment_words = positive_count + negative_count
        sentiment_score = ((positive_count - negative_count) / max(total_sentiment_words, 1)) * 100
        
        return {
            "sentiment_score": max(-100, min(100, int(sentiment_score))),
            "positive_words": positive_count,
            "negative_words": negative_count,
            "technical_words": technical_count,
            "tone": "Positivo" if sentiment_score > 20 else "Neutro" if sentiment_score > -20 else "Técnico/Neutro"
        }
    
    def detect_main_themes(self, content, sections, top_n=10):
        """Detecta los temas principales del documento"""
        # Extraer palabras importantes (excluyendo stopwords básicas)
        stopwords = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 
                    'of', 'with', 'by', 'is', 'are', 'was', 'were', 'be', 'been', 'have',
                    'has', 'had', 'do', 'does', 'did', 'will', 'would', 'should', 'could',
                    'may', 'might', 'can', 'this', 'that', 'these', 'those', 'it', 'its',
                    'el', 'la', 'los', 'las', 'un', 'una', 'de', 'del', 'en', 'con', 'por'}
        
        # Extraer palabras de títulos y contenido
        all_text = content.lower()
        words = re.findall(r'\b[a-z]{4,}\b', all_text)
        words = [w for w in words if w not in stopwords]
        
        # Contar frecuencia
        word_freq = Counter(words)
        
        # Agrupar por temas relacionados
        themes = {}
        for word, count in word_freq.most_common(50):
            # Buscar palabras relacionadas
            related = [w for w in word_freq.keys() if word in w or w in word or 
                      (len(word) > 5 and len(w) > 5 and word[:4] == w[:4])]
            if related:
                theme_key = word
                if theme_key not in themes:
                    themes[theme_key] = count
                else:
                    themes[theme_key] += count
        
        # Ordenar y tomar top N
        sorted_themes = sorted(themes.items(), key=lambda x: x[1], reverse=True)[:top_n]
        
        return {
            "main_themes": dict(sorted_themes),
            "theme_count": len(sorted_themes),
            "most_common_words": dict(word_freq.most_common(20))
        }
    
    def analyze_coherence(self, sections):
        """Analiza la coherencia y estructura del documento"""
        if not sections:
            return {"coherence_score": 0, "issues": []}
        
        issues = []
        score = 100
        
        # Verificar niveles jerárquicos
        levels = [s["level"] for s in sections if s["level"] > 0]
        if levels:
            max_level = max(levels)
            # Penalizar si hay saltos grandes de nivel
            for i in range(len(levels) - 1):
                if abs(levels[i] - levels[i+1]) > 2:
                    issues.append(f"Salto grande de nivel entre secciones: {levels[i]} -> {levels[i+1]}")
                    score -= 5
        
        # Verificar secciones vacías
        empty_sections = sum(1 for s in sections if not s["content"].strip())
        if empty_sections > 0:
            issues.append(f"{empty_sections} secciones sin contenido")
            score -= empty_sections * 2
        
        # Verificar distribución de contenido
        section_lengths = [len(s["content"]) for s in sections]
        if section_lengths:
            avg_length = sum(section_lengths) / len(section_lengths)
            # Penalizar si hay mucha variación
            variance = sum((l - avg_length) ** 2 for l in section_lengths) / len(section_lengths)
            if variance > avg_length * 10:
                issues.append("Distribución desigual de contenido entre secciones")
                score -= 10
        
        return {
            "coherence_score": max(0, score),
            "issues": issues,
            "structure_quality": "Excelente" if score > 80 else "Buena" if score > 60 else "Regular"
        }
    
    def generate_auto_summary(self, content, sections, max_sentences=10):
        """Genera un resumen automático del documento"""
        # Extraer oraciones importantes (las que están en títulos o tienen palabras clave)
        sentences = re.split(r'[.!?]+', content)
        sentences = [s.strip() for s in sentences if len(s.strip()) > 20]
        
        # Priorizar oraciones con palabras clave
        keywords = ['automation', 'workflow', 'pipeline', 'deployment', 'configuration',
                   'orchestration', 'monitoring', 'testing', 'security', 'optimization']
        
        scored_sentences = []
        for sentence in sentences[:100]:  # Limitar para eficiencia
            score = 0
            sentence_lower = sentence.lower()
            # Puntos por palabras clave
            for keyword in keywords:
                if keyword in sentence_lower:
                    score += 3
            # Puntos por longitud (oraciones medianas son mejores)
            if 50 <= len(sentence) <= 200:
                score += 2
            # Puntos por estar cerca del inicio
            score += 1
            
            scored_sentences.append((score, sentence))
        
        # Ordenar y tomar las mejores
        scored_sentences.sort(reverse=True)
        summary_sentences = [s[1] for s in scored_sentences[:max_sentences]]
        
        return {
            "summary": " ".join(summary_sentences),
            "key_points": [s["title"] for s in sections[:5] if s["title"]],
            "total_sentences_analyzed": len(sentences),
            "summary_length": len(" ".join(summary_sentences))
        }
    
    def analyze_code_complexity(self, content):
        """Analiza la complejidad del código en el documento"""
        code_blocks = re.findall(r'```[\s\S]*?```', content)
        
        total_lines = 0
        total_functions = 0
        total_classes = 0
        languages = Counter()
        
        for block in code_blocks:
            lines = block.split('\n')
            # Detectar lenguaje
            first_line = lines[0].strip()
            if first_line.startswith('```'):
                lang = first_line[3:].strip()
                if lang:
                    languages[lang] += 1
            
            code_content = '\n'.join(lines[1:-1])
            total_lines += len([l for l in code_content.split('\n') if l.strip()])
            total_functions += len(re.findall(r'\bdef\s+\w+|function\s+\w+|const\s+\w+\s*=', code_content))
            total_classes += len(re.findall(r'\bclass\s+\w+|export\s+class\s+\w+', code_content))
        
        return {
            "code_blocks": len(code_blocks),
            "total_code_lines": total_lines,
            "functions_detected": total_functions,
            "classes_detected": total_classes,
            "languages_used": dict(languages),
            "avg_lines_per_block": total_lines / max(len(code_blocks), 1),
            "complexity_level": "Alto" if total_lines > 500 else "Medio" if total_lines > 100 else "Bajo"
        }
    
    def analyze_dependencies(self, sections):
        """Analiza dependencias y referencias entre secciones"""
        dependencies = []
        references = defaultdict(list)
        
        for i, section in enumerate(sections):
            content = section["content"].lower()
            title = section["title"].lower()
            
            # Buscar referencias a otras secciones
            for j, other_section in enumerate(sections):
                if i != j and other_section["title"]:
                    other_title_words = other_section["title"].lower().split()
                    # Buscar palabras clave del título en el contenido
                    matches = sum(1 for word in other_title_words if len(word) > 4 and word in content)
                    if matches >= 2:  # Al menos 2 palabras coinciden
                        references[i].append({
                            "target_section": j,
                            "target_title": other_section["title"],
                            "match_score": matches
                        })
        
        return {
            "total_dependencies": sum(len(refs) for refs in references.values()),
            "sections_with_dependencies": len([i for i in references.keys() if references[i]]),
            "dependency_map": {str(i): refs for i, refs in references.items() if refs},
            "most_referenced_sections": sorted(
                [(i, len(refs)) for i, refs in references.items()],
                key=lambda x: x[1],
                reverse=True
            )[:5]
        }
    
    def create_advanced_analysis_report(self, doc_data, output_name):
        """Crea reporte de análisis avanzado"""
        print(f"Creando análisis avanzado para {output_name}...")
        
        content = self.read_markdown(doc_data["path"]) if isinstance(doc_data["path"], (str, Path)) else ""
        sections = doc_data["sections"]
        
        # Realizar todos los análisis
        sentiment = self.analyze_sentiment_tone(content)
        themes = self.detect_main_themes(content, sections)
        coherence = self.analyze_coherence(sections)
        summary = self.generate_auto_summary(content, sections)
        code_analysis = self.analyze_code_complexity(content)
        dependencies = self.analyze_dependencies(sections)
        
        # Crear documento Word con análisis avanzado
        doc = Document()
        
        # Título
        title = doc.add_heading(f'Análisis Avanzado: {output_name}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Fecha
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run(f"Generado el {datetime.now().strftime('%d de %B de %Y a las %H:%M')}")
        
        doc.add_page_break()
        
        # 1. Resumen Automático
        doc.add_heading('Resumen Automático', level=1)
        doc.add_paragraph(summary["summary"])
        doc.add_paragraph()
        doc.add_paragraph(f"Puntos clave:")
        for point in summary["key_points"]:
            doc.add_paragraph(point, style='List Bullet')
        
        doc.add_page_break()
        
        # 2. Análisis de Sentimiento y Tono
        doc.add_heading('Análisis de Sentimiento y Tono', level=1)
        para = doc.add_paragraph()
        para.add_run(f"Score de Sentimiento: {sentiment['sentiment_score']}/100").bold = True
        doc.add_paragraph(f"Tono general: {sentiment['tone']}")
        doc.add_paragraph(f"Palabras positivas detectadas: {sentiment['positive_words']}")
        doc.add_paragraph(f"Palabras negativas detectadas: {sentiment['negative_words']}")
        doc.add_paragraph(f"Palabras técnicas detectadas: {sentiment['technical_words']}")
        
        doc.add_page_break()
        
        # 3. Temas Principales
        doc.add_heading('Temas Principales Detectados', level=1)
        themes_table = doc.add_table(rows=1, cols=2)
        themes_table.style = 'Light Grid Accent 1'
        hdr_cells = themes_table.rows[0].cells
        hdr_cells[0].text = 'Tema'
        hdr_cells[1].text = 'Frecuencia'
        hdr_cells[0].paragraphs[0].runs[0].font.bold = True
        hdr_cells[1].paragraphs[0].runs[0].font.bold = True
        
        for theme, count in list(themes["main_themes"].items())[:15]:
            row_cells = themes_table.add_row().cells
            row_cells[0].text = theme
            row_cells[1].text = str(count)
        
        doc.add_page_break()
        
        # 4. Análisis de Coherencia
        doc.add_heading('Análisis de Coherencia y Estructura', level=1)
        para = doc.add_paragraph()
        para.add_run(f"Score de Coherencia: {coherence['coherence_score']}/100").bold = True
        doc.add_paragraph(f"Calidad de estructura: {coherence['structure_quality']}")
        if coherence['issues']:
            doc.add_paragraph("Problemas detectados:")
            for issue in coherence['issues']:
                doc.add_paragraph(issue, style='List Bullet')
        
        doc.add_page_break()
        
        # 5. Análisis de Código
        doc.add_heading('Análisis de Complejidad de Código', level=1)
        doc.add_paragraph(f"Bloques de código: {code_analysis['code_blocks']}")
        doc.add_paragraph(f"Líneas de código totales: {code_analysis['total_code_lines']}")
        doc.add_paragraph(f"Funciones detectadas: {code_analysis['functions_detected']}")
        doc.add_paragraph(f"Clases detectadas: {code_analysis['classes_detected']}")
        doc.add_paragraph(f"Nivel de complejidad: {code_analysis['complexity_level']}")
        if code_analysis['languages_used']:
            doc.add_paragraph("Lenguajes utilizados:")
            for lang, count in code_analysis['languages_used'].items():
                doc.add_paragraph(f"  - {lang}: {count} bloques", style='List Bullet')
        
        doc.add_page_break()
        
        # 6. Análisis de Dependencias
        doc.add_heading('Análisis de Dependencias entre Secciones', level=1)
        doc.add_paragraph(f"Total de dependencias detectadas: {dependencies['total_dependencies']}")
        doc.add_paragraph(f"Secciones con dependencias: {dependencies['sections_with_dependencies']}")
        if dependencies['most_referenced_sections']:
            doc.add_paragraph("Secciones más referenciadas:")
            for section_idx, ref_count in dependencies['most_referenced_sections']:
                if section_idx < len(sections):
                    doc.add_paragraph(f"  - {sections[section_idx]['title']}: {ref_count} referencias", 
                                    style='List Bullet')
        
        # Guardar
        output_path = self.analysis_dir / f"{output_name}_ADVANCED_ANALYSIS.docx"
        doc.save(str(output_path))
        print(f"✓ Análisis avanzado guardado: {output_path}")
        
        # Guardar también en JSON
        analysis_data = {
            "metadata": {
                "document": output_name,
                "generated_at": datetime.now().isoformat()
            },
            "sentiment_analysis": sentiment,
            "theme_analysis": themes,
            "coherence_analysis": coherence,
            "auto_summary": summary,
            "code_analysis": code_analysis,
            "dependency_analysis": dependencies
        }
        
        json_path = self.analysis_dir / f"{output_name}_ADVANCED_ANALYSIS.json"
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(analysis_data, f, indent=2, ensure_ascii=False)
        
        print(f"✓ JSON de análisis avanzado guardado: {json_path}")
        
        return output_path, json_path
    
    def create_advanced_charts(self, all_analyses, output_name="ADVANCED"):
        """Crea gráficas avanzadas de análisis"""
        charts = []
        
        if not all_analyses:
            return charts
        
        # 1. Gráfica de sentimiento
        fig, ax = plt.subplots(figsize=(12, 7))
        doc_names = [a["name"] for a in all_analyses]
        sentiment_scores = [a["sentiment"]["sentiment_score"] for a in all_analyses]
        
        colors_sent = ['#2ecc71' if s > 0 else '#e74c3c' if s < 0 else '#f39c12' for s in sentiment_scores]
        bars = ax.bar(doc_names, sentiment_scores, color=colors_sent, edgecolor='black', linewidth=1.5)
        ax.axhline(y=0, color='black', linestyle='--', linewidth=1)
        ax.set_title('Análisis de Sentimiento por Documento', fontsize=16, fontweight='bold', pad=20)
        ax.set_ylabel('Score de Sentimiento (-100 a 100)', fontweight='bold')
        ax.set_xlabel('Documento', fontweight='bold')
        ax.grid(axis='y', alpha=0.3, linestyle='--')
        
        for bar, score in zip(bars, sentiment_scores):
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2., height,
                   f'{score}', ha='center', va='bottom' if height > 0 else 'top', fontweight='bold')
        
        plt.xticks(rotation=45, ha='right')
        plt.tight_layout()
        chart_path = self.temp_dir / f"{output_name}_sentiment.png"
        plt.savefig(chart_path, dpi=300, bbox_inches='tight', facecolor='white')
        plt.close()
        charts.append(("Análisis de Sentimiento", chart_path))
        
        # 2. Gráfica de coherencia
        fig, ax = plt.subplots(figsize=(12, 7))
        coherence_scores = [a["coherence"]["coherence_score"] for a in all_analyses]
        
        bars = ax.bar(doc_names, coherence_scores, color='#3498db', edgecolor='black', linewidth=1.5)
        ax.set_title('Análisis de Coherencia por Documento', fontsize=16, fontweight='bold', pad=20)
        ax.set_ylabel('Score de Coherencia (0-100)', fontweight='bold')
        ax.set_xlabel('Documento', fontweight='bold')
        ax.set_ylim(0, 100)
        ax.grid(axis='y', alpha=0.3, linestyle='--')
        
        for bar, score in zip(bars, coherence_scores):
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2., height,
                   f'{score}', ha='center', va='bottom', fontweight='bold')
        
        plt.xticks(rotation=45, ha='right')
        plt.tight_layout()
        chart_path = self.temp_dir / f"{output_name}_coherence.png"
        plt.savefig(chart_path, dpi=300, bbox_inches='tight', facecolor='white')
        plt.close()
        charts.append(("Análisis de Coherencia", chart_path))
        
        # 3. Gráfica de complejidad de código
        fig, ax = plt.subplots(figsize=(12, 7))
        code_complexity = [a["code"]["complexity_level"] for a in all_analyses]
        complexity_map = {"Alto": 3, "Medio": 2, "Bajo": 1}
        complexity_values = [complexity_map.get(c, 1) for c in code_complexity]
        
        bars = ax.bar(doc_names, complexity_values, color='#9b59b6', edgecolor='black', linewidth=1.5)
        ax.set_title('Nivel de Complejidad de Código', fontsize=16, fontweight='bold', pad=20)
        ax.set_ylabel('Nivel (1=Bajo, 2=Medio, 3=Alto)', fontweight='bold')
        ax.set_xlabel('Documento', fontweight='bold')
        ax.set_ylim(0, 3.5)
        ax.set_yticks([1, 2, 3])
        ax.set_yticklabels(['Bajo', 'Medio', 'Alto'])
        ax.grid(axis='y', alpha=0.3, linestyle='--')
        
        plt.xticks(rotation=45, ha='right')
        plt.tight_layout()
        chart_path = self.temp_dir / f"{output_name}_code_complexity.png"
        plt.savefig(chart_path, dpi=300, bbox_inches='tight', facecolor='white')
        plt.close()
        charts.append(("Complejidad de Código", chart_path))
        
        return charts

def main():
    """Función principal"""
    base_dir = Path("/Users/adan/Documents/documentos_blatam")
    production_dir = base_dir / "truthgpt_collected/integration_code/production_code"
    
    documents = [
        {
            "path": base_dir / "airflow_automation_prompt.md",
            "name": "Automation_Expert_Prompt"
        },
        {
            "path": production_dir / "ARCHITECTURE_IMPROVEMENTS.md",
            "name": "Architecture_Improvements"
        },
        {
            "path": production_dir / "REFACTORING_PLAN.md",
            "name": "Refactoring_Plan"
        }
    ]
    
    converter = AdvancedAnalysisConverter()
    
    print("=" * 70)
    print("ANÁLISIS AVANZADO DE DOCUMENTOS")
    print("=" * 70)
    print()
    
    all_analyses = []
    
    for doc in documents:
        if not doc["path"].exists():
            continue
        
        print(f"\n📄 Analizando: {doc['name']}")
        print("-" * 70)
        
        try:
            content = converter.read_markdown(doc["path"])
            sections = converter.parse_markdown_sections(content)
            
            doc_data = {
                "name": doc["name"],
                "path": doc["path"],
                "sections": sections
            }
            
            # Realizar análisis avanzado
            sentiment = converter.analyze_sentiment_tone(content)
            themes = converter.detect_main_themes(content, sections)
            coherence = converter.analyze_coherence(sections)
            summary = converter.generate_auto_summary(content, sections)
            code_analysis = converter.analyze_code_complexity(content)
            dependencies = converter.analyze_dependencies(sections)
            
            all_analyses.append({
                "name": doc["name"],
                "sentiment": sentiment,
                "themes": themes,
                "coherence": coherence,
                "summary": summary,
                "code": code_analysis,
                "dependencies": dependencies
            })
            
            # Crear reporte
            converter.create_advanced_analysis_report(doc_data, doc["name"])
            
        except Exception as e:
            print(f"❌ Error: {str(e)}")
            import traceback
            traceback.print_exc()
    
    # Crear gráficas comparativas avanzadas
    if len(all_analyses) >= 2:
        print("\n" + "=" * 70)
        print("CREANDO GRÁFICAS AVANZADAS")
        print("=" * 70)
        charts = converter.create_advanced_charts(all_analyses)
        print(f"✓ {len(charts)} gráficas avanzadas creadas")
    
    print("\n" + "=" * 70)
    print("ANÁLISIS AVANZADO COMPLETADO")
    print("=" * 70)
    print(f"\nArchivos generados en: {converter.analysis_dir}")
    print("\n✅ Análisis avanzado completado!")

if __name__ == "__main__":
    main()



