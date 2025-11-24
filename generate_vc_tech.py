import docx
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_tech_brief(filename):
    doc = docx.Document()
    
    # Styles
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Courier New'
    h1.font.size = Pt(14)
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.font.bold = True
    
    body = doc.styles['Normal']
    body.font.name = 'Calibri'

    # Header
    p = doc.add_paragraph("TECHNICAL DUE DILIGENCE BRIEF")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(18)
    
    doc.add_paragraph("Version 1.2 | Status: Production | Compliance: SOC2 Readiness").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("_" * 80)

    # 1. Architecture Overview
    doc.add_heading("1. SYSTEM ARCHITECTURE", 1)
    doc.add_paragraph("Cloud-native microservices architecture hosted on AWS (us-east-1).")
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    rows = table.rows
    rows[0].cells[0].text = "Frontend"
    rows[0].cells[1].text = "React 18, TypeScript, TailwindCSS (Hosted on Vercel)"
    rows[1].cells[0].text = "Backend API"
    rows[1].cells[1].text = "Python FastAPI (Async), running on AWS ECS (Fargate)"
    rows[2].cells[0].text = "AI Inference"
    rows[2].cells[1].text = "Ray Serve clusters on GPU instances (g5.xlarge)"
    rows[3].cells[0].text = "Database"
    rows[3].cells[1].text = "PostgreSQL (RDS) for relational data, Pinecone for Vector DB"
    rows[4].cells[0].text = "Queue/Async"
    rows[4].cells[1].text = "Redis & Celery for long-running generation tasks"

    # 2. AI Pipeline Details
    doc.add_heading("2. AI MODEL PIPELINE (The 'Secret Sauce')", 1)
    doc.add_paragraph("Our proprietary pipeline consists of three stages:")
    
    doc.add_paragraph("1. Intent Classification:", style='List Number').bold = True
    doc.add_paragraph("   BERT-based model classifies user intent (e.g., 'Sales Email', 'Insta Caption').")
    
    doc.add_paragraph("2. Context Retrieval (RAG):", style='List Number').bold = True
    doc.add_paragraph("   Retrieves relevant cultural context and brand guidelines from Vector DB.")
    
    doc.add_paragraph("3. Generation & Polishing:", style='List Number').bold = True
    doc.add_paragraph("   Fine-tuned Llama-3-70b generates content. A secondary 'Critic Model' evaluates output against local dialect rules (e.g., 'voseo' in Argentina).")

    # 3. Security & Data Privacy
    doc.add_heading("3. SECURITY & COMPLIANCE", 1)
    doc.add_paragraph("• Data Encryption: TLS 1.3 in transit, AES-256 at rest.")
    doc.add_paragraph("• PII Handling: Automatic redaction of Personally Identifiable Information before sending to LLMs.")
    doc.add_paragraph("• SOC2: Currently in observation period (Drata integration active). Target certification: Q2 2025.")

    # 4. Scalability
    doc.add_heading("4. SCALABILITY METRICS", 1)
    doc.add_paragraph("• Current Throughput: 50k requests/day.")
    doc.add_paragraph("• Stress Test Limit: 1M requests/day (validated).")
    doc.add_paragraph("• Auto-scaling: KEDA policies scale pods based on queue latency.")

    doc.save(filename)
    print(f"Tech Brief created: {filename}")

if __name__ == "__main__":
    create_tech_brief("Ventura_Capital_Tech_Brief.docx")


