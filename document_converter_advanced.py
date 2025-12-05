#!/usr/bin/env python3
"""
Sistema Avanzado de Conversión de Documentos
=============================================

Módulo completo para convertir documentos a PDF, Word y Excel
con todas las librerías disponibles y mejoradas.

Librerías soportadas:
- PDF: reportlab, fpdf, weasyprint, pdfkit, pypdf, PyMuPDF, xhtml2pdf
- Word: python-docx, docx2pdf, python-docx2txt, mammoth
- Excel: openpyxl, xlsxwriter, pandas, xlrd, xlwt, pyexcel

Autor: Sistema de Conversión Avanzado
Versión: 3.0.0
"""

import os
import sys
import subprocess
import shutil
from pathlib import Path
from typing import Dict, List, Optional, Union, Any, Tuple
from datetime import datetime
import json
import logging
from enum import Enum

# Configurar logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class FormatType(Enum):
    """Tipos de formato soportados"""
    PDF = "pdf"
    WORD = "word"
    EXCEL = "excel"
    HTML = "html"
    MARKDOWN = "markdown"


class DocumentConverterAdvanced:
    """
    Convertidor avanzado de documentos con múltiples librerías
    """
    
    def __init__(self, output_dir: Optional[str] = None):
        """
        Inicializa el convertidor
        
        Args:
            output_dir: Directorio de salida (opcional)
        """
        self.output_dir = output_dir or "output_documents"
        Path(self.output_dir).mkdir(parents=True, exist_ok=True)
        
        # Detectar librerías disponibles
        self.available_libraries = self._detect_libraries()
        logger.info(f"Librerías detectadas: {self.available_libraries}")
    
    def _detect_libraries(self) -> Dict[str, bool]:
        """Detecta qué librerías están instaladas"""
        libraries = {
            # PDF
            'reportlab': False,
            'fpdf': False,
            'weasyprint': False,
            'pdfkit': False,
            'pypdf': False,
            'PyMuPDF': False,
            'xhtml2pdf': False,
            'docx2pdf': False,
            # Word
            'python-docx': False,
            'mammoth': False,
            # Excel
            'openpyxl': False,
            'xlsxwriter': False,
            'pandas': False,
            'xlrd': False,
            'xlwt': False,
            'pyexcel': False,
            # Utilidades
            'markdown': False,
            'Pillow': False,
            'matplotlib': False,
            'pandas': False,
        }
        
        for lib in libraries.keys():
            try:
                if lib == 'python-docx':
                    import docx
                    libraries[lib] = True
                elif lib == 'PyMuPDF':
                    import fitz
                    libraries[lib] = True
                elif lib == 'Pillow':
                    import PIL
                    libraries[lib] = True
                else:
                    __import__(lib.replace('-', '_'))
                    libraries[lib] = True
            except ImportError:
                pass
        
        return libraries
    
    def convert_to_pdf(self, 
                       input_file: str,
                       output_file: Optional[str] = None,
                       method: Optional[str] = None) -> Optional[str]:
        """
        Convierte un archivo a PDF usando múltiples métodos
        
        Args:
            input_file: Archivo de entrada
            output_file: Archivo de salida (opcional)
            method: Método específico a usar (opcional)
        
        Returns:
            Ruta del archivo PDF creado o None
        """
        input_path = Path(input_file)
        if not input_path.exists():
            logger.error(f"Archivo no encontrado: {input_file}")
            return None
        
        if output_file is None:
            output_file = str(Path(self.output_dir) / f"{input_path.stem}.pdf")
        
        # Determinar método basado en extensión
        ext = input_path.suffix.lower()
        
        # Intentar diferentes métodos según disponibilidad
        methods = []
        
        if ext == '.docx':
            methods = ['docx2pdf', 'weasyprint', 'reportlab', 'libreoffice']
        elif ext == '.html':
            methods = ['weasyprint', 'pdfkit', 'xhtml2pdf', 'reportlab']
        elif ext == '.md':
            methods = ['weasyprint', 'pdfkit', 'reportlab']
        else:
            methods = ['reportlab', 'fpdf', 'weasyprint']
        
        # Si se especifica un método, usarlo primero
        if method and method in methods:
            methods = [method] + [m for m in methods if m != method]
        
        for method_name in methods:
            try:
                result = self._convert_to_pdf_method(input_file, output_file, method_name)
                if result:
                    logger.info(f"PDF creado exitosamente usando {method_name}: {output_file}")
                    return result
            except Exception as e:
                logger.warning(f"Error con método {method_name}: {e}")
                continue
        
        logger.error("No se pudo convertir a PDF con ningún método disponible")
        return None
    
    def _convert_to_pdf_method(self, 
                               input_file: str,
                               output_file: str,
                               method: str) -> Optional[str]:
        """Convierte usando un método específico"""
        
        if method == 'docx2pdf':
            if not self.available_libraries.get('docx2pdf'):
                raise ImportError("docx2pdf no está instalado")
            from docx2pdf import convert
            convert(input_file, output_file)
            return output_file
        
        elif method == 'weasyprint':
            if not self.available_libraries.get('weasyprint'):
                raise ImportError("weasyprint no está instalado")
            from weasyprint import HTML
            if Path(input_file).suffix == '.html':
                HTML(filename=input_file).write_pdf(output_file)
            else:
                # Convertir a HTML primero
                html_content = self._to_html(input_file)
                HTML(string=html_content).write_pdf(output_file)
            return output_file
        
        elif method == 'reportlab':
            if not self.available_libraries.get('reportlab'):
                raise ImportError("reportlab no está instalado")
            return self._convert_with_reportlab(input_file, output_file)
        
        elif method == 'pdfkit':
            if not self.available_libraries.get('pdfkit'):
                raise ImportError("pdfkit no está instalado")
            import pdfkit
            if Path(input_file).suffix == '.html':
                pdfkit.from_file(input_file, output_file)
            else:
                html_content = self._to_html(input_file)
                pdfkit.from_string(html_content, output_file)
            return output_file
        
        elif method == 'xhtml2pdf':
            if not self.available_libraries.get('xhtml2pdf'):
                raise ImportError("xhtml2pdf no está instalado")
            from xhtml2pdf import pisa
            html_content = self._to_html(input_file)
            with open(output_file, 'wb') as f:
                pisa.CreatePDF(html_content, dest=f)
            return output_file
        
        elif method == 'libreoffice':
            # Usar LibreOffice desde línea de comandos
            if shutil.which('libreoffice'):
                subprocess.run([
                    'libreoffice', '--headless', '--convert-to', 'pdf',
                    '--outdir', str(Path(output_file).parent),
                    input_file
                ], check=True)
                # LibreOffice genera el nombre automáticamente
                expected_file = Path(output_file).parent / f"{Path(input_file).stem}.pdf"
                if expected_file.exists():
                    if expected_file != Path(output_file):
                        shutil.move(str(expected_file), output_file)
                    return output_file
            raise FileNotFoundError("LibreOffice no está instalado")
        
        elif method == 'fpdf':
            if not self.available_libraries.get('fpdf'):
                raise ImportError("fpdf no está instalado")
            return self._convert_with_fpdf(input_file, output_file)
        
        else:
            raise ValueError(f"Método desconocido: {method}")
    
    def _convert_with_reportlab(self, input_file: str, output_file: str) -> str:
        """Convierte usando ReportLab"""
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
        
        doc = SimpleDocTemplate(output_file, pagesize=A4)
        story = []
        styles = getSampleStyleSheet()
        
        # Leer contenido
        content = self._read_file(input_file)
        
        # Procesar contenido
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                story.append(Spacer(1, 0.2*inch))
            elif line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                style = styles['Heading' + str(min(level, 6))]
                text = line.lstrip('#').strip()
                story.append(Paragraph(text, style))
            else:
                story.append(Paragraph(line, styles['Normal']))
        
        doc.build(story)
        return output_file
    
    def _convert_with_fpdf(self, input_file: str, output_file: str) -> str:
        """Convierte usando FPDF"""
        from fpdf import FPDF
        
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        
        content = self._read_file(input_file)
        for line in content.split('\n'):
            pdf.cell(0, 10, txt=line, ln=1)
        
        pdf.output(output_file)
        return output_file
    
    def convert_to_word(self,
                       input_file: str,
                       output_file: Optional[str] = None,
                       style: str = 'professional') -> Optional[str]:
        """
        Convierte un archivo a Word (.docx)
        
        Args:
            input_file: Archivo de entrada
            output_file: Archivo de salida (opcional)
            style: Estilo del documento ('professional', 'simple', 'premium')
        
        Returns:
            Ruta del archivo Word creado o None
        """
        if not self.available_libraries.get('python-docx'):
            logger.error("python-docx no está instalado. Instale con: pip install python-docx")
            return None
        
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        
        input_path = Path(input_file)
        if not input_path.exists():
            logger.error(f"Archivo no encontrado: {input_file}")
            return None
        
        if output_file is None:
            output_file = str(Path(self.output_dir) / f"{input_path.stem}.docx")
        
        try:
            doc = Document()
            
            # Configurar estilos según el tipo
            if style == 'professional':
                self._setup_professional_word(doc)
            elif style == 'premium':
                self._setup_premium_word(doc)
            else:
                self._setup_simple_word(doc)
            
            # Leer y procesar contenido
            content = self._read_file(input_file)
            
            # Procesar según extensión
            ext = input_path.suffix.lower()
            if ext == '.md':
                self._process_markdown_to_word(doc, content)
            elif ext == '.html':
                self._process_html_to_word(doc, content)
            else:
                self._process_text_to_word(doc, content)
            
            # Guardar
            doc.save(output_file)
            logger.info(f"Documento Word creado: {output_file}")
            return output_file
        
        except Exception as e:
            logger.error(f"Error creando Word: {e}", exc_info=True)
            return None
    
    def _setup_professional_word(self, doc):
        """Configura estilos profesionales para Word"""
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Configurar márgenes
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3)
            section.right_margin = Cm(3)
        
        # Crear estilos personalizados
        styles = doc.styles
        
        # Título principal
        if 'Custom Title' not in [s.name for s in styles]:
            title_style = styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
            title_format = title_style.font
            title_format.name = 'Calibri'
            title_format.size = Pt(24)
            title_format.bold = True
            title_format.color = RGBColor(0, 51, 102)
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)
    
    def _setup_premium_word(self, doc):
        """Configura estilos premium para Word"""
        self._setup_professional_word(doc)
        # Agregar portada, tabla de contenidos, etc.
    
    def _setup_simple_word(self, doc):
        """Configura estilos simples para Word"""
        pass
    
    def _process_markdown_to_word(self, doc, content: str):
        """Procesa Markdown y lo convierte a Word"""
        if self.available_libraries.get('markdown'):
            import markdown
            html = markdown.markdown(content, extensions=['tables', 'fenced_code'])
            self._process_html_to_word(doc, html)
        else:
            # Procesamiento básico sin markdown
            self._process_text_to_word(doc, content)
    
    def _process_html_to_word(self, doc, html: str):
        """Procesa HTML y lo convierte a Word"""
        from bs4 import BeautifulSoup
        try:
            soup = BeautifulSoup(html, 'html.parser')
            for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'table']):
                if element.name.startswith('h'):
                    level = int(element.name[1])
                    heading = doc.add_heading(element.get_text(), level=level)
                elif element.name == 'p':
                    doc.add_paragraph(element.get_text())
                elif element.name in ['ul', 'ol']:
                    for li in element.find_all('li'):
                        doc.add_paragraph(li.get_text(), style='List Bullet' if element.name == 'ul' else 'List Number')
        except ImportError:
            # Fallback si BeautifulSoup no está disponible
            self._process_text_to_word(doc, html)
    
    def _process_text_to_word(self, doc, content: str):
        """Procesa texto plano y lo convierte a Word"""
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                continue
            elif line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                doc.add_heading(line.lstrip('#').strip(), level=min(level, 6))
            else:
                doc.add_paragraph(line)
    
    def convert_to_excel(self,
                        input_file: str,
                        output_file: Optional[str] = None,
                        style: str = 'professional') -> Optional[str]:
        """
        Convierte un archivo a Excel (.xlsx)
        
        Args:
            input_file: Archivo de entrada
            output_file: Archivo de salida (opcional)
            style: Estilo del documento ('professional', 'simple', 'premium')
        
        Returns:
            Ruta del archivo Excel creado o None
        """
        # Intentar openpyxl primero
        if self.available_libraries.get('openpyxl'):
            return self._convert_to_excel_openpyxl(input_file, output_file, style)
        elif self.available_libraries.get('xlsxwriter'):
            return self._convert_to_excel_xlsxwriter(input_file, output_file, style)
        elif self.available_libraries.get('pandas'):
            return self._convert_to_excel_pandas(input_file, output_file, style)
        else:
            logger.error("No hay librerías de Excel disponibles. Instale: pip install openpyxl")
            return None
    
    def _convert_to_excel_openpyxl(self,
                                   input_file: str,
                                   output_file: Optional[str],
                                   style: str) -> Optional[str]:
        """Convierte usando openpyxl"""
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
        
        input_path = Path(input_file)
        if output_file is None:
            output_file = str(Path(self.output_dir) / f"{input_path.stem}.xlsx")
        
        try:
            wb = Workbook()
            ws = wb.active
            ws.title = "Contenido"
            
            # Leer contenido
            content = self._read_file(input_file)
            
            # Procesar contenido
            if input_path.suffix.lower() == '.md':
                self._process_markdown_to_excel(ws, content, style)
            else:
                self._process_text_to_excel(ws, content, style)
            
            # Ajustar ancho de columnas
            for column in ws.columns:
                max_length = 0
                column_letter = get_column_letter(column[0].column)
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                ws.column_dimensions[column_letter].width = adjusted_width
            
            wb.save(output_file)
            logger.info(f"Documento Excel creado: {output_file}")
            return output_file
        
        except Exception as e:
            logger.error(f"Error creando Excel: {e}", exc_info=True)
            return None
    
    def _convert_to_excel_xlsxwriter(self,
                                    input_file: str,
                                    output_file: Optional[str],
                                    style: str) -> Optional[str]:
        """Convierte usando xlsxwriter"""
        import xlsxwriter
        
        input_path = Path(input_file)
        if output_file is None:
            output_file = str(Path(self.output_dir) / f"{input_path.stem}.xlsx")
        
        try:
            workbook = xlsxwriter.Workbook(output_file)
            worksheet = workbook.add_worksheet('Contenido')
            
            content = self._read_file(input_file)
            row = 0
            
            for line in content.split('\n'):
                line = line.strip()
                if line:
                    worksheet.write(row, 0, line)
                    row += 1
            
            workbook.close()
            logger.info(f"Documento Excel creado: {output_file}")
            return output_file
        
        except Exception as e:
            logger.error(f"Error creando Excel: {e}", exc_info=True)
            return None
    
    def _convert_to_excel_pandas(self,
                                 input_file: str,
                                 output_file: Optional[str],
                                 style: str) -> Optional[str]:
        """Convierte usando pandas"""
        import pandas as pd
        
        input_path = Path(input_file)
        if output_file is None:
            output_file = str(Path(self.output_dir) / f"{input_path.stem}.xlsx")
        
        try:
            content = self._read_file(input_file)
            lines = [line.strip() for line in content.split('\n') if line.strip()]
            df = pd.DataFrame({'Contenido': lines})
            df.to_excel(output_file, index=False, engine='openpyxl')
            logger.info(f"Documento Excel creado: {output_file}")
            return output_file
        
        except Exception as e:
            logger.error(f"Error creando Excel: {e}", exc_info=True)
            return None
    
    def _process_markdown_to_excel(self, ws, content: str, style: str):
        """Procesa Markdown y lo convierte a Excel"""
        from openpyxl.styles import Font, PatternFill, Alignment
        
        row = 1
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                continue
            
            cell = ws.cell(row=row, column=1, value=line)
            
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                cell.value = line.lstrip('#').strip()
                cell.font = Font(bold=True, size=14 - level)
                if style == 'professional':
                    cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
                    cell.font = Font(bold=True, color="FFFFFF", size=14 - level)
            
            row += 1
    
    def _process_text_to_excel(self, ws, content: str, style: str):
        """Procesa texto y lo convierte a Excel"""
        row = 1
        for line in content.split('\n'):
            line = line.strip()
            if line:
                ws.cell(row=row, column=1, value=line)
                row += 1
    
    def convert_all_formats(self,
                           input_file: str,
                           formats: List[str] = ['pdf', 'word', 'excel'],
                           **kwargs) -> Dict[str, Optional[str]]:
        """
        Convierte un archivo a múltiples formatos
        
        Args:
            input_file: Archivo de entrada
            formats: Lista de formatos ['pdf', 'word', 'excel']
            **kwargs: Argumentos adicionales para cada formato
        
        Returns:
            Diccionario con rutas de archivos creados
        """
        results = {}
        
        for fmt in formats:
            try:
                if fmt.lower() == 'pdf':
                    results['pdf'] = self.convert_to_pdf(input_file, **kwargs.get('pdf', {}))
                elif fmt.lower() in ['word', 'docx']:
                    results['word'] = self.convert_to_word(input_file, **kwargs.get('word', {}))
                elif fmt.lower() in ['excel', 'xlsx']:
                    results['excel'] = self.convert_to_excel(input_file, **kwargs.get('excel', {}))
            except Exception as e:
                logger.error(f"Error convirtiendo a {fmt}: {e}")
                results[fmt] = None
        
        return results
    
    def _to_html(self, input_file: str) -> str:
        """Convierte un archivo a HTML"""
        content = self._read_file(input_file)
        ext = Path(input_file).suffix.lower()
        
        if ext == '.md':
            if self.available_libraries.get('markdown'):
                import markdown
                return markdown.markdown(content, extensions=['tables', 'fenced_code'])
            else:
                # Conversión básica
                html = f"<html><body><pre>{content}</pre></body></html>"
                return html
        elif ext == '.html':
            return content
        else:
            # Texto plano
            return f"<html><body><pre>{content}</pre></body></html>"
    
    def _read_file(self, file_path: str) -> str:
        """Lee un archivo con diferentes codificaciones"""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        raise ValueError(f"No se pudo leer el archivo: {file_path}")
    
    def get_requirements(self) -> List[str]:
        """Obtiene lista de todas las librerías recomendadas"""
        return [
            # PDF
            'reportlab>=4.0.0',
            'fpdf2>=2.7.0',
            'weasyprint>=60.0',
            'pdfkit>=1.0.0',
            'pypdf>=3.0.0',
            'PyMuPDF>=1.23.0',
            'xhtml2pdf>=0.2.11',
            'docx2pdf>=0.1.8',
            # Word
            'python-docx>=1.1.0',
            'mammoth>=1.6.0',
            # Excel
            'openpyxl>=3.1.0',
            'xlsxwriter>=3.1.0',
            'pandas>=2.0.0',
            'xlrd>=2.0.1',
            'xlwt>=1.3.0',
            # Utilidades
            'markdown>=3.4.0',
            'Pillow>=10.0.0',
            'matplotlib>=3.7.0',
            'beautifulsoup4>=4.12.0',
        ]
    
    def install_requirements(self):
        """Instala todas las librerías recomendadas"""
        requirements = self.get_requirements()
        for req in requirements:
            try:
                subprocess.check_call([sys.executable, '-m', 'pip', 'install', req])
                logger.info(f"Instalado: {req}")
            except subprocess.CalledProcessError as e:
                logger.warning(f"Error instalando {req}: {e}")


def main():
    """Función principal para uso desde línea de comandos"""
    import argparse
    
    parser = argparse.ArgumentParser(description='Convertidor avanzado de documentos')
    parser.add_argument('input_file', help='Archivo de entrada')
    parser.add_argument('-f', '--formats', nargs='+', 
                       choices=['pdf', 'word', 'excel', 'all'],
                       default=['all'],
                       help='Formatos de salida')
    parser.add_argument('-o', '--output-dir', default='output_documents',
                       help='Directorio de salida')
    parser.add_argument('--style', choices=['simple', 'professional', 'premium'],
                       default='professional',
                       help='Estilo del documento')
    parser.add_argument('--install-deps', action='store_true',
                       help='Instalar dependencias automáticamente')
    
    args = parser.parse_args()
    
    # Instalar dependencias si se solicita
    if args.install_deps:
        converter = DocumentConverterAdvanced()
        converter.install_requirements()
        return
    
    # Crear convertidor
    converter = DocumentConverterAdvanced(output_dir=args.output_dir)
    
    # Determinar formatos
    if 'all' in args.formats:
        formats = ['pdf', 'word', 'excel']
    else:
        formats = args.formats
    
    # Convertir
    results = converter.convert_all_formats(
        args.input_file,
        formats=formats,
        word={'style': args.style},
        excel={'style': args.style}
    )
    
    # Mostrar resultados
    print("\n" + "="*70)
    print("RESULTADOS DE CONVERSIÓN")
    print("="*70)
    for fmt, path in results.items():
        if path:
            print(f"✅ {fmt.upper()}: {path}")
        else:
            print(f"❌ {fmt.upper()}: Error en la conversión")
    print("="*70)


if __name__ == "__main__":
    main()

