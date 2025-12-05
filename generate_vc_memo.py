import docx
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_investment_memo(filename):
    doc = docx.Document()
    
    # Layout
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    
    # Styles
    style_h1 = doc.styles['Heading 1']
    style_h1.font.name = 'Calibri'
    style_h1.font.size = Pt(18)
    style_h1.font.color.rgb = RGBColor(32, 55, 100)
    style_h1.font.bold = True

    style_h2 = doc.styles['Heading 2']
    style_h2.font.name = 'Calibri'
    style_h2.font.size = Pt(14)
    style_h2.font.color.rgb = RGBColor(47, 117, 181)
    style_h2.font.bold = True

    # Header
    p = doc.add_paragraph()
    run = p.add_run("INVESTMENT MEMORANDUM")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(128, 128, 128)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph("To: Investment Committee\nFrom: Deal Team\nDate: December 2024\nSubject: Series A Investment in Ventura Capital").bold = True
    doc.add_paragraph("_" * 80)

    # 1. Executive Summary
    doc.add_heading("1. Executive Summary", 1)
    doc.add_paragraph("Ventura Capital es una plataforma SaaS B2B de IA Generativa especializada en la creación de contenido de marketing para el mercado latinoamericano. La compañía ha demostrado una fuerte tracción temprana con $139k de ARR, creciendo 20% mes a mes, y unit economics saludables (LTV/CAC 9:1).")
    doc.add_paragraph("Recomendación: Inversión de $2.0M a una valuación pre-money de $8.0M para capturar el liderazgo en un mercado de $2.8B.")

    # 2. Market Opportunity
    doc.add_heading("2. Market Opportunity", 1)
    doc.add_heading("The Problem", 2)
    doc.add_paragraph("Las herramientas actuales (Jasper, Copy.ai) son anglocéntricas y carecen de matices culturales locales, resultando en contenido genérico que no convierte en LATAM. Las agencias tradicionales son lentas y costosas.")
    
    doc.add_heading("The Solution", 2)
    doc.add_paragraph("Modelos de lenguaje (LLMs) fine-tuned con datasets propietarios de dialectos regionales (español mexicano, rioplatense, andino, etc.) y contextos culturales específicos.")

    doc.add_heading("Market Size", 2)
    doc.add_paragraph("• TAM: $2.8B (Gasto en MarTech y Contenido Digital en LATAM)")
    doc.add_paragraph("• SAM: $280M (Empresas Digitales y E-commerce)")
    doc.add_paragraph("• SOM: $28M (Objetivo 3 años)")

    # 3. Product & Technology
    doc.add_heading("3. Product & Technology", 1)
    doc.add_paragraph("Tech Stack: Python/FastAPI backend, React frontend, PyTorch para inferencia de modelos. Infraestructura en AWS (GPU instances).")
    doc.add_paragraph("Moat Tecnológico: Dataset propietario de 50M+ ejemplos de copy de alta conversión en LATAM, inalcanzable para competidores globales sin esfuerzo manual masivo.")

    # 4. Go-to-Market Strategy
    doc.add_heading("4. Go-to-Market Strategy", 1)
    doc.add_paragraph("• Direct Sales: Equipo de ventas outbound enfocado en Agencias de Marketing (ticket alto).")
    doc.add_paragraph("• Product-Led Growth (PLG): Freemium tier para solopreneurs y creadores de contenido.")
    doc.add_paragraph("• Partnerships: Integración con TiendaNube y MercadoLibre para adquirir e-commerce merchants.")

    # 5. Financials
    doc.add_heading("5. Financial Plan", 1)
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Light Shading Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "2024 (Act)"
    hdr[2].text = "2025 (Proj)"
    hdr[3].text = "2026 (Proj)"
    
    vals = [
        ("ARR", "$139k", "$1.46M", "$3.6M"),
        ("EBITDA", "($865k)", "($350k)", "$1.6M"),
        ("Customers", "500", "2,500", "8,000")
    ]
    
    for m, v1, v2, v3 in vals:
        row = table.add_row().cells
        row[0].text = m
        row[1].text = v1
        row[2].text = v2
        row[3].text = v3

    # 6. Risks
    doc.add_heading("6. Key Risks & Mitigations", 1)
    doc.add_paragraph("• Riesgo de Plataforma: Dependencia de modelos base (GPT/Llama). Mitigación: Desarrollo de modelos propios verticales.")
    doc.add_paragraph("• Competencia: Entrada agresiva de jugadores US. Mitigación: Fuerte efecto de red de datos locales y alianzas regionales.")

    doc.save(filename)
    print(f"Investment Memo created: {filename}")

if __name__ == "__main__":
    create_investment_memo("Ventura_Capital_Investment_Memo.docx")







