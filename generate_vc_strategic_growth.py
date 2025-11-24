import docx
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_strategic_docs():
    # --- Document 1: Exit Strategy & M&A ---
    create_exit_strategy("Ventura_Capital_Exit_Strategy.docx")
    
    # --- Document 2: GTM Playbook ---
    create_gtm_playbook("Ventura_Capital_GTM_Playbook.docx")

def setup_styles(doc):
    # Title Style
    styles = doc.styles
    if 'VC Title' not in styles:
        style = styles.add_style('VC Title', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(26)
        font.bold = True
        font.color.rgb = RGBColor(32, 55, 100) # Navy

    if 'VC Heading 1' not in styles:
        h1 = styles.add_style('VC Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        h1.base_style = styles['Heading 1']
        font = h1.font
        font.name = 'Calibri'
        font.size = Pt(18)
        font.bold = True
        font.color.rgb = RGBColor(47, 117, 181) # Blue

def create_exit_strategy(filename):
    doc = docx.Document()
    setup_styles(doc)
    
    # Header
    doc.add_paragraph("VENTURA CAPITAL", style='VC Title').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("EXIT STRATEGY & M&A LANDSCAPE", style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("_" * 80).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. Strategic Rationale
    doc.add_paragraph("1. Why We Are an Attractive Target", style='VC Heading 1')
    doc.add_paragraph("Ventura Capital representa un activo estratégico único para compradores globales que buscan penetración inmediata en el mercado latinoamericano. Nuestros 'Data Moats' culturales son difíciles de replicar sin años de recolección de datos.")
    
    # 2. Potential Acquirers
    doc.add_paragraph("2. Buyer Universe (Tier 1)", style='VC Heading 1')
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light List Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = "Acquirer Category"
    hdr[1].text = "Potential Buyers"
    hdr[2].text = "Strategic Rationale"
    
    buyers = [
        ("Global MarTech", "HubSpot, Salesforce, Adobe", "Necesidad de localizar sus nubes de marketing para capturar el mercado SMB de LATAM."),
        ("Regional Tech Giants", "MercadoLibre, Globant", "Integración vertical de herramientas para sus ecosistemas de vendedores (MeLi) o clientes enterprise."),
        ("Generative AI Platforms", "Jasper, Canva", "Expansión geográfica rápida y adquisición de talento/datos locales.")
    ]
    
    for cat, buyer, rat in buyers:
        row = table.add_row().cells
        row[0].text = cat
        row[1].text = buyer
        row[2].text = rat

    # 3. Comparable Transactions
    doc.add_paragraph("3. Comparable Transactions (Comps)", style='VC Heading 1')
    doc.add_paragraph("• Writesonic (2023): Valued at ~$100M (20x ARR). Focus on SEO content.")
    doc.add_paragraph("• Copy.ai (2022): Series A at $10M ARR. Strong PLG motion similar to Ventura.")
    
    # 4. Return Analysis
    doc.add_paragraph("4. Projected Returns (Series A Investors)", style='VC Heading 1')
    doc.add_paragraph("Based on a $10M Post-Money Valuation today:")
    
    doc.add_paragraph("Scenario A: Base Case Exit (2027)", style='Heading 3')
    doc.add_paragraph("• Exit Value: $50M (5x Revenue)")
    doc.add_paragraph("• MOIC: 5.0x")
    doc.add_paragraph("• IRR: ~70%")

    doc.add_paragraph("Scenario B: Bull Case Exit (2028)", style='Heading 3')
    doc.add_paragraph("• Exit Value: $150M (10x Revenue)")
    doc.add_paragraph("• MOIC: 15.0x")
    doc.add_paragraph("• IRR: ~110%")

    doc.save(filename)
    print(f"Created: {filename}")

def create_gtm_playbook(filename):
    doc = docx.Document()
    setup_styles(doc)

    # Header
    doc.add_paragraph("VENTURA CAPITAL", style='VC Title').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("GTM & SALES PLAYBOOK 2025", style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("_" * 80).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. Strategy Overview
    doc.add_paragraph("1. The 'Pincer' Strategy", style='VC Heading 1')
    doc.add_paragraph("Nuestra estrategia de Go-To-Market ataca el mercado desde dos frentes simultáneos para maximizar la cobertura y eficiencia del CAC.")
    
    doc.add_paragraph("Top-Down (Direct Sales):", style='Heading 3')
    doc.add_paragraph("Enfoque en Agencias Digitales y E-commerce Enterprise (>100 SKUs). Ticket promedio: $12k/año.")
    
    doc.add_paragraph("Bottom-Up (PLG):", style='Heading 3')
    doc.add_paragraph("Freemium para creadores y PyMEs. Viralidad a través de marcas de agua y 'shareable assets'. Conversión a Pro Plan ($29/mo).")

    # 2. Sales Org Structure
    doc.add_paragraph("2. Sales Organization Ramp (Use of Funds)", style='VC Heading 1')
    
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    rows = table.rows
    rows[0].cells[0].text = "Role"
    rows[0].cells[1].text = "Q1 2025 Hires"
    rows[0].cells[2].text = "Key Responsibility"
    
    rows[1].cells[0].text = "Head of Sales"
    rows[1].cells[1].text = "1"
    rows[1].cells[2].text = "Build playbook, hire SDRs, close first 10 Enterprise deals."
    
    rows[2].cells[0].text = "SDRs (Outbound)"
    rows[2].cells[1].text = "2"
    rows[2].cells[2].text = "Cold outreach to Agencias Marketing (List building + Email/LinkedIn)."
    
    rows[3].cells[0].text = "Growth Marketer"
    rows[3].cells[1].text = "1"
    rows[3].cells[2].text = "Optimize PLG funnel, SEO content, and Paid Ads (Meta/LinkedIn)."

    # 3. Channel Metrics
    doc.add_paragraph("3. Unit Economics Targets", style='VC Heading 1')
    doc.add_paragraph("• Enterprise CAC Target: $2,500")
    doc.add_paragraph("• Enterprise LTV Target: $30,000 (12:1 ratio)")
    doc.add_paragraph("• PLG CAC Target: $25")
    doc.add_paragraph("• PLG LTV Target: $350 (14:1 ratio)")

    doc.save(filename)
    print(f"Created: {filename}")

if __name__ == "__main__":
    create_strategic_docs()


