#!/usr/bin/env python3
"""
Generador Premium de Documentos Importantes
===========================================
Convierte los documentos más importantes a PDF, Word y Excel
con gráficas de alta calidad y formato profesional.

Documentos procesados:
- airflow_automation_prompt.md
- ARCHITECTURE_IMPROVEMENTS.md
- REFACTORING_PLAN.md
"""

import os
import sys
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Tuple
import re
import json

# Importar librerías de conversión
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Instalando python-docx...")
    os.system("pip install python-docx")
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    import openpyxl
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.chart import BarChart, LineChart, PieChart, Reference
    from openpyxl.drawing.image import Image
except ImportError:
    print("Instalando openpyxl...")
    os.system("pip install openpyxl")
    import openpyxl
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.chart import BarChart, LineChart, PieChart, Reference

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image as RLImage
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
except ImportError:
    print("Instalando reportlab...")
    os.system("pip install reportlab")
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image as RLImage
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY

try:
    import matplotlib
    matplotlib.use('Agg')  # Backend sin GUI
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    from matplotlib.backends.backend_pdf import PdfPages
    import seaborn as sns
    sns.set_style("whitegrid")
    sns.set_palette("husl")
except ImportError:
    print("Instalando matplotlib y seaborn...")
    os.system("pip install matplotlib seaborn")
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import seaborn as sns
    sns.set_style("whitegrid")
    sns.set_palette("husl")

try:
    import markdown
    from markdown.extensions import codehilite, tables, fenced_code
except ImportError:
    print("Instalando markdown...")
    os.system("pip install markdown")
    import markdown
    from markdown.extensions import codehilite, tables, fenced_code

WEASYPRINT_AVAILABLE = False
try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except Exception:
    # WeasyPrint requiere librerías del sistema (GTK, Pango, etc.)
    # Si no están disponibles, simplemente no lo usamos
    pass


class DocumentGeneratorPremium:
    """Generador premium de documentos con gráficas"""
    
    def __init__(self, output_dir: str = "documentos_importantes_premium"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.graphs_dir = self.output_dir / "graficas"
        self.graphs_dir.mkdir(exist_ok=True)
        
        # Configurar matplotlib
        plt.rcParams['figure.figsize'] = (12, 8)
        plt.rcParams['font.size'] = 10
        plt.rcParams['axes.labelsize'] = 12
        plt.rcParams['axes.titlesize'] = 14
        plt.rcParams['xtick.labelsize'] = 10
        plt.rcParams['ytick.labelsize'] = 10
        plt.rcParams['legend.fontsize'] = 10
        
    def analyze_document(self, file_path: Path) -> Dict:
        """Analiza un documento y extrae estadísticas"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Estadísticas básicas
        lines = content.split('\n')
        words = content.split()
        paragraphs = [p for p in content.split('\n\n') if p.strip()]
        
        # Contar secciones
        sections = len(re.findall(r'^#+\s+', content, re.MULTILINE))
        code_blocks = len(re.findall(r'```', content))
        tables = len(re.findall(r'\|.*\|', content))
        links = len(re.findall(r'\[.*?\]\(.*?\)', content))
        
        # Extraer estructura
        structure = []
        for line in lines:
            if line.strip().startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                title = line.lstrip('#').strip()
                structure.append({'level': level, 'title': title})
        
        return {
            'file_path': str(file_path),
            'name': file_path.stem,
            'total_lines': len(lines),
            'total_words': len(words),
            'total_paragraphs': len(paragraphs),
            'sections': sections,
            'code_blocks': code_blocks,
            'tables': tables,
            'links': links,
            'structure': structure[:20],  # Primeras 20 secciones
            'content': content
        }
    
    def create_statistics_graphs(self, doc_stats: Dict) -> List[str]:
        """Crea gráficas de estadísticas del documento"""
        graph_files = []
        
        # Gráfica 1: Distribución de contenido
        fig, ax = plt.subplots(figsize=(10, 6))
        categories = ['Líneas', 'Palabras', 'Párrafos', 'Secciones', 'Bloques Código', 'Tablas', 'Enlaces']
        values = [
            doc_stats['total_lines'],
            doc_stats['total_words'],
            doc_stats['total_paragraphs'],
            doc_stats['sections'],
            doc_stats['code_blocks'],
            doc_stats['tables'],
            doc_stats['links']
        ]
        
        # Normalizar valores para visualización
        max_val = max(values) if values else 1
        normalized = [v / max_val * 100 if max_val > 0 else 0 for v in values]
        
        bars = ax.barh(categories, normalized, color=sns.color_palette("husl", len(categories)))
        ax.set_xlabel('Valor Normalizado (%)', fontsize=12, fontweight='bold')
        ax.set_title(f'Análisis de Contenido: {doc_stats["name"]}', fontsize=14, fontweight='bold')
        ax.grid(axis='x', alpha=0.3)
        
        # Agregar valores en las barras
        for i, (bar, val) in enumerate(zip(bars, values)):
            width = bar.get_width()
            ax.text(width + 1, bar.get_y() + bar.get_height()/2, 
                   f'{val:,}', ha='left', va='center', fontweight='bold')
        
        plt.tight_layout()
        graph_file = self.graphs_dir / f"{doc_stats['name']}_estadisticas.png"
        plt.savefig(graph_file, dpi=300, bbox_inches='tight')
        plt.close()
        graph_files.append(str(graph_file))
        
        # Gráfica 2: Estructura jerárquica
        if doc_stats['structure']:
            fig, ax = plt.subplots(figsize=(12, 8))
            
            # Contar niveles
            level_counts = {}
            for item in doc_stats['structure']:
                level = item['level']
                level_counts[level] = level_counts.get(level, 0) + 1
            
            if level_counts:
                levels = sorted(level_counts.keys())
                counts = [level_counts[l] for l in levels]
                labels = [f'Nivel {l}' for l in levels]
                
                colors_pie = sns.color_palette("Set3", len(levels))
                wedges, texts, autotexts = ax.pie(counts, labels=labels, autopct='%1.1f%%',
                                                  colors=colors_pie, startangle=90)
                
                for autotext in autotexts:
                    autotext.set_color('black')
                    autotext.set_fontweight('bold')
                
                ax.set_title(f'Distribución de Secciones por Nivel: {doc_stats["name"]}', 
                           fontsize=14, fontweight='bold')
                
                plt.tight_layout()
                graph_file = self.graphs_dir / f"{doc_stats['name']}_estructura.png"
                plt.savefig(graph_file, dpi=300, bbox_inches='tight')
                plt.close()
                graph_files.append(str(graph_file))
        
        # Gráfica 3: Comparativa de documentos (si hay múltiples)
        return graph_files
    
    def create_comparison_graph(self, all_stats: List[Dict]) -> str:
        """Crea gráfica comparativa de todos los documentos"""
        if len(all_stats) < 2:
            return None
        
        fig, axes = plt.subplots(2, 2, figsize=(14, 10))
        fig.suptitle('Comparativa de Documentos Importantes', fontsize=16, fontweight='bold')
        
        names = [s['name'] for s in all_stats]
        
        # Gráfica 1: Líneas y palabras
        ax1 = axes[0, 0]
        x = range(len(names))
        width = 0.35
        lines = [s['total_lines'] for s in all_stats]
        words = [s['total_words'] for s in all_stats]
        
        # Normalizar para comparación
        max_lines = max(lines) if lines else 1
        max_words = max(words) if words else 1
        lines_norm = [l / max_lines * 100 for l in lines]
        words_norm = [w / max_words * 100 for w in words]
        
        ax1.bar([i - width/2 for i in x], lines_norm, width, label='Líneas (normalizado)', alpha=0.8)
        ax1.bar([i + width/2 for i in x], words_norm, width, label='Palabras (normalizado)', alpha=0.8)
        ax1.set_xlabel('Documentos', fontweight='bold')
        ax1.set_ylabel('Valor Normalizado (%)', fontweight='bold')
        ax1.set_title('Líneas vs Palabras', fontweight='bold')
        ax1.set_xticks(x)
        ax1.set_xticklabels(names, rotation=45, ha='right')
        ax1.legend()
        ax1.grid(axis='y', alpha=0.3)
        
        # Gráfica 2: Secciones y bloques de código
        ax2 = axes[0, 1]
        sections = [s['sections'] for s in all_stats]
        code_blocks = [s['code_blocks'] for s in all_stats]
        
        max_sec = max(sections) if sections else 1
        max_code = max(code_blocks) if code_blocks else 1
        sections_norm = [s / max_sec * 100 for s in sections]
        code_norm = [c / max_code * 100 for c in code_blocks]
        
        ax2.bar([i - width/2 for i in x], sections_norm, width, label='Secciones (normalizado)', alpha=0.8)
        ax2.bar([i + width/2 for i in x], code_norm, width, label='Bloques Código (normalizado)', alpha=0.8)
        ax2.set_xlabel('Documentos', fontweight='bold')
        ax2.set_ylabel('Valor Normalizado (%)', fontweight='bold')
        ax2.set_title('Secciones vs Bloques de Código', fontweight='bold')
        ax2.set_xticks(x)
        ax2.set_xticklabels(names, rotation=45, ha='right')
        ax2.legend()
        ax2.grid(axis='y', alpha=0.3)
        
        # Gráfica 3: Complejidad (tablas + enlaces)
        ax3 = axes[1, 0]
        tables = [s['tables'] for s in all_stats]
        links = [s['links'] for s in all_stats]
        
        max_tab = max(tables) if tables else 1
        max_link = max(links) if links else 1
        tables_norm = [t / max_tab * 100 if max_tab > 0 else 0 for t in tables]
        links_norm = [l / max_link * 100 if max_link > 0 else 0 for l in links]
        
        ax3.bar([i - width/2 for i in x], tables_norm, width, label='Tablas (normalizado)', alpha=0.8)
        ax3.bar([i + width/2 for i in x], links_norm, width, label='Enlaces (normalizado)', alpha=0.8)
        ax3.set_xlabel('Documentos', fontweight='bold')
        ax3.set_ylabel('Valor Normalizado (%)', fontweight='bold')
        ax3.set_title('Tablas vs Enlaces', fontweight='bold')
        ax3.set_xticks(x)
        ax3.set_xticklabels(names, rotation=45, ha='right')
        ax3.legend()
        ax3.grid(axis='y', alpha=0.3)
        
        # Gráfica 4: Resumen total
        ax4 = axes[1, 1]
        totals = [s['total_words'] for s in all_stats]
        colors_bar = sns.color_palette("husl", len(names))
        bars = ax4.barh(names, totals, color=colors_bar, alpha=0.8)
        ax4.set_xlabel('Total de Palabras', fontweight='bold')
        ax4.set_title('Total de Contenido por Documento', fontweight='bold')
        ax4.grid(axis='x', alpha=0.3)
        
        # Agregar valores
        for bar, val in zip(bars, totals):
            width = bar.get_width()
            ax4.text(width + max(totals) * 0.01, bar.get_y() + bar.get_height()/2,
                    f'{val:,}', ha='left', va='center', fontweight='bold')
        
        plt.tight_layout()
        graph_file = self.graphs_dir / "comparativa_documentos.png"
        plt.savefig(graph_file, dpi=300, bbox_inches='tight')
        plt.close()
        
        return str(graph_file)
    
    def create_pdf(self, doc_stats: Dict, graph_files: List[str]) -> str:
        """Crea un PDF profesional del documento"""
        pdf_file = self.output_dir / f"{doc_stats['name']}_premium.pdf"
        
        doc = SimpleDocTemplate(str(pdf_file), pagesize=A4,
                               rightMargin=72, leftMargin=72,
                               topMargin=72, bottomMargin=72)
        
        # Estilos
        styles = getSampleStyleSheet()
        
        # Estilo personalizado para título
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1a1a1a'),
            spaceAfter=30,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        # Estilo para subtítulos
        subtitle_style = ParagraphStyle(
            'CustomSubtitle',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#2c3e50'),
            spaceAfter=12,
            spaceBefore=20,
            fontName='Helvetica-Bold'
        )
        
        # Estilo para texto normal
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            textColor=colors.HexColor('#333333'),
            spaceAfter=12,
            leading=14,
            alignment=TA_JUSTIFY
        )
        
        # Construir contenido
        story = []
        
        # Portada
        story.append(Paragraph(doc_stats['name'].replace('_', ' ').title(), title_style))
        story.append(Spacer(1, 0.5*inch))
        story.append(Paragraph(f"Generado: {datetime.now().strftime('%d de %B de %Y, %H:%M')}", 
                              styles['Normal']))
        story.append(PageBreak())
        
        # Resumen ejecutivo
        story.append(Paragraph("Resumen Ejecutivo", subtitle_style))
        story.append(Spacer(1, 0.2*inch))
        
        summary_data = [
            ['Métrica', 'Valor'],
            ['Total de Líneas', f"{doc_stats['total_lines']:,}"],
            ['Total de Palabras', f"{doc_stats['total_words']:,}"],
            ['Total de Párrafos', f"{doc_stats['total_paragraphs']:,}"],
            ['Secciones', f"{doc_stats['sections']}"],
            ['Bloques de Código', f"{doc_stats['code_blocks']}"],
            ['Tablas', f"{doc_stats['tables']}"],
            ['Enlaces', f"{doc_stats['links']}"]
        ]
        
        summary_table = Table(summary_data, colWidths=[3*inch, 2*inch])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#34495e')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#ecf0f1')),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#bdc3c7')),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8f9fa')])
        ]))
        story.append(summary_table)
        story.append(Spacer(1, 0.3*inch))
        
        # Agregar gráficas
        for graph_file in graph_files:
            if Path(graph_file).exists():
                img = RLImage(graph_file, width=6*inch, height=4*inch)
                story.append(img)
                story.append(Spacer(1, 0.2*inch))
        
        story.append(PageBreak())
        
        # Contenido del documento
        story.append(Paragraph("Contenido del Documento", subtitle_style))
        story.append(Spacer(1, 0.2*inch))
        
        # Procesar contenido markdown básico
        content = doc_stats['content']
        lines = content.split('\n')
        
        # Limitar a 1000 líneas para evitar PDFs muy grandes
        max_lines = min(1000, len(lines))
        
        for i, line in enumerate(lines[:max_lines]):
            line = line.strip()
            if not line:
                if i < max_lines - 1:  # No agregar espacio al final
                    story.append(Spacer(1, 0.1*inch))
                continue
            
            try:
                if line.startswith('#'):
                    level = len(line) - len(line.lstrip('#'))
                    text = line.lstrip('#').strip()
                    # Limpiar caracteres especiales
                    text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    if level == 1:
                        story.append(Paragraph(text, subtitle_style))
                    elif level == 2:
                        story.append(Paragraph(text, styles['Heading2']))
                    else:
                        story.append(Paragraph(text, styles['Heading3']))
                elif line.startswith('-') or line.startswith('*'):
                    text = line.lstrip('-*').strip()
                    text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    story.append(Paragraph(f"• {text}", normal_style))
                elif line.startswith('|') and '|' in line[1:]:
                    # Es una tabla, simplificarla
                    cells = [c.strip() for c in line.split('|')[1:-1]]
                    if cells and not all(c.startswith('-') for c in cells):
                        text = ' | '.join(cells)
                        text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        story.append(Paragraph(text, normal_style))
                else:
                    # Limpiar texto para ReportLab
                    text = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    if text:
                        story.append(Paragraph(text, normal_style))
            except Exception as e:
                # Si hay error procesando una línea, continuar
                logger.warning(f"Error procesando línea {i}: {e}")
                continue
        
        # Construir PDF
        doc.build(story)
        return str(pdf_file)
    
    def create_word(self, doc_stats: Dict, graph_files: List[str]) -> str:
        """Crea un documento Word profesional"""
        doc_file = self.output_dir / f"{doc_stats['name']}_premium.docx"
        doc = Document()
        
        # Configurar estilos
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Portada
        title = doc.add_heading(doc_stats['name'].replace('_', ' ').title(), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        para = doc.add_paragraph(f"Generado: {datetime.now().strftime('%d de %B de %Y, %H:%M')}")
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # Resumen ejecutivo
        doc.add_heading('Resumen Ejecutivo', 1)
        
        # Tabla de resumen
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Encabezado
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Métrica'
        hdr_cells[1].text = 'Valor'
        hdr_cells[0].paragraphs[0].runs[0].font.bold = True
        hdr_cells[1].paragraphs[0].runs[0].font.bold = True
        
        # Datos
        metrics = [
            ('Total de Líneas', f"{doc_stats['total_lines']:,}"),
            ('Total de Palabras', f"{doc_stats['total_words']:,}"),
            ('Total de Párrafos', f"{doc_stats['total_paragraphs']:,}"),
            ('Secciones', f"{doc_stats['sections']}"),
            ('Bloques de Código', f"{doc_stats['code_blocks']}"),
            ('Tablas', f"{doc_stats['tables']}"),
            ('Enlaces', f"{doc_stats['links']}")
        ]
        
        for metric, value in metrics:
            row_cells = table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = value
        
        doc.add_paragraph()
        
        # Agregar gráficas
        for graph_file in graph_files:
            if Path(graph_file).exists():
                doc.add_paragraph()
                doc.add_picture(graph_file, width=Inches(6))
                doc.add_paragraph()
        
        doc.add_page_break()
        
        # Contenido
        doc.add_heading('Contenido del Documento', 1)
        
        content = doc_stats['content']
        lines = content.split('\n')
        
        # Limitar a 2000 líneas para evitar documentos muy grandes
        max_lines = min(2000, len(lines))
        
        for i, line in enumerate(lines[:max_lines]):
            line = line.strip()
            if not line:
                if i < max_lines - 1:
                    doc.add_paragraph()
                continue
            
            try:
                if line.startswith('#'):
                    level = len(line) - len(line.lstrip('#'))
                    text = line.lstrip('#').strip()
                    # Limitar nivel a 9 (máximo de Word)
                    level = min(level, 9)
                    doc.add_heading(text, level)
                elif line.startswith('-') or line.startswith('*'):
                    text = line.lstrip('-*').strip()
                    if text:
                        doc.add_paragraph(text, style='List Bullet')
                elif line.startswith('|') and '|' in line[1:]:
                    # Es una tabla, simplificarla como texto
                    cells = [c.strip() for c in line.split('|')[1:-1]]
                    if cells and not all(c.startswith('-') for c in cells):
                        text = ' | '.join(cells)
                        doc.add_paragraph(text)
                else:
                    if line:
                        doc.add_paragraph(line)
            except Exception as e:
                # Si hay error procesando una línea, continuar
                logger.warning(f"Error procesando línea {i}: {e}")
                continue
        
        doc.save(str(doc_file))
        return str(doc_file)
    
    def create_excel(self, doc_stats: Dict, all_stats: List[Dict]) -> str:
        """Crea un archivo Excel con análisis y gráficas"""
        excel_file = self.output_dir / f"{doc_stats['name']}_premium.xlsx"
        wb = Workbook()
        
        # Hoja 1: Resumen
        ws1 = wb.active
        ws1.title = "Resumen"
        
        # Estilos
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF", size=12)
        title_font = Font(bold=True, size=14)
        border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # Título
        ws1['A1'] = doc_stats['name'].replace('_', ' ').title()
        ws1['A1'].font = title_font
        ws1.merge_cells('A1:B1')
        
        ws1['A2'] = f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
        ws1['A2'].font = Font(italic=True)
        
        # Tabla de resumen
        ws1['A4'] = 'Métrica'
        ws1['B4'] = 'Valor'
        ws1['A4'].fill = header_fill
        ws1['A4'].font = header_font
        ws1['B4'].fill = header_fill
        ws1['B4'].font = header_font
        
        metrics = [
            ('Total de Líneas', doc_stats['total_lines']),
            ('Total de Palabras', doc_stats['total_words']),
            ('Total de Párrafos', doc_stats['total_paragraphs']),
            ('Secciones', doc_stats['sections']),
            ('Bloques de Código', doc_stats['code_blocks']),
            ('Tablas', doc_stats['tables']),
            ('Enlaces', doc_stats['links'])
        ]
        
        for i, (metric, value) in enumerate(metrics, start=5):
            ws1[f'A{i}'] = metric
            ws1[f'B{i}'] = value
            ws1[f'A{i}'].border = border
            ws1[f'B{i}'].border = border
            ws1[f'B{i}'].alignment = Alignment(horizontal='right')
        
        # Ajustar ancho de columnas
        ws1.column_dimensions['A'].width = 25
        ws1.column_dimensions['B'].width = 15
        
        # Hoja 2: Comparativa (si hay múltiples documentos)
        if len(all_stats) > 1:
            ws2 = wb.create_sheet("Comparativa")
            
            # Encabezados
            headers = ['Documento', 'Líneas', 'Palabras', 'Párrafos', 'Secciones', 'Código', 'Tablas', 'Enlaces']
            for col, header in enumerate(headers, start=1):
                cell = ws2.cell(row=1, column=col, value=header)
                cell.fill = header_fill
                cell.font = header_font
                cell.border = border
            
            # Datos
            for row, stats in enumerate(all_stats, start=2):
                ws2.cell(row=row, column=1, value=stats['name'])
                ws2.cell(row=row, column=2, value=stats['total_lines'])
                ws2.cell(row=row, column=3, value=stats['total_words'])
                ws2.cell(row=row, column=4, value=stats['total_paragraphs'])
                ws2.cell(row=row, column=5, value=stats['sections'])
                ws2.cell(row=row, column=6, value=stats['code_blocks'])
                ws2.cell(row=row, column=7, value=stats['tables'])
                ws2.cell(row=row, column=8, value=stats['links'])
                
                for col in range(1, 9):
                    ws2.cell(row=row, column=col).border = border
            
            # Gráfica comparativa
            chart = BarChart()
            chart.type = "col"
            chart.style = 10
            chart.title = "Comparativa de Documentos"
            chart.y_axis.title = 'Valor'
            chart.x_axis.title = 'Documentos'
            
            data = Reference(ws2, min_col=2, min_row=1, max_col=8, max_row=len(all_stats)+1)
            cats = Reference(ws2, min_col=1, min_row=2, max_row=len(all_stats)+1)
            chart.add_data(data, titles_from_data=True)
            chart.set_categories(cats)
            chart.height = 10
            chart.width = 15
            
            ws2.add_chart(chart, "J2")
            
            # Ajustar ancho de columnas
            for col in range(1, 9):
                ws2.column_dimensions[chr(64 + col)].width = 15
        
        # Hoja 3: Estructura del documento
        ws3 = wb.create_sheet("Estructura")
        ws3['A1'] = 'Nivel'
        ws3['B1'] = 'Título'
        ws3['A1'].fill = header_fill
        ws3['A1'].font = header_font
        ws3['B1'].fill = header_fill
        ws3['B1'].font = header_font
        
        for i, item in enumerate(doc_stats['structure'], start=2):
            ws3.cell(row=i, column=1, value=item['level'])
            ws3.cell(row=i, column=2, value=item['title'])
            ws3.cell(row=i, column=1).border = border
            ws3.cell(row=i, column=2).border = border
        
        ws3.column_dimensions['A'].width = 10
        ws3.column_dimensions['B'].width = 50
        
        wb.save(str(excel_file))
        return str(excel_file)
    
    def process_documents(self, doc_files: List[str]) -> Dict[str, List[str]]:
        """Procesa múltiples documentos"""
        results = {}
        all_stats = []
        
        for doc_file in doc_files:
            file_path = Path(doc_file)
            if not file_path.exists():
                print(f"⚠️  Archivo no encontrado: {doc_file}")
                continue
            
            try:
                print(f"📄 Procesando: {file_path.name}")
                
                # Analizar documento
                stats = self.analyze_document(file_path)
                all_stats.append(stats)
                
                # Crear gráficas
                print("  📊 Generando gráficas...")
                try:
                    graph_files = self.create_statistics_graphs(stats)
                except Exception as e:
                    print(f"  ⚠️  Error generando gráficas: {e}")
                    graph_files = []
                
                # Crear PDF
                print("  📑 Generando PDF...")
                try:
                    pdf_file = self.create_pdf(stats, graph_files)
                except Exception as e:
                    print(f"  ⚠️  Error generando PDF: {e}")
                    pdf_file = None
                
                # Crear Word
                print("  📝 Generando Word...")
                try:
                    word_file = self.create_word(stats, graph_files)
                except Exception as e:
                    print(f"  ⚠️  Error generando Word: {e}")
                    word_file = None
                
                # Crear Excel
                print("  📊 Generando Excel...")
                try:
                    excel_file = self.create_excel(stats, all_stats)
                except Exception as e:
                    print(f"  ⚠️  Error generando Excel: {e}")
                    excel_file = None
                
                results[stats['name']] = {
                    'pdf': pdf_file,
                    'word': word_file,
                    'excel': excel_file,
                    'graphs': graph_files
                }
                
                print(f"  ✅ Completado: {stats['name']}\n")
                
            except Exception as e:
                print(f"  ❌ Error procesando {file_path.name}: {e}\n")
                continue
        
        # Crear gráfica comparativa
        if len(all_stats) > 1:
            print("📊 Generando gráfica comparativa...")
            try:
                comparison_graph = self.create_comparison_graph(all_stats)
                if comparison_graph:
                    print(f"  ✅ Gráfica comparativa: {comparison_graph}\n")
            except Exception as e:
                print(f"  ⚠️  Error generando gráfica comparativa: {e}\n")
        
        return results


def main():
    """Función principal"""
    # Documentos importantes a procesar
    base_dir = Path(__file__).parent
    important_docs = [
        str(base_dir / "airflow_automation_prompt.md"),
        str(base_dir / "truthgpt_collected/integration_code/production_code/ARCHITECTURE_IMPROVEMENTS.md"),
        str(base_dir / "truthgpt_collected/integration_code/production_code/REFACTORING_PLAN.md")
    ]
    
    print("=" * 60)
    print("GENERADOR PREMIUM DE DOCUMENTOS IMPORTANTES")
    print("=" * 60)
    print()
    
    generator = DocumentGeneratorPremium()
    results = generator.process_documents(important_docs)
    
    print("=" * 60)
    print("RESUMEN DE ARCHIVOS GENERADOS")
    print("=" * 60)
    print()
    
    for doc_name, files in results.items():
        print(f"📄 {doc_name}:")
        print(f"   PDF:  {files['pdf']}")
        print(f"   Word: {files['word']}")
        print(f"   Excel: {files['excel']}")
        print(f"   Gráficas: {len(files['graphs'])} archivos")
        print()
    
    print(f"✅ Todos los archivos se han generado en: {generator.output_dir}")
    print()


if __name__ == "__main__":
    main()

