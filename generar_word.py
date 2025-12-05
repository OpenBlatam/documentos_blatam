from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_word_doc(filename):
    doc = Document()
    
    # Title
    title = doc.add_heading('Solución de Problemas en la Tarjeta Madre', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Experiencias Prácticas y Soluciones')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('') # Spacer

    # Introduction
    intro = doc.add_paragraph('Este documento presenta una guía visual y práctica sobre los problemas más comunes encontrados en tarjetas madre de computadoras de escritorio, sus causas y soluciones efectivas.')
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph('')

    components = [
        {
            "title": "1. Memoria RAM (Ranuras DIMM)",
            "problem": "Problema: El equipo enciende, los ventiladores giran, pero no hay señal de video y se escuchan pitidos.",
            "cause": "Causa: Acumulación de polvo o sulfatación en los contactos dorados de la memoria.",
            "solution": "Solución: Limpiar los contactos con una goma de borrar blanca suave y retirar el polvo de las ranuras con aire comprimido.",
            "image": "ram_word.png"
        },
        {
            "title": "2. Batería CMOS (Pila CR2032)",
            "problem": "Problema: La fecha y hora se desconfiguran cada vez que se desconecta el equipo de la corriente.",
            "cause": "Causa: La batería tipo botón ha agotado su carga (voltaje inferior a 3V).",
            "solution": "Solución: Reemplazar la batería por una nueva modelo CR2032, respetando la polaridad positiva (+).",
            "image": "battery_word.png"
        },
        {
            "title": "3. Panel Frontal (Conexiones)",
            "problem": "Problema: El botón de encendido no responde o los LEDs de actividad no funcionan tras un mantenimiento.",
            "cause": "Causa: Conexión incorrecta de los cables (Power SW, Reset SW, HDD LED) en los pines de la placa.",
            "solution": "Solución: Consultar el diagrama en el manual de la placa madre y reconectar verificando la polaridad.",
            "image": "panel_word.png"
        }
    ]

    for comp in components:
        # Section Title
        h = doc.add_heading(comp['title'], level=1)
        
        # Details
        p = doc.add_paragraph()
        run = p.add_run(comp['problem'])
        run.bold = True
        p.add_run('\n')
        p.add_run(comp['cause'])
        p.add_run('\n')
        p.add_run(comp['solution'])
        
        # Image
        if os.path.exists(comp['image']):
            try:
                doc.add_picture(comp['image'], width=Inches(4.0))
                last_paragraph = doc.paragraphs[-1] 
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                doc.add_paragraph(f"[Error al insertar imagen: {e}]")
        
        doc.add_paragraph('') # Spacer
        doc.add_paragraph('─' * 20) # Separator
        doc.add_paragraph('') # Spacer

    # Footer / Credits
    footer = doc.sections[0].footer
    p_footer = footer.paragraphs[0]
    p_footer.text = "Elaborado por: Adán Muñoz Pablo | Entregable 1"
    p_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.save(filename)
    print(f"Documento guardado como {filename}")

if __name__ == "__main__":
    create_word_doc("Adan_Munoz_Pablo_Entregable_1.docx")







