import sys
import subprocess
import importlib.util
import os

def install(package):
    subprocess.check_call([sys.executable, "-m", "pip", "install", package])

def check_and_install(package_name, import_name=None):
    if import_name is None:
        import_name = package_name
    
    if importlib.util.find_spec(import_name) is None:
        print(f"Installing {package_name}...")
        try:
            install(package_name)
        except Exception as e:
            print(f"Failed to install {package_name}: {e}")

# Check and install dependencies
check_and_install("python-docx", "docx")
check_and_install("fpdf")

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Failed to import python-docx. Word document generation skipped.")
    Document = None

try:
    from fpdf import FPDF
except ImportError:
    print("Failed to import fpdf. PDF generation skipped.")
    FPDF = None

# Content Data in Spanish matching the template
TITLE = "Análisis de factibilidad del negocio"
SUBTITLE = "Implementación de Portal Web para Biciosos Workshop"
INFO_GENERAL = "Datos Generales:\nEmpresa: Biciosos Workshop\nGiro: Venta y servicios de ciclismo\nProyecto: Portal Web con Citas y Presupuestos"

SECTIONS = [
    {
        "title": "1. Introducción",
        "content": [
            ("", "Biciosos Workshop es una empresa con 5 años de antigüedad, que cuenta con una sucursal, un gerente y dos empleados. Su operación actual es manual y presencial. La estrategia planteada por la gerencia es implementar un portal web que permita exhibir productos, contactar clientes, gestionar citas y generar presupuestos automatizados."),
            ("", "El objetivo de este análisis es evaluar 3 alternativas de software libre para determinar la más viable técnica, económica y operativamente, cubriendo aspectos críticos como hosting, seguridad, costos y funcionalidades específicas.")
        ]
    },
    {
        "title": "2. Análisis de alternativas",
        "content": [
            ("", "Se analizaron tres plataformas Open Source (Software Libre) para tiendas virtuales:"),
            ("Alternativa 1: WooCommerce (WordPress)", "Plataforma líder mundial. Es un plugin que transforma WordPress en una tienda. Ideal para contenido + venta.\nFuncionalidades deseables no incluidas nativamente: El sistema de citas y la generación de presupuestos PDF requieren la instalación de plugins adicionales (gratuitos o de pago)."),
            ("Alternativa 2: PrestaShop", "Software dedicado 100% al comercio electrónico. Muy robusto para manejo de miles de productos.\nFuncionalidades deseables no incluidas nativamente: No cuenta con sistema de agenda/citas ni generador de presupuestos formal; requiere módulos costosos para estas funciones."),
            ("Alternativa 3: OpenCart", "Solución ligera y rápida. Requiere pocos recursos de servidor.\nFuncionalidades deseables no incluidas nativamente: La gestión de servicios (citas) es muy limitada y compleja de adaptar sin desarrollo a medida.")
        ]
    },
    {
        "title": "3. Estudio económico",
        "content": [
            ("Costo y Licencias", [
                "WooCommerce: Licencia GPL (Gratuita). Inversión inicial estimada: $150 USD (Hosting + Dominio + Tema Premium).",
                "PrestaShop: Licencia OSL (Gratuita). Inversión inicial estimada: $300 USD (Requiere hosting más potente + Módulos esenciales).",
                "OpenCart: Licencia GPL (Gratuita). Inversión inicial estimada: $100 USD (Hosting básico)."
            ]),
            ("Gastos de Operación", "Los gastos recurrentes principales son el Hosting anual y la renovación de dominio. No hay pago mensual por uso del software base en ninguna de las opciones.")
        ]
    },
    {
        "title": "4. Estudio técnico",
        "content": [
            ("Hosting y Base de Datos", [
                "Hosting: Todas requieren servidor LAMP (Linux, Apache, MySQL, PHP). WooCommerce y OpenCart funcionan bien en hosting compartido estándar. PrestaShop se beneficia de servidores VPS.",
                "Base de Datos: Todas utilizan MySQL o MariaDB, estándares de la industria fáciles de migrar."
            ]),
            ("Facilidades de almacenamiento y respaldo", [
                "Almacenamiento: Depende del proveedor de hosting contratado (se recomienda mínimo 10GB SSD).",
                "Respaldo: Las tres permiten respaldos manuales de la base de datos y archivos. Existen plugins de automatización de backups para WordPress (WooCommerce) que facilitan esta tarea."
            ]),
            ("Seguridad e Integración de Cobro", [
                "Seguridad: Requieren certificado SSL (HTTPS). WooCommerce y PrestaShop lanzan parches de seguridad frecuentes.",
                "Cobros: Las tres tienen integración nativa o vía plugins oficiales con VISA, Mastercard (vía pasarelas como Stripe/MercadoPago) y PayPal."
            ])
        ]
    },
    {
        "title": "5. Estudio Operativo",
        "content": [
            ("Tiempo de Implementación", [
                "WooCommerce: 2 semanas (Configuración rápida).",
                "PrestaShop: 4 semanas (Curva de aprendizaje más alta).",
                "OpenCart: 3 semanas."
            ]),
            ("Recursos Humanos", "El personal actual (2 empleados) puede administrar WooCommerce con una capacitación mínima (5 horas). PrestaShop requeriría una capacitación técnica más extensa (15+ horas)."),
            ("Funcionalidad de Citas y Presupuestos", "Este es el requerimiento crítico. WooCommerce es el único que permite integrar un sistema de 'Booking' y plugins de facturación/presupuestos de manera fluida en el mismo entorno.")
        ]
    },
    {
        "title": "6. Resultados",
        "content": [
            ("", "Cuadro comparativo de los principales aspectos identificados:"),
            ("table", [
                ["Aspecto", "WooCommerce", "PrestaShop", "OpenCart"],
                ["Licencia", "Gratuita (GPL)", "Gratuita (OSL)", "Gratuita (GPL)"],
                ["Costo Implementación", "Bajo", "Medio", "Bajo"],
                ["Tiempo Implem.", "Corto (2 sem)", "Medio (4 sem)", "Medio (3 sem)"],
                ["Base de Datos", "MySQL", "MySQL", "MySQL"],
                ["Hosting Requerido", "Estándar", "Optimizado (VPS)", "Básico"],
                ["Seguridad", "Alta", "Alta", "Media"],
                ["Pagos (Visa/Paypal)", "Integración Total", "Integración Total", "Integración Total"],
                ["Citas y Presupuestos", "Excelente (Plugins)", "Difícil (Módulos)", "Muy Limitado"],
                ["Funcionalidad Faltante", "Requiere plugins", "Nativa robusta", "Básica"]
            ])
        ]
    },
    {
        "title": "7. Conclusiones",
        "content": [
            ("", "Basado en el análisis, la mejor alternativa para Biciosos Workshop es WooCommerce."),
            ("", "Esta decisión se sustenta en que cumple mejor con el requerimiento estratégico del gerente: 'proporcionar citas y generar presupuestos'. Aunque PrestaShop es una tienda más potente, WooCommerce ofrece la flexibilidad necesaria para gestionar servicios (taller) y productos en una sola interfaz amigable para los 2 empleados actuales. Además, su costo de implementación y mantenimiento es el más bajo, asegurando la viabilidad económica.")
        ]
    },
    {
        "title": "8. Bibliografía",
        "content": [
            ("", "1. Automattic. (2023). WooCommerce Documentation. Recuperado de woocommerce.com\n2. PrestaShop S.A. (2023). PrestaShop User Guide. Recuperado de prestashop.com\n3. Laudon, K. C., & Laudon, J. P. (2016). Sistemas de información gerencial. Pearson Educación.")
        ]
    }
]

# --- Generate DOCX ---
def generate_docx():
    if Document is None:
        return
    
    doc = Document()
    
    # Title Style
    title = doc.add_heading(TITLE, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph(SUBTITLE)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    info = doc.add_paragraph(INFO_GENERAL)
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    
    # Content
    for section in SECTIONS:
        doc.add_heading(section["title"], level=1)
        
        for sub_title, text in section["content"]:
            if sub_title == "table":
                # Create Table
                table_data = text
                table = doc.add_table(rows=1, cols=len(table_data[0]))
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(table_data[0]):
                    hdr_cells[i].text = header
                    # Bold header
                    for paragraph in hdr_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                
                for row_data in table_data[1:]:
                    row_cells = table.add_row().cells
                    for i, item in enumerate(row_data):
                        row_cells[i].text = item
                doc.add_paragraph() # Spacing after table
            elif sub_title:
                doc.add_heading(sub_title, level=2)
                if isinstance(text, list):
                    for item in text:
                        doc.add_paragraph(item, style='List Bullet')
                else:
                    doc.add_paragraph(text)
            else:
                # Content without subtitle
                if isinstance(text, list):
                    for item in text:
                        doc.add_paragraph(item, style='List Bullet')
                else:
                    doc.add_paragraph(text)
        
        doc.add_paragraph() # Spacing
    
    output_path = "Analisis_Factibilidad_Biciosos.docx"
    doc.save(output_path)
    print(f"Generated Word document: {output_path}")

# --- Generate PDF ---
class PDF(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 10)
        self.cell(0, 10, 'Biciosos Workshop - Analisis de Factibilidad', 0, 1, 'R')
        self.ln(5)

    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, f'Pagina {self.page_no()}', 0, 0, 'C')

    def chapter_title(self, title):
        self.set_font('Arial', 'B', 14)
        try:
            title = title.encode('latin-1', 'replace').decode('latin-1')
        except:
            pass
        self.cell(0, 10, title, 0, 1, 'L')
        self.ln(2)

    def chapter_subtitle(self, subtitle):
        self.set_font('Arial', 'B', 12)
        try:
            subtitle = subtitle.encode('latin-1', 'replace').decode('latin-1')
        except:
            pass
        self.cell(0, 10, subtitle, 0, 1, 'L')
        self.ln(1)

    def chapter_body(self, body):
        self.set_font('Arial', '', 11)
        try:
            body = body.encode('latin-1', 'replace').decode('latin-1')
        except:
            pass
        self.multi_cell(0, 6, body)
        self.ln()

    def bullet_points(self, points):
        self.set_font('Arial', '', 11)
        for point in points:
            try:
                point = point.encode('latin-1', 'replace').decode('latin-1')
            except:
                pass
            self.cell(5)
            self.multi_cell(0, 6, f"- {point}")
        self.ln()

    def create_table(self, data):
        self.set_font('Arial', 'B', 9) # Smaller font for table
        # Calculate widths
        # 4 columns: Aspect, Woo, Presta, Open
        col_widths = [45, 45, 45, 45] # Total 180
        line_height = 6
        
        # Header
        for i, item in enumerate(data[0]):
            try:
                item = item.encode('latin-1', 'replace').decode('latin-1')
            except:
                pass
            self.cell(col_widths[i], line_height, item, border=1)
        self.ln(line_height)
        
        # Rows
        self.set_font('Arial', '', 9)
        for row in data[1:]:
            for i, item in enumerate(row):
                try:
                    item = item.encode('latin-1', 'replace').decode('latin-1')
                except:
                    pass
                # Use multi_cell for cells that might wrap? simplified to cell for now as FPDF tables are tricky
                # Truncate if too long to avoid break
                if len(item) > 25:
                     item = item[:22] + "..."
                self.cell(col_widths[i], line_height, item, border=1)
            self.ln(line_height)
        self.ln()

def generate_pdf():
    if FPDF is None:
        return

    pdf = PDF()
    pdf.add_page()
    
    # Title
    pdf.set_font('Arial', 'B', 16)
    try:
        t = TITLE.encode('latin-1', 'replace').decode('latin-1')
        st = SUBTITLE.encode('latin-1', 'replace').decode('latin-1')
    except:
        t = TITLE
        st = SUBTITLE
        
    pdf.cell(0, 10, t, 0, 1, 'C')
    pdf.ln(5)
    
    pdf.set_font('Arial', '', 12)
    pdf.multi_cell(0, 6, st, 0, 'C')
    pdf.ln(10)
    
    for section in SECTIONS:
        pdf.chapter_title(section["title"])
        
        for sub_title, text in section["content"]:
            if sub_title == "table":
                pdf.create_table(text)
            elif sub_title:
                pdf.chapter_subtitle(sub_title)
                if isinstance(text, list):
                    pdf.bullet_points(text)
                else:
                    pdf.chapter_body(text)
            else:
                if isinstance(text, list):
                    pdf.bullet_points(text)
                else:
                    pdf.chapter_body(text)
    
    output_path = "Analisis_Factibilidad_Biciosos.pdf"
    pdf.output(output_path)
    print(f"Generated PDF document: {output_path}")

if __name__ == "__main__":
    generate_docx()
    generate_pdf()
