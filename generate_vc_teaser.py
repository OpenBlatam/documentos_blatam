import docx
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_teaser(filename):
    doc = docx.Document()
    
    # --- Setup One-Pager Layout ---
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(1.5) # Narrow margins for one-pager
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # --- Header ---
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    header_table.columns[0].width = Cm(12)
    header_table.columns[1].width = Cm(6)
    
    c1 = header_table.rows[0].cells[0]
    p = c1.paragraphs[0]
    run = p.add_run("VENTURA CAPITAL")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(32, 55, 100)
    p.add_run("\nSaaS IA Marketing para LATAM").font.size = Pt(12)

    c2 = header_table.rows[0].cells[1]
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run2 = p2.add_run("INVESTMENT TEASER\nSeries A - $2M Ask")
    run2.bold = True
    run2.font.color.rgb = RGBColor(47, 117, 181)

    doc.add_paragraph("_" * 90).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Content Grid ---
    # 2 Columns: Left (Main), Right (Highlights)
    
    # We'll simulate columns with a table for layout stability
    layout_table = doc.add_table(rows=1, cols=2)
    layout_table.autofit = False
    layout_table.columns[0].width = Cm(11.5)
    layout_table.columns[1].width = Cm(6.5)
    
    # Left Column (Narrative)
    left_cell = layout_table.rows[0].cells[0]
    
    def add_section(cell, title, text):
        p = cell.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.color.rgb = RGBColor(32, 55, 100)
        run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(12)
        
        p_text = cell.add_paragraph()
        p_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_text = p_text.add_run(text)
        run_text.font.size = Pt(10)

    add_section(left_cell, "THE OPPORTUNITY", 
                "El mercado de marketing digital en LATAM sufre de ineficiencia crónica. Las agencias son costosas y lentas. Las herramientas de IA globales no entienden el contexto cultural local. Ventura AI es la primera plataforma de generación de contenido nativa para la región.")
    
    add_section(left_cell, "PRODUCT & TECHNOLOGY", 
                "Nuestra plataforma propietaria utiliza LLMs fine-tuned con dialectos regionales (MX, AR, CO, CL). Genera campañas completas (Copy + Imagen) en segundos. Integraciones nativas con MercadoLibre y WhatsApp Business.")
    
    add_section(left_cell, "TRACTION", 
                "• $139k ARR actual, creciendo 20% MoM\n• 500+ Clientes pagados\n• CAC: $150 | LTV: $1,350 (9x LTV/CAC)\n• Churn < 5%")
    
    add_section(left_cell, "USE OF FUNDS", 
                "Buscamos $2M para acelerar la adquisición de usuarios en Brasil y México, y finalizar el desarrollo de nuestro modelo propietario 'Ventura-LLM-v2'.")

    # Right Column (Key Stats & Team)
    right_cell = layout_table.rows[0].cells[1]
    
    # Highlights Box style
    right_cell.paragraphs[0].text = "" # Clear first empty paragraph
    
    def add_stat_box(cell, label, value):
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        
        r_val = p.add_run(value + "\n")
        r_val.bold = True
        r_val.font.size = Pt(18)
        r_val.font.color.rgb = RGBColor(47, 117, 181)
        
        r_lbl = p.add_run(label)
        r_lbl.font.size = Pt(9)
        r_lbl.font.color.rgb = RGBColor(128, 128, 128)

    add_stat_box(right_cell, "ARR Proyectado 2026", "$3.6M")
    add_stat_box(right_cell, "Gross Margin", "85%")
    add_stat_box(right_cell, "TAM LATAM", "$2.8B")
    
    # Team Section in Right Column
    p = right_cell.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run("TEAM")
    run.bold = True
    run.font.color.rgb = RGBColor(32, 55, 100)
    
    team_p = right_cell.add_paragraph()
    team_run = team_p.add_run("CEO: Ex-Unicornio Founder\nCTO: PhD AI Stanford\nCOO: Ex-McKinsey")
    team_run.font.size = Pt(9)

    # Footer
    doc.add_paragraph("_" * 90).alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Confidencial | Para Inversionistas Acreditados | contact@venturacapital.ai")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)

    doc.save(filename)
    print(f"Teaser created: {filename}")

if __name__ == "__main__":
    create_teaser("Ventura_Capital_Teaser.docx")

