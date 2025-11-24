#!/usr/bin/env python3
"""
Script para convertir documentos Markdown importantes a Word y Excel
con gráficas e imágenes.
"""

import os
import re
import json
import statistics
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Tuple, Any
import markdown
from markdown.extensions import codehilite, tables, fenced_code

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import openpyxl
    from openpyxl import Workbook
    from openpyxl.chart import BarChart, LineChart, PieChart, Reference
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    import matplotlib.pyplot as plt
    import matplotlib
    matplotlib.use('Agg')  # Para no requerir display
    from matplotlib.backends.backend_agg import FigureCanvasAgg
    import seaborn as sns
    sns.set_style("whitegrid")
except ImportError as e:
    print(f"Error: Faltan dependencias. Instala con: pip install python-docx openpyxl matplotlib seaborn markdown")
    print(f"Error específico: {e}")
    exit(1)


class MarkdownToWordExcel:
    """Convierte documentos Markdown a Word y Excel con gráficas."""
    
    def __init__(self, output_dir: str = "exports"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        self.stats = {}
        
    def parse_markdown(self, file_path: str) -> Dict[str, Any]:
        """Parsea un archivo Markdown y extrae información estructurada."""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Parsear con markdown
        md = markdown.Markdown(extensions=['codehilite', 'tables', 'fenced_code'])
        html = md.convert(content)
        
        # Extraer estadísticas
        lines = content.split('\n')
        stats = {
            'total_lines': len(lines),
            'total_words': len(content.split()),
            'total_chars': len(content),
            'headers': [],
            'code_blocks': [],
            'tables': [],
            'links': [],
            'images': [],
            'sections': []
        }
        
        # Extraer headers
        for line in lines:
            if line.strip().startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                text = line.lstrip('#').strip()
                stats['headers'].append({'level': level, 'text': text})
        
        # Extraer código
        code_pattern = r'```[\s\S]*?```'
        code_blocks = re.findall(code_pattern, content)
        stats['code_blocks'] = code_blocks
        
        # Extraer tablas (básico)
        table_pattern = r'\|.*\|'
        tables = re.findall(table_pattern, content, re.MULTILINE)
        stats['tables'] = tables
        
        # Extraer links
        link_pattern = r'\[([^\]]+)\]\(([^\)]+)\)'
        links = re.findall(link_pattern, content)
        stats['links'] = links
        
        # Extraer imágenes
        img_pattern = r'!\[([^\]]*)\]\(([^\)]+)\)'
        images = re.findall(img_pattern, content)
        stats['images'] = images
        
        # Identificar secciones principales
        current_section = None
        for line in lines:
            if line.strip().startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                if level <= 2:  # Solo secciones principales
                    current_section = line.lstrip('#').strip()
                    stats['sections'].append({
                        'title': current_section,
                        'level': level,
                        'line': lines.index(line)
                    })
        
        return {
            'content': content,
            'html': html,
            'stats': stats,
            'filename': Path(file_path).stem
        }
    
    def create_word_document(self, parsed_data: Dict[str, Any], output_path: str):
        """Crea un documento Word desde datos parseados."""
        doc = Document()
        
        # Título principal
        title = doc.add_heading(parsed_data['filename'].replace('_', ' ').title(), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Fecha
        date_para = doc.add_paragraph(f"Generado el: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.runs[0].font.size = Pt(10)
        date_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_paragraph()  # Espacio
        
        # Estadísticas
        stats = parsed_data['stats']
        stats_heading = doc.add_heading('📊 Estadísticas del Documento', 1)
        
        stats_table = doc.add_table(rows=1, cols=2)
        stats_table.style = 'Light Grid Accent 1'
        hdr_cells = stats_table.rows[0].cells
        hdr_cells[0].text = 'Métrica'
        hdr_cells[1].text = 'Valor'
        
        metrics = [
            ('Total de líneas', stats['total_lines']),
            ('Total de palabras', stats['total_words']),
            ('Total de caracteres', stats['total_chars']),
            ('Encabezados', len(stats['headers'])),
            ('Bloques de código', len(stats['code_blocks'])),
            ('Tablas', len(stats['tables'])),
            ('Enlaces', len(stats['links'])),
            ('Imágenes', len(stats['images'])),
            ('Secciones principales', len(stats['sections']))
        ]
        
        for metric, value in metrics:
            row_cells = stats_table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = str(value)
        
        doc.add_paragraph()  # Espacio
        
        # Secciones principales
        if stats['sections']:
            sections_heading = doc.add_heading('📑 Índice de Secciones', 1)
            sections_list = doc.add_paragraph(style='List Bullet')
            for section in stats['sections'][:20]:  # Limitar a 20
                sections_list.add_run(f"{'  ' * (section['level'] - 1)}{section['title']}\n")
        
        doc.add_page_break()
        
        # Contenido principal (simplificado)
        content_heading = doc.add_heading('📄 Contenido', 1)
        
        # Procesar contenido línea por línea
        lines = parsed_data['content'].split('\n')
        for i, line in enumerate(lines[:500]):  # Limitar para no hacer el doc muy grande
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
            
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                text = line.lstrip('#').strip()
                doc.add_heading(text, min(level, 9))
            elif line.startswith('-') or line.startswith('*'):
                p = doc.add_paragraph(line, style='List Bullet')
            elif line.startswith('```'):
                # Código - omitir por ahora
                continue
            elif '|' in line and line.count('|') >= 2:
                # Tabla - procesar básicamente
                continue
            else:
                doc.add_paragraph(line)
        
        # Guardar
        doc.save(output_path)
        print(f"✅ Documento Word creado: {output_path}")
    
    def create_excel_workbook(self, parsed_data: Dict[str, Any], output_path: str):
        """Crea un libro Excel con datos estructurados y gráficas."""
        wb = Workbook()
        
        # Hoja 1: Estadísticas
        ws_stats = wb.active
        ws_stats.title = "Estadísticas"
        
        # Título
        ws_stats['A1'] = parsed_data['filename'].replace('_', ' ').title()
        ws_stats['A1'].font = Font(size=16, bold=True)
        ws_stats['A2'] = f"Generado el: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
        ws_stats['A2'].font = Font(size=10, italic=True)
        
        # Datos estadísticos
        stats = parsed_data['stats']
        row = 4
        ws_stats[f'A{row}'] = 'Métrica'
        ws_stats[f'B{row}'] = 'Valor'
        ws_stats[f'A{row}'].font = Font(bold=True)
        ws_stats[f'B{row}'].font = Font(bold=True)
        
        metrics_data = [
            ('Total de líneas', stats['total_lines']),
            ('Total de palabras', stats['total_words']),
            ('Total de caracteres', stats['total_chars']),
            ('Encabezados', len(stats['headers'])),
            ('Bloques de código', len(stats['code_blocks'])),
            ('Tablas', len(stats['tables'])),
            ('Enlaces', len(stats['links'])),
            ('Imágenes', len(stats['images'])),
            ('Secciones principales', len(stats['sections']))
        ]
        
        for metric, value in metrics_data:
            row += 1
            ws_stats[f'A{row}'] = metric
            ws_stats[f'B{row}'] = value
        
        # Ajustar columnas
        ws_stats.column_dimensions['A'].width = 25
        ws_stats.column_dimensions['B'].width = 15
        
        # Gráfica de estadísticas
        chart = BarChart()
        chart.type = "col"
        chart.style = 10
        chart.title = "Estadísticas del Documento"
        chart.y_axis.title = 'Cantidad'
        chart.x_axis.title = 'Métricas'
        
        data = Reference(ws_stats, min_col=2, min_row=4, max_row=row)
        cats = Reference(ws_stats, min_col=1, min_row=5, max_row=row)
        chart.add_data(data, titles_from_data=False)
        chart.set_categories(cats)
        ws_stats.add_chart(chart, "D4")
        
        # Hoja 2: Encabezados
        ws_headers = wb.create_sheet("Encabezados")
        ws_headers['A1'] = 'Nivel'
        ws_headers['B1'] = 'Texto'
        ws_headers['A1'].font = Font(bold=True)
        ws_headers['B1'].font = Font(bold=True)
        
        for i, header in enumerate(stats['headers'][:100], start=2):  # Limitar a 100
            ws_headers[f'A{i}'] = header['level']
            ws_headers[f'B{i}'] = header['text']
        
        ws_headers.column_dimensions['A'].width = 10
        ws_headers.column_dimensions['B'].width = 50
        
        # Gráfica de distribución de niveles
        if stats['headers']:
            level_counts = {}
            for h in stats['headers']:
                level_counts[h['level']] = level_counts.get(h['level'], 0) + 1
            
            ws_levels = wb.create_sheet("Distribución Niveles")
            ws_levels['A1'] = 'Nivel'
            ws_levels['B1'] = 'Cantidad'
            ws_levels['A1'].font = Font(bold=True)
            ws_levels['B1'].font = Font(bold=True)
            
            row = 2
            for level, count in sorted(level_counts.items()):
                ws_levels[f'A{row}'] = f"Nivel {level}"
                ws_levels[f'B{row}'] = count
                row += 1
            
            # Gráfica de pastel
            pie_chart = PieChart()
            pie_chart.title = "Distribución de Niveles de Encabezados"
            data = Reference(ws_levels, min_col=2, min_row=1, max_row=row-1)
            labels = Reference(ws_levels, min_col=1, min_row=2, max_row=row-1)
            pie_chart.add_data(data, titles_from_data=False)
            pie_chart.set_categories(labels)
            ws_levels.add_chart(pie_chart, "D2")
        
        # Hoja 3: Enlaces
        if stats['links']:
            ws_links = wb.create_sheet("Enlaces")
            ws_links['A1'] = 'Texto'
            ws_links['B1'] = 'URL'
            ws_links['A1'].font = Font(bold=True)
            ws_links['B1'].font = Font(bold=True)
            
            for i, (text, url) in enumerate(stats['links'][:200], start=2):  # Limitar a 200
                ws_links[f'A{i}'] = text[:50]  # Limitar texto
                ws_links[f'B{i}'] = url
            
            ws_links.column_dimensions['A'].width = 30
            ws_links.column_dimensions['B'].width = 60
        
        # Hoja 4: Secciones
        if stats['sections']:
            ws_sections = wb.create_sheet("Secciones")
            ws_sections['A1'] = 'Nivel'
            ws_sections['B1'] = 'Título'
            ws_sections['C1'] = 'Línea'
            ws_sections['A1'].font = Font(bold=True)
            ws_sections['B1'].font = Font(bold=True)
            ws_sections['C1'].font = Font(bold=True)
            
            for i, section in enumerate(stats['sections'], start=2):
                ws_sections[f'A{i}'] = section['level']
                ws_sections[f'B{i}'] = section['title']
                ws_sections[f'C{i}'] = section['line']
            
            ws_sections.column_dimensions['A'].width = 10
            ws_sections.column_dimensions['B'].width = 50
            ws_sections.column_dimensions['C'].width = 10
        
        # Guardar
        wb.save(output_path)
        print(f"✅ Libro Excel creado: {output_path}")
    
    def create_visualizations(self, parsed_data: Dict[str, Any], output_dir: Path):
        """Crea visualizaciones y las guarda como imágenes."""
        stats = parsed_data['stats']
        filename = parsed_data['filename']
        
        # Gráfica 1: Distribución de contenido
        fig, axes = plt.subplots(2, 2, figsize=(12, 10))
        fig.suptitle(f'Análisis Visual: {filename}', fontsize=16, fontweight='bold')
        
        # Subplot 1: Distribución de niveles de encabezados
        if stats['headers']:
            level_counts = {}
            for h in stats['headers']:
                level_counts[h['level']] = level_counts.get(h['level'], 0) + 1
            
            axes[0, 0].bar(level_counts.keys(), level_counts.values(), color='steelblue')
            axes[0, 0].set_title('Distribución de Niveles de Encabezados')
            axes[0, 0].set_xlabel('Nivel')
            axes[0, 0].set_ylabel('Cantidad')
            axes[0, 0].grid(True, alpha=0.3)
        
        # Subplot 2: Métricas principales
        metrics = ['Líneas', 'Palabras', 'Encabezados', 'Enlaces']
        values = [
            stats['total_lines'],
            stats['total_words'],
            len(stats['headers']),
            len(stats['links'])
        ]
        axes[0, 1].barh(metrics, values, color='coral')
        axes[0, 1].set_title('Métricas Principales')
        axes[0, 1].grid(True, alpha=0.3)
        
        # Subplot 3: Tipos de contenido
        content_types = ['Código', 'Tablas', 'Enlaces', 'Imágenes']
        content_counts = [
            len(stats['code_blocks']),
            len(stats['tables']),
            len(stats['links']),
            len(stats['images'])
        ]
        axes[1, 0].pie(content_counts, labels=content_types, autopct='%1.1f%%', startangle=90)
        axes[1, 0].set_title('Distribución de Tipos de Contenido')
        
        # Subplot 4: Palabras por sección (aproximado)
        if stats['sections']:
            section_names = [s['title'][:20] for s in stats['sections'][:10]]
            section_sizes = [100] * len(section_names)  # Placeholder
            axes[1, 1].barh(range(len(section_names)), section_sizes, color='lightgreen')
            axes[1, 1].set_yticks(range(len(section_names)))
            axes[1, 1].set_yticklabels(section_names)
            axes[1, 1].set_title('Secciones Principales')
            axes[1, 1].grid(True, alpha=0.3)
        
        plt.tight_layout()
        img_path = output_dir / f"{filename}_analisis.png"
        plt.savefig(img_path, dpi=300, bbox_inches='tight')
        plt.close()
        
        print(f"✅ Visualización creada: {img_path}")
        return img_path
    
    def convert(self, markdown_file: str):
        """Convierte un archivo Markdown a Word y Excel."""
        print(f"\n🔄 Procesando: {markdown_file}")
        
        # Parsear
        parsed_data = self.parse_markdown(markdown_file)
        filename = parsed_data['filename']
        
        # Crear Word
        word_path = self.output_dir / f"{filename}.docx"
        self.create_word_document(parsed_data, str(word_path))
        
        # Crear Excel
        excel_path = self.output_dir / f"{filename}.xlsx"
        self.create_excel_workbook(parsed_data, str(excel_path))
        
        # Crear visualizaciones
        img_path = self.create_visualizations(parsed_data, self.output_dir)
        
        # Guardar estadísticas
        self.stats[filename] = parsed_data['stats']
        
        print(f"✅ Conversión completada para: {filename}")
        return {
            'word': str(word_path),
            'excel': str(excel_path),
            'image': str(img_path)
        }


def main():
    """Función principal."""
    # Archivos importantes a convertir
    base_dir = Path("/Users/adan/Documents/documentos_blatam")
    important_files = [
        base_dir / "01_Marketing" / "dm_linkedin_INDICE_MAESTRO.md",
        base_dir / "01_marketing" / "FAQ_OUTREACH.md",
        base_dir / "README.md"
    ]
    
    converter = MarkdownToWordExcel(output_dir=str(base_dir / "exports"))
    
    results = []
    for file_path in important_files:
        if file_path.exists():
            try:
                result = converter.convert(str(file_path))
                results.append(result)
            except Exception as e:
                print(f"❌ Error procesando {file_path}: {e}")
        else:
            print(f"⚠️  Archivo no encontrado: {file_path}")
    
    # Resumen
    print("\n" + "="*60)
    print("📊 RESUMEN DE CONVERSIÓN")
    print("="*60)
    for result in results:
        print(f"\n✅ {Path(result['word']).stem}:")
        print(f"   - Word: {result['word']}")
        print(f"   - Excel: {result['excel']}")
        print(f"   - Imagen: {result['image']}")
    
    print("\n✨ ¡Conversión completada!")


if __name__ == "__main__":
    main()



