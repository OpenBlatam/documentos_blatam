#!/usr/bin/env python3
"""
Genera versiones PDF y Word con formato profesional de un archivo Python.
Incluye resumen ejecutivo del código, numeración de líneas y metadatos.
"""

from __future__ import annotations

import argparse
import re
import sys
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm, inch
from reportlab.pdfgen import canvas
from reportlab.platypus import (
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
@dataclass
class CodeMetrics:
    total_lines: int
    non_empty_lines: int
    functions: int
    async_functions: int
    classes: int
    imports: int
    comments: int
    docstrings: int


def analyze_code(content: str) -> CodeMetrics:
    """Calcula métricas básicas del código para incluirlas en los reportes."""
    lines = content.splitlines()
    total_lines = len(lines)
    non_empty_lines = sum(1 for line in lines if line.strip())

    async_functions = len(re.findall(r"^\s*async\s+def\s+", content, re.MULTILINE))
    functions = len(re.findall(r"^\s*def\s+", content, re.MULTILINE))
    classes = len(re.findall(r"^\s*class\s+", content, re.MULTILINE))
    imports = len(re.findall(r"^\s*(from\s+\S+\s+import|import\s+\S+)", content, re.MULTILINE))
    comments = len(re.findall(r"^\s*#", content, re.MULTILINE))
    docstrings = len(re.findall(r'("""[\s\S]*?"""|\'\'\'[\s\S]*?\'\'\')', content))

    return CodeMetrics(
        total_lines=total_lines,
        non_empty_lines=non_empty_lines,
        functions=functions - async_functions,
        async_functions=async_functions,
        classes=classes,
        imports=imports,
        comments=comments,
        docstrings=docstrings,
    )

def highlight_python_syntax(line):
    """Aplica colores básicos de sintaxis Python"""
    # Palabras clave
    keywords = ['def', 'class', 'import', 'from', 'if', 'elif', 'else', 'for', 'while', 
                'try', 'except', 'finally', 'with', 'as', 'return', 'yield', 'async', 'await',
                'True', 'False', 'None', 'and', 'or', 'not', 'in', 'is', 'pass', 'break', 'continue']
    
    # Strings
    line = re.sub(r'("""[\s\S]*?""")', r'<font color="#008000">\1</font>', line)
    line = re.sub(r"('''[\s\S]*?''')", r'<font color="#008000">\1</font>', line)
    line = re.sub(r'(".*?")', r'<font color="#008000">\1</font>', line)
    line = re.sub(r"('.*?')", r'<font color="#008000">\1</font>', line)
    
    # Comentarios
    line = re.sub(r'(#.*)', r'<font color="#808080">\1</font>', line)
    
    # Palabras clave
    for keyword in keywords:
        pattern = r'\b(' + keyword + r')\b'
        line = re.sub(pattern, r'<font color="#0000FF"><b>\1</b></font>', line)
    
    # Números
    line = re.sub(r'\b(\d+\.?\d*)\b', r'<font color="#FF0000">\1</font>', line)
    
    return line


def create_pdf(input_path: Path, output_path: Path, content: str, metrics: CodeMetrics) -> None:
    """Crea un PDF mejorado del archivo Python."""
    lines = content.splitlines()
    
    # Crear documento PDF con header/footer personalizado
    class NumberedCanvas(canvas.Canvas):
        def __init__(self, *args, **kwargs):
            canvas.Canvas.__init__(self, *args, **kwargs)
            self._saved_page_states = []
        
        def showPage(self):
            self._saved_page_states.append(dict(self.__dict__))
            self._startPage()
        
        def save(self):
            num_pages = len(self._saved_page_states)
            for state in self._saved_page_states:
                self.__dict__.update(state)
                self.draw_page_number(num_pages)
                canvas.Canvas.showPage(self)
            canvas.Canvas.save(self)
        
        def draw_page_number(self, page_count):
            self.saveState()
            self.setFont("Helvetica", 9)
            page_text = f"Página {self._pageNumber} de {page_count}"
            self.drawRightString(19.5*cm, 1*cm, page_text)
            self.restoreState()
    
    doc = SimpleDocTemplate(
        output_file,
        pagesize=A4,
        rightMargin=2*cm,
        leftMargin=2*cm,
        topMargin=3*cm,
        bottomMargin=2.5*cm
    )
    
    # Estilos mejorados
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=20,
        textColor=colors.HexColor('#1a1a1a'),
        spaceAfter=12,
        spaceBefore=0,
        alignment=TA_LEFT,
        fontName='Helvetica-Bold'
    )
    
    info_style = ParagraphStyle(
        'InfoStyle',
        parent=styles['Normal'],
        fontSize=9,
        textColor=colors.HexColor('#666666'),
        spaceAfter=15,
        alignment=TA_LEFT,
        fontName='Helvetica'
    )
    
    code_style = ParagraphStyle(
        'CodeStyle',
        parent=styles['Code'],
        fontSize=7.5,
        fontName='Courier',
        leading=9,
        leftIndent=0,
        rightIndent=0,
        alignment=TA_LEFT,
        textColor=colors.HexColor('#000000'),
        backColor=colors.HexColor('#f5f5f5')
    )
    
    # Contenido
    story = []
    
    # Título
    title = file_path.name
    story.append(Paragraph(f"<b>{title}</b>", title_style))
    
    # Información del archivo
    file_info = f"""
    <b>Archivo:</b> {file_path.absolute()}<br/>
    <b>Líneas de código:</b> {total_lines}<br/>
    <b>Fecha de generación:</b> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
    """
    story.append(Paragraph(file_info, info_style))

    summary_data = [
        ["Líneas no vacías", str(metrics.non_empty_lines)],
        ["Funciones síncronas", str(metrics.functions)],
        ["Funciones asíncronas", str(metrics.async_functions)],
        ["Clases", str(metrics.classes)],
        ["Importaciones", str(metrics.imports)],
        ["Comentarios", str(metrics.comments)],
        ["Docstrings", str(metrics.docstrings)],
    ]
    summary_table = Table(summary_data, colWidths=[doc.width * 0.55, doc.width * 0.45])
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#ededed')),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f9f9f9')]),
        ('BOX', (0, 0), (-1, -1), 0.75, colors.HexColor('#cccccc')),
        ('INNERGRID', (0, 0), (-1, -1), 0.25, colors.HexColor('#d5d5d5')),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('LEFTPADDING', (0, 0), (-1, -1), 6),
        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
    ]))
    story.append(summary_table)
    story.append(Spacer(1, 0.3*inch))
    
    # Código con numeración de líneas - dividir en bloques
    lines_per_block = 50  # Líneas por bloque para evitar tablas muy grandes
    
    for block_start in range(0, len(lines), lines_per_block):
        block_end = min(block_start + lines_per_block, len(lines))
        block_lines = lines[block_start:block_end]
        
        # Crear texto con numeración
        code_text = "\n".join(
            [f"{i:4d} | {line}" for i, line in enumerate(block_lines, start=block_start + 1)]
        )
        
        # Crear tabla para cada bloque
        code_table = Table([[Preformatted(code_text, code_style, maxLineLength=120)]], 
                           colWidths=[doc.width])
        code_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f8f8f8')),
            ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#cccccc')),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ]))
        
        story.append(code_table)
        
        # Agregar espacio entre bloques (excepto el último)
        if block_end < len(lines):
            story.append(Spacer(1, 0.1*inch))
    
    # Construir PDF con canvas personalizado
    doc.build(story, canvasmaker=NumberedCanvas)
    print(f"✓ PDF mejorado creado: {output_file}")


def create_word(input_file, output_file):
    """Crea un documento Word mejorado del archivo Python"""
    # Leer el archivo
    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.splitlines()
    file_path = input_path
    total_lines = metrics.total_lines
    
    # Crear documento Word
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Título
    title = doc.add_heading(file_path.name, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Información del archivo
    info_para = doc.add_paragraph()
    info_para.add_run('Archivo: ').bold = True
    info_para.add_run(str(file_path.absolute()))
    info_para.add_run('\nLíneas de código: ').bold = True
    info_para.add_run(str(total_lines))
    info_para.add_run('\nFecha de generación: ').bold = True
    info_para.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    
    doc.add_paragraph()  # Espacio
    
    # Crear tabla para el código con fondo
    table = doc.add_table(rows=len(lines), cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Configurar ancho de columnas
    table.columns[0].width = Inches(0.6)  # Columna de números
    table.columns[1].width = Inches(6.4)  # Columna de código
    
    # Agregar código línea por línea
    for i, line in enumerate(lines):
        row = table.rows[i]
        
        # Número de línea
        cell_num = row.cells[0]
        para_num = cell_num.paragraphs[0]
        run_num = para_num.add_run(f"{i+1:4d}")
        run_num.font.name = 'Courier New'
        run_num.font.size = Pt(8)
        run_num.font.color.rgb = RGBColor(128, 128, 128)
        para_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Código
        cell_code = row.cells[1]
        para_code = cell_code.paragraphs[0]
        run_code = para_code.add_run(line)
        run_code.font.name = 'Courier New'
        run_code.font.size = Pt(9)
        
        # Aplicar colores básicos de sintaxis
        line_stripped = line.strip()
        if line_stripped.startswith('#'):
            # Comentarios grises
            run_code.font.color.rgb = RGBColor(128, 128, 128)
        elif line_stripped.startswith('"""') or line_stripped.startswith("'''"):
            # Docstrings verdes
            run_code.font.color.rgb = RGBColor(0, 128, 0)
        elif '"' in line or "'" in line:
            # Strings verdes (simple)
            if line.count('"') >= 2 or line.count("'") >= 2:
                run_code.font.color.rgb = RGBColor(0, 128, 0)
    
    # Guardar
    doc.save(output_file)
    print(f"✓ Word mejorado creado: {output_file}")


if __name__ == "__main__":
    input_file = "api_unified.py"
    
    if not Path(input_file).exists():
        print(f"Error: No se encontró el archivo {input_file}")
        sys.exit(1)
    
    # Crear PDF
    pdf_output = input_file.replace('.py', '.pdf')
    create_pdf(input_file, pdf_output)
    
    # Crear Word
    word_output = input_file.replace('.py', '.docx')
    create_word(input_file, word_output)
    
    print(f"\n✓ Documentos generados exitosamente:")
    print(f"  - {pdf_output}")
    print(f"  - {word_output}")

