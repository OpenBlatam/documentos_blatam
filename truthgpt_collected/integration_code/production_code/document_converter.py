#!/usr/bin/env python3
"""
Document Converter - Conversor Universal de Documentos
======================================================

Módulo completo para convertir documentos a PDF, Word y Excel
con múltiples librerías y opciones avanzadas.

Soporta:
- PDF: reportlab, fpdf, weasyprint, pdfkit, PyPDF2, PyMuPDF
- Word: python-docx, docx2pdf, python-docx-template
- Excel: openpyxl, xlsxwriter, pandas, xlswriter
"""

__version__ = '1.0.0'

import os
import io
import json
from pathlib import Path
from typing import Dict, Any, List, Optional, Union, Tuple
from datetime import datetime
from enum import Enum

from core.utils import setup_logger

logger = setup_logger(__name__)


class ExportFormat(Enum):
    """Formatos de exportación soportados"""
    PDF = "pdf"
    WORD = "docx"
    EXCEL = "xlsx"


class DocumentConverter:
    """Conversor universal de documentos con múltiples librerías"""
    
    def __init__(self):
        self.logger = logger
        self._check_libraries()
    
    def _check_libraries(self):
        """Verifica qué librerías están disponibles"""
        self.available_libs = {
            'pdf': {
                'reportlab': self._check_import('reportlab'),
                'fpdf': self._check_import('fpdf'),
                'weasyprint': self._check_import('weasyprint'),
                'pdfkit': self._check_import('pdfkit'),
                'PyPDF2': self._check_import('PyPDF2'),
                'PyMuPDF': self._check_import('fitz'),  # PyMuPDF se importa como fitz
            },
            'word': {
                'python-docx': self._check_import('docx'),
                'docx2pdf': self._check_import('docx2pdf'),
            },
            'excel': {
                'openpyxl': self._check_import('openpyxl'),
                'xlsxwriter': self._check_import('xlsxwriter'),
                'pandas': self._check_import('pandas'),
            }
        }
        
        # Log de librerías disponibles
        for format_type, libs in self.available_libs.items():
            available = [name for name, available in libs.items() if available]
            if available:
                self.logger.info(f"Librerías {format_type} disponibles: {', '.join(available)}")
            else:
                self.logger.warning(f"No hay librerías {format_type} disponibles")
    
    def _check_import(self, module_name: str) -> bool:
        """Verifica si un módulo puede ser importado"""
        try:
            __import__(module_name)
            return True
        except ImportError:
            return False
    
    def convert_to_pdf(
        self,
        data: Union[Dict, List, str],
        output_path: str,
        method: str = "auto",
        title: str = "Documento",
        **kwargs
    ) -> Dict[str, Any]:
        """
        Convierte datos a PDF usando múltiples métodos.
        
        Args:
            data: Datos a convertir (dict, list, o string)
            output_path: Ruta del archivo PDF de salida
            method: Método a usar ('auto', 'reportlab', 'fpdf', 'weasyprint', 'pdfkit')
            title: Título del documento
            **kwargs: Opciones adicionales
        
        Returns:
            Dict con resultado de la conversión
        """
        if method == "auto":
            # Seleccionar mejor método disponible
            if self.available_libs['pdf']['reportlab']:
                method = "reportlab"
            elif self.available_libs['pdf']['fpdf']:
                method = "fpdf"
            elif self.available_libs['pdf']['weasyprint']:
                method = "weasyprint"
            elif self.available_libs['pdf']['pdfkit']:
                method = "pdfkit"
            else:
                raise ImportError("No hay librerías PDF disponibles. Instala: pip install reportlab fpdf2 weasyprint pdfkit")
        
        try:
            if method == "reportlab":
                return self._convert_to_pdf_reportlab(data, output_path, title, **kwargs)
            elif method == "fpdf":
                return self._convert_to_pdf_fpdf(data, output_path, title, **kwargs)
            elif method == "weasyprint":
                return self._convert_to_pdf_weasyprint(data, output_path, title, **kwargs)
            elif method == "pdfkit":
                return self._convert_to_pdf_pdfkit(data, output_path, title, **kwargs)
            else:
                raise ValueError(f"Método PDF no soportado: {method}")
        except Exception as e:
            self.logger.error(f"Error convirtiendo a PDF con {method}: {e}", exc_info=True)
            raise
    
    def _convert_to_pdf_reportlab(
        self,
        data: Union[Dict, List, str],
        output_path: str,
        title: str,
        **kwargs
    ) -> Dict[str, Any]:
        """Convierte a PDF usando ReportLab"""
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        
        doc = SimpleDocTemplate(output_path, pagesize=A4)
        story = []
        styles = getSampleStyleSheet()
        
        # Título
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=24,
            textColor=colors.HexColor('#1a1a1a'),
            spaceAfter=30,
            alignment=1  # Centrado
        )
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.3 * inch))
        
        # Procesar datos
        if isinstance(data, str):
            story.append(Paragraph(data, styles['Normal']))
        elif isinstance(data, dict):
            # Crear tabla de datos
            table_data = [['Campo', 'Valor']]
            for key, value in data.items():
                table_data.append([str(key), str(value)])
            
            table = Table(table_data, colWidths=[2*inch, 4*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#366092')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
            ]))
            story.append(table)
        elif isinstance(data, list):
            if data and isinstance(data[0], dict):
                # Lista de diccionarios - tabla
                if data:
                    headers = list(data[0].keys())
                    table_data = [headers]
                    for item in data:
                        table_data.append([str(item.get(h, '')) for h in headers])
                    
                    table = Table(table_data, repeatRows=1)
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#366092')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 10),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ]))
                    story.append(table)
            else:
                # Lista simple
                for item in data:
                    story.append(Paragraph(str(item), styles['Normal']))
                    story.append(Spacer(1, 0.1 * inch))
        
        # Pie de página
        story.append(Spacer(1, 0.5 * inch))
        story.append(Paragraph(
            f"Generado el {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            styles['Normal']
        ))
        
        doc.build(story)
        
        return {
            "success": True,
            "output_path": output_path,
            "method": "reportlab",
            "file_size": os.path.getsize(output_path)
        }
    
    def _convert_to_pdf_fpdf(
        self,
        data: Union[Dict, List, str],
        output_path: str,
        title: str,
        **kwargs
    ) -> Dict[str, Any]:
        """Convierte a PDF usando FPDF"""
        from fpdf import FPDF
        
        class PDF(FPDF):
            def header(self):
                self.set_font('Arial', 'B', 15)
                self.cell(0, 10, title, 0, 1, 'C')
                self.ln(10)
            
            def footer(self):
                self.set_y(-15)
                self.set_font('Arial', 'I', 8)
                self.cell(0, 10, f'Página {self.page_no()}', 0, 0, 'C')
        
        pdf = PDF()
        pdf.add_page()
        pdf.set_font('Arial', '', 12)
        
        if isinstance(data, str):
            pdf.multi_cell(0, 10, data)
        elif isinstance(data, dict):
            for key, value in data.items():
                pdf.set_font('Arial', 'B', 12)
                pdf.cell(0, 10, f"{key}:", 0, 1)
                pdf.set_font('Arial', '', 12)
                pdf.multi_cell(0, 10, str(value))
                pdf.ln(5)
        elif isinstance(data, list):
            for item in data:
                pdf.multi_cell(0, 10, str(item))
                pdf.ln(5)
        
        pdf.output(output_path)
        
        return {
            "success": True,
            "output_path": output_path,
            "method": "fpdf",
            "file_size": os.path.getsize(output_path)
        }
    
    def _convert_to_pdf_weasyprint(
        self,
        data: Union[Dict, List, str],
        output_path: str,
        title: str,
        **kwargs
    ) -> Dict[str, Any]:
        """Convierte a PDF usando WeasyPrint"""
        from weasyprint import HTML, CSS
        from weasyprint.text.fonts import FontConfiguration
        
        # Generar HTML
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>{title}</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; }}
                h1 {{ color: #366092; text-align: center; }}
                table {{ width: 100%; border-collapse: collapse; margin: 20px 0; }}
                th {{ background-color: #366092; color: white; padding: 10px; text-align: left; }}
                td {{ padding: 8px; border: 1px solid #ddd; }}
                tr:nth-child(even) {{ background-color: #f2f2f2; }}
            </style>
        </head>
        <body>
            <h1>{title}</h1>
        """
        
        if isinstance(data, str):
            html_content += f"<p>{data}</p>"
        elif isinstance(data, dict):
            html_content += "<table><tr><th>Campo</th><th>Valor</th></tr>"
            for key, value in data.items():
                html_content += f"<tr><td>{key}</td><td>{value}</td></tr>"
            html_content += "</table>"
        elif isinstance(data, list):
            if data and isinstance(data[0], dict):
                if data:
                    headers = list(data[0].keys())
                    html_content += f"<table><tr>{''.join(f'<th>{h}</th>' for h in headers)}</tr>"
                    for item in data:
                        html_content += f"<tr>{''.join(f'<td>{item.get(h, \"\")}</td>' for h in headers)}</tr>"
                    html_content += "</table>"
            else:
                html_content += "<ul>"
                for item in data:
                    html_content += f"<li>{item}</li>"
                html_content += "</ul>"
        
        html_content += f"""
            <p style="margin-top: 40px; font-size: 10px; color: #666;">
                Generado el {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
            </p>
        </body>
        </html>
        """
        
        HTML(string=html_content).write_pdf(output_path)
        
        return {
            "success": True,
            "output_path": output_path,
            "method": "weasyprint",
            "file_size": os.path.getsize(output_path)
        }
    
    def _convert_to_pdf_pdfkit(
        self,
        data: Union[Dict, List, str],
        output_path: str,
        title: str,
        **kwargs
    ) -> Dict[str, Any]:
        """Convierte a PDF usando pdfkit (requiere wkhtmltopdf)"""
        import pdfkit
        
        # Similar a weasyprint pero usando pdfkit
        html_content = self._generate_html(data, title)
        
        options = {
            'page-size': 'A4',
            'margin-top': '0.75in',
            'margin-right': '0.75in',
            'margin-bottom': '0.75in',
            'margin-left': '0.75in',
            'encoding': "UTF-8",
            'no-outline': None
        }
        
        pdfkit.from_string(html_content, output_path, options=options)
        
        return {
            "success": True,
            "output_path": output_path,
            "method": "pdfkit",
            "file_size": os.path.getsize(output_path)
        }
    
    def _generate_html(self, data: Union[Dict, List, str], title: str) -> str:
        """Genera HTML a partir de datos"""
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>{title}</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; }}
                h1 {{ color: #366092; text-align: center; }}
                table {{ width: 100%; border-collapse: collapse; margin: 20px 0; }}
                th {{ background-color: #366092; color: white; padding: 10px; text-align: left; }}
                td {{ padding: 8px; border: 1px solid #ddd; }}
                tr:nth-child(even) {{ background-color: #f2f2f2; }}
            </style>
        </head>
        <body>
            <h1>{title}</h1>
        """
        
        if isinstance(data, str):
            html_content += f"<p>{data}</p>"
        elif isinstance(data, dict):
            html_content += "<table><tr><th>Campo</th><th>Valor</th></tr>"
            for key, value in data.items():
                html_content += f"<tr><td>{key}</td><td>{value}</td></tr>"
            html_content += "</table>"
        elif isinstance(data, list):
            if data and isinstance(data[0], dict):
                if data:
                    headers = list(data[0].keys())
                    html_content += f"<table><tr>{''.join(f'<th>{h}</th>' for h in headers)}</tr>"
                    for item in data:
                        html_content += f"<tr>{''.join(f'<td>{item.get(h, \"\")}</td>' for h in headers)}</tr>"
                    html_content += "</table>"
            else:
                html_content += "<ul>"
                for item in data:
                    html_content += f"<li>{item}</li>"
                html_content += "</ul>"
        
        html_content += f"""
            <p style="margin-top: 40px; font-size: 10px; color: #666;">
                Generado el {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
            </p>
        </body>
        </html>
        """
        return html_content
    
    def convert_to_word(
        self,
        data: Union[Dict, List, str],
        output_path: str,
        title: str = "Documento",
        **kwargs
    ) -> Dict[str, Any]:
        """
        Convierte datos a Word (.docx).
        
        Args:
            data: Datos a convertir
            output_path: Ruta del archivo Word de salida
            title: Título del documento
            **kwargs: Opciones adicionales
        
        Returns:
            Dict con resultado de la conversión
        """
        if not self.available_libs['word']['python-docx']:
            raise ImportError("python-docx no está disponible. Instala con: pip install python-docx")
        
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            
            doc = Document()
            
            # Configurar estilos
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(11)
            
            # Título
            title_para = doc.add_heading(title, 0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.runs[0]
            title_run.font.color.rgb = RGBColor(54, 96, 146)
            title_run.font.size = Pt(24)
            
            doc.add_paragraph()  # Espacio
            
            # Procesar datos
            if isinstance(data, str):
                doc.add_paragraph(data)
            elif isinstance(data, dict):
                # Crear tabla
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Light Grid Accent 1'
                
                # Encabezados
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Campo'
                hdr_cells[1].text = 'Valor'
                
                # Estilo encabezados
                for cell in hdr_cells:
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                    shading_elm = cell._element.get_or_add_tcPr()
                    shading = shading_elm.get_or_add_shd()
                    shading.set(qn('w:fill'), '366092')
                
                # Datos
                for key, value in data.items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(key)
                    row_cells[1].text = str(value)
            elif isinstance(data, list):
                if data and isinstance(data[0], dict):
                    # Lista de diccionarios - tabla
                    if data:
                        headers = list(data[0].keys())
                        table = doc.add_table(rows=1, cols=len(headers))
                        table.style = 'Light Grid Accent 1'
                        
                        # Encabezados
                        hdr_cells = table.rows[0].cells
                        for i, header in enumerate(headers):
                            hdr_cells[i].text = str(header)
                            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                            hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                            shading_elm = hdr_cells[i]._element.get_or_add_tcPr()
                            shading = shading_elm.get_or_add_shd()
                            shading.set(qn('w:fill'), '366092')
                        
                        # Datos
                        for item in data:
                            row_cells = table.add_row().cells
                            for i, header in enumerate(headers):
                                row_cells[i].text = str(item.get(header, ''))
                else:
                    # Lista simple
                    for item in data:
                        doc.add_paragraph(str(item), style='List Bullet')
            
            # Pie de página
            doc.add_paragraph()
            para = doc.add_paragraph(f"Generado el {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            para.runs[0].font.size = Pt(9)
            para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
            
            # Guardar
            doc.save(output_path)
            
            return {
                "success": True,
                "output_path": output_path,
                "method": "python-docx",
                "file_size": os.path.getsize(output_path)
            }
        except Exception as e:
            self.logger.error(f"Error convirtiendo a Word: {e}", exc_info=True)
            raise
    
    def convert_to_excel(
        self,
        data: Union[Dict, List, str],
        output_path: str,
        sheet_name: str = "Datos",
        method: str = "auto",
        **kwargs
    ) -> Dict[str, Any]:
        """
        Convierte datos a Excel (.xlsx).
        
        Args:
            data: Datos a convertir
            output_path: Ruta del archivo Excel de salida
            sheet_name: Nombre de la hoja
            method: Método a usar ('auto', 'openpyxl', 'xlsxwriter', 'pandas')
            **kwargs: Opciones adicionales
        
        Returns:
            Dict con resultado de la conversión
        """
        if method == "auto":
            if self.available_libs['excel']['openpyxl']:
                method = "openpyxl"
            elif self.available_libs['excel']['xlsxwriter']:
                method = "xlsxwriter"
            elif self.available_libs['excel']['pandas']:
                method = "pandas"
            else:
                raise ImportError("No hay librerías Excel disponibles. Instala: pip install openpyxl xlsxwriter pandas")
        
        try:
            if method == "openpyxl":
                return self._convert_to_excel_openpyxl(data, output_path, sheet_name, **kwargs)
            elif method == "xlsxwriter":
                return self._convert_to_excel_xlsxwriter(data, output_path, sheet_name, **kwargs)
            elif method == "pandas":
                return self._convert_to_excel_pandas(data, output_path, sheet_name, **kwargs)
            else:
                raise ValueError(f"Método Excel no soportado: {method}")
        except Exception as e:
            self.logger.error(f"Error convirtiendo a Excel con {method}: {e}", exc_info=True)
            raise
    
    def _convert_to_excel_openpyxl(
        self,
        data: Union[Dict, List, str],
        output_path: str,
        sheet_name: str,
        **kwargs
    ) -> Dict[str, Any]:
        """Convierte a Excel usando openpyxl"""
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
        
        wb = Workbook()
        ws = wb.active
        ws.title = sheet_name
        
        # Estilos
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF", size=12)
        border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        row = 1
        
        if isinstance(data, str):
            ws['A1'] = data
        elif isinstance(data, dict):
            # Encabezados
            ws['A1'] = 'Campo'
            ws['B1'] = 'Valor'
            ws['A1'].fill = header_fill
            ws['A1'].font = header_font
            ws['B1'].fill = header_fill
            ws['B1'].font = header_font
            ws['A1'].border = border
            ws['B1'].border = border
            
            # Datos
            row = 2
            for key, value in data.items():
                ws[f'A{row}'] = str(key)
                ws[f'B{row}'] = str(value)
                ws[f'A{row}'].border = border
                ws[f'B{row}'].border = border
                if row % 2 == 0:
                    ws[f'A{row}'].fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
                    ws[f'B{row}'].fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
                row += 1
        elif isinstance(data, list):
            if data and isinstance(data[0], dict):
                # Lista de diccionarios
                if data:
                    headers = list(data[0].keys())
                    # Encabezados
                    for col_idx, header in enumerate(headers, 1):
                        cell = ws.cell(row=1, column=col_idx, value=str(header))
                        cell.fill = header_fill
                        cell.font = header_font
                        cell.border = border
                    
                    # Datos
                    for row_idx, item in enumerate(data, 2):
                        for col_idx, header in enumerate(headers, 1):
                            cell = ws.cell(row=row_idx, column=col_idx, value=str(item.get(header, '')))
                            cell.border = border
                            if row_idx % 2 == 0:
                                cell.fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
            else:
                # Lista simple
                for idx, item in enumerate(data, 1):
                    ws[f'A{idx}'] = str(item)
        
        # Ajustar ancho de columnas
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
        
        wb.save(output_path)
        
        return {
            "success": True,
            "output_path": output_path,
            "method": "openpyxl",
            "file_size": os.path.getsize(output_path)
        }
    
    def _convert_to_excel_xlsxwriter(
        self,
        data: Union[Dict, List, str],
        output_path: str,
        sheet_name: str,
        **kwargs
    ) -> Dict[str, Any]:
        """Convierte a Excel usando xlsxwriter"""
        import xlsxwriter
        
        workbook = xlsxwriter.Workbook(output_path)
        worksheet = workbook.add_worksheet(sheet_name)
        
        # Formatos
        header_format = workbook.add_format({
            'bold': True,
            'bg_color': '#366092',
            'font_color': '#FFFFFF',
            'border': 1,
            'align': 'left',
            'valign': 'vcenter'
        })
        
        cell_format = workbook.add_format({
            'border': 1,
            'align': 'left',
            'valign': 'vcenter'
        })
        
        row = 0
        
        if isinstance(data, str):
            worksheet.write(row, 0, data)
        elif isinstance(data, dict):
            # Encabezados
            worksheet.write(row, 0, 'Campo', header_format)
            worksheet.write(row, 1, 'Valor', header_format)
            
            # Datos
            row = 1
            for key, value in data.items():
                worksheet.write(row, 0, str(key), cell_format)
                worksheet.write(row, 1, str(value), cell_format)
                row += 1
        elif isinstance(data, list):
            if data and isinstance(data[0], dict):
                if data:
                    headers = list(data[0].keys())
                    # Encabezados
                    for col_idx, header in enumerate(headers):
                        worksheet.write(row, col_idx, str(header), header_format)
                    
                    # Datos
                    row = 1
                    for item in data:
                        for col_idx, header in enumerate(headers):
                            worksheet.write(row, col_idx, str(item.get(header, '')), cell_format)
                        row += 1
            else:
                for idx, item in enumerate(data):
                    worksheet.write(idx, 0, str(item), cell_format)
        
        workbook.close()
        
        return {
            "success": True,
            "output_path": output_path,
            "method": "xlsxwriter",
            "file_size": os.path.getsize(output_path)
        }
    
    def _convert_to_excel_pandas(
        self,
        data: Union[Dict, List, str],
        output_path: str,
        sheet_name: str,
        **kwargs
    ) -> Dict[str, Any]:
        """Convierte a Excel usando pandas"""
        import pandas as pd
        
        if isinstance(data, str):
            df = pd.DataFrame({'Contenido': [data]})
        elif isinstance(data, dict):
            df = pd.DataFrame([data])
        elif isinstance(data, list):
            if data and isinstance(data[0], dict):
                df = pd.DataFrame(data)
            else:
                df = pd.DataFrame({'Item': data})
        else:
            df = pd.DataFrame({'Datos': [str(data)]})
        
        # Crear Excel con formato
        with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name=sheet_name, index=False)
            
            # Aplicar formato
            workbook = writer.book
            worksheet = writer.sheets[sheet_name]
            
            from openpyxl.styles import Font, PatternFill
            
            # Formato de encabezados
            header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            header_font = Font(bold=True, color="FFFFFF", size=12)
            
            for cell in worksheet[1]:
                cell.fill = header_fill
                cell.font = header_font
        
        return {
            "success": True,
            "output_path": output_path,
            "method": "pandas",
            "file_size": os.path.getsize(output_path)
        }
    
    def convert_multiple_formats(
        self,
        data: Union[Dict, List, str],
        base_output_path: str,
        formats: List[str] = None,
        title: str = "Documento",
        **kwargs
    ) -> Dict[str, Any]:
        """
        Convierte datos a múltiples formatos simultáneamente.
        
        Args:
            data: Datos a convertir
            base_output_path: Ruta base (sin extensión)
            formats: Lista de formatos ['pdf', 'docx', 'xlsx']
            title: Título del documento
            **kwargs: Opciones adicionales
        
        Returns:
            Dict con resultados de todas las conversiones
        """
        if formats is None:
            formats = ['pdf', 'docx', 'xlsx']
        
        results = {}
        base_path = Path(base_output_path)
        
        for fmt in formats:
            try:
                if fmt == 'pdf':
                    output_path = str(base_path.with_suffix('.pdf'))
                    results['pdf'] = self.convert_to_pdf(data, output_path, title=title, **kwargs)
                elif fmt == 'docx':
                    output_path = str(base_path.with_suffix('.docx'))
                    results['docx'] = self.convert_to_word(data, output_path, title=title, **kwargs)
                elif fmt == 'xlsx':
                    output_path = str(base_path.with_suffix('.xlsx'))
                    results['xlsx'] = self.convert_to_excel(data, output_path, **kwargs)
            except Exception as e:
                results[fmt] = {
                    "success": False,
                    "error": str(e)
                }
                self.logger.error(f"Error convirtiendo a {fmt}: {e}")
        
        return results
    
    def get_available_libraries(self) -> Dict[str, Dict[str, bool]]:
        """Retorna información sobre librerías disponibles"""
        return self.available_libs.copy()


# Instancia global
_converter_instance = None

def get_document_converter() -> DocumentConverter:
    """Obtiene la instancia global del conversor"""
    global _converter_instance
    if _converter_instance is None:
        _converter_instance = DocumentConverter()
    return _converter_instance

