#!/usr/bin/env python3
"""
Script para convertir el documento ultra completo a Word (.docx) y PDF
"""

import os
import subprocess
import sys
from pathlib import Path

def install_requirements():
    """Instala las dependencias necesarias"""
    try:
        import markdown
        import docx
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        print("✅ Todas las dependencias ya están instaladas")
    except ImportError:
        print("📦 Instalando dependencias necesarias...")
        subprocess.check_call([
            sys.executable, "-m", "pip", "install", 
            "markdown", "python-docx", "reportlab", "beautifulsoup4"
        ])
        print("✅ Dependencias instaladas correctamente")

def markdown_to_html(md_file):
    """Convierte Markdown a HTML"""
    import markdown
    from markdown.extensions import tables, toc, codehilite
    
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Configurar extensiones de Markdown
    md = markdown.Markdown(
        extensions=[
            'tables',
            'toc',
            'codehilite',
            'fenced_code',
            'attr_list'
        ],
        extension_configs={
            'toc': {
                'permalink': True,
                'permalink_title': "Enlace permanente a esta sección"
            }
        }
    )
    
    html_content = md.convert(md_content)
    
    # Crear HTML completo con estilos
    full_html = f"""
    <!DOCTYPE html>
    <html lang="es">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Frases Populares de Búsqueda en Google 2024 - Versión Ultra Completa</title>
        <style>
            body {{
                font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                line-height: 1.6;
                max-width: 1200px;
                margin: 0 auto;
                padding: 20px;
                color: #333;
                background-color: #f9f9f9;
            }}
            .container {{
                background: white;
                padding: 40px;
                border-radius: 10px;
                box-shadow: 0 0 20px rgba(0,0,0,0.1);
            }}
            h1 {{
                color: #2c3e50;
                border-bottom: 3px solid #3498db;
                padding-bottom: 10px;
                text-align: center;
                font-size: 2.5em;
                margin-bottom: 30px;
            }}
            h2 {{
                color: #34495e;
                border-left: 4px solid #3498db;
                padding-left: 15px;
                margin-top: 30px;
                margin-bottom: 15px;
                font-size: 1.8em;
            }}
            h3 {{
                color: #2c3e50;
                margin-top: 25px;
                margin-bottom: 10px;
                font-size: 1.4em;
            }}
            h4 {{
                color: #34495e;
                margin-top: 20px;
                margin-bottom: 8px;
                font-size: 1.2em;
            }}
            table {{
                width: 100%;
                border-collapse: collapse;
                margin: 20px 0;
                background: white;
                box-shadow: 0 2px 10px rgba(0,0,0,0.1);
                border-radius: 8px;
                overflow: hidden;
            }}
            th, td {{
                padding: 12px 15px;
                text-align: left;
                border-bottom: 1px solid #ddd;
            }}
            th {{
                background-color: #3498db;
                color: white;
                font-weight: bold;
                text-transform: uppercase;
                font-size: 0.9em;
                letter-spacing: 0.5px;
            }}
            tr:nth-child(even) {{
                background-color: #f8f9fa;
            }}
            tr:hover {{
                background-color: #e8f4f8;
            }}
            ul, ol {{
                margin: 15px 0;
                padding-left: 30px;
            }}
            li {{
                margin: 8px 0;
                line-height: 1.5;
            }}
            strong {{
                color: #2c3e50;
                font-weight: 600;
            }}
            code {{
                background-color: #f4f4f4;
                padding: 2px 6px;
                border-radius: 3px;
                font-family: 'Courier New', monospace;
                color: #e74c3c;
            }}
            pre {{
                background-color: #f8f9fa;
                padding: 15px;
                border-radius: 5px;
                border-left: 4px solid #3498db;
                overflow-x: auto;
                margin: 15px 0;
            }}
            blockquote {{
                border-left: 4px solid #3498db;
                margin: 20px 0;
                padding: 10px 20px;
                background-color: #f8f9fa;
                font-style: italic;
            }}
            .toc {{
                background-color: #f8f9fa;
                padding: 20px;
                border-radius: 8px;
                margin: 20px 0;
                border: 1px solid #e9ecef;
            }}
            .toc h2 {{
                margin-top: 0;
                color: #2c3e50;
            }}
            .toc ul {{
                list-style-type: none;
                padding-left: 0;
            }}
            .toc li {{
                margin: 5px 0;
            }}
            .toc a {{
                text-decoration: none;
                color: #3498db;
                font-weight: 500;
            }}
            .toc a:hover {{
                text-decoration: underline;
            }}
            .highlight {{
                background-color: #fff3cd;
                padding: 15px;
                border-radius: 5px;
                border-left: 4px solid #ffc107;
                margin: 15px 0;
            }}
            .emoji {{
                font-size: 1.2em;
                margin-right: 8px;
            }}
            @media print {{
                body {{
                    background: white;
                    padding: 0;
                }}
                .container {{
                    box-shadow: none;
                    padding: 0;
                }}
                h1, h2, h3, h4 {{
                    page-break-after: avoid;
                }}
                table {{
                    page-break-inside: avoid;
                }}
            }}
        </style>
    </head>
    <body>
        <div class="container">
            {html_content}
        </div>
    </body>
    </html>
    """
    
    return full_html

def html_to_pdf_reportlab(html_content, output_file):
    """Convierte HTML a PDF usando ReportLab"""
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from bs4 import BeautifulSoup
    import re
    
    # Crear documento PDF
    doc = SimpleDocTemplate(output_file, pagesize=A4, 
                          rightMargin=72, leftMargin=72, 
                          topMargin=72, bottomMargin=18)
    
    # Obtener estilos
    styles = getSampleStyleSheet()
    
    # Crear estilos personalizados
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
        alignment=1,  # Centrado
        textColor=colors.HexColor('#2c3e50')
    )
    
    heading2_style = ParagraphStyle(
        'CustomHeading2',
        parent=styles['Heading2'],
        fontSize=18,
        spaceBefore=20,
        spaceAfter=12,
        textColor=colors.HexColor('#34495e')
    )
    
    heading3_style = ParagraphStyle(
        'CustomHeading3',
        parent=styles['Heading3'],
        fontSize=14,
        spaceBefore=15,
        spaceAfter=8,
        textColor=colors.HexColor('#2c3e50')
    )
    
    heading4_style = ParagraphStyle(
        'CustomHeading4',
        parent=styles['Heading4'],
        fontSize=12,
        spaceBefore=10,
        spaceAfter=6,
        textColor=colors.HexColor('#34495e')
    )
    
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=6
    )
    
    # Parsear HTML
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Crear lista de elementos para el PDF
    story = []
    
    # Procesar elementos
    for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'ul', 'ol', 'table']):
        if element.name == 'h1':
            text = element.get_text().strip()
            if text:
                story.append(Paragraph(text, title_style))
                story.append(Spacer(1, 12))
        
        elif element.name == 'h2':
            text = element.get_text().strip()
            if text:
                story.append(Paragraph(text, heading2_style))
                story.append(Spacer(1, 6))
        
        elif element.name == 'h3':
            text = element.get_text().strip()
            if text:
                story.append(Paragraph(text, heading3_style))
                story.append(Spacer(1, 4))
        
        elif element.name == 'h4':
            text = element.get_text().strip()
            if text:
                story.append(Paragraph(text, heading4_style))
                story.append(Spacer(1, 3))
        
        elif element.name == 'p':
            text = element.get_text().strip()
            if text and not text.startswith('---'):
                # Procesar texto en negrita
                text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
                story.append(Paragraph(text, normal_style))
                story.append(Spacer(1, 6))
        
        elif element.name in ['ul', 'ol']:
            for li in element.find_all('li'):
                text = li.get_text().strip()
                if text:
                    # Procesar texto en negrita
                    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
                    if element.name == 'ul':
                        text = f"• {text}"
                    story.append(Paragraph(text, normal_style))
                    story.append(Spacer(1, 3))
            story.append(Spacer(1, 6))
        
        elif element.name == 'table':
            # Crear tabla
            rows = []
            for tr in element.find_all('tr'):
                row = []
                for td in tr.find_all(['td', 'th']):
                    cell_text = td.get_text().strip()
                    # Procesar texto en negrita
                    cell_text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', cell_text)
                    row.append(cell_text)
                if row:
                    rows.append(row)
            
            if rows:
                # Crear tabla con ReportLab
                table = Table(rows)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3498db')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('FONTSIZE', (0, 1), (-1, -1), 9),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))
                story.append(table)
                story.append(Spacer(1, 12))
    
    # Construir PDF
    try:
        doc.build(story)
        print(f"✅ PDF creado exitosamente: {output_file}")
    except Exception as e:
        print(f"❌ Error al crear PDF: {e}")

def markdown_to_docx_improved(md_file, output_file):
    """Convierte Markdown a Word (.docx) - Versión mejorada y robusta"""
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    import re
    
    # Leer el archivo Markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Crear documento Word
    doc = Document()
    
    # Configurar estilos
    styles = doc.styles
    
    # Estilo para títulos principales
    if 'Título 1' not in [style.name for style in styles]:
        title1_style = styles.add_style('Título 1', WD_STYLE_TYPE.PARAGRAPH)
        title1_font = title1_style.font
        title1_font.name = 'Calibri'
        title1_font.size = Pt(24)
        title1_font.bold = True
        title1_font.color.rgb = None
    
    # Estilo para títulos secundarios
    if 'Título 2' not in [style.name for style in styles]:
        title2_style = styles.add_style('Título 2', WD_STYLE_TYPE.PARAGRAPH)
        title2_font = title2_style.font
        title2_font.name = 'Calibri'
        title2_font.size = Pt(18)
        title2_font.bold = True
        title2_font.color.rgb = None
    
    # Estilo para títulos terciarios
    if 'Título 3' not in [style.name for style in styles]:
        title3_style = styles.add_style('Título 3', WD_STYLE_TYPE.PARAGRAPH)
        title3_font = title3_style.font
        title3_font.name = 'Calibri'
        title3_font.size = Pt(14)
        title3_font.bold = True
        title3_font.color.rgb = None
    
    # Procesar el contenido línea por línea
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Verificar que la línea no esté vacía
        if not line:
            i += 1
            continue
        
        try:
            if line.startswith('# '):
                # Título principal
                title = line[2:].strip()
                p = doc.add_heading(title, level=1)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif line.startswith('## '):
                # Título secundario
                title = line[3:].strip()
                doc.add_heading(title, level=2)
            elif line.startswith('### '):
                # Título terciario
                title = line[4:].strip()
                doc.add_heading(title, level=3)
            elif line.startswith('#### '):
                # Título cuaternario
                title = line[5:].strip()
                doc.add_heading(title, level=4)
            elif line.startswith('|') and '|' in line:
                # Tabla
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i].strip())
                    i += 1
                i -= 1  # Retroceder una línea
                
                if len(table_lines) > 1:
                    # Crear tabla
                    rows = []
                    for table_line in table_lines:
                        if '---' not in table_line:  # Saltar líneas de separación
                            cells = [cell.strip() for cell in table_line.split('|')[1:-1]]
                            rows.append(cells)
                    
                    if rows:
                        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                        table.style = 'Table Grid'
                        
                        for row_idx, row_data in enumerate(rows):
                            for col_idx, cell_data in enumerate(row_data):
                                if col_idx < len(table.rows[row_idx].cells):
                                    table.rows[row_idx].cells[col_idx].text = cell_data
                                    # Hacer la primera fila en negrita
                                    if row_idx == 0:
                                        for paragraph in table.rows[row_idx].cells[col_idx].paragraphs:
                                            for run in paragraph.runs:
                                                run.bold = True
            elif line.startswith('- ') or line.startswith('* '):
                # Lista con viñetas
                items = []
                while i < len(lines) and (lines[i].strip().startswith('- ') or lines[i].strip().startswith('* ')):
                    item = lines[i].strip()[2:].strip()
                    # Procesar texto en negrita
                    item = re.sub(r'\*\*(.*?)\*\*', r'\1', item)
                    items.append(item)
                    i += 1
                i -= 1  # Retroceder una línea
                
                for item in items:
                    p = doc.add_paragraph(item, style='List Bullet')
            elif len(line) > 0 and line[0].isdigit() and '. ' in line:
                # Lista numerada
                items = []
                while i < len(lines) and len(lines[i].strip()) > 0 and lines[i].strip()[0].isdigit() and '. ' in lines[i].strip():
                    item = lines[i].strip()
                    # Extraer el número y el texto
                    match = re.match(r'^\d+\.\s*(.*)', item)
                    if match:
                        items.append(match.group(1))
                    i += 1
                i -= 1  # Retroceder una línea
                
                for item in items:
                    p = doc.add_paragraph(item, style='List Number')
            elif line.startswith('**') and line.endswith('**'):
                # Texto en negrita
                text = line[2:-2]
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
            elif line and not line.startswith('---'):
                # Párrafo normal
                # Procesar texto en negrita
                text = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
                p = doc.add_paragraph(text)
        except Exception as e:
            print(f"⚠️ Error procesando línea {i+1}: {e}")
            # Continuar con la siguiente línea
            pass
        
        i += 1
    
    # Guardar documento
    try:
        doc.save(output_file)
        print(f"✅ Documento Word creado exitosamente: {output_file}")
    except Exception as e:
        print(f"❌ Error al guardar documento Word: {e}")

def main():
    """Función principal"""
    print("🚀 Iniciando conversión del documento ultra completo...")
    
    # Instalar dependencias
    install_requirements()
    
    # Archivos de entrada y salida
    md_file = "/Users/adan/frontier/frases-populares-busqueda-google-2024-ultra-completo.md"
    html_file = "/Users/adan/frontier/frases-populares-busqueda-google-2024-ultra-completo.html"
    docx_file = "/Users/adan/frontier/frases-populares-busqueda-google-2024-ultra-completo.docx"
    pdf_file = "/Users/adan/frontier/frases-populares-busqueda-google-2024-ultra-completo.pdf"
    
    # Verificar que el archivo Markdown existe
    if not os.path.exists(md_file):
        print(f"❌ Error: No se encontró el archivo {md_file}")
        return
    
    try:
        # Convertir a HTML
        print("📄 Convirtiendo Markdown a HTML...")
        html_content = markdown_to_html(md_file)
        
        with open(html_file, 'w', encoding='utf-8') as f:
            f.write(html_content)
        print(f"✅ HTML creado: {html_file}")
        
        # Convertir a PDF usando ReportLab
        print("📄 Convirtiendo HTML a PDF usando ReportLab...")
        html_to_pdf_reportlab(html_content, pdf_file)
        
        # Convertir a Word
        print("📄 Convirtiendo Markdown a Word...")
        markdown_to_docx_improved(md_file, docx_file)
        
        print("\n🎉 ¡Conversión completada exitosamente!")
        print(f"📁 Archivos generados:")
        print(f"   • HTML: {html_file}")
        print(f"   • PDF: {pdf_file}")
        print(f"   • Word: {docx_file}")
        
        # Verificar que los archivos se crearon
        files_created = []
        if os.path.exists(html_file):
            files_created.append(f"✅ HTML: {os.path.getsize(html_file)} bytes")
        if os.path.exists(pdf_file):
            files_created.append(f"✅ PDF: {os.path.getsize(pdf_file)} bytes")
        if os.path.exists(docx_file):
            files_created.append(f"✅ Word: {os.path.getsize(docx_file)} bytes")
        
        print("\n📊 Resumen de archivos creados:")
        for file_info in files_created:
            print(f"   {file_info}")
        
    except Exception as e:
        print(f"❌ Error durante la conversión: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()








