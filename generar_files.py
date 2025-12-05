
from fpdf import FPDF
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- Contenido del Documento ---
titulo = "Propuesta de Optimización: Capa de Enlace de Datos"
autor = "Adan (AD)"
curso = "Redes de Computadoras"
entregable = "U2_E2_AD"

texto_intro = (
    "A continuación se presenta la propuesta de solución para los problemas de colisión, "
    "congestión y errores de transmisión detectados en la red local (LAN) de la empresa."
)

seccion_1_titulo = "1. Identificación de Causas y Diagnóstico"
seccion_1_cuerpo = (
    "Diagnóstico de la Situación Actual:\n"
    "Los problemas reportados (colisiones constantes, congestión y retransmisiones) indican deficiencias "
    "estructurales en la implementación de las capas físicas y de enlace de datos.\n\n"
    "Causas Identificadas:\n"
    "1. Colisiones y Dominio de Colisión: La red opera bajo un dominio de colisión excesivamente amplio. "
    "Esto suele ocurrir por el uso de concentradores (Hubs) o topologías de bus compartidas donde todos los dispositivos "
    "compiten por el medio al mismo tiempo. Al transmitir simultáneamente, las señales chocan y se corrompen.\n"
    "2. Congestión y Dominio de Transmisión (Broadcast): La red carece de segmentación lógica. El tráfico de difusión (broadcast) "
    "inunda todos los puertos, saturando el ancho de banda y los recursos de procesamiento de los dispositivos.\n\n"
    "Solución Conceptual:\n"
    "- Reducir el dominio de colisión mediante el uso de Switches (micro-segmentación).\n"
    "- Reducir el dominio de transmisión mediante la implementación de VLANs o Routers."
)

seccion_2_titulo = "2. Plan de Implementación: Control de Flujo y Errores"
seccion_2_cuerpo = (
    "Para optimizar la eficiencia en la Capa de Enlace de Datos, se propone la siguiente estrategia algorítmica:\n\n"
    "A) Método de Control de Flujo: Ventana Deslizante (Sliding Window)\n"
    "Se reemplazará cualquier mecanismo de 'Parar y Esperar' por el protocolo de Ventana Deslizante.\n"
    "- Funcionamiento: Permite al emisor enviar múltiples tramas (hasta un tamaño de ventana N) sin esperar confirmación inmediata de cada una.\n"
    "- Beneficio: Maximiza el uso del canal ('pipelining'), evitando que el medio físico permanezca ocioso esperando ACKs, lo cual es crítico para reducir la latencia percibida.\n\n"
    "B) Mecanismo de Control de Errores: CRC + ARQ (Repetición Selectiva)\n"
    "- Detección (CRC): Se implementará una Comprobación de Redundancia Cíclica (CRC-32) en el tráiler de cada trama para detectar corrupción de bits con alta precisión.\n"
    "- Corrección (Repetición Selectiva): Se configurará el protocolo ARQ para solicitar la retransmisión ÚNICAMENTE de las tramas que lleguen corruptas o se pierdan, en lugar de reenviar toda la secuencia (como haría Go-Back-N)."
)

seccion_3_titulo = "3. Estrategias de Mejora en la Eficiencia"
seccion_3_cuerpo = (
    "La implementación conjunta de estas tecnologías generará las siguientes mejoras:\n\n"
    "1. Eliminación de Colisiones: Al segmentar la red, cada dispositivo obtiene su propio ancho de banda dedicado (Full-Duplex), eliminando la contienda por el medio.\n"
    "2. Optimización del 'Goodput': La Repetición Selectiva asegura que el ancho de banda solo se consuma en datos nuevos o correcciones necesarias, no en retransmisiones redundantes.\n"
    "3. Integridad de Datos: El uso estricto de CRC garantiza que la información procesada por las capas superiores esté libre de errores, mejorando la confiabilidad general del sistema."
)

# --- Generación de PDF ---
def generar_pdf():
    pdf = FPDF()
    pdf.add_page()
    
    # Título
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(0, 10, titulo, 0, 1, 'C')
    pdf.set_font("Arial", 'I', 10)
    pdf.cell(0, 10, f"Entregable: {entregable} | Autor: {autor}", 0, 1, 'C')
    pdf.ln(5)
    
    # Cuerpo
    pdf.set_font("Arial", '', 11)
    pdf.multi_cell(0, 6, texto_intro)
    pdf.ln(5)
    
    # Secciones
    for tit, cuerpo in [(seccion_1_titulo, seccion_1_cuerpo), 
                        (seccion_2_titulo, seccion_2_cuerpo), 
                        (seccion_3_titulo, seccion_3_cuerpo)]:
        pdf.set_font("Arial", 'B', 12)
        pdf.set_fill_color(230, 230, 230)
        pdf.cell(0, 8, tit, 0, 1, 'L', 1)
        pdf.ln(2)
        
        pdf.set_font("Arial", '', 11)
        pdf.multi_cell(0, 6, cuerpo)
        pdf.ln(5)
        
    filename = f"{entregable}.pdf"
    pdf.output(filename)
    print(f"PDF generado: {filename}")

# --- Generación de Word ---
def generar_word():
    doc = Document()
    
    # Título
    t = doc.add_heading(titulo, 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Entregable: {entregable} | Autor: {autor}")
    r.italic = True
    
    doc.add_paragraph(texto_intro)
    
    # Secciones
    for tit, cuerpo in [(seccion_1_titulo, seccion_1_cuerpo), 
                        (seccion_2_titulo, seccion_2_cuerpo), 
                        (seccion_3_titulo, seccion_3_cuerpo)]:
        doc.add_heading(tit, level=1)
        doc.add_paragraph(cuerpo)
        
    filename = f"{entregable}.docx"
    doc.save(filename)
    print(f"Word generado: {filename}")

if __name__ == "__main__":
    try:
        generar_pdf()
        generar_word()
    except Exception as e:
        print(f"Error: {e}")







